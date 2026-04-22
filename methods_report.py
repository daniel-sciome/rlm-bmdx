"""
methods_report.py — Structured Materials & Methods section for NIEHS 5dTox reports.

Replicates the exact structure from NIEHS Report 10 (PFHxSAm study):
  Materials and Methods
  ├── Study Design
  ├── Dose Selection Rationale
  ├── Chemistry
  ├── Clinical Examinations and Sample Collection
  │   ├── Clinical Observations
  │   ├── Body and Organ Weights
  │   ├── Clinical Pathology
  │   └── Internal Dose Assessment       (conditional: tissue_conc in pool)
  ├── Transcriptomics                     (conditional: gene_expression in pool)
  │   ├── Sample Collection for Transcriptomics
  │   ├── RNA Isolation, Library Creation, and Sequencing
  │   ├── Sequence Data Processing
  │   ├── Sequencing Quality Checks and Outlier Removal
  │   └── Data Normalization
  ├── Data Analysis
  │   ├── Statistical Analysis of Body Weights, Organ Weights, and Clinical Pathology
  │   ├── Benchmark Dose Analysis of Body Weights, Organ Weights, and Clinical Pathology
  │   ├── Benchmark Dose Analysis of Transcriptomics Data
  │   ├── Empirical False Discovery Rate Determination for Genomic Dose-response Modeling
  │   └── Data Accessibility
  └── [Table 1: Final Sample Counts for BMD Analysis of Transcriptomics Data]

Approach: Hybrid data + LLM.
  - Programmatically extract study metadata (doses, sample counts, domains,
    BMDExpress analysis parameters) from fingerprints, animal_report, and .bm2
    analysisInfo.notes.
  - Feed the structured context to an LLM prompt that generates prose for
    each subsection.
  - Subsections are CONDITIONAL — only included when the file pool has the
    relevant data domain.

This module is imported by background_server.py for the /api/generate-methods
endpoint and the /api/export-docx DOCX builder.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field

# base_domain is no longer needed — platform strings are used directly.
# Kept as a no-op import guard in case downstream code still references it.


logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Constants — heading hierarchy and subsection keys
# ---------------------------------------------------------------------------

# The canonical subsection ordering and heading levels.
# level 2 = H2 ("Materials and Methods" — added by caller)
# level 3 = H3 (major subsections: Study Design, Chemistry, etc.)
# level 4 = H4 (sub-subsections: Clinical Observations, Body Weights, etc.)
#
# Each tuple: (subsection_key, heading_text, heading_level, condition_field)
# condition_field is the MethodsContext bool field that must be True for the
# subsection to be included.  None means always included.
SUBSECTION_SKELETON: list[tuple[str, str, int, str | None]] = [
    ("study_design",            "Study Design",                                                  3, None),
    ("dose_selection",          "Dose Selection Rationale",                                      3, None),
    ("chemistry",               "Chemistry",                                                     3, None),
    ("clinical_exams",          "Clinical Examinations and Sample Collection",                   3, None),
    ("clinical_obs",            "Clinical Observations",                                         4, None),
    ("body_organ_weights",      "Body and Organ Weights",                                        4, "has_body_weight"),
    ("clinical_pathology",      "Clinical Pathology",                                            4, "has_clin_path"),
    ("internal_dose",           "Internal Dose Assessment",                                      4, "has_tissue_conc"),
    ("transcriptomics",         "Transcriptomics",                                               3, "has_gene_expression"),
    ("txomics_sample",          "Sample Collection for Transcriptomics",                         4, "has_gene_expression"),
    ("txomics_rna",             "RNA Isolation, Library Creation, and Sequencing",               4, "has_gene_expression"),
    ("txomics_seq_processing",  "Sequence Data Processing",                                      4, "has_gene_expression"),
    ("txomics_qc",              "Sequencing Quality Checks and Outlier Removal",                 4, "has_gene_expression"),
    ("txomics_normalization",   "Data Normalization",                                            4, "has_gene_expression"),
    ("data_analysis",           "Data Analysis",                                                 3, None),
    ("stat_analysis",           "Statistical Analysis of Body Weights, Organ Weights, and Clinical Pathology",  4, None),
    ("bmd_apical",              "Benchmark Dose Analysis of Body Weights, Organ Weights, and Clinical Pathology", 4, None),
    ("bmd_genomics",            "Benchmark Dose Analysis of Transcriptomics Data",               4, "has_gene_expression"),
    ("efdr",                    "Empirical False Discovery Rate Determination for Genomic Dose-response Modeling", 4, "has_gene_expression"),
    ("data_accessibility",      "Data Accessibility",                                            4, None),
]


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class MethodsContext:
    """
    All study metadata extracted from the file pool.

    Used to:
      1. Inform the LLM prompt with actual study parameters.
      2. Decide which conditional subsections to include.
      3. Build Table 1 (genomics sample counts) programmatically.
    """
    # --- Chemical identity ---
    chemical_name: str = ""
    casrn: str = ""
    dtxsid: str = ""

    # --- Study design (from fingerprints + animal_report) ---
    # species includes strain when available, e.g. "Hsd:Sprague Dawley® SD®"
    species: str = "Sprague Dawley"
    dose_groups: list[float] = field(default_factory=list)
    dose_unit: str = "mg/kg"
    n_per_group: int = 5
    n_control: int = 10
    # How many animals per sex per dose for internal dose assessment (biosampling)
    n_biosampling: int = 0
    sexes: list[str] = field(default_factory=list)
    vehicle: str = "corn oil"
    route: str = "gavage"
    duration_days: int = 5

    # --- Domain presence flags (drive conditional subsections) ---
    has_body_weight: bool = False
    has_organ_weights: bool = False
    has_clin_chem: bool = False
    has_hematology: bool = False
    has_hormones: bool = False
    has_tissue_conc: bool = False
    has_gene_expression: bool = False

    @property
    def has_clin_path(self) -> bool:
        """True if any clinical pathology domain exists (clin_chem, hematology, or hormones)."""
        return self.has_clin_chem or self.has_hematology or self.has_hormones

    # --- Endpoint names by domain (from fingerprints) ---
    organ_weight_endpoints: list[str] = field(default_factory=list)
    clin_chem_endpoints: list[str] = field(default_factory=list)
    hematology_endpoints: list[str] = field(default_factory=list)
    hormone_endpoints: list[str] = field(default_factory=list)
    # Organs with gene expression data (e.g. ["Liver", "Kidney"])
    ge_organs: list[str] = field(default_factory=list)

    # --- Biosampling / pharmacokinetic context (for Abstract-Methods) ---
    # Dose groups that had biosampling animals dedicated to internal dose
    # assessment (blood/plasma collection).  Extracted from sidecar files
    # (tissue_conc or body_weight) by scanning for rows with selection
    # containing "biosampling".  Reference report writes e.g.:
    #   "Blood was collected from animals dedicated for internal dose
    #    assessment in the 4 and 37 mg/kg groups."
    biosampling_doses: list[float] = field(default_factory=list)

    # --- Pharmacokinetics (for Abstract-Results PK sentence) ---
    # Aggregated plasma concentration means per sex × dose × timepoint,
    # plus calculated half-lives per sex × dose.  Built from tissue
    # concentration sidecars (the only domain that uses biosampling
    # animals).  Half-lives use the standard two-point formula:
    #   t½ = ln(2) × Δt / ln(C₁/C₂)
    # where C₁ is the early timepoint concentration and C₂ the later one.
    # Reference report writes:
    #   "Average PFHxSAm plasma concentrations at 2 and 24 hours postdose
    #    were lower in male rats than in female rats. Half-lives ... were
    #    78.2 and 25.6 hours for the 4 and 37 mg/kg groups, respectively..."
    #
    # Schema:
    #   pk_concentrations: {sex: {dose: {hour: mean_value}}}
    #   pk_half_lives:     {sex: {dose: hours_float}}
    #   pk_timepoints:     sorted list of timepoints in hours (e.g., [2, 24])
    pk_concentrations: dict | None = None
    pk_half_lives: dict | None = None
    pk_timepoints: list[int] = field(default_factory=list)

    # --- Genomics assay identification (for Abstract-Methods) ---
    # Human-readable assay name (e.g., "TempO-Seq", "Affymetrix", "RNA-seq")
    # and the chip/probe-set name (e.g., "S1500+").  Extracted from the
    # integrated BMDProject's gene-expression experiments via chip.name
    # and chip.chipId — S1500 in the chip name implies TempO-Seq.
    # Reference report writes e.g.:
    #   "...assayed in gene expression studies using the TempO-Seq assay."
    genomics_assay: str | None = None
    genomics_chip: str | None = None

    # --- BMDExpress / BMDS metadata (from .bm2 analysisInfo.notes) ---
    bmdexpress_version: str | None = None
    bmds_version: str | None = None
    bmr_type: str | None = None
    bmr_factor: float | None = None
    models_fit: list[str] | None = None
    constant_variance: bool | None = None
    # Pre-filter method for transcriptomics (Williams or CurveFit)
    prefilter_method: str | None = None
    prefilter_pvalue: float | None = None
    fold_change_filter: float | None = None

    # --- Table 1: sample counts for genomics BMD analysis ---
    # Structure: {organ: {sex: {dose_float: count}}}
    # Built from animal_report or gene_expression fingerprints
    genomics_sample_counts: dict | None = None

    def to_dict(self) -> dict:
        """Serialize for JSON persistence.  Converts dataclass to a plain dict."""
        d = {}
        for f in self.__dataclass_fields__:
            d[f] = getattr(self, f)
        # Include the computed property too
        d["has_clin_path"] = self.has_clin_path
        return d

    @classmethod
    def from_dict(cls, d: dict) -> MethodsContext:
        """Reconstruct from a JSON-serialized dict."""
        # Filter to only fields that exist on the dataclass
        valid_fields = set(cls.__dataclass_fields__)
        filtered = {k: v for k, v in d.items() if k in valid_fields}
        return cls(**filtered)


@dataclass
class MethodsSection:
    """
    A single M&M subsection with heading and content.

    heading:    The subsection heading text (e.g. "Study Design").
    level:      Heading depth — 3 = H3, 4 = H4 in the DOCX.
    key:        Subsection key matching SUBSECTION_SKELETON (e.g. "study_design").
    paragraphs: Prose paragraphs (user-editable in the frontend).
    table:      Optional table data for programmatic tables like Table 1.
                Format: {"caption": str, "headers": [...], "rows": [[...]], "footnotes": [...]}
    """
    heading: str
    level: int
    key: str
    paragraphs: list[str] = field(default_factory=list)
    table: dict | None = None

    def to_dict(self) -> dict:
        return {
            "heading": self.heading,
            "level": self.level,
            "key": self.key,
            "paragraphs": self.paragraphs,
            "table": self.table,
        }

    @classmethod
    def from_dict(cls, d: dict) -> MethodsSection:
        return cls(
            heading=d["heading"],
            level=d["level"],
            key=d["key"],
            paragraphs=d.get("paragraphs", []),
            table=d.get("table"),
        )


@dataclass
class MethodsReport:
    """
    Complete structured M&M output.

    sections:  Ordered list of MethodsSection objects (heading hierarchy preserved).
    context:   The MethodsContext used to generate this report — retained so the
               DOCX builder can access study parameters for Table 1 generation.
    """
    sections: list[MethodsSection] = field(default_factory=list)
    context: MethodsContext = field(default_factory=MethodsContext)

    def to_dict(self) -> dict:
        return {
            "sections": [s.to_dict() for s in self.sections],
            "context": self.context.to_dict(),
        }

    @classmethod
    def from_dict(cls, d: dict) -> MethodsReport:
        return cls(
            sections=[MethodsSection.from_dict(s) for s in d.get("sections", [])],
            context=MethodsContext.from_dict(d.get("context", {})),
        )


# ---------------------------------------------------------------------------
# Extraction: parse bm2 analysisInfo.notes into structured metadata
# ---------------------------------------------------------------------------

def _parse_bm2_analysis_info(notes_list: list[str]) -> dict:
    """
    Parse the analysisInfo.notes list from a .bm2 file's bMDResult,
    williamsTrendResults, curveFitPrefilterResults, or categoryAnalysisResults.

    BMDExpress 3 stores analysis parameters as a flat list of "Key: Value"
    strings in each analysis node's analysisInfo.notes.  This function
    extracts the subset we need for the M&M report.

    Args:
        notes_list: List of strings like ["BMDExpress3 Version: BMDExpress 3.20.0156 BETA",
                    "Models fit: hill, power, exponential 3, exponential 5", ...]

    Returns:
        Dict with parsed fields.  Missing fields are absent (not None).
        Possible keys: bmdexpress_version, bmds_version, bmr_type, bmr_factor,
        models_fit, constant_variance, prefilter_method, prefilter_pvalue,
        fold_change_filter.
    """
    result = {}
    for note in notes_list:
        # Most notes follow "Key: Value" format
        if ": " not in note:
            # Some notes are just labels like "Williams Trend Test" or "Benchmark Dose Analyses"
            # Use these to identify the prefilter method
            lower = note.strip().lower()
            if "williams" in lower:
                result["prefilter_method"] = "Williams Trend Test"
            elif "curve fit" in lower:
                result["prefilter_method"] = "Curve Fit Prefilter"
            continue

        key, _, value = note.partition(": ")
        key = key.strip()
        value = value.strip()

        if key == "BMDExpress3 Version":
            result["bmdexpress_version"] = value
        elif key == "BMDS Major Version":
            result["bmds_version"] = value
        elif key == "BMR Type":
            result["bmr_type"] = value
        elif key == "BMR Factor":
            try:
                result["bmr_factor"] = float(value)
            except ValueError:
                result["bmr_factor_str"] = value
        elif key == "Models fit" or key == "Models Used":
            # "hill, power, exponential 3, exponential 5"
            # or "hill: Hill EPA BMDS MLE ToxicR,power: Power EPA BMDS MLE ToxicR,..."
            # Normalize to clean model names
            models = []
            for m in value.split(","):
                # Strip the "Hill EPA BMDS MLE ToxicR" suffix if present
                name = m.split(":")[0].strip()
                if name:
                    models.append(name)
            result["models_fit"] = models
        elif key == "Constant Variance":
            result["constant_variance"] = value in ("1", "true", "True")
        elif key == "Unadjusted P-Value Cutoff":
            try:
                result["prefilter_pvalue"] = float(value)
            except ValueError:
                pass
        elif key == "NOTEL/LOTEL Fold Change Threshold":
            try:
                result["fold_change_filter"] = float(value)
            except ValueError:
                pass

    return result


def _collect_bm2_analysis_metadata(bm2_json: dict) -> dict:
    """
    Collect analysis metadata from ALL analysis nodes in a .bm2 file.

    Merges notes from williamsTrendResults (prefilter), curveFitPrefilterResults,
    bMDResult (BMD modeling params), and categoryAnalysisResults.
    Later entries overwrite earlier ones for the same key, so the BMD result
    (most specific) takes priority.

    Args:
        bm2_json: The full deserialized BMDProject dict from Java export / LMDB cache.

    Returns:
        Merged dict of analysis parameters (same keys as _parse_bm2_analysis_info).
    """
    merged = {}

    # Parse in order of specificity: prefilter < BMD < category
    # Each overwrites the previous for shared keys like bmdexpress_version
    for section_key in ("williamsTrendResults", "curveFitPrefilterResults", "bMDResult", "categoryAnalysisResults"):
        items = bm2_json.get(section_key, [])
        if not items:
            continue
        # Only parse the first item — all experiments in a section share params
        first = items[0] if isinstance(items, list) else items
        if not isinstance(first, dict):
            continue
        notes = first.get("analysisInfo", {}).get("notes", [])
        if notes:
            parsed = _parse_bm2_analysis_info(notes)
            merged.update(parsed)

    return merged


# ---------------------------------------------------------------------------
# Extraction: build MethodsContext from file pool data
# ---------------------------------------------------------------------------

def extract_methods_context(
    identity: dict,
    fingerprints: dict,
    animal_report: dict | None = None,
    study_params: dict | None = None,
    bm2_jsons: dict | None = None,
    session_dir: str | None = None,
    integrated: dict | None = None,
) -> MethodsContext:
    """
    Build a MethodsContext from all available data sources.

    This is the main entry point for extracting study metadata that drives
    both the LLM prompt and the conditional subsection logic.

    Args:
        identity:      Chemical identity dict from the frontend (name, casrn, dtxsid, ...).
        fingerprints:  Dict of {file_id: FileFingerprint-as-dict-or-object} from the server's
                       _pool_fingerprints[dtxsid] cache.  Each fingerprint has domain, sexes,
                       dose_groups, endpoint_names, organ, etc.
        animal_report: Optional dict from animal_report.json (dose_design, domain_coverage, etc.).
        study_params:  Optional user-provided overrides: vehicle, route, duration_days, species.
        bm2_jsons:     Optional dict of {file_id: bm2_json_dict} for extracting BMDExpress
                       analysis metadata from analysisInfo.notes.
        session_dir:   Optional path to the session directory.  Used to scan sidecar files
                       for biosampling dose groups (animals with selection="biosampling").
        integrated:    Optional integrated BMDProject dict.  Used to extract the genomics
                       assay/chip from doseResponseExperiments[].chip — e.g., chip name
                       containing "S1500" identifies TempO-Seq.

    Returns:
        Populated MethodsContext with all available study metadata.
    """
    ctx = MethodsContext()
    study_params = study_params or {}
    bm2_jsons = bm2_jsons or {}

    # --- Chemical identity ---
    ctx.chemical_name = identity.get("name", "the test chemical")
    ctx.casrn = identity.get("casrn", "")
    ctx.dtxsid = identity.get("dtxsid", "")

    # --- Study params (user-provided overrides) ---
    # These are NIEHS 5-day protocol defaults.  The dose_design from the
    # animal report contains TOTAL animals per group (core + biosampling),
    # which inflates the counts.  The actual core group sizes (5/10) are
    # protocol constants, so we don't override from dose_design.
    ctx.vehicle = study_params.get("vehicle", "corn oil")
    ctx.route = study_params.get("route", "gavage")
    ctx.duration_days = study_params.get("duration_days", 5)
    ctx.species = study_params.get("species", "Sprague Dawley")
    ctx.n_per_group = study_params.get("n_per_group", 5)
    ctx.n_control = study_params.get("n_control", 10)

    # --- Scan fingerprints for platform presence and collect metadata ---
    all_doses: set[float] = set()
    all_sexes: set[str] = set()
    dose_unit_found = None

    for fid, fp in fingerprints.items():
        # Support both dict and object-style access
        _get = fp.get if isinstance(fp, dict) else lambda k, d=None: getattr(fp, k, d)

        # Use platform directly — no suffix stripping needed.
        # data_type "gene_expression" is checked separately since
        # gene expression files have platform=None.
        platform = _get("platform")
        data_type = _get("data_type")

        if not platform and data_type != "gene_expression":
            continue

        # Set platform presence flags using human-readable platform strings.
        if platform == "Body Weight":
            ctx.has_body_weight = True
        elif platform == "Organ Weights":
            ctx.has_organ_weights = True
            eps = _get("endpoint_names", [])
            ctx.organ_weight_endpoints = list(set(ctx.organ_weight_endpoints + eps))
        elif platform == "Clinical Chemistry":
            ctx.has_clin_chem = True
            eps = _get("endpoint_names", [])
            ctx.clin_chem_endpoints = list(set(ctx.clin_chem_endpoints + eps))
        elif platform == "Hematology":
            ctx.has_hematology = True
            eps = _get("endpoint_names", [])
            ctx.hematology_endpoints = list(set(ctx.hematology_endpoints + eps))
        elif platform == "Hormones":
            ctx.has_hormones = True
            eps = _get("endpoint_names", [])
            ctx.hormone_endpoints = list(set(ctx.hormone_endpoints + eps))
        elif platform == "Tissue Concentration":
            ctx.has_tissue_conc = True
        if data_type == "gene_expression":
            ctx.has_gene_expression = True
            organ = _get("organ")
            if organ and organ not in ctx.ge_organs:
                ctx.ge_organs.append(organ)

        # Collect doses and sexes from all fingerprints
        doses = _get("dose_groups", [])
        if doses:
            all_doses.update(float(d) for d in doses)
        sexes = _get("sexes", [])
        if sexes:
            all_sexes.update(sexes)
        du = _get("dose_unit")
        if du:
            dose_unit_found = du
        # Species from fingerprints (LLM-inferred) — only if user didn't override
        sp = _get("species")
        if sp and not study_params.get("species"):
            ctx.species = sp

    if all_doses:
        ctx.dose_groups = sorted(all_doses)
    if all_sexes:
        ctx.sexes = sorted(all_sexes)
    if dose_unit_found:
        ctx.dose_unit = dose_unit_found

    # --- Animal report: domain coverage and biosampling count ---
    if animal_report:
        # Biosampling count from animal_report
        ctx.n_biosampling = animal_report.get("biosampling_count", 0)

        # Fill dose_groups from animal_report if fingerprints didn't have them
        if not ctx.dose_groups and animal_report.get("dose_groups"):
            ctx.dose_groups = [float(d) for d in animal_report["dose_groups"]]

        # Domain coverage can confirm platform presence — keys are now
        # platform strings (e.g., "Body Weight", "Hematology").
        dc = animal_report.get("domain_coverage", {})
        for plat in dc:
            if plat == "Body Weight":
                ctx.has_body_weight = True
            elif plat == "Organ Weights":
                ctx.has_organ_weights = True
            elif plat == "Clinical Chemistry":
                ctx.has_clin_chem = True
            elif plat == "Hematology":
                ctx.has_hematology = True
            elif plat == "Hormones":
                ctx.has_hormones = True
            elif plat == "Tissue Concentration":
                ctx.has_tissue_conc = True
            elif plat == "Gene Expression":
                ctx.has_gene_expression = True

    # --- BMDExpress analysis metadata from .bm2 files ---
    for fid, bm2_json in bm2_jsons.items():
        if not isinstance(bm2_json, dict):
            continue
        meta = _collect_bm2_analysis_metadata(bm2_json)
        if meta:
            # Apply to context (first non-None wins for each field)
            if meta.get("bmdexpress_version") and not ctx.bmdexpress_version:
                ctx.bmdexpress_version = meta["bmdexpress_version"]
            if meta.get("bmds_version") and not ctx.bmds_version:
                ctx.bmds_version = meta["bmds_version"]
            if meta.get("bmr_type") and not ctx.bmr_type:
                ctx.bmr_type = meta["bmr_type"]
            if meta.get("bmr_factor") is not None and ctx.bmr_factor is None:
                ctx.bmr_factor = meta["bmr_factor"]
            if meta.get("models_fit") and not ctx.models_fit:
                ctx.models_fit = meta["models_fit"]
            if meta.get("constant_variance") is not None and ctx.constant_variance is None:
                ctx.constant_variance = meta["constant_variance"]
            if meta.get("prefilter_method") and not ctx.prefilter_method:
                ctx.prefilter_method = meta["prefilter_method"]
            if meta.get("prefilter_pvalue") is not None and ctx.prefilter_pvalue is None:
                ctx.prefilter_pvalue = meta["prefilter_pvalue"]
            if meta.get("fold_change_filter") is not None and ctx.fold_change_filter is None:
                ctx.fold_change_filter = meta["fold_change_filter"]

    # --- Build genomics sample counts for Table 1 ---
    # Structure: {organ: {sex: {dose: count}}}
    # Source: gene_expression fingerprints' n_animals_by_dose, grouped by organ and sex
    if ctx.has_gene_expression:
        ctx.genomics_sample_counts = _build_genomics_sample_counts(
            fingerprints, ctx.dose_groups,
        )

    # --- Biosampling dose groups (for Abstract-Methods) ---
    # Scan sidecar files for rows where selection contains "biosampling".
    # The reference report writes: "Blood was collected from animals
    # dedicated for internal dose assessment in the 4 and 37 mg/kg groups."
    if session_dir:
        ctx.biosampling_doses = _extract_biosampling_doses(session_dir)

    # --- Pharmacokinetics (for Abstract-Results) ---
    # Aggregate plasma concentrations + half-lives from tissue conc
    # sidecars.  Used by build_abstract_results_pk() to produce the
    # cross-sex comparison sentence (e.g., "Half-lives ... were 78.2
    # and 25.6 hours for the 4 and 37 mg/kg groups...").
    if session_dir:
        concs, t_half, timepoints = _extract_pk_data(session_dir)
        if concs:
            ctx.pk_concentrations = concs
        if t_half:
            ctx.pk_half_lives = t_half
        if timepoints:
            ctx.pk_timepoints = timepoints

    # --- Genomics assay and chip (for Abstract-Methods) ---
    # Extract from integrated.doseResponseExperiments[].chip.  The chip name
    # or ID identifies the assay — "S1500" implies TempO-Seq, "HG_U133" implies
    # Affymetrix, etc.
    if integrated and ctx.has_gene_expression:
        assay, chip = _extract_genomics_assay(integrated)
        ctx.genomics_assay = assay
        ctx.genomics_chip = chip

    return ctx


def _extract_biosampling_doses(session_dir: str) -> list[float]:
    """
    Scan session sidecar JSON files for biosampling animals and collect
    the set of dose groups they belong to.

    Biosampling animals are tagged in sidecar data via the per-animal
    `selection` field containing "biosampling".  Usually only 2 of the
    study's 10 doses have biosampling animals (the two chosen for
    pharmacokinetic/internal dose assessment).

    Returns a sorted list of dose values.  Empty list if no biosampling
    animals are found.
    """
    import os
    import json

    doses: set[float] = set()
    files_dir = os.path.join(session_dir, "files")
    if not os.path.isdir(files_dir):
        return []

    for fname in os.listdir(files_dir):
        if not fname.endswith(".sidecar.json"):
            continue
        try:
            with open(os.path.join(files_dir, fname)) as f:
                sc = json.load(f)
        except (json.JSONDecodeError, OSError):
            continue

        for _aid, rec in sc.get("animals", {}).items():
            selection = str(rec.get("selection", ""))
            if "biosampling" in selection.lower():
                dose = rec.get("dose")
                if dose is not None:
                    try:
                        doses.add(float(dose))
                    except (TypeError, ValueError):
                        pass

    return sorted(doses)


def _extract_pk_data(session_dir: str) -> tuple[dict, dict, list[int]]:
    """
    Extract pharmacokinetic data from tissue concentration sidecars.

    Reads any tissue_conc_*.sidecar.json files in the session, scans
    biosampling animal records, and aggregates plasma concentrations
    by (sex, dose, timepoint).  Then computes per-(sex, dose) plasma
    half-lives using the standard two-point formula:

        t½ = ln(2) × Δt / ln(C_early / C_late)

    Half-life is only computed when:
      - Exactly two timepoints exist with positive mean concentrations
      - The early concentration is greater than the late concentration
        (monotonic decay — required for the log-linear assumption)

    Returns:
        (concentrations, half_lives, timepoints)
        concentrations: {sex: {dose: {hour_int: mean_value}}}
        half_lives:     {sex: {dose: hours_float}}
        timepoints:     sorted list of unique hour integers seen across
                        all observations (e.g., [2, 24])
    """
    import os
    import json
    import math
    import re

    files_dir = os.path.join(session_dir, "files")
    if not os.path.isdir(files_dir):
        return {}, {}, []

    # Per (sex, dose, hour) → list of concentration values
    raw_values: dict[tuple[str, float, int], list[float]] = {}
    timepoints_seen: set[int] = set()

    # Pattern to extract the timepoint hours from endpoint names like
    # "Plasma 2 Hour Perfluorohexanesulfonamide Concentration".
    _HOUR_RE = re.compile(r"Plasma\s+(\d+)\s+Hour", re.IGNORECASE)

    for fname in os.listdir(files_dir):
        if not fname.startswith("tissue_conc") or not fname.endswith(".sidecar.json"):
            continue
        try:
            with open(os.path.join(files_dir, fname)) as f:
                sc = json.load(f)
        except (json.JSONDecodeError, OSError):
            continue

        # The sidecar's "sex" field carries the per-file sex (e.g., "Male"),
        # which we trust over per-animal sex (often None in tissue conc).
        sex = sc.get("sex") or ""
        if not sex:
            continue

        for _aid, rec in sc.get("animals", {}).items():
            selection = str(rec.get("selection", ""))
            if "biosampling" not in selection.lower():
                continue
            dose = rec.get("dose")
            if dose is None:
                continue
            try:
                dose = float(dose)
            except (TypeError, ValueError):
                continue

            for obs in rec.get("observations", []):
                ep = obs.get("endpoint", "")
                if "Concentration" not in ep:
                    continue  # skip LOQ rows
                m = _HOUR_RE.search(ep)
                if not m:
                    continue
                hour = int(m.group(1))
                val = obs.get("value")
                if val is None:
                    continue
                try:
                    v = float(val)
                except (TypeError, ValueError):
                    continue
                raw_values.setdefault((sex, dose, hour), []).append(v)
                timepoints_seen.add(hour)

    # Aggregate to means
    concentrations: dict[str, dict[float, dict[int, float]]] = {}
    for (sex, dose, hour), vs in raw_values.items():
        if not vs:
            continue
        mean = sum(vs) / len(vs)
        concentrations.setdefault(sex, {}).setdefault(dose, {})[hour] = mean

    # Compute half-lives where we have two-timepoint monotonic decay
    half_lives: dict[str, dict[float, float]] = {}
    for sex, by_dose in concentrations.items():
        for dose, by_hour in by_dose.items():
            tps = sorted(by_hour.keys())
            if len(tps) < 2:
                continue
            # Use the first and last available timepoints
            t_early, t_late = tps[0], tps[-1]
            c_early = by_hour[t_early]
            c_late = by_hour[t_late]
            if c_early <= 0 or c_late <= 0 or c_early <= c_late:
                continue  # require monotonic decay for log-linear half-life
            try:
                t_half = math.log(2) * (t_late - t_early) / math.log(c_early / c_late)
            except (ValueError, ZeroDivisionError):
                continue
            half_lives.setdefault(sex, {})[dose] = t_half

    return concentrations, half_lives, sorted(timepoints_seen)


def _extract_genomics_assay(integrated: dict) -> tuple[str | None, str | None]:
    """
    Identify the genomics assay platform from the integrated BMDProject.

    Scans doseResponseExperiments for gene-expression experiments (those with
    a non-generic chip) and reads chip.name / chip.chipId.  Maps known chip
    identifiers to their canonical assay names:

      - "S1500", "S1500+"        → TempO-Seq
      - "BioSpyder"              → TempO-Seq
      - "HG-U133", "HT_MG_..."   → Affymetrix GeneChip
      - "Illumina"               → Illumina BeadChip / RNA-seq

    Returns (assay_name, chip_name) tuple — either may be None if not
    identifiable.  For unknown chips, returns (None, chip.name) so the
    caller can at least report the raw chip identifier.
    """
    experiments = integrated.get("doseResponseExperiments", [])
    for e in experiments:
        chip = e.get("chip")
        # Skip refs (int) and None — we need a full dict with name/chipId
        if not isinstance(chip, dict):
            continue
        chip_id = str(chip.get("chipId", "") or "")
        chip_name = str(chip.get("name", "") or "")
        # Skip placeholder "generic" chips attached to apical experiments
        if chip_id.lower() in ("generic", "") and chip_name.lower() in ("generic", ""):
            continue
        # Compare case-insensitively against known signatures
        probe = (chip_id + " " + chip_name).lower()
        if "s1500" in probe or "tempo" in probe or "biospyder" in probe:
            return ("TempO-Seq", chip_name or chip_id)
        if "affy" in probe or "hg-u133" in probe or "ht_mg" in probe:
            return ("Affymetrix GeneChip", chip_name or chip_id)
        if "illumina" in probe:
            return ("Illumina", chip_name or chip_id)
        # Unknown real chip — return its name so the caller can still report it
        return (None, chip_name or chip_id)

    return (None, None)


def _build_genomics_sample_counts(
    fingerprints: dict,
    dose_groups: list[float],
) -> dict | None:
    """
    Build the Table 1 sample-count matrix from gene_expression fingerprints.

    Each gene_expression fingerprint represents one organ × sex combination.
    We extract n_animals_by_dose from each to build:
        {organ: {sex: {dose: count}}}

    Args:
        fingerprints: Dict of {file_id: fingerprint_dict_or_object}.
        dose_groups:  Sorted list of all dose values in the study.

    Returns:
        Nested dict of sample counts, or None if no GE fingerprints found.
    """
    counts: dict[str, dict[str, dict[float, int]]] = {}

    for fid, fp in fingerprints.items():
        _get = fp.get if isinstance(fp, dict) else lambda k, d=None: getattr(fp, k, d)

        if _get("data_type") != "gene_expression":
            continue

        organ = _get("organ", "Unknown")
        sexes = _get("sexes", [])
        n_by_dose = _get("n_animals_by_dose", {})

        if not n_by_dose:
            continue

        # Each GE fingerprint is typically one sex (inferred by LLM),
        # but may have multiple sexes if the experiment combines them
        sex_label = sexes[0] if sexes else "Unknown"

        if organ not in counts:
            counts[organ] = {}
        if sex_label not in counts[organ]:
            counts[organ][sex_label] = {}

        for dose_str, count in n_by_dose.items():
            dose_val = float(dose_str)
            # Take the max if we see the same organ/sex/dose from multiple files
            existing = counts[organ][sex_label].get(dose_val, 0)
            counts[organ][sex_label][dose_val] = max(existing, count)

    return counts if counts else None


# ---------------------------------------------------------------------------
# Subsection skeleton: which subsections to include given the context
# ---------------------------------------------------------------------------

def build_subsection_skeleton(ctx: MethodsContext) -> list[tuple[str, str, int]]:
    """
    Determine which M&M subsections to include based on domain presence.

    Returns a list of (subsection_key, heading_text, heading_level) tuples
    in the correct order for both the LLM prompt and the DOCX output.

    Args:
        ctx: MethodsContext with domain presence flags set.

    Returns:
        Filtered list of subsections that apply to this study.
    """
    skeleton = []
    for key, heading, level, condition in SUBSECTION_SKELETON:
        if condition is None:
            # Always included
            skeleton.append((key, heading, level))
        else:
            # Check the condition field on MethodsContext
            if getattr(ctx, condition, False):
                skeleton.append((key, heading, level))
    return skeleton


# ---------------------------------------------------------------------------
# LLM prompt builder
# ---------------------------------------------------------------------------

def build_methods_prompt(ctx: MethodsContext) -> tuple[str, str]:
    """
    Build the structured LLM prompt for M&M generation.

    Returns (system_prompt, user_prompt) strings.

    The user prompt includes all extracted study context and asks the LLM
    to return a JSON object keyed by subsection key.  Only keys for
    present domains are requested.

    Args:
        ctx: Populated MethodsContext.

    Returns:
        Tuple of (system_prompt, user_prompt).
    """
    system = (
        "You are a toxicology report writer specializing in NTP/NIEHS-style "
        "technical reports for 5-day genomic dose-response studies. "
        "Write precise, formal scientific text. "
        "Return ONLY valid JSON with no markdown formatting."
    )

    # Build the study context block with actual data
    dose_str = ", ".join(str(d) for d in ctx.dose_groups) if ctx.dose_groups else "not specified"
    sexes_str = " and ".join(ctx.sexes) if ctx.sexes else "male and female"

    context_block = f"""## Study Context
- Chemical: {ctx.chemical_name} (CASRN: {ctx.casrn}, DTXSID: {ctx.dtxsid})
- Species: {ctx.species} rats
- Sexes: {sexes_str}
- Route of administration: oral {ctx.route}
- Vehicle: {ctx.vehicle}
- Duration: {ctx.duration_days} days
- Dose groups ({ctx.dose_unit}): {dose_str}
- Animals per treatment group: {ctx.n_per_group} {sexes_str} per dose
- Animals in vehicle control: {ctx.n_control} {sexes_str}"""

    if ctx.n_biosampling > 0:
        context_block += f"\n- Biosampling animals: {ctx.n_biosampling} total"

    # Domain-specific context
    if ctx.has_organ_weights and ctx.organ_weight_endpoints:
        context_block += f"\n- Organs weighed: {', '.join(ctx.organ_weight_endpoints)}"
    if ctx.has_clin_chem and ctx.clin_chem_endpoints:
        context_block += f"\n- Clinical chemistry parameters: {', '.join(ctx.clin_chem_endpoints)}"
    if ctx.has_hematology and ctx.hematology_endpoints:
        context_block += f"\n- Hematology parameters: {', '.join(ctx.hematology_endpoints)}"
    if ctx.has_hormones and ctx.hormone_endpoints:
        context_block += f"\n- Hormone parameters: {', '.join(ctx.hormone_endpoints)}"
    if ctx.has_gene_expression and ctx.ge_organs:
        context_block += f"\n- Transcriptomics organs: {', '.join(ctx.ge_organs)}"
    if ctx.has_tissue_conc:
        context_block += "\n- Internal dose assessment: tissue concentration data available"

    # BMDExpress metadata
    if ctx.bmdexpress_version:
        context_block += f"\n- BMDExpress version: {ctx.bmdexpress_version}"
    if ctx.bmds_version:
        context_block += f"\n- BMDS version: {ctx.bmds_version}"
    if ctx.models_fit:
        context_block += f"\n- BMD models fit: {', '.join(ctx.models_fit)}"
    if ctx.bmr_type:
        context_block += f"\n- BMR type: {ctx.bmr_type}"
    if ctx.bmr_factor is not None:
        context_block += f"\n- BMR factor: {ctx.bmr_factor}"
    if ctx.prefilter_method:
        context_block += f"\n- Pre-filter method: {ctx.prefilter_method}"
    if ctx.prefilter_pvalue is not None:
        context_block += f"\n- Pre-filter p-value cutoff: {ctx.prefilter_pvalue}"
    if ctx.constant_variance is not None:
        context_block += f"\n- Constant variance: {'yes' if ctx.constant_variance else 'no'}"

    # Build subsection requirements
    skeleton = build_subsection_skeleton(ctx)
    subsection_keys = [key for key, _, _ in skeleton]

    # Per-subsection guidance for the LLM
    guidelines = _build_subsection_guidelines(ctx, skeleton)

    user_prompt = f"""Generate the Materials and Methods section for a {ctx.duration_days}-day genomic dose-response study report.

{context_block}

## Required Subsections
Generate content for ONLY the following subsections (in order).
Return your response as a JSON object where each key is a subsection identifier
and the value is a string containing 1-3 paragraphs of prose.

Keys to generate: {subsection_keys}

{guidelines}

## Important formatting notes:
- Do NOT include section headers in the text — they will be added by the template.
- Write in past tense, third person, formal NIEHS/NTP technical report style.
- Reference exact study parameters from the context above (dose groups, species, etc.).
- Each value should be a single string. Use \\n\\n to separate multiple paragraphs within a subsection.
- Return ONLY the JSON object, no markdown code fences."""

    return system, user_prompt


def _build_subsection_guidelines(
    ctx: MethodsContext,
    skeleton: list[tuple[str, str, int]],
) -> str:
    """
    Build per-subsection prose guidelines for the LLM prompt.

    Each guideline tells the LLM exactly what to cover in that subsection
    and references the actual study parameters from the context.

    Args:
        ctx:      MethodsContext with study parameters.
        skeleton: The filtered subsection skeleton.

    Returns:
        Formatted string of guidelines.
    """
    dose_str = ", ".join(str(d) for d in ctx.dose_groups) if ctx.dose_groups else "not specified"
    sexes_str = " and ".join(ctx.sexes) if ctx.sexes else "male and female"

    guides: dict[str, str] = {
        "study_design": (
            f"One paragraph covering: animal source and quarantine (7 days), "
            f"randomization by body weight, dose groups ({dose_str} {ctx.dose_unit}), "
            f"dosing schedule ({ctx.duration_days} consecutive days by oral {ctx.route}), "
            f"sample sizes ({ctx.n_per_group} {sexes_str} per treatment group, "
            f"{ctx.n_control} {sexes_str} vehicle control), "
            f"{'biosampling animals for internal dose assessment, ' if ctx.n_biosampling > 0 else ''}"
            f"necropsy timing (approximately 24 hours after the final dose)."
        ),
        "dose_selection": (
            "One paragraph. Reference LD50 predictions if available, "
            "explain the dose spacing rationale (half-log dose spacing), "
            "rationale for top dose selection, and target of no overt systemic toxicity."
        ),
        "chemistry": (
            f"First paragraph: {ctx.chemical_name} source (typically synthesized by contract lab or obtained from commercial supplier), "
            f"lot number (state that lot details are in the full study protocol), "
            f"identity confirmation methods (IR, NMR, or mass spectrometry), purity (typically >95%%), storage conditions. "
            f"Second paragraph: dose formulation preparation in {ctx.vehicle}, "
            f"concentration verification by analytical chemistry, formulation QC analysis."
        ),
        "clinical_exams": (
            "Brief introductory paragraph describing the scope of clinical examinations "
            "performed during the study: clinical observations, body/organ weight collection, "
            "clinical pathology sampling, and (if applicable) tissue collection for internal dose assessment."
        ),
        "clinical_obs": (
            "One paragraph. Standard boilerplate: animals were observed twice daily for morbidity/mortality, "
            "clinical observations were recorded once daily, "
            "cage-side observations included assessment of general appearance, behavior, and respiratory patterns."
        ),
        "body_organ_weights": (
            f"One paragraph. Body weights recorded on study day 0 and at necropsy. "
            f"{'Organ weights: ' + ', '.join(ctx.organ_weight_endpoints[:10]) + '.' if ctx.organ_weight_endpoints else 'Organs were weighed at necropsy.'} "
            f"Absolute and relative (organ-to-body-weight ratio) weights calculated."
        ),
        "clinical_pathology": (
            f"One paragraph covering blood collection method (retroorbital or cardiac puncture under anesthesia), "
            f"sample processing. "
            f"{'Clinical chemistry parameters: ' + ', '.join(ctx.clin_chem_endpoints[:10]) + '. ' if ctx.clin_chem_endpoints else ''}"
            f"{'Hematology parameters: ' + ', '.join(ctx.hematology_endpoints[:10]) + '. ' if ctx.hematology_endpoints else ''}"
            f"{'Hormone parameters: ' + ', '.join(ctx.hormone_endpoints[:5]) + '. ' if ctx.hormone_endpoints else ''}"
            f"Reference the analyzer instruments (e.g., Advia 120 for hematology, AU680 for clinical chemistry)."
        ),
        "internal_dose": (
            "One paragraph. Describe tissue collection for internal dose assessment: "
            "blood/serum and tissue samples collected at necropsy, stored at -80°C, "
            "analyzed by LC-MS/MS for parent compound and/or metabolite concentrations."
        ),
        "transcriptomics": (
            f"Brief introductory paragraph: transcriptomic profiling was performed on "
            f"{', '.join(ctx.ge_organs) if ctx.ge_organs else 'target organs'} "
            f"using the TempO-Seq targeted RNA sequencing platform (BioSpyder Technologies)."
        ),
        "txomics_sample": (
            "One paragraph. RNA was isolated from flash-frozen tissue samples. "
            "Describe tissue homogenization, RNA extraction method (e.g., RNeasy), "
            "RNA quality assessment (RIN values)."
        ),
        "txomics_rna": (
            "One paragraph. TempO-Seq S1500+ platform (BioSpyder Technologies), "
            "library preparation with detector oligos, ligation, amplification, "
            "multiplexed sequencing on Illumina platform."
        ),
        "txomics_seq_processing": (
            "One paragraph. Demultiplexing, alignment to probe reference sequences, "
            "count table generation using the TempO-SeqR pipeline."
        ),
        "txomics_qc": (
            "One paragraph. Outlier detection by principal component analysis and "
            "total read count thresholds. Samples below quality thresholds were excluded "
            "from downstream analysis."
        ),
        "txomics_normalization": (
            "One paragraph. DESeq2 median-of-ratios normalization or TMM normalization. "
            "Describe the normalization method used for count data."
        ),
        "data_analysis": (
            "Brief introductory paragraph: overview of the statistical and dose-response "
            "modeling approaches applied to both apical and transcriptomic endpoints."
        ),
        "stat_analysis": (
            "One paragraph. Standard NIEHS statistical pipeline: "
            "Jonckheere trend test for monotonic dose-response, followed by "
            "Williams test (if monotonic) or Dunnett test (if non-monotonic) for pairwise comparisons. "
            "Shirley nonparametric test used when data did not meet normality/homogeneity assumptions. "
            "Dunn test for non-parametric pairwise comparisons. "
            "Significance threshold: p < 0.05."
        ),
        "bmd_apical": (
            f"One paragraph. Benchmark dose modeling of apical endpoints using "
            f"{'BMDS ' + ctx.bmds_version if ctx.bmds_version else 'EPA BMDS software'}. "
            f"{'Models fit: ' + ', '.join(ctx.models_fit) + '. ' if ctx.models_fit else ''}"
            f"{'BMR: ' + str(ctx.bmr_factor) + ' ' + (ctx.bmr_type or 'standard deviation') + ' from the control mean. ' if ctx.bmr_factor is not None else ''}"
            f"{'Constant variance assumed. ' if ctx.constant_variance else ''}"
            f"Model selection by lowest AIC or Laplace model averaging."
        ),
        "bmd_genomics": (
            f"One paragraph. Benchmark dose modeling of transcriptomic data using "
            f"{ctx.bmdexpress_version or 'BMDExpress'} with "
            f"{'BMDS ' + ctx.bmds_version if ctx.bmds_version else 'EPA BMDS'}. "
            f"{'Pre-filter: ' + ctx.prefilter_method + ' (p < ' + str(ctx.prefilter_pvalue) + ')' if ctx.prefilter_method and ctx.prefilter_pvalue else ''}"
            f"{'Models fit: ' + ', '.join(ctx.models_fit) + '. ' if ctx.models_fit else ''}"
            f"Describe gene-level BMD derivation, GO term / pathway enrichment using "
            f"category analysis, and filtering criteria (goodness-of-fit p > 0.1, "
            f"BMDU/BMDL ratio ≤ 40, minimum 5%% gene set coverage)."
        ),
        "efdr": (
            "One paragraph. Empirical false discovery rate estimation by permutation: "
            "dose labels are randomly shuffled to create null datasets, "
            "the full BMD analysis pipeline is run on permuted data, "
            "and the fraction of false positives at each BMD threshold is estimated. "
            "Reference the number of permutations used."
        ),
        "data_accessibility": (
            "One paragraph. Data availability statement: "
            "all raw and processed data are available through the NIEHS Chemical Effects "
            "in Biological Systems (CEBS) database or the study's data repository. "
            "Provide placeholder for DOI or accession number."
        ),
    }

    lines = ["## Guidelines per subsection:"]
    for key, heading, level in skeleton:
        if key in guides:
            lines.append(f"- {key}: {guides[key]}")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Table 1 builder
# ---------------------------------------------------------------------------

def build_table1_data(ctx: MethodsContext) -> dict | None:
    """
    Build Table 1: Final Sample Counts for BMD Analysis of Transcriptomics Data.

    This table shows the number of samples per organ × sex × dose group
    that passed QC and were included in the BMD analysis.

    Args:
        ctx: MethodsContext with genomics_sample_counts populated.

    Returns:
        Dict with keys: caption, headers, rows, footnotes.
        Or None if no genomics data.
    """
    if not ctx.genomics_sample_counts:
        return None

    # Headers: empty corner cell + dose group columns
    dose_headers = []
    for d in ctx.dose_groups:
        # Format: "0 mg/kg", "0.15 mg/kg", etc.
        if d == int(d):
            dose_headers.append(f"{int(d)} {ctx.dose_unit}")
        else:
            dose_headers.append(f"{d} {ctx.dose_unit}")

    headers = [""] + dose_headers

    # Rows: grouped by sex, then by organ
    rows = []
    all_sexes = sorted({
        sex
        for organ_data in ctx.genomics_sample_counts.values()
        for sex in organ_data
    })
    all_organs = sorted(ctx.genomics_sample_counts.keys())

    for sex in all_sexes:
        # Sex header row (bold — indicated by leading **)
        rows.append([f"**{sex}**"] + [""] * len(ctx.dose_groups))

        for organ in all_organs:
            sex_data = ctx.genomics_sample_counts.get(organ, {}).get(sex, {})
            row = [f"  {organ}"]
            for dose in ctx.dose_groups:
                count = sex_data.get(dose, 0)
                if count > 0:
                    row.append(str(count))
                else:
                    # Dash indicates no samples (mortality, exclusion, etc.)
                    row.append("–")
            rows.append(row)

    return {
        "caption": "Final Sample Counts for BMD Analysis of Transcriptomics Data",
        "headers": headers,
        "rows": rows,
        "footnotes": [],
    }


# ---------------------------------------------------------------------------
# Abstract → Methods paragraph builder
# ---------------------------------------------------------------------------
# The Abstract section has 4 labeled paragraphs (Background, Methods, Results,
# Summary).  The Methods paragraph is deterministic — every fact it reports is
# already in MethodsContext, so we can generate it from a template without an
# LLM call.  This keeps the abstract methods faithful to the data and fast.

def _format_dose_list(doses: list[float], dose_unit: str) -> str:
    """
    Format a list of dose values into a NIEHS-style comma-separated string.

    Per NIEHS Report 10 convention:
      - Preserve the exact dose values (including leading zero and decimals)
      - Separate with commas, with "and" before the last value
      - Append "mg/kg body weight [mg/kg]" style unit annotation

    Examples:
      _format_dose_list([0, 0.15, 0.5], "mg/kg")
        → "0, 0.15, and 0.5"
      _format_dose_list([4, 37], "mg/kg")
        → "4 and 37"
    """
    if not doses:
        return "(doses not specified)"

    # Format each dose: strip trailing ".0" for integers (with NIEHS-style
    # thousand-separator comma for >=1000), keep decimals otherwise.
    def _fmt(d: float) -> str:
        if d == int(d):
            return f"{int(d):,}"
        return f"{d:g}"

    formatted = [_fmt(d) for d in doses]

    if len(formatted) == 1:
        return formatted[0]
    if len(formatted) == 2:
        return f"{formatted[0]} and {formatted[1]}"
    return ", ".join(formatted[:-1]) + ", and " + formatted[-1]


def _format_organ_list(organs: list[str]) -> str:
    """
    Format a list of organ names (e.g., ["Liver", "Kidney"]) into natural
    English: "liver and kidney" (two), "liver, kidney, and spleen" (three+),
    "liver" (one).  Lowercases organ names since they appear mid-sentence.
    """
    lc = [o.lower() for o in organs]
    if not lc:
        return ""
    if len(lc) == 1:
        return lc[0]
    if len(lc) == 2:
        return f"{lc[0]} and {lc[1]}"
    return ", ".join(lc[:-1]) + ", and " + lc[-1]


def build_abstract_methods(ctx: MethodsContext) -> str:
    """
    Build the Abstract → Methods paragraph from a MethodsContext.

    Mirrors the structure of NIEHS Report 10's Abstract-Methods paragraph:

      A short-term in vivo biological potency study on {TA} in adult {sexes}
      {Species} (Strain) rats was conducted.  {TA} was formulated in
      {vehicle} and administered once daily for {duration} consecutive days
      by {route} (study days 0–{duration-1}).  {TA} was administered at
      {N_doses} doses ({dose_list} {dose_unit}).  Blood was collected from
      animals dedicated for internal dose assessment in the
      {biosampling_doses} {dose_unit} groups.  On study day {duration}, the
      day after the final dose was administered, animals were euthanized,
      standard toxicological measures were assessed, and the {ge_organs}
      were assayed in gene expression studies using the {assay} assay.
      Modeling was conducted to identify the benchmark doses (BMDs)
      associated with apical toxicological endpoints and transcriptional
      changes in the {ge_organs}.  A benchmark response of {bmr} was used
      to model all endpoints.

    All data comes from MethodsContext — no LLM call.  Conditional clauses
    are omitted when the corresponding data isn't present (e.g., no
    biosampling animals → skip the blood collection sentence; no gene
    expression → skip the transcriptomics clauses).
    """
    # --- Test article display name ---
    # Use abbreviation if present, else full name
    ta_name = ctx.chemical_name or "the test article"

    # --- Sexes ("adult male and female") ---
    sexes = [s.lower() for s in ctx.sexes] if ctx.sexes else ["male", "female"]
    sexes_str = " and ".join(sexes)

    # --- Species ("Sprague Dawley (Hsd:Sprague Dawley® SD®)") ---
    # If the species string contains the strain in parens, use as-is;
    # otherwise fall back to "Sprague Dawley".
    species = ctx.species or "Sprague Dawley"

    # --- Vehicle ---
    vehicle = ctx.vehicle or "corn oil"

    # --- Duration + study days range ---
    duration = ctx.duration_days or 5
    last_day = duration - 1
    sacrifice_day = duration

    # --- Route ---
    route = ctx.route or "gavage"

    # --- Dose list ---
    n_doses = len(ctx.dose_groups)
    dose_list = _format_dose_list(ctx.dose_groups, ctx.dose_unit)
    # NIEHS convention: "mg/kg body weight [mg/kg]" on first mention
    dose_unit_full = f"{ctx.dose_unit} body weight [{ctx.dose_unit}]"

    # --- Biosampling sentence (conditional) ---
    biosampling_sentence = ""
    if ctx.biosampling_doses:
        bio_list = _format_dose_list(ctx.biosampling_doses, ctx.dose_unit)
        biosampling_sentence = (
            f" Blood was collected from animals dedicated for internal "
            f"dose assessment in the {bio_list} {ctx.dose_unit} groups."
        )

    # --- Transcriptomics clauses (conditional) ---
    assay_clause = ""
    bmd_clause = (
        " Modeling was conducted to identify the benchmark doses (BMDs) "
        "associated with apical toxicological endpoints"
    )
    if ctx.has_gene_expression and ctx.ge_organs:
        organs_str = _format_organ_list(ctx.ge_organs)
        assay_name = ctx.genomics_assay or "gene expression"
        assay_clause = (
            f", and the {organs_str} were assayed in gene expression "
            f"studies using the {assay_name} assay"
        )
        bmd_clause += f" and transcriptional changes in the {organs_str}"
    bmd_clause += "."

    # --- BMR description ---
    # Default to NIEHS protocol: "one standard deviation"
    bmr_desc = "one standard deviation"
    if ctx.bmr_factor is not None and ctx.bmr_type:
        # e.g., bmr_factor=1.0, bmr_type="Std. Dev." → "one standard deviation"
        factor_words = {1.0: "one", 2.0: "two", 0.5: "one half of"}
        word = factor_words.get(ctx.bmr_factor, str(ctx.bmr_factor))
        type_clean = ctx.bmr_type.lower().replace("std. dev.", "standard deviation").strip()
        bmr_desc = f"{word} {type_clean}" if word else type_clean

    # --- Assemble the paragraph ---
    paragraph = (
        f"A short-term in vivo biological potency study on {ta_name} in "
        f"adult {sexes_str} {species} rats was conducted. "
        f"{ta_name} was formulated in {vehicle} and administered once "
        f"daily for {duration} consecutive days by {route} (study days "
        f"0–{last_day}). "
        f"{ta_name} was administered at {n_doses} doses "
        f"({dose_list} {dose_unit_full})."
        f"{biosampling_sentence}"
        f" On study day {sacrifice_day}, the day after the final dose "
        f"was administered, animals were euthanized, standard "
        f"toxicological measures were assessed{assay_clause}."
        f"{bmd_clause}"
        f" A benchmark response of {bmr_desc} was used to model all "
        f"endpoints."
    )

    return paragraph


# ---------------------------------------------------------------------------
# Abstract → Results paragraph builder
# ---------------------------------------------------------------------------
# The Abstract → Results paragraph summarizes the apical-endpoint findings
# (significant changes + BMD/BMDL values) per sex.  Like the Methods
# paragraph, every fact is already in the BMD summary, so we generate it
# deterministically without an LLM.  Genomics and pharmacokinetics sub-
# paragraphs follow the same pattern and can be added in future iterations.

# Direction words used to convert UP/DOWN flags into adjectives that read
# naturally when prefixed to an endpoint name.
_DIRECTION_WORDS = {
    "UP":   "increased",
    "DOWN": "decreased",
    "":     "altered",  # fallback if direction wasn't recorded
}


def _is_reliable_bmd(entry: dict) -> bool:
    """
    Decide whether a BMD summary entry has a reliable, parseable BMD.

    NIEHS reference convention: report "no reliable BMD" when the value
    is missing, the dash placeholder ("—"), an explicit failure marker
    ("NVM" = no viable model), or when bmd_status is something other
    than "viable".
    """
    bmd = entry.get("bmd")
    if bmd is None:
        return False
    s = str(bmd).strip()
    if s in ("", "—", "NVM", "ND", "NA"):
        return False
    # Must be parseable as float
    try:
        float(s)
    except (TypeError, ValueError):
        return False
    if entry.get("bmd_status") and entry["bmd_status"] != "viable":
        return False
    return True


# Threshold for the anomalous-BMD heuristic.  When the curve-fit BMD is
# this many times lower than the statistically-observed NOEL/LOEL, the
# BMD is treated as a model artifact rather than a true potency estimate.
#
# The NIEHS reference report quantifies the anomaly callouts as
# "approximately 75- to 230-fold", "140- to 430-fold", and "25- to
# 80-fold" lower than NOEL/LOEL — implying the threshold for excluding
# from the Abstract effects list is somewhere around 10× (any larger
# discrepancy would be reported in the body Results as anomalous).
ANOMALY_RATIO_THRESHOLD = 10.0


def _is_anomalous_bmd(entry: dict, threshold: float = ANOMALY_RATIO_THRESHOLD) -> bool:
    """
    Detect whether a BMD entry has an anomalously LOW model-derived BMD
    compared to the statistically-observed NOEL/LOEL.

    Rationale: pairwise statistical tests identify the lowest dose where
    the effect is statistically significant (LOEL) and the highest dose
    where it isn't (NOEL).  A curve-fit BMD substantially below the NOEL
    means the model extrapolates an "effect" at doses where the actual
    measurements showed none — almost always a model artifact (e.g.,
    poor model fit, control variability, BMR set too tight relative to
    the data).  The reference report flags such BMDs as anomalous and
    excludes them from the Abstract's "effects included..." list.

    Decision rule:
      - If NOEL is present and BMD < NOEL / threshold  → anomalous.
      - If NOEL is None but LOEL is present and
        BMD < LOEL / threshold → anomalous (weaker signal).
      - Otherwise → not flagged.

    Args:
        entry: BMD summary dict with keys 'bmd', 'noel', 'loel'.
        threshold: ratio at which BMD is considered too low to trust
                   (default 10, i.e., BMD < NOEL/10 is anomalous).

    Returns:
        True if the BMD should be filtered out of the Abstract effects list.
    """
    bmd_raw = entry.get("bmd")
    if bmd_raw is None:
        return False
    try:
        bmd = float(str(bmd_raw).strip())
    except (TypeError, ValueError):
        return False
    if bmd <= 0:
        return False

    noel = entry.get("noel")
    loel = entry.get("loel")

    # Prefer NOEL when available (it's the strongest evidence of "no
    # statistical effect at this dose"); fall back to LOEL otherwise.
    reference_dose = None
    if noel is not None:
        try:
            reference_dose = float(noel)
        except (TypeError, ValueError):
            pass
    if reference_dose is None and loel is not None:
        try:
            reference_dose = float(loel)
        except (TypeError, ValueError):
            pass

    if reference_dose is None or reference_dose <= 0:
        return False

    return bmd < (reference_dose / threshold)


def _normalize_endpoint_name(endpoint: str) -> str:
    """
    Lowercase a multi-word endpoint name for mid-sentence use, while
    preserving short ALL-CAPS acronyms (e.g., "ALT", "RBC") that should
    stay uppercase.

    Examples:
      "Total Thyroxine"           → "total thyroxine"
      "Aspartate Aminotransferase" → "aspartate aminotransferase"
      "Thyroid Stimulating Hormone" → "thyroid stimulating hormone"
      "ALT"                        → "ALT"
      "Hgb"                        → "hgb"
    """
    if not endpoint:
        return ""

    words = endpoint.strip().split()
    out_words: list[str] = []
    for w in words:
        # Acronym heuristic: all-uppercase, 2-5 chars, no digits → keep
        if 2 <= len(w) <= 5 and w.isupper() and w.isalpha():
            out_words.append(w)
        else:
            out_words.append(w.lower())
    return " ".join(out_words)


def _format_endpoint_phrase(endpoint: str, direction: str, platform: str) -> str:
    """
    Build the descriptor phrase used in the abstract sentence:
      "{direction word} {endpoint name lowercased}{contextual suffix}"

    Hormone endpoints conventionally include "concentration" (e.g.,
    "decreased total thyroxine concentration"); organ weight endpoints
    append "weight" (e.g., "increased liver weight"); other platforms
    use the bare endpoint name.
    """
    direction_word = _DIRECTION_WORDS.get(direction or "", "altered")
    endpoint_lower = _normalize_endpoint_name(endpoint) or "(unknown endpoint)"

    suffix = ""
    if platform == "Hormones":
        suffix = " concentration"
    elif platform == "Organ Weight":
        suffix = " weight"

    return f"{direction_word} {endpoint_lower}{suffix}"


def _format_bmd_pair(bmd, bmdl) -> str:
    """
    Format one BMD (BMDL) value pair for the listing sentence.

    Strategy: when the source value parses as a clean decimal, preserve
    the source's precision — the upstream BMDExpress output already
    rounds to a sensible number of significant figures, and matching
    that display avoids spurious trailing zeros (e.g., "8.54" not
    "8.540").  Falls back to "—" for missing/non-numeric values.
    """
    def _fmt(v) -> str:
        if v is None:
            return "—"
        s = str(v).strip()
        if s in ("", "—", "NVM", "ND", "NA"):
            return "—"
        # Verify it's parseable, but return the original string to
        # preserve source-side rounding (e.g., "8.54" not "8.5400")
        try:
            float(s)
        except (TypeError, ValueError):
            return s
        return s

    return f"{_fmt(bmd)} ({_fmt(bmdl)})"


def _join_oxford(items: list[str]) -> str:
    """
    Join a list with Oxford comma and "and" before the last element.
    "" → "", ["a"] → "a", ["a","b"] → "a and b",
    ["a","b","c"] → "a, b, and c".
    """
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return ", ".join(items[:-1]) + ", and " + items[-1]


def build_abstract_results_apical(
    apical_bmd_summary: list[dict],
    sexes: list[str] | None = None,
) -> str:
    """
    Build the apical-findings portion of the Abstract → Results paragraph.

    Mirrors NIEHS Report 10's per-sex apical summary:

      Several clinical pathology and organ weight measurements showed
      dose-related changes from which BMD values were calculated. In
      male rats, the effects included {direction} {endpoint}, ... The
      BMDs and benchmark dose lower confidence limits (BMDLs) were
      {bmd1} ({bmdl1}), ..., respectively. In female rats, there were
      no apical endpoints for which a BMD value could be reliably
      estimated.

    Strategy:
      1. Lead sentence reports which platform categories had reliable
         BMDs (e.g., "clinical pathology and organ weight" or just
         "organ weight"), or omits the lead if no reliable BMDs exist
         in either sex.
      2. For each sex (in canonical Male, Female order), filter to
         entries with reliable BMDs and a known direction, sort by
         ascending BMD, and emit a "the effects included ..." sentence
         followed by the BMD/BMDL listing.
      3. If a sex has no reliable BMDs, emit the standard "no apical
         endpoints" fallback sentence.

    Args:
        apical_bmd_summary: list of {endpoint, sex, platform, bmd, bmdl,
                            bmd_status, loel, noel, direction} dicts —
                            same structure as data["bmd_summary"]["endpoints"].
        sexes: optional list of sex labels to include (defaults to
               ["Male", "Female"]).  Order is preserved in output.

    Returns:
        A single paragraph string ready to insert into the Abstract.
    """
    if sexes is None:
        sexes = ["Male", "Female"]

    # --- De-duplicate ---
    # The cache may carry duplicate rows where one has BMD values and
    # the other has "—" placeholders.  Keep only the rows with reliable
    # BMD values, dropping the placeholder duplicates.
    reliable = [e for e in apical_bmd_summary if _is_reliable_bmd(e)]

    # --- Anomalous-BMD filter ---
    # Exclude entries whose curve-fit BMD is implausibly low compared to
    # the statistically-observed NOEL/LOEL.  These are model artifacts,
    # not real potency estimates — the reference report excludes them
    # from the Abstract effects list and instead calls them out as
    # anomalous in the body Results section.
    reliable = [e for e in reliable if not _is_anomalous_bmd(e)]

    # --- Lead sentence: which platform categories had findings? ---
    # Group reliable entries by platform → human-readable category name
    PLATFORM_TO_CATEGORY = {
        "Clinical Chemistry": "clinical pathology",
        "Hematology":         "clinical pathology",
        "Hormones":           "clinical pathology",
        "Organ Weight":       "organ weight",
        "Body Weight":        "body weight",
    }
    categories: list[str] = []
    seen: set[str] = set()
    # Order categories by first appearance among reliable entries sorted
    # by sex then platform — gives stable output across runs.
    for entry in sorted(reliable, key=lambda e: (e.get("sex", ""), e.get("platform", ""))):
        cat = PLATFORM_TO_CATEGORY.get(entry.get("platform", ""))
        if cat and cat not in seen:
            categories.append(cat)
            seen.add(cat)

    sentences: list[str] = []
    if categories:
        cat_phrase = _join_oxford(categories)
        sentences.append(
            f"Several {cat_phrase} measurements showed dose-related changes "
            f"from which BMD values were calculated."
        )

    # --- Per-sex apical findings ---
    for sex in sexes:
        sex_entries = [e for e in reliable if e.get("sex") == sex]
        # Filter to entries with a known direction (non-empty)
        sex_entries = [e for e in sex_entries if e.get("direction")]

        if not sex_entries:
            sentences.append(
                f"In {sex.lower()} rats, there were no apical endpoints "
                f"for which a BMD value could be reliably estimated."
            )
            continue

        # Sort by ascending BMD (most sensitive first — matches NIEHS reference)
        sex_entries.sort(key=lambda e: float(e["bmd"]))

        # Build the descriptor phrases and the parallel BMD (BMDL) list
        descriptors = [
            _format_endpoint_phrase(
                e.get("endpoint", ""),
                e.get("direction", ""),
                e.get("platform", ""),
            )
            for e in sex_entries
        ]
        bmd_pairs = [_format_bmd_pair(e["bmd"], e.get("bmdl")) for e in sex_entries]

        # Differentiate "BMD" vs "BMDs" if singular
        plural = len(sex_entries) > 1
        sentences.append(
            f"In {sex.lower()} rats, the effects included "
            f"significantly {_join_oxford(descriptors)}. "
            f"The BMD{'s' if plural else ''} and benchmark dose lower "
            f"confidence limit{'s' if plural else ''} (BMDL{'s' if plural else ''}) "
            f"{'were' if plural else 'was'} {_join_oxford(bmd_pairs)}, "
            f"{'respectively' if plural else ''}".rstrip(", ") + "."
        )

    return " ".join(sentences)


# ---------------------------------------------------------------------------
# Genomics: most sensitive gene sets and genes per organ × sex
# ---------------------------------------------------------------------------

def _format_rat_gene_symbol(symbol: str) -> str:
    """
    Convert an uppercase gene symbol (e.g., "PRLR") into NIEHS rat-gene
    convention (e.g., "Prlr") — first letter capitalized, rest lowercase,
    with embedded slashes and digits preserved.

    Examples:
      "PRLR"                  → "Prlr"
      "GSTA2"                 → "Gsta2"
      "LOC100911545/A2M"      → "Loc100911545/A2m"
      "CYP7A1"                → "Cyp7a1"
    """
    if not symbol:
        return symbol
    # Slash-separated multi-symbols: format each segment independently
    if "/" in symbol:
        return "/".join(_format_rat_gene_symbol(s) for s in symbol.split("/"))
    s = symbol.strip()
    if not s:
        return s
    return s[0].upper() + s[1:].lower()


def _stat_display_name(stat_key: str) -> str:
    """
    Convert a BMD aggregation stat key (e.g., "median", "fifth_pct") into
    the prose word used in the abstract sentence ("median", "fifth percentile").
    """
    return {
        "median":     "median",
        "mean":       "mean",
        "fifth_pct":  "fifth percentile",
        "minimum":    "minimum",
        "maximum":    "maximum",
    }.get(stat_key, stat_key.replace("_", " "))


def _format_dose_value(v) -> str:
    """
    Format a dose/BMD value preserving sensible precision.  Drops trailing
    zeros while keeping at most 3 decimal places (matches NIEHS reference
    e.g., "0.520", "5.725", "1,000").  NaN/inf values render as the dash
    placeholder used elsewhere for missing data.
    """
    import math

    if v is None:
        return "—"
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    if math.isnan(f) or math.isinf(f):
        return "—"
    if f >= 1000 and f == int(f):
        return f"{int(f):,}"
    if f == int(f):
        return str(int(f))
    # Round to 3 decimal places, strip trailing zeros, restore at least one
    s = f"{f:.3f}".rstrip("0").rstrip(".")
    return s if s else "0"


def _picks_above_lle(items: list[dict], lle: float, n: int) -> list[dict]:
    """
    Return up to N items sorted by ascending BMD, filtered to those with
    BMD >= lower-limit-of-extrapolation (LLE).  The reference report
    excludes anything below LLE from the "most sensitive" lists.

    Items with NaN/inf BMD or BMDL are treated as unreliable and dropped
    — these come from failed model fits where curve coefficients didn't
    converge.
    """
    import math

    reliable: list[dict] = []
    for item in items:
        bmd = item.get("bmd")
        bmdl = item.get("bmdl")
        if bmd is None:
            continue
        try:
            bmd_f = float(bmd)
        except (TypeError, ValueError):
            continue
        if math.isnan(bmd_f) or math.isinf(bmd_f) or bmd_f < lle:
            continue
        # Also filter NaN BMDLs — a NaN BMDL means the lower confidence
        # limit didn't converge, which the reference treats as unreliable.
        try:
            bmdl_f = float(bmdl) if bmdl is not None else None
            if bmdl_f is not None and (math.isnan(bmdl_f) or math.isinf(bmdl_f)):
                continue
        except (TypeError, ValueError):
            pass
        reliable.append({**item, "_bmd_float": bmd_f})

    reliable.sort(key=lambda x: x["_bmd_float"])
    return reliable[:n]


def _build_gene_sets_sentence(
    sets: list[dict], sex: str, stat_label: str, dose_unit: str,
) -> str:
    """
    Build one "most sensitive gene sets" sentence for a given sex.

    Reference pattern:
      "The most sensitive gene sets in {sex} rats for which a reliable
       estimate of the BMD could be made were {GO terms with Oxford
       comma} with {stat} BMDs of {values} {unit} and {stat} BMDLs of
       {values} {unit}, respectively."

    Returns "" when the input list is empty (caller decides whether to
    emit a "no reliable BMD" fallback).
    """
    if not sets:
        return ""

    terms = [s.get("go_term", s.get("go_id", "(unknown)")) for s in sets]
    bmds = [_format_dose_value(s.get("bmd")) for s in sets]
    bmdls = [_format_dose_value(s.get("bmdl")) for s in sets]

    plural = len(sets) > 1
    return (
        f"The most sensitive gene set{'s' if plural else ''} in "
        f"{sex.lower()} rats for which a reliable estimate of the "
        f"BMD could be made {'were' if plural else 'was'} "
        f"{_join_oxford(terms)} with {stat_label} BMD"
        f"{'s' if plural else ''} of {_join_oxford(bmds)} {dose_unit} "
        f"and {stat_label} BMDL{'s' if plural else ''} of "
        f"{_join_oxford(bmdls)} {dose_unit}, respectively."
    )


def _build_top_genes_sentence(
    genes: list[dict], sex: str, direction: str, dose_unit: str,
) -> str:
    """
    Build one "most sensitive up/down-regulated genes" sentence.

    Reference pattern:
      "The most sensitive {up/down}regulated genes in {sex} rats with
       reliable BMD estimates included {gene symbols} with BMDs (BMDLs)
       of {pairs}, respectively."

    Returns "" when no genes match the direction filter.
    """
    if not genes:
        return ""

    symbols = [_format_rat_gene_symbol(g.get("gene_symbol", "")) for g in genes]
    pairs = [
        f"{_format_dose_value(g.get('bmd'))} ({_format_dose_value(g.get('bmdl'))})"
        for g in genes
    ]

    plural = len(genes) > 1
    return (
        f"The most sensitive {direction}regulated gene"
        f"{'s' if plural else ''} in {sex.lower()} rats with reliable "
        f"BMD estimate{'s' if plural else ''} "
        f"{'included' if plural else 'was'} {_join_oxford(symbols)} "
        f"with BMD{'s' if plural else ''} (BMDL{'s' if plural else ''}) "
        f"of {_join_oxford(pairs)} {dose_unit}, respectively."
    )


def build_abstract_results_genomics(
    genomics_sections: dict,
    dose_groups: list[float],
    dose_unit: str = "mg/kg",
    bmd_stat: str | None = None,
    n_top_sets: int = 3,
    n_top_genes: int = 8,
    sexes: list[str] | None = None,
) -> str:
    """
    Build the genomics portion of the Abstract → Results paragraph.

    For each organ (alphabetical order, matching reference convention),
    emits:
      1. A lower-limit-of-extrapolation summary sentence: either "no GO
         process or individual genes had BMD median values below
         <{LLE} {unit}>" or a count of items below.
      2. Per sex: the most sensitive gene sets sentence (up to n_top_sets).
      3. Per sex: the most sensitive up-regulated genes sentence
         (up to n_top_genes).
      4. Per sex: the most sensitive down-regulated genes sentence
         (up to n_top_genes).

    Empty sentences are omitted (e.g., a sex with no reliable gene sets
    silently drops that sentence).

    Args:
        genomics_sections: dict keyed by "{organ}_{sex}" (e.g., "liver_male"),
                           each value with gene_sets_by_stat and top_genes.
        dose_groups:       The full study dose list — used to compute the
                           lower limit of extrapolation (lowest non-zero / 3).
        dose_unit:         Display unit (default "mg/kg").
        bmd_stat:          Which stat to use ("median", "fifth_pct", etc.).
                           Defaults to whichever stat is present in the data.
        n_top_sets:        How many top gene sets to list per sex.
        n_top_genes:       How many top up/down-regulated genes per sex.
        sexes:             Sex order (defaults to ["Male", "Female"]).

    Returns:
        Paragraph string, or empty string if no genomics data is present.
    """
    if not genomics_sections:
        return ""

    sexes = sexes or ["Male", "Female"]

    # Lower limit of extrapolation = lowest non-zero dose / 3 (BMDExpress convention)
    nonzero_doses = [d for d in dose_groups if d and d > 0]
    if not nonzero_doses:
        return ""
    lle = min(nonzero_doses) / 3.0
    lle_str = _format_dose_value(lle)

    # Identify all organs from the section keys (e.g. "liver_male" → "liver")
    organs: set[str] = set()
    for key in genomics_sections:
        if "_" in key:
            organs.add(key.split("_", 1)[0])
    organs_sorted = sorted(organs)

    sentences: list[str] = []

    for organ in organs_sorted:
        # --- Pick which BMD stat to read (default to whatever exists) ---
        sample_section = None
        for sex in sexes:
            sec = genomics_sections.get(f"{organ}_{sex.lower()}")
            if sec and sec.get("gene_sets_by_stat"):
                sample_section = sec
                break
        if not sample_section:
            continue

        available_stats = list(sample_section["gene_sets_by_stat"].keys())
        chosen_stat = bmd_stat if bmd_stat in available_stats else (available_stats[0] if available_stats else None)
        if not chosen_stat:
            continue
        stat_label = _stat_display_name(chosen_stat)

        # --- Lower-limit-of-extrapolation summary across both sexes ---
        # Count how many gene sets and individual genes have BMD < LLE
        # across all sexes for this organ.
        below_lle_sets = 0
        below_lle_genes = 0
        for sex in sexes:
            sec = genomics_sections.get(f"{organ}_{sex.lower()}", {})
            if not sec:
                continue
            sets = sec.get("gene_sets_by_stat", {}).get(chosen_stat, [])
            for s in sets:
                bmd = s.get("bmd")
                try:
                    if bmd is not None and float(bmd) < lle:
                        below_lle_sets += 1
                except (TypeError, ValueError):
                    pass
            for g in sec.get("top_genes", []):
                bmd = g.get("bmd")
                try:
                    if bmd is not None and float(bmd) < lle:
                        below_lle_genes += 1
                except (TypeError, ValueError):
                    pass

        sex_phrase = " and ".join(s.lower() for s in sexes) + " rats"
        if below_lle_sets == 0 and below_lle_genes == 0:
            sentences.append(
                f"In the {organ} of {sex_phrase}, no Gene Ontology "
                f"biological process or individual genes had BMD "
                f"{stat_label} values below the lower limit of "
                f"extrapolation (<{lle_str} {dose_unit})."
            )
        else:
            # Plural-aware count phrase
            sets_word = "gene set" if below_lle_sets == 1 else "gene sets"
            genes_word = "gene" if below_lle_genes == 1 else "genes"
            sentences.append(
                f"In the {organ} of {sex_phrase}, "
                f"{below_lle_sets} Gene Ontology biological process "
                f"{sets_word} and {below_lle_genes} individual "
                f"{genes_word} had BMD {stat_label} values below the "
                f"lower limit of extrapolation (<{lle_str} {dose_unit})."
            )

        # --- Per-sex gene sets and top genes ---
        for sex in sexes:
            sec = genomics_sections.get(f"{organ}_{sex.lower()}", {})
            if not sec:
                continue

            # Top gene sets (above LLE, sorted by BMD ascending)
            sets = sec.get("gene_sets_by_stat", {}).get(chosen_stat, [])
            top_sets = _picks_above_lle(sets, lle, n_top_sets)
            sets_sentence = _build_gene_sets_sentence(top_sets, sex, stat_label, dose_unit)
            if sets_sentence:
                sentences.append(sets_sentence)

            # Top up-regulated and down-regulated genes
            top_genes = sec.get("top_genes", [])
            up_genes = [g for g in top_genes if str(g.get("direction", "")).lower() == "up"]
            down_genes = [g for g in top_genes if str(g.get("direction", "")).lower() == "down"]
            up_top = _picks_above_lle(up_genes, lle, n_top_genes)
            down_top = _picks_above_lle(down_genes, lle, n_top_genes)

            up_sentence = _build_top_genes_sentence(up_top, sex, "up", dose_unit)
            if up_sentence:
                sentences.append(up_sentence)
            down_sentence = _build_top_genes_sentence(down_top, sex, "down", dose_unit)
            if down_sentence:
                sentences.append(down_sentence)

    return " ".join(sentences)


def build_abstract_results_pk(
    chemical_name: str,
    pk_concentrations: dict | None,
    pk_half_lives: dict | None,
    pk_timepoints: list[int],
    dose_unit: str = "mg/kg",
    sexes: list[str] | None = None,
) -> str:
    """
    Build the pharmacokinetic portion of the Abstract → Results paragraph.

    Reference pattern:
      "Average {TA} plasma concentrations at {2 and 24} hours postdose
       were {lower/higher} in {sex_a} rats than in {sex_b} rats. Half-
       lives estimated using the two time points were {longer/shorter}
       in {sex_b} rats ({78.2 and 25.6} hours for the {4 and 37} {unit}
       groups, respectively) than in {sex_a} rats ({40.1 and 15.1}
       hours for the {4 and 37} {unit} groups, respectively)."

    Strategy:
      1. Compute total exposure (sum of mean concentrations across all
         timepoints and biosampling doses) per sex to decide the sentence
         polarity (which sex had lower/higher concentrations).
      2. Format the timepoint list ("at 2 and 24 hours postdose").
      3. Build the half-life comparison sentence with values aligned to
         the same dose order.

    Returns "" when there are insufficient data (no two-timepoint half-
    lives, or only one sex represented).

    Args:
        chemical_name:    The test article name (used in the concentration sentence).
        pk_concentrations: {sex: {dose: {hour: mean_value}}} from MethodsContext.
        pk_half_lives:    {sex: {dose: hours_float}} from MethodsContext.
        pk_timepoints:    Sorted list of timepoint hours seen in the data.
        dose_unit:        Display unit (default "mg/kg").
        sexes:            Sex labels (defaults to ["Male", "Female"]).

    Returns:
        A paragraph string, or empty string if data is insufficient.
    """
    if not pk_concentrations or not pk_half_lives:
        return ""
    if not pk_timepoints:
        return ""

    sexes = sexes or ["Male", "Female"]

    # Need both sexes for a comparison sentence
    sexes_present = [s for s in sexes if s in pk_concentrations and s in pk_half_lives]
    if len(sexes_present) < 2:
        return ""

    sex_a, sex_b = sexes_present[0], sexes_present[1]

    # --- Total mean concentrations per sex (sum across dose × timepoint) ---
    # Used only to decide which sex had "lower" vs "higher" concentrations.
    def _total(sex: str) -> float:
        total = 0.0
        for dose, by_hour in pk_concentrations.get(sex, {}).items():
            for v in by_hour.values():
                total += v
        return total

    total_a = _total(sex_a)
    total_b = _total(sex_b)
    if total_a < total_b:
        conc_lower, conc_higher = sex_a, sex_b
        conc_polarity = "lower"
    else:
        conc_lower, conc_higher = sex_b, sex_a
        conc_polarity = "lower"  # phrasing always uses "lower"

    # --- Timepoints sentence ("at 2 and 24 hours postdose") ---
    tp_phrase = _join_oxford([str(t) for t in pk_timepoints])

    # --- Half-life comparison ---
    # Find the union of doses where both sexes have a half-life
    common_doses = sorted(
        set(pk_half_lives[sex_a].keys()) & set(pk_half_lives[sex_b].keys())
    )
    if not common_doses:
        return ""

    # Decide polarity: which sex has the longer mean half-life?
    def _mean_half_life(sex: str) -> float:
        vals = [pk_half_lives[sex][d] for d in common_doses if d in pk_half_lives[sex]]
        return sum(vals) / len(vals) if vals else 0.0

    if _mean_half_life(sex_a) > _mean_half_life(sex_b):
        hl_longer, hl_shorter = sex_a, sex_b
    else:
        hl_longer, hl_shorter = sex_b, sex_a

    # Format dose group list and the matched half-life lists
    dose_list = _join_oxford([_format_dose_value(d) for d in common_doses])

    def _hl_list(sex: str) -> str:
        # NIEHS reference uses one decimal for half-lives (78.2, 25.6, etc.)
        vals = [pk_half_lives[sex][d] for d in common_doses]
        formatted = [f"{v:.1f}" for v in vals]
        return _join_oxford(formatted)

    hl_longer_str = _hl_list(hl_longer)
    hl_shorter_str = _hl_list(hl_shorter)

    # --- Assemble ---
    # The chemical name appears in the first sentence; subsequent sentences
    # use "rats" alone since the test article is implicit.
    return (
        f"Average {chemical_name} plasma concentrations at "
        f"{tp_phrase} hours postdose were {conc_polarity} in "
        f"{conc_lower.lower()} rats than in {conc_higher.lower()} rats. "
        f"Half-lives estimated using the two time points were longer in "
        f"{hl_longer.lower()} rats ({hl_longer_str} hours for the "
        f"{dose_list} {dose_unit} groups, respectively) than in "
        f"{hl_shorter.lower()} rats ({hl_shorter_str} hours for the "
        f"{dose_list} {dose_unit} groups, respectively)."
    )


# ---------------------------------------------------------------------------
# Body Results: Gene Set BMD Analysis prose
# ---------------------------------------------------------------------------
# The body Results section has two genomics blocks (Gene Set, Gene BMD)
# each preceded by ~2 paragraphs of boilerplate framing followed by
# per-organ paragraphs of findings.  These are deterministic — values
# come from MethodsContext + the genomics cache and the prose follows a
# fixed NIEHS-Report-10 sentence skeleton.

def _format_organ_phrase(organs: list[str]) -> str:
    """
    "liver and kidney" / "liver, kidney, and spleen" / "liver".

    Lowercases organ names since they appear mid-sentence.  Identical
    behavior to _format_organ_list — kept as a separate helper so the
    body-narrative builders can evolve independently.
    """
    return _format_organ_list(organs)


def _normalize_organ_name(organ: str) -> str:
    """
    Normalize an organ key (e.g., 'liver') into title case for headings
    (e.g., 'Liver'), and into lowercase for mid-sentence prose.  This
    helper just returns the lowercase form; callers that need title
    case should capitalize themselves.
    """
    return (organ or "").lower()


def _format_paired_bmd_pairs(items: list[dict]) -> str:
    """
    Format a list of {bmd, bmdl} entries as a NIEHS-style "BMD (BMDL)"
    sequence joined with the Oxford comma:
      "0.520 (0.160) and 0.750 (0.186)"
      "5.725 (1.686), 7.423 (5.757), and 8.417 (7.129)"
    """
    pairs = [
        f"{_format_dose_value(it.get('bmd'))} ({_format_dose_value(it.get('bmdl'))})"
        for it in items
    ]
    return _join_oxford(pairs)


def build_gene_set_body_intro(
    chemical_name: str,
    ge_organs: list[str],
    table_numbers: list[int] | None = None,
) -> list[str]:
    """
    Build the two boilerplate intro paragraphs for the Gene Set BMD
    Analysis section, matching NIEHS Report 10 page 19.

    Paragraph 1 — Methodology framing:
      "Chemical-induced alterations in {organs} gene transcript expression
       were examined to determine those gene sets most sensitive to {TA}
       exposure. To that end, BMD analysis of transcripts and gene sets
       (Gene Ontology [GO] biological process) was conducted to determine
       the potency of the chemical to elicit gene expression changes in
       the {organs}. This analysis used transcript-level BMD data to
       assess an aggregate score of gene set potency (median transcript
       BMD) and enrichment."

    Paragraph 2 — Interpretation caveat:
      "The 'active' gene sets in the {organs} with the lowest BMD median
       values are shown in {Table N} and {Table N+1}, respectively. The
       gene sets in {Tables} should be interpreted with caution from the
       standpoint of the underlying biological mechanism..."

    Args:
        chemical_name:  Test article name, used in paragraph 1.
        ge_organs:      Organ list from MethodsContext.ge_organs.
        table_numbers:  Auto-assigned table numbers for the per-organ
                        gene set tables (e.g., [9, 10] for Tables 9 & 10).
                        When omitted or short, uses generic "the tables".

    Returns:
        Two-paragraph list ready to inject into data.gene_set_narrative.
    """
    organs_phrase = _format_organ_phrase(ge_organs) or "the assayed tissues"

    # Methodology paragraph
    p1 = (
        f"Chemical-induced alterations in {organs_phrase} gene transcript "
        f"expression were examined to determine those gene sets most "
        f"sensitive to {chemical_name} exposure. To that end, BMD analysis "
        f"of transcripts and gene sets (Gene Ontology [GO] biological "
        f"process) was conducted to determine the potency of the chemical "
        f"to elicit gene expression changes in the {organs_phrase}. This "
        f"analysis used transcript-level BMD data to assess an aggregate "
        f"score of gene set potency (median transcript BMD) and enrichment."
    )

    # Interpretation caveat — table refs use the auto-assigned numbers
    # when available, else a generic "the tables below" fallback.
    if table_numbers and len(table_numbers) >= 1:
        table_refs = [f"Table {n}" for n in table_numbers]
        if len(table_refs) == 1:
            tables_str = table_refs[0]
            tables_str_repeat = table_refs[0]
        else:
            tables_str = _join_oxford(table_refs) + ", respectively"
            tables_str_repeat = _join_oxford(table_refs)
    else:
        tables_str = "the tables below"
        tables_str_repeat = "the tables"

    p2 = (
        f"The “active” gene sets in the {organs_phrase} with the "
        f"lowest BMD median values are shown in {tables_str}. The gene "
        f"sets in {tables_str_repeat} should be interpreted with caution "
        f"from the standpoint of the underlying biological mechanism and "
        f"any relationship to toxicity or toxic agents referenced in the "
        f"GO term definitions. The data primarily should be considered a "
        f"metric of potency for chemical-induced transcriptional changes "
        f"(i.e., a concerted biological change) that could serve as a "
        f"surrogate of estimated biological potency and, by extension, "
        f"toxicological potency when more definitive toxicological data "
        f"are unavailable."
    )

    return [p1, p2]


def build_gene_set_body_findings(
    genomics_sections: dict,
    dose_groups: list[float],
    dose_unit: str = "mg/kg",
    bmd_stat: str | None = None,
    n_top: int = 2,
    sexes: list[str] | None = None,
) -> list[str]:
    """
    Build the per-organ "findings" paragraphs for the Gene Set BMD
    Analysis section body.

    Per organ (alphabetical), one paragraph composed of:
      - Lower-limit-of-extrapolation check, scoped to gene sets only:
          "No gene sets in the {organ} of male or female rats had
           estimated BMD median values <{LLE} {unit}."
      - One sub-clause per sex describing the most sensitive gene sets,
        with GO IDs in parens, BMDs and BMDLs paired:
          "In male rats, the most sensitive GO biological processes for
           which a BMD value could be reliably calculated were
           {GO term} ({GO ID}) and {GO term} ({GO ID}) with median BMDs
           (BMDLs) of {bmd1} ({bmdl1}) and {bmd2} ({bmdl2}) {unit},
           respectively."

    Differs from build_abstract_results_genomics in:
      - Single-organ scope per paragraph (vs. abstract's combined per-organ)
      - Includes GO IDs in parens after each gene set name
      - Uses paired "BMDs (BMDLs) of X (Y) and Z (W) mg/kg" format
        instead of separated "BMDs of X and Z and BMDLs of Y and W"
    """
    if not genomics_sections:
        return []

    sexes = sexes or ["Male", "Female"]

    nonzero = [d for d in (dose_groups or []) if d and d > 0]
    if not nonzero:
        return []
    lle = min(nonzero) / 3.0
    lle_str = _format_dose_value(lle)

    # Pick the BMD stat — same logic as abstract genomics builder
    chosen_stat = bmd_stat
    if not chosen_stat:
        for sec in genomics_sections.values():
            if sec and sec.get("gene_sets_by_stat"):
                stats = list(sec["gene_sets_by_stat"].keys())
                if stats:
                    chosen_stat = stats[0]
                    break
    if not chosen_stat:
        return []
    stat_label = _stat_display_name(chosen_stat)

    # Walk organs alphabetically
    organs = sorted({k.split("_", 1)[0] for k in genomics_sections if "_" in k})
    paragraphs: list[str] = []

    for organ in organs:
        sentences: list[str] = []

        # LLE-scoped-to-gene-sets check across both sexes for this organ
        below_lle = 0
        for sex in sexes:
            sec = genomics_sections.get(f"{organ}_{sex.lower()}", {})
            if not sec:
                continue
            sets = sec.get("gene_sets_by_stat", {}).get(chosen_stat, [])
            for s in sets:
                bmd = s.get("bmd")
                try:
                    if bmd is not None and float(bmd) < lle:
                        below_lle += 1
                except (TypeError, ValueError):
                    pass

        sex_phrase = " and ".join(s.lower() for s in sexes) + " rats"
        if below_lle == 0:
            sentences.append(
                f"No gene sets in the {organ} of {sex_phrase} had "
                f"estimated BMD {stat_label} values <{lle_str} {dose_unit}."
            )
        else:
            word = "gene set" if below_lle == 1 else "gene sets"
            sentences.append(
                f"In the {organ} of {sex_phrase}, {below_lle} {word} had "
                f"estimated BMD {stat_label} values <{lle_str} {dose_unit}."
            )

        # Per-sex findings clauses
        for sex in sexes:
            sec = genomics_sections.get(f"{organ}_{sex.lower()}", {})
            if not sec:
                continue
            sets = sec.get("gene_sets_by_stat", {}).get(chosen_stat, [])
            top = _picks_above_lle(sets, lle, n_top)
            if not top:
                sentences.append(
                    f"In {sex.lower()} rats, no GO biological processes "
                    f"had a reliable BMD estimate above the lower limit of "
                    f"extrapolation."
                )
                continue

            # "{name} ({GO ID})" descriptors, joined with Oxford.
            # Omit the parens entirely when the GO ID is missing — never
            # leave a trailing "(" or ")".
            descriptors = []
            for s in top:
                term = s.get("go_term", "(unknown)")
                go_id = (s.get("go_id") or "").strip()
                if go_id:
                    descriptors.append(f"{term} ({go_id})")
                else:
                    descriptors.append(term)
            pairs = _format_paired_bmd_pairs(top)

            plural = len(top) > 1
            sentences.append(
                f"In {sex.lower()} rats, the most sensitive GO biological "
                f"process{'es' if plural else ''} for which a BMD value "
                f"could be reliably calculated "
                f"{'were' if plural else 'was'} {_join_oxford(descriptors)} "
                f"with {stat_label} BMD{'s' if plural else ''} (BMDL"
                f"{'s' if plural else ''}) of {pairs} {dose_unit}"
                f"{', respectively' if plural else ''}."
            )

        paragraphs.append(" ".join(sentences))

    return paragraphs


# ---------------------------------------------------------------------------
# Body Results: Gene BMD Analysis prose
# ---------------------------------------------------------------------------

def build_gene_body_intro(
    ge_organs: list[str],
    table_numbers: list[int] | None = None,
    fold_change_filter: float | None = None,
    bmdu_bmdl_ratio: float | None = 40.0,
    fit_pvalue_threshold: float | None = 0.1,
) -> list[str]:
    """
    Build the two boilerplate intro paragraphs for the Gene BMD Analysis
    section body, matching NIEHS Report 10 page 26.

    Paragraph 1 — Methodology + filter values:
      "The top 10 genes based on BMD potency in the {organs} (fold change
       >|{fc}|, significant Williams trend test, global goodness-of-fit p
       value >{p}, and BMDU/BMDL ≤{ratio}) are shown in {Table N} and
       {Table N+1}."

    Paragraph 2 — Interpretation caveat:
      "As with the GO analysis, the biological or toxicological
       significance of the changes in gene expression shown in {Tables}
       should be interpreted with caution. The data primarily should be
       considered a metric of potency..."

    Filter values default to the NIEHS Report 10 reference settings
    (|2|, p > 0.1, BMDU/BMDL ≤ 40) when not provided by MethodsContext.
    """
    organs_phrase = _format_organ_phrase(ge_organs) or "the assayed tissues"

    fc = fold_change_filter if fold_change_filter is not None else 2
    p = fit_pvalue_threshold if fit_pvalue_threshold is not None else 0.1
    ratio = bmdu_bmdl_ratio if bmdu_bmdl_ratio is not None else 40

    # Format filter values: drop unnecessary decimals
    def _fmt(v: float) -> str:
        if v == int(v):
            return str(int(v))
        return f"{v:g}"

    if table_numbers and len(table_numbers) >= 1:
        table_refs = [f"Table {n}" for n in table_numbers]
        tables_str = _join_oxford(table_refs)
    else:
        tables_str = "the tables below"

    p1 = (
        f"The top 10 genes based on BMD potency in the {organs_phrase} "
        f"(fold change >|{_fmt(fc)}|, significant Williams trend test, "
        f"global goodness-of-fit p value >{_fmt(p)}, and BMDU/BMDL "
        f"≤{_fmt(ratio)}) are shown in {tables_str}."
    )

    p2 = (
        f"As with the GO analysis, the biological or toxicological "
        f"significance of the changes in gene expression shown in "
        f"{tables_str} should be interpreted with caution. The data "
        f"primarily should be considered a metric of potency for "
        f"chemical-induced transcriptional changes that could serve as a "
        f"conservative surrogate of estimated biological potency, and by "
        f"extension toxicological potency, when more definitive "
        f"toxicological data are unavailable."
    )

    return [p1, p2]


def build_gene_body_findings(
    genomics_sections: dict,
    dose_groups: list[float],
    dose_unit: str = "mg/kg",
    n_top: int = 8,
    sexes: list[str] | None = None,
) -> list[str]:
    """
    Build the per-organ × per-sex "findings" paragraphs for the Gene BMD
    Analysis section body.

    Per organ (alphabetical), one paragraph composed of:
      - Lower-limit-of-extrapolation check, scoped to genes only.
      - For each sex, separate clauses for upregulated and downregulated
        most-sensitive genes:
          "In male rats, the most sensitive upregulated genes with a
           calculated BMD were {Gsta2}, {Gsta5}, ... with BMDs (BMDLs) of
           {x} ({y}), {x} ({y}), ... {unit}, respectively."
          "The most sensitive genes exhibiting a decrease in expression
           were {Egr1}, ... with BMDs (BMDLs) of {x} ({y}), ..."

    Gene name expansions ("Gsta2 (glutathione S-transferase alpha 2)")
    require an external annotation source not present in the genomics
    cache, so this function emits bare gene symbols.  Future enhancement:
    look up gene names from integrated.json's referenceGeneAnnotations
    or from bmdx.duckdb.
    """
    if not genomics_sections:
        return []

    sexes = sexes or ["Male", "Female"]

    nonzero = [d for d in (dose_groups or []) if d and d > 0]
    if not nonzero:
        return []
    lle = min(nonzero) / 3.0
    lle_str = _format_dose_value(lle)

    organs = sorted({k.split("_", 1)[0] for k in genomics_sections if "_" in k})
    paragraphs: list[str] = []

    for organ in organs:
        sentences: list[str] = []

        # LLE-scoped-to-genes check across both sexes
        below_lle = 0
        for sex in sexes:
            sec = genomics_sections.get(f"{organ}_{sex.lower()}", {})
            if not sec:
                continue
            for g in sec.get("top_genes", []):
                bmd = g.get("bmd")
                try:
                    if bmd is not None and float(bmd) < lle:
                        below_lle += 1
                except (TypeError, ValueError):
                    pass

        sex_phrase = " and ".join(s.lower() for s in sexes) + " rats"
        if below_lle == 0:
            sentences.append(
                f"No {organ} genes in {sex_phrase} had estimated BMD "
                f"median values <{lle_str} {dose_unit}."
            )
        else:
            word = "gene" if below_lle == 1 else "genes"
            sentences.append(
                f"In the {organ} of {sex_phrase}, {below_lle} {word} had "
                f"estimated BMD median values <{lle_str} {dose_unit}."
            )

        # Per-sex up/down clauses
        for sex in sexes:
            sec = genomics_sections.get(f"{organ}_{sex.lower()}", {})
            if not sec:
                continue
            top_all = sec.get("top_genes", [])
            up = [g for g in top_all if str(g.get("direction", "")).lower() == "up"]
            down = [g for g in top_all if str(g.get("direction", "")).lower() == "down"]
            up_top = _picks_above_lle(up, lle, n_top)
            down_top = _picks_above_lle(down, lle, n_top)

            for direction_kind, items in (("upregulated", up_top), ("decrease in expression", down_top)):
                if not items:
                    continue
                symbols = [_format_rat_gene_symbol(g.get("gene_symbol", "")) for g in items]
                pairs = _format_paired_bmd_pairs(items)
                plural = len(items) > 1

                if direction_kind == "upregulated":
                    sentences.append(
                        f"In {sex.lower()} rats, the most sensitive "
                        f"upregulated gene{'s' if plural else ''} with a "
                        f"calculated BMD "
                        f"{'were' if plural else 'was'} "
                        f"{_join_oxford(symbols)} with BMD"
                        f"{'s' if plural else ''} (BMDL"
                        f"{'s' if plural else ''}) of {pairs} {dose_unit}"
                        f"{', respectively' if plural else ''}."
                    )
                else:
                    sentences.append(
                        f"The most sensitive gene{'s' if plural else ''} "
                        f"exhibiting a decrease in expression in "
                        f"{sex.lower()} rats "
                        f"{'were' if plural else 'was'} "
                        f"{_join_oxford(symbols)} with BMD"
                        f"{'s' if plural else ''} (BMDL"
                        f"{'s' if plural else ''}) of {pairs} {dose_unit}"
                        f"{', respectively' if plural else ''}."
                    )

        paragraphs.append(" ".join(sentences))

    return paragraphs


def build_abstract_summary(
    apical_bmd_summary: list[dict] | None = None,
    genomics_sections: dict | None = None,
    dose_groups: list[float] | None = None,
    dose_unit: str = "mg/kg",
    bmd_stat: str | None = None,
    sexes: list[str] | None = None,
) -> str:
    """
    Build the Abstract → Summary paragraph.

    Reference pattern (NIEHS Report 10):
      "Taken together, in male rats, the most sensitive gene set BMD
       (BMDL) median, individual gene BMD (BMDL), and apical endpoint
       BMD (BMDL) values that could be reliably determined occurred at
       0.520 (0.160), 0.510 (0.212), and 7.264 (5.024) mg/kg,
       respectively. In female rats, the most sensitive gene set BMD
       (BMDL) median and individual gene BMD (BMDL) values that could
       be reliably determined occurred at 10.324 (7.461) and 1.163
       (0.179) mg/kg, respectively. There were no apical endpoints in
       female rats for which a BMD value could be reliably estimated."

    Strategy: for each sex, find the lowest reliable BMD across all
    organs in three categories — gene sets, individual genes, apical
    endpoints — and assemble a sentence whose category list is gated
    on availability.  When apical fails entirely for a sex, we add the
    "no apical endpoints" fallback sentence.

    Args:
        apical_bmd_summary: list of BMD summary entry dicts.
        genomics_sections:  dict keyed by "{organ}_{sex}" with gene_sets and top_genes.
        dose_groups:        Full study dose list (for lower-limit-of-extrapolation).
        dose_unit:          Display unit (default "mg/kg").
        bmd_stat:           Which stat to use for gene sets ("median", "fifth_pct").
        sexes:              Optional sex order (defaults to ["Male", "Female"]).

    Returns:
        Paragraph string, or empty string if no reliable BMDs exist anywhere.
    """
    sexes = sexes or ["Male", "Female"]

    # Lower limit of extrapolation (LLE) — reuse same convention as Results
    nonzero_doses = [d for d in (dose_groups or []) if d and d > 0]
    lle = (min(nonzero_doses) / 3.0) if nonzero_doses else 0.0

    # Determine which BMD stat to read from gene_sets_by_stat.  Default
    # to whatever the data carries (matches the Results paragraph logic).
    chosen_stat = bmd_stat
    if not chosen_stat and genomics_sections:
        for sec in genomics_sections.values():
            if sec and sec.get("gene_sets_by_stat"):
                stats = list(sec["gene_sets_by_stat"].keys())
                if stats:
                    chosen_stat = stats[0]
                    break
    stat_label = _stat_display_name(chosen_stat) if chosen_stat else ""

    # --- Per-sex lowest-BMD lookups ---
    # Each helper returns {bmd_str, bmdl_str} or None if no reliable
    # value exists for the sex × category combination.
    def _lowest_geneset(sex: str) -> dict | None:
        if not genomics_sections or not chosen_stat:
            return None
        candidates: list[dict] = []
        for key, sec in genomics_sections.items():
            if not key.endswith(f"_{sex.lower()}"):
                continue
            sets = sec.get("gene_sets_by_stat", {}).get(chosen_stat, [])
            candidates.extend(_picks_above_lle(sets, lle, n=1))
        if not candidates:
            return None
        # Among per-organ winners, pick the one with the lowest BMD overall
        candidates.sort(key=lambda x: x["_bmd_float"])
        winner = candidates[0]
        return {
            "bmd": _format_dose_value(winner.get("bmd")),
            "bmdl": _format_dose_value(winner.get("bmdl")),
        }

    def _lowest_gene(sex: str) -> dict | None:
        if not genomics_sections:
            return None
        candidates: list[dict] = []
        for key, sec in genomics_sections.items():
            if not key.endswith(f"_{sex.lower()}"):
                continue
            genes = sec.get("top_genes", [])
            candidates.extend(_picks_above_lle(genes, lle, n=1))
        if not candidates:
            return None
        candidates.sort(key=lambda x: x["_bmd_float"])
        winner = candidates[0]
        return {
            "bmd": _format_dose_value(winner.get("bmd")),
            "bmdl": _format_dose_value(winner.get("bmdl")),
        }

    def _lowest_apical(sex: str) -> dict | None:
        if not apical_bmd_summary:
            return None
        # Apply same reliability + anomaly filters as the Results paragraph
        reliable = [
            e for e in apical_bmd_summary
            if e.get("sex") == sex
            and _is_reliable_bmd(e)
            and not _is_anomalous_bmd(e)
            and e.get("direction")
        ]
        if not reliable:
            return None
        reliable.sort(key=lambda e: float(e["bmd"]))
        winner = reliable[0]
        return {
            "bmd": _format_dose_value(winner.get("bmd")),
            "bmdl": _format_dose_value(winner.get("bmdl")),
        }

    # --- Build per-sex sentences ---
    # Each sex gets one main sentence listing all available categories,
    # plus an optional fallback sentence when apical is missing.
    sentences: list[str] = []
    has_any_content = False

    for sex in sexes:
        gs = _lowest_geneset(sex)
        gene = _lowest_gene(sex)
        apical = _lowest_apical(sex)

        # Build the (label, value_string) pairs in NIEHS reference order:
        # gene set → individual gene → apical endpoint.
        category_phrases: list[tuple[str, str]] = []
        if gs:
            label = (
                f"the most sensitive gene set BMD (BMDL) {stat_label}"
                if stat_label else
                "the most sensitive gene set BMD (BMDL)"
            )
            category_phrases.append((label, f"{gs['bmd']} ({gs['bmdl']})"))
        if gene:
            category_phrases.append((
                "individual gene BMD (BMDL)",
                f"{gene['bmd']} ({gene['bmdl']})",
            ))
        if apical:
            category_phrases.append((
                "apical endpoint BMD (BMDL)",
                f"{apical['bmd']} ({apical['bmdl']})",
            ))

        if not category_phrases:
            # Nothing reliable for this sex at all — skip the sentence,
            # but still emit the apical-missing fallback if relevant.
            if apical_bmd_summary:
                sentences.append(
                    f"There were no apical endpoints in {sex.lower()} rats "
                    f"for which a BMD value could be reliably estimated."
                )
            continue

        has_any_content = True
        labels = [p[0] for p in category_phrases]
        values = [p[1] for p in category_phrases]
        plural = len(category_phrases) > 1

        # Sentence start: the very first sex sentence is preceded by
        # the "Taken together, " connective (which uses lowercase "in"
        # because it comes after a comma); subsequent sex sentences are
        # standalone and start with capital "In".
        if not sentences:
            opener = f"Taken together, in {sex.lower()} rats,"
        else:
            opener = f"In {sex.lower()} rats,"

        sentences.append(
            f"{opener} "
            f"{_join_oxford(labels)} value{'s' if plural else ''} that could "
            f"be reliably determined occurred at {_join_oxford(values)} "
            f"{dose_unit}{', respectively' if plural else ''}."
        )

        # If apical was missing for this sex but genomics existed, add
        # the standard fallback sentence about apical specifically.
        if apical_bmd_summary and not apical:
            sentences.append(
                f"There were no apical endpoints in {sex.lower()} rats "
                f"for which a BMD value could be reliably estimated."
            )

    if not has_any_content:
        return ""

    return " ".join(sentences)


def build_abstract_results(
    apical_bmd_summary: list[dict] | None = None,
    genomics_sections: dict | None = None,
    dose_groups: list[float] | None = None,
    dose_unit: str = "mg/kg",
    bmd_stat: str | None = None,
    sexes: list[str] | None = None,
    methods_ctx: dict | None = None,
) -> str:
    """
    Build the full Abstract → Results paragraph.

    Combines, in NIEHS Report 10 order:
      1. Apical findings per sex (build_abstract_results_apical)
      2. Pharmacokinetic findings — plasma concentrations + half-lives
         (build_abstract_results_pk), driven by methods_ctx.pk_*
      3. Genomics findings per organ × sex (build_abstract_results_genomics)

    Args:
        apical_bmd_summary: list of BMD summary entry dicts.
        genomics_sections:  dict keyed by "{organ}_{sex}" with gene_sets and top_genes.
        dose_groups:        Full study dose list (for lower-limit-of-extrapolation).
        dose_unit:          Display unit (default "mg/kg").
        bmd_stat:           Which stat to use for gene sets ("median", "fifth_pct", ...).
        sexes:              Optional sex order (defaults to ["Male", "Female"]).
        methods_ctx:        Optional MethodsContext-as-dict.  When present, its
                            pk_concentrations / pk_half_lives / pk_timepoints fields
                            drive the pharmacokinetics sentence.

    Returns:
        A single paragraph string for the Abstract Results section.
    """
    parts: list[str] = []

    if apical_bmd_summary:
        ap = build_abstract_results_apical(apical_bmd_summary, sexes=sexes)
        if ap:
            parts.append(ap)

    # PK paragraph: only when MethodsContext carries pk_* aggregates
    if methods_ctx:
        pk = build_abstract_results_pk(
            chemical_name=methods_ctx.get("chemical_name", "the test article"),
            pk_concentrations=methods_ctx.get("pk_concentrations"),
            pk_half_lives=methods_ctx.get("pk_half_lives"),
            pk_timepoints=methods_ctx.get("pk_timepoints", []),
            dose_unit=methods_ctx.get("dose_unit", dose_unit),
            sexes=sexes,
        )
        if pk:
            parts.append(pk)

    if genomics_sections and dose_groups:
        gn = build_abstract_results_genomics(
            genomics_sections=genomics_sections,
            dose_groups=dose_groups,
            dose_unit=dose_unit,
            bmd_stat=bmd_stat,
            sexes=sexes,
        )
        if gn:
            parts.append(gn)

    return " ".join(parts)


# ---------------------------------------------------------------------------
# DOCX generation: add structured M&M to a python-docx Document
# ---------------------------------------------------------------------------

def add_methods_to_doc(
    doc,
    methods_report: MethodsReport,
    start_table_num: int = 1,
) -> int:
    """
    Add the structured Materials and Methods section to a python-docx Document.

    Generates the full NIEHS-style M&M with hierarchical headings (H2, H3, H4),
    prose paragraphs per subsection, and Table 1 (genomics sample counts).

    Follows the same DOCX formatting conventions as add_animal_report_to_doc():
    - Calibri 11pt for body text
    - Calibri 9pt for table cells and captions
    - "Light Shading Accent 1" table style
    - Sequential table numbering

    Args:
        doc:             python-docx Document object.
        methods_report:  MethodsReport with sections and context populated.
        start_table_num: Table number to start from (for sequential numbering).

    Returns:
        Next table number (for subsequent sections to continue from).
    """
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table_num = start_table_num

    # --- Top-level heading ---
    doc.add_heading("Materials and Methods", level=2)

    # --- Render each subsection ---
    for section in methods_report.sections:
        # Add heading at the appropriate level
        doc.add_heading(section.heading, level=section.level)

        # Add prose paragraphs
        for para_text in section.paragraphs:
            if not para_text.strip():
                continue
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

        # Add table if present (e.g. Table 1)
        if section.table:
            table_num = _add_methods_table(
                doc, section.table, table_num,
            )

    # --- Add Table 1 at the end if genomics data exists ---
    # Table 1 is appended after all prose subsections, matching the
    # NIEHS report layout where it follows the Data Analysis section.
    table1_data = build_table1_data(methods_report.context)
    if table1_data:
        table_num = _add_methods_table(doc, table1_data, table_num)

    return table_num


def _add_methods_table(doc, table_data: dict, table_num: int) -> int:
    """
    Add a formatted table to the DOCX document.

    Handles the caption-above-table pattern, bold sex-header rows,
    and footnotes below the table.

    Args:
        doc:        python-docx Document.
        table_data: Dict with caption, headers, rows, footnotes.
        table_num:  Current table number for caption.

    Returns:
        Next table number.
    """
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    headers = table_data["headers"]
    rows = table_data["rows"]
    caption_text = table_data.get("caption", "")
    footnotes = table_data.get("footnotes", [])

    n_cols = len(headers)
    n_rows = len(rows)

    # Create the table
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Caption paragraph (inserted before the table)
    caption = doc.add_paragraph()
    run = caption.add_run(f"Table {table_num}. {caption_text}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.italic = True
    caption.paragraph_format.space_after = Pt(4)
    # Move caption before the table element in the document XML
    table._element.addprevious(caption._element)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = "Calibri"

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[1 + row_idx]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            # Check for bold marker (sex header rows: "**Male**")
            is_bold = val.startswith("**") and val.endswith("**")
            clean_val = val.strip("*").strip() if is_bold else val.strip()
            cell.text = clean_val
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Calibri"
                    if is_bold:
                        r.bold = True

    # Footnotes below the table
    for fn in footnotes:
        p = doc.add_paragraph()
        run = p.add_run(fn)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.italic = True
        p.paragraph_format.space_after = Pt(2)

    return table_num + 1
