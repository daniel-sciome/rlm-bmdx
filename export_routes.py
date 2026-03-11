"""
Export and style-profile routes for 5dToxReport.

Provides endpoints for exporting the complete NIEHS Report 10-structured
document in .docx and PDF/UA-1 formats, plus managing the global writing
style profile that the LLM uses to match user preferences.

Extracted from background_server.py (Phase 4) to keep the main server
file focused on app creation, middleware, and router mounting.

Endpoints:
  GET    /api/style-profile          — Retrieve the global style profile
  DELETE /api/style-profile/{idx}    — Delete a style rule by index
  POST   /api/export-docx            — Export full report to .docx
  POST   /api/export-pdf             — Export full report to tagged PDF/UA-1
  GET    /api/export-pdf-scaffold    — Generate scaffold PDF with placeholders
"""

import asyncio
import logging
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Request
from fastapi.responses import FileResponse, JSONResponse

from build_docx import add_heading, add_para, fmt
from apical_report import (
    build_table_data_from_bm2,
    add_apical_tables_to_doc,
    add_bmd_summary_table_to_doc,
    add_gene_set_bmd_tables_to_doc,
    add_gene_bmd_tables_to_doc,
)
from animal_report import add_animal_report_to_doc
from session_store import safe_filename
from style_learning import (
    load_style_profile, save_style_profile,
)
from server_state import get_bm2_uploads

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Router — mounted by background_server.py as app.include_router(...)
# ---------------------------------------------------------------------------
router = APIRouter()


# ---------------------------------------------------------------------------
# Helper: superscript reference markers in .docx paragraphs
# ---------------------------------------------------------------------------

def _add_text_with_superscript_refs(paragraph, text: str) -> None:
    """
    Add text to a paragraph, converting [N] reference markers to superscript.

    Splits the text on [N] patterns, adds normal text as regular runs and
    reference numbers as superscript runs. Uses Calibri 11pt as the base font.
    """
    from docx.oxml.ns import qn

    # Split on reference markers like [1], [2,3], [1-3], etc.
    parts = re.split(r'(\[\d+(?:[,\-–]\d+)*\])', text)

    for part in parts:
        if re.match(r'\[\d+(?:[,\-–]\d+)*\]', part):
            # This is a reference marker — make it superscript
            run = paragraph.add_run(part)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.superscript = True
        else:
            # Normal text
            if part:
                run = paragraph.add_run(part)
                run.font.size = Pt(11)
                run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# GET /api/style-profile — retrieve the global style profile
# ---------------------------------------------------------------------------

@router.get("/api/style-profile")
async def api_style_profile():
    """
    Return the global style profile (learned writing preferences).

    Returns the full profile JSON including version, updated_at, and rules
    array.  If no profile exists yet, returns an empty structure with zero
    rules so the client can always expect the same shape.
    """
    profile = load_style_profile()
    return JSONResponse(profile)


# ---------------------------------------------------------------------------
# DELETE /api/style-profile/{idx} — delete a specific style rule by index
# ---------------------------------------------------------------------------

@router.delete("/api/style-profile/{idx}")
async def api_delete_style_rule(idx: int):
    """
    Delete a style rule at the given index from the global profile.

    The index is 0-based and corresponds to the rule's position in the
    rules array.  After deletion, the profile is rewritten to disk.

    Returns the updated profile so the client can re-render immediately.
    """
    profile = load_style_profile()
    rules = profile.get("rules", [])

    if idx < 0 or idx >= len(rules):
        return JSONResponse(
            {"error": f"Rule index {idx} out of range (0..{len(rules) - 1})"},
            status_code=404,
        )

    removed = rules.pop(idx)
    save_style_profile(profile)
    logger.info("Deleted style rule #%d: %s", idx, removed.get("rule", ""))

    return JSONResponse(profile)


# ---------------------------------------------------------------------------
# POST /api/export-docx — export full report to .docx
# ---------------------------------------------------------------------------

@router.post("/api/export-docx")
async def api_export_docx(request: Request):
    """
    Export the full 5dToxReport to a .docx file.

    Builds a document in NIEHS report order:
      1. Title + subtitle
      2. Background section (paragraphs + references)
      3. Materials and Methods (if approved)
      4. Results heading
         4a. Apical endpoint tables (from .bm2 files)
         4b. Apical Endpoint BMD Summary table (if approved)
         4c. Gene Set BMD Analysis tables (if genomics approved)
         4d. Gene BMD Analysis tables (if genomics approved)
      5. Summary (if approved)
      6. References

    Input JSON:
      {
        "paragraphs": ["paragraph 1...", ...],
        "references": ["[1] Ref text...", ...],
        "chemical_name": "1,2-Dichlorobenzene",
        "apical_sections": [...],
        "methods_paragraphs": ["Methods paragraph 1...", ...],
        "bmd_summary_endpoints": [{endpoint, bmd, ...}, ...],
        "genomics_sections": [
          {"organ": "liver", "sex": "male",
           "gene_sets": [...], "top_genes": [...], "dose_unit": "mg/kg"},
          ...
        ],
        "summary_paragraphs": ["Summary paragraph 1...", ...]
      }

    All new fields are optional — if omitted, those sections are skipped
    (backwards-compatible with the original two-section report).

    Returns the .docx file as a downloadable attachment.
    """
    body = await request.json()
    paragraphs = body.get("paragraphs", [])
    references = body.get("references", [])
    chemical_name = body.get("chemical_name", "Chemical")
    apical_sections = body.get("apical_sections", [])

    # New NIEHS sections (all optional)
    # methods_data is the new structured format (dict with sections + context).
    # methods_paragraphs is the legacy flat format (list of strings).
    # We accept either for backwards compatibility.
    methods_data = body.get("methods_data")
    methods_paragraphs = body.get("methods_paragraphs", [])
    animal_report_data = body.get("animal_report")
    bmd_summary_endpoints = body.get("bmd_summary_endpoints", [])
    genomics_sections = body.get("genomics_sections", [])
    summary_paragraphs = body.get("summary_paragraphs", [])

    # At least the background or one results section must be present
    if not paragraphs and not apical_sections:
        return JSONResponse(
            {"error": "No paragraphs or apical sections provided"},
            status_code=400,
        )

    # Build the .docx document
    doc = Document()

    # ===== 1. Title + subtitle =====
    doc.add_heading(
        f"5 Day Genomic Dose Response: {chemical_name}",
        level=1,
    )
    add_para(doc,
        "5 Day Genomic Dose Response in Sprague-Dawley Rats",
        italic=True, size=12,
    )

    # ===== 2. Background section =====
    if paragraphs:
        doc.add_heading("Background", level=2)
        for para_text in paragraphs:
            p = doc.add_paragraph()
            _add_text_with_superscript_refs(p, para_text)
            p.paragraph_format.space_after = Pt(6)

    # Sequential table numbering across all report sections
    next_table_num = 1

    # ===== 3. Materials and Methods (NEW — if approved) =====
    # Supports two formats:
    #   - Structured (methods_data): dict with "sections" and "context" from
    #     the new NIEHS-style generation.  Uses add_methods_to_doc() for full
    #     heading hierarchy, conditional subsections, and Table 1.
    #   - Legacy (methods_paragraphs): flat list of paragraph strings from the
    #     old 5-paragraph generation.  Rendered as simple body text.
    if methods_data and methods_data.get("sections"):
        from methods_report import MethodsReport, add_methods_to_doc
        report = MethodsReport.from_dict(methods_data)
        next_table_num = add_methods_to_doc(doc, report, start_table_num=next_table_num)
    elif methods_paragraphs:
        doc.add_heading("Materials and Methods", level=2)
        for para_text in methods_paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

    # ===== 3b. Animal Report (NEW — if approved) =====
    # Inserted between Materials & Methods and Results because it documents
    # the study's animal roster and data coverage — context needed before
    # reading the dose-response results.
    #
    # animal_report_data is a dict matching report_to_dict() output from
    # animal_report.py.  We reconstruct an AnimalReport object from it
    # to reuse the add_animal_report_to_doc() DOCX builder.
    if animal_report_data:
        from animal_report import AnimalReport, AnimalRecord
        # Reconstruct AnimalReport from serialized dict
        ar = AnimalReport()
        ar.study_number = animal_report_data.get("study_number")
        ar.total_animals = animal_report_data.get("total_animals", 0)
        ar.core_count = animal_report_data.get("core_count", 0)
        ar.biosampling_count = animal_report_data.get("biosampling_count", 0)
        ar.dose_groups = animal_report_data.get("dose_groups", [])
        # Reconstruct dose_design — keys were stringified for JSON
        raw_dd = animal_report_data.get("dose_design", {})
        ar.dose_design = {float(k): v for k, v in raw_dd.items()}
        ar.domain_coverage = animal_report_data.get("domain_coverage", {})
        ar.completeness = animal_report_data.get("completeness", {})
        ar.consistency_issues = animal_report_data.get("consistency_issues", [])
        # Reconstruct animal records
        for aid, rec_dict in animal_report_data.get("animals", {}).items():
            ar.animals[aid] = AnimalRecord(
                animal_id=aid,
                sex=rec_dict.get("sex"),
                dose=rec_dict.get("dose"),
                selection=rec_dict.get("selection"),
                domain_presence=rec_dict.get("domain_presence", {}),
            )
        # Reconstruct attrition objects
        from animal_report import DomainAttrition
        for domain, att_dict in animal_report_data.get("attrition", {}).items():
            ar.attrition[domain] = DomainAttrition(
                domain=domain,
                xlsx_ids=set(att_dict.get("xlsx_ids", [])),
                txt_csv_ids=set(att_dict.get("txt_csv_ids", [])),
                bm2_ids=set(att_dict.get("bm2_ids", [])),
                excluded_xlsx_to_txt=set(att_dict.get("excluded_xlsx_to_txt", [])),
                excluded_txt_to_bm2=set(att_dict.get("excluded_txt_to_bm2", [])),
                exclusion_reasons=att_dict.get("exclusion_reasons", {}),
            )
        next_table_num = add_animal_report_to_doc(doc, ar, start_table_num=next_table_num)

    # ===== 4. Results =====
    # Add a "Results" heading if we have any results sections
    has_results = bool(apical_sections or bmd_summary_endpoints or genomics_sections)
    if has_results:
        doc.add_heading("Results", level=2)

    # --- 4a. Apical endpoint tables (existing .bm2 tables) ---
    # Looks up cached table_data from the upload store by bm2_id.
    _bm2_uploads = get_bm2_uploads()
    for section in apical_sections:
        bm2_id = section.get("bm2_id", "")
        section_title = section.get(
            "section_title",
            "Animal Condition, Body Weights, and Organ Weights",
        )
        table_caption = section.get(
            "table_caption_template",
            "Summary of Body Weights and Organ Weights "
            "of {sex} Rats Administered {compound} for Five Days",
        )
        compound = section.get("compound_name", chemical_name)
        dose_unit = section.get("dose_unit", "mg/kg")

        upload = _bm2_uploads.get(bm2_id)
        if not upload:
            continue

        table_data = upload.get("table_data")
        if table_data is None:
            try:
                loop = asyncio.get_running_loop()
                table_data, _, bm2_json = await loop.run_in_executor(
                    None, build_table_data_from_bm2, upload["temp_path"],
                )
                upload["table_data"] = table_data
                upload["bm2_json"] = bm2_json
            except Exception:
                continue

        narrative_paras = section.get("narrative_paragraphs")
        if narrative_paras is None and upload:
            narrative_paras = upload.get("narrative")

        next_table_num = add_apical_tables_to_doc(
            doc,
            table_data,
            section_title=section_title,
            compound_name=compound,
            dose_unit=dose_unit,
            table_caption_template=table_caption,
            start_table_num=next_table_num,
            narrative_paragraphs=narrative_paras,
        )

    # --- 4b. Apical Endpoint BMD Summary table (NEW) ---
    if bmd_summary_endpoints:
        dose_unit = body.get("dose_unit", "mg/kg")
        next_table_num = add_bmd_summary_table_to_doc(
            doc,
            bmd_summary_endpoints,
            table_num=next_table_num,
            dose_unit=dose_unit,
        )

    # --- 4c & 4d. Gene Set BMD and Gene BMD tables (NEW) ---
    for gs_section in genomics_sections:
        organ = gs_section.get("organ", "")
        sex = gs_section.get("sex", "")
        gene_sets = gs_section.get("gene_sets", [])
        top_genes = gs_section.get("top_genes", [])
        dose_unit = gs_section.get("dose_unit", "mg/kg")

        # Gene Set BMD Analysis table (4c) — use the stat label from
        # the section (e.g. "5th %ile") so the column header matches
        # the statistic selected in the settings panel.
        stat_label = gs_section.get("bmd_stat_label", "Median")
        if gene_sets:
            next_table_num = add_gene_set_bmd_tables_to_doc(
                doc, gene_sets, organ, sex,
                table_num=next_table_num,
                dose_unit=dose_unit,
                bmd_stat_label=stat_label,
            )

        # Gene BMD Analysis table (4d)
        if top_genes:
            next_table_num = add_gene_bmd_tables_to_doc(
                doc, top_genes, organ, sex,
                table_num=next_table_num,
                dose_unit=dose_unit,
            )

    # ===== 5. Summary (NEW — if approved) =====
    if summary_paragraphs:
        doc.add_heading("Summary", level=2)
        for para_text in summary_paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

    # ===== 6. References (moved to end for NIEHS ordering) =====
    if references:
        doc.add_heading("References", level=2)
        for ref_line in references:
            add_para(doc, ref_line, size=10)

    # Save to a temporary file and return as download
    safe_name = safe_filename(chemical_name)
    filename = f"5dToxReport_{safe_name}.docx"

    tmp = tempfile.NamedTemporaryFile(
        delete=False, suffix=".docx", prefix="5dtox_",
    )
    doc.save(tmp.name)
    tmp.close()

    return FileResponse(
        tmp.name,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# POST /api/export-pdf — export full report to tagged PDF/UA-1
# ---------------------------------------------------------------------------

@router.post("/api/export-pdf")
async def api_export_pdf(request: Request):
    """
    Export the full 5dToxReport to a PDF/UA-1 compliant PDF file.

    Accepts the same JSON payload as /api/export-docx for consistency.
    Marshals the data into the Typst template schema and compiles it
    using the Typst compiler (embedded Rust binary via typst-py).

    The output PDF has:
      - Full StructTreeRoot with H1-H3, P, Table/TH/TD/TR tags
      - PDF/UA-1 identifier in XMP metadata
      - /Lang set to "en" in the document catalog
      - /MarkInfo << /Marked true >>
      - Proper heading hierarchy matching NIEHS Report 10
      - Running header and page numbers marked as Artifacts

    Returns the PDF file as a downloadable attachment.
    """
    from report_pdf import build_report_pdf, marshal_export_data

    body = await request.json()

    # Resolve table_data for apical sections that reference uploaded .bm2 files.
    # The web UI may send bm2_id references instead of inline table_data,
    # so we need to look up the cached data from the upload store.
    _bm2_uploads = get_bm2_uploads()
    for sec in body.get("apical_sections", []):
        if "table_data" not in sec or not sec["table_data"]:
            bm2_id = sec.get("bm2_id", "")
            upload = _bm2_uploads.get(bm2_id)
            if upload and upload.get("table_data"):
                sec["table_data"] = upload["table_data"]

    # Also inject narrative from server cache if not provided inline
    for sec in body.get("apical_sections", []):
        if not sec.get("narrative_paragraphs"):
            bm2_id = sec.get("bm2_id", "")
            upload = _bm2_uploads.get(bm2_id)
            if upload and upload.get("narrative"):
                sec["narrative_paragraphs"] = upload["narrative"]

    try:
        # Marshal the DOCX-format payload into the Typst template schema
        report_data = marshal_export_data(body)

        # Compile to PDF/UA-1 — sub-second for typical report sizes
        pdf_bytes = build_report_pdf(report_data)
    except Exception as e:
        logging.exception("PDF export failed")
        return JSONResponse(
            {"error": f"PDF generation failed: {e}"},
            status_code=500,
        )

    # Write to a temp file for FileResponse
    chemical_name = body.get("chemical_name", "Chemical")
    safe_name = safe_filename(chemical_name)
    filename = f"5dToxReport_{safe_name}.pdf"

    tmp = tempfile.NamedTemporaryFile(
        delete=False, suffix=".pdf", prefix="5dtox_",
    )
    tmp.write(pdf_bytes)
    tmp.close()

    return FileResponse(
        tmp.name,
        filename=filename,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# GET /api/export-pdf-scaffold — generate a scaffold PDF showing all sections
# ---------------------------------------------------------------------------

@router.get("/api/export-pdf-scaffold")
async def api_export_pdf_scaffold(
    chemical_name: str = "Test Article",
    casrn: str = "000-00-0",
    dtxsid: str = "DTXSID0000000",
):
    """
    Generate a complete scaffold PDF with placeholder content in every section.

    This endpoint produces a full NIEHS Report 10-structured PDF with all
    sections populated by clearly-marked placeholder text (wrapped in angle
    quotes).  The purpose is to show the exact page flow, typography, table
    layout, landscape pages, and pagination that the final report will have.

    Every template code path is exercised: title page, roman-numeral front
    matter (foreword, TOC, tables list, about, peer review, pub details,
    acknowledgments, abstract), arabic body pages (background, M&M with
    Table 1, results with landscape dose-response tables, internal dose
    portrait table, BMD summary sex-grouped table, genomics gene set
    and gene tables with GO/gene descriptions), summary, and references.

    Query parameters allow customizing the chemical identity on the
    title page and throughout the document:
      ?chemical_name=Perfluorohexanesulfonamide&casrn=41997-13-1&dtxsid=DTXSID50469320

    Returns the PDF as a downloadable attachment.
    """
    from report_pdf import build_report_pdf, scaffold_report_data

    try:
        data = scaffold_report_data(
            chemical_name=chemical_name,
            casrn=casrn,
            dtxsid=dtxsid,
        )
        pdf_bytes = build_report_pdf(data)
    except Exception as e:
        logging.exception("Scaffold PDF generation failed")
        return JSONResponse(
            {"error": f"Scaffold PDF generation failed: {e}"},
            status_code=500,
        )

    # Clean filename from chemical name
    safe_name = safe_filename(chemical_name)
    filename = f"5dToxReport_Scaffold_{safe_name}.pdf"

    tmp = tempfile.NamedTemporaryFile(
        delete=False, suffix=".pdf", prefix="5dtox_scaffold_",
    )
    tmp.write(pdf_bytes)
    tmp.close()

    return FileResponse(
        tmp.name,
        filename=filename,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# GET /api/export-bm2/{dtxsid} — download enriched .bm2 file
# ---------------------------------------------------------------------------

@router.get("/api/export-bm2/{dtxsid}")
async def api_export_bm2(dtxsid: str):
    """
    Download the metadata-enriched .bm2 file for a session.

    Reads the session's integrated.json (which contains LLM-inferred and
    user-approved ExperimentDescription metadata) and converts it to the
    canonical .bm2 format (Java ObjectOutputStream) via JsonToBm2.

    The resulting file is a standard BMDExpress 3 project file — opening
    it in BMDExpress 3 shows experiments with metadata pre-filled (species,
    sex, organ, test article, study duration, etc.).

    If an integrated.bm2 already exists and is newer than integrated.json,
    it's served directly without re-export.

    Returns the .bm2 file as a download attachment.
    """
    from session_store import session_dir
    from pool_integrator import export_integrated_bm2

    sess_path = session_dir(dtxsid)
    json_path = sess_path / "integrated.json"
    bm2_path = sess_path / "integrated.bm2"

    if not json_path.exists():
        return JSONResponse(
            {"error": "No integrated data found — run integration first"},
            status_code=404,
        )

    # Re-export if .bm2 is missing or older than the JSON source.
    # This handles the case where the user edits metadata (re-runs
    # integration) and then downloads — always gets the latest.
    needs_export = (
        not bm2_path.exists()
        or bm2_path.stat().st_mtime < json_path.stat().st_mtime
    )
    if needs_export:
        try:
            loop = asyncio.get_running_loop()
            await loop.run_in_executor(
                None,
                export_integrated_bm2,
                str(json_path),
                str(bm2_path),
            )
        except Exception as e:
            logger.exception("Failed to export enriched .bm2 for %s", dtxsid)
            return JSONResponse(
                {"error": f"Export failed: {e}"},
                status_code=500,
            )

    # Derive a human-readable filename from the session identity
    identity_path = sess_path / "identity.json"
    filename = f"{dtxsid}_integrated.bm2"
    if identity_path.exists():
        try:
            import json
            identity = json.loads(identity_path.read_text(encoding="utf-8"))
            chem_name = identity.get("preferredName", "")
            if chem_name:
                safe_name = safe_filename(chem_name)
                filename = f"{safe_name}_integrated.bm2"
        except Exception:
            pass

    return FileResponse(
        str(bm2_path),
        filename=filename,
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
