"""
build_docx.py — Shared DOCX formatting utilities for report generation.

Provides reusable functions for adding formatted headings, paragraphs,
bullet lists, bold-lead paragraphs, and tables to python-docx Document
objects.  Used by background_server.py and interpret.py for DOCX export.
"""

from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def fmt(paragraph, size=11, bold=False, italic=False, font="Calibri", color=None,
        space_after=Pt(6), space_before=Pt(0), align=None):
    """Apply formatting to all runs in a paragraph."""
    pf = paragraph.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if align:
        pf.alignment = align
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.name = font
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    """Add a heading at the specified level."""
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text="", bold=False, italic=False, size=11):
    """Add a formatted paragraph with a single run."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bold_lead(doc, lead, rest):
    """Add paragraph with bold lead text followed by normal text."""
    p = doc.add_paragraph()
    r1 = p.add_run(lead)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=""):
    """Add a bullet-list paragraph with optional bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row and optional column widths."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Calibri"

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Calibri"

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table
