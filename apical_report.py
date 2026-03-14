"""
Generate NTP-style apical endpoint report sections from BMDExpress 3 .bm2 files.

Originally built for the "Animal Condition, Body Weights, and Organ Weights"
section, but parameterized so it can also produce "Clinical Pathology" or any
other apical endpoint section — just pass --section-title and --table-caption.

This script is the bridge between BMDExpress category analysis results and
the NTP-standard statistical tests (apical_stats.py).  It:

  1. Exports raw dose-response and category analysis data from a .bm2 file
     using the BMDExpress 3 Java CLI
  2. Runs NTP-standard statistical tests (Jonckheere → Williams/Dunnett)
     on the raw dose-response data
  3. Cross-references category analysis results for BMD/BMDL values
  4. Applies the business rule: BMD is only reported if BOTH pairwise
     significance AND trend significance are met
  5. Generates a .docx report section with tables split by sex, matching
     the NTP subchronic study report format (PFHxSAm prototype)

Table format (per sex):
  - First row: "n" — sample sizes per dose group
  - Subsequent rows: one per endpoint (e.g., "Terminal Body Wt.", "Liver
    Absolute", "Liver Relative")
  - Columns: Endpoint label, then mean ± SE for each dose group (with */**
    significance markers), then BMD1Std and BMDL1Std columns
  - Control row gets * / ** for the trend test (Jonckheere)
  - Treatment rows get * / ** for pairwise test (Williams or Dunnett)
  - Endpoints not significant for BOTH trend + pairwise get "ND" for BMD/BMDL
  - Dose groups with no surviving animals get a dash and footnote

Usage:
    python apical_report.py <path_to.bm2> [--output report.docx]

    # Or import and use programmatically:
    from apical_report import generate_report
    generate_report("path/to/file.bm2", "output.docx")
"""

import argparse
import csv
import json
import math
import os
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

import bm2_cache
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from apical_stats import analyze_endpoint, EndpointStats, run_java_prefilter


# ---------------------------------------------------------------------------
# Constants — paths and configuration
# ---------------------------------------------------------------------------

# Java classpath, project paths, and helper directory are centralized in
# java_bridge.py to avoid duplication across apical_report, apical_stats,
# and pool_integrator.
from java_bridge import build_classpath as _build_classpath
from java_bridge import JAVA_HELPER_DIR

# Font size for table cells (small to fit many dose columns)
TABLE_FONT_SIZE = 8

# Font size for table header cells
HEADER_FONT_SIZE = 8


def export_bm2(bm2_path: str, output_json: str, output_tsv: str) -> None:
    """
    Export a .bm2 file to JSON + category TSV in a single JVM launch.

    Uses the pre-compiled ExportBm2.class which deserializes the .bm2 file
    (Java ObjectInputStream → BMDProject), then:
      1. Writes the full BMDProject as JSON via Jackson
      2. Writes category analysis results as TSV via BMDExpress's own
         DataCombinerService + ProjectNavigationService

    This is ~30x faster than the old 3-JVM approach because JVM startup
    is the bottleneck, not the actual work.

    Args:
        bm2_path:    Path to the input .bm2 file.
        output_json: Path where the JSON export will be written.
        output_tsv:  Path for category TSV, or "NONE" to skip.
    """
    cp = _build_classpath()
    helper_dir = str(JAVA_HELPER_DIR)

    subprocess.run(
        [
            "java", "-cp", f"{cp}:{helper_dir}",
            "ExportBm2", bm2_path, output_json, output_tsv,
        ],
        check=True, capture_output=True,
    )


def export_genomics(bm2_path: str, output_json: str) -> dict:
    """
    Extract genomics results (gene-level BMD + GO BP categories) from a .bm2 file.

    Uses the pre-compiled ExportGenomics.class which deserializes the .bm2 via
    BMDExpress 3's native Java API, then extracts:
      1. Per-probe BMD/BMDL/BMDU, direction, rSquared, fold change
      2. GO Biological Process category analysis with aggregated BMD stats

    Results are grouped by organ × sex experiment.  This is the authoritative
    source — BMDExpress's own prefilter → curve fit → BMD pipeline ran on the
    data, and we're reading its results directly.

    Args:
        bm2_path:    Path to the gene expression .bm2 file.
        output_json: Path where the genomics JSON will be written.

    Returns:
        The parsed JSON dict with structure:
        {
            "experiments": [
                {
                    "name": str, "organ": str, "sex": str,
                    "total_probes": int,
                    "genes": [{ "probe_id", "gene_symbol", "bmd", "bmdl",
                                "bmdu", "direction", "r_squared", ... }],
                    "go_bp": [{ "go_id", "go_term", "bmd_median", "bmdl_median",
                                "n_genes", "n_passed", "direction", ... }]
                }
            ]
        }
    """
    cp = _build_classpath()
    helper_dir = str(JAVA_HELPER_DIR)

    subprocess.run(
        [
            "java", "-cp", f"{cp}:{helper_dir}",
            "ExportGenomics", bm2_path, output_json,
        ],
        check=True, capture_output=True,
    )

    with open(output_json, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Category analysis BMD extraction — via native Java API
# ---------------------------------------------------------------------------

def export_categories(bm2_path: str, output_json: str) -> dict:
    """
    Extract all category analysis BMD values from a .bm2 file.

    Uses the pre-compiled ExportCategories.class which reads
    CategoryAnalysisResults directly from the BMDProject via BMDExpress's
    native Java API.  This replaces the broken DataCombinerService TSV
    export path (which crashes with IndexOutOfBoundsException).

    Exports ALL BMD statistics per endpoint: mean, median, minimum,
    weighted mean, 5th/10th percentile, 95% CI bounds — so the user
    can choose which statistic to display in report tables.

    Args:
        bm2_path:    Path to the .bm2 file.
        output_json: Path where the categories JSON will be written.

    Returns:
        The parsed JSON dict with structure:
        {
            "analysis_count": int,
            "total_endpoints": int,
            "analyses": [
                {
                    "name": str,
                    "experiment_prefix": str,
                    "category_type": str,
                    "results": [
                        {
                            "category_id": str,
                            "category_title": str,
                            "bmd": { "mean", "median", "minimum", ... },
                            "bmdl": { ... },
                            "bmdu": { ... },
                            "genes_passed": int,
                            ...
                        }
                    ]
                }
            ]
        }
    """
    cp = _build_classpath()
    helper_dir = str(JAVA_HELPER_DIR)

    subprocess.run(
        [
            "java", "-cp", f"{cp}:{helper_dir}",
            "ExportCategories", bm2_path, output_json,
        ],
        check=True, capture_output=True,
    )

    with open(output_json, "rb") as f:
        import orjson
        return orjson.loads(f.read())


def build_category_lookup(
    categories_json: dict,
    bmd_stat: str = "mean",
    experiment_names: list[str] | None = None,
) -> dict[tuple[str, str], dict]:
    """
    Build the category lookup dict from ExportCategories output.

    Transforms the structured JSON from export_categories() into the
    (experiment_name, endpoint_name) → {bmd, bmdl, bmdu, ...} lookup
    that build_table_data() uses to populate BMD/BMDL columns.

    BMDExpress appends pipeline suffixes to experiment names in category
    analysis results (e.g., "female_clin_chem" becomes
    "female_clin_chem_williams_0.05_NOMTC_nofoldfilter").  When
    experiment_names is provided, this function resolves each suffixed
    prefix back to its raw experiment name so build_table_data() can
    look up BMD values using the original experiment name.

    The bmd_stat parameter controls which BMD aggregate statistic is used.
    For single-gene apical endpoints (1 gene per category), all statistics
    are identical.  For GO categories in genomics data, the choice matters
    (e.g., median vs 5th percentile can differ substantially).

    Args:
        categories_json:  Output from export_categories().
        bmd_stat:         Which BMD statistic to use. One of:
                          "mean", "median", "minimum", "weighted_mean",
                          "fifth_pct", "tenth_pct", "lower95", "upper95".
                          Defaults to "mean" for backward compatibility.
        experiment_names: Optional list of raw experiment names from the
                          BMDProject's doseResponseExperiments.  When
                          provided, category prefixes are resolved to the
                          longest matching experiment name, fixing the
                          key mismatch caused by BMDExpress pipeline
                          suffixes.

    Returns:
        Dict mapping (experiment_name, endpoint_name) → {bmd, bmdl, bmdu, ...}
    """
    # Pre-sort experiment names by length (longest first) so we match the
    # most specific name when a prefix starts with multiple candidates.
    # E.g., "female_clin_chem" should match before "female_clin".
    sorted_exp_names = (
        sorted(experiment_names, key=len, reverse=True)
        if experiment_names
        else []
    )

    lookup = {}

    for analysis in categories_json.get("analyses", []):
        prefix = analysis.get("experiment_prefix", "")

        # Resolve the suffixed prefix to a raw experiment name.
        # BMDExpress category analysis names look like:
        #   female_clin_chem_williams_0.05_NOMTC_nofoldfilter_BMD_null_DEFINED-...
        # ExportCategories splits on "_BMD" → prefix is the part before it:
        #   female_clin_chem_williams_0.05_NOMTC_nofoldfilter
        # We need to map this back to the raw experiment name:
        #   female_clin_chem
        resolved_prefix = prefix
        for exp_name in sorted_exp_names:
            # The prefix must start with the experiment name, and the
            # character after (if any) must be '_' to avoid partial matches
            # (e.g., "female_clin" should not match "female_clin_chem").
            if prefix == exp_name:
                resolved_prefix = exp_name
                break
            if prefix.startswith(exp_name) and prefix[len(exp_name):len(exp_name) + 1] == "_":
                resolved_prefix = exp_name
                break

        for result in analysis.get("results", []):
            endpoint = result.get("category_title", "")

            # Pull the selected statistic from the nested bmd/bmdl/bmdu blocks
            bmd_block = result.get("bmd", {})
            bmdl_block = result.get("bmdl", {})
            bmdu_block = result.get("bmdu", {})

            entry = {
                "bmd": bmd_block.get(bmd_stat, bmd_block.get("mean", "")),
                "bmdl": bmdl_block.get(bmd_stat, bmdl_block.get("mean", "")),
                "bmdu": bmdu_block.get(bmd_stat, bmdu_block.get("mean", "")),
                "category_id": result.get("category_id", ""),
                "direction": result.get("direction", ""),
                "fold_change": result.get("max_fold_change", ""),
                # Preserve the full stat blocks for downstream use
                "bmd_stats": bmd_block,
                "bmdl_stats": bmdl_block,
                "bmdu_stats": bmdu_block,
            }

            # Store under the resolved prefix (raw experiment name).
            # Also store under the full suffixed prefix for callers that
            # already use the suffixed key (backward compatibility).
            lookup[(resolved_prefix, endpoint)] = entry
            if resolved_prefix != prefix:
                lookup[(prefix, endpoint)] = entry

    return lookup


# ---------------------------------------------------------------------------
# Data assembly — combine stats + category analysis into table-ready data
# ---------------------------------------------------------------------------

@dataclass
class TableRow:
    """
    One row of the final report table — represents a single endpoint
    (e.g., "Liver Weight Relative") for one sex.

    The reference NIEHS report splits apical data into two table types:
      - Domain tables (Tables 2-7): mean ± SE per dose group, significance
        markers, sample sizes.  NO BMD columns.
      - BMD summary (Table 8): Endpoint, BMD, BMDL, LOEL, NOEL, Direction.
        Only includes endpoints with significant findings.

    This dataclass carries all fields needed for both table types.

    Attributes:
        label:          Display label for the endpoint column
        values_by_dose: Dict mapping dose → formatted string ("mean ± SE**")
        n_by_dose:      Dict mapping dose → sample size (for the n row)
        bmd_str:        Formatted BMD value, "NVM", "UREP", "—"
        bmdl_str:       Formatted BMDL value, "NVM", "UREP", "—"
        trend_marker:   "*" or "**" or "" — for the control column
        bmd_status:     NIEHS classification: "viable", "NVM", "NR", "UREP",
                        "failure", or None (not modeled).
        loel:           Lowest observed effect level (mg/kg), or None
        noel:           No observed effect level (mg/kg), or None
        direction:      "UP" or "DOWN" — direction of change vs control
        responsive:     True if both Jonckheere trend AND Dunnett pairwise tests
                        are significant (p ≤ 0.05).  Used to filter main body
                        tables (Tables 2-7) per NIEHS business rule: only
                        responsive endpoints appear in domain tables.  BMD
                        summary (Table 8) uses its own inclusion gate.
    """
    label: str = ""
    values_by_dose: dict = field(default_factory=dict)
    n_by_dose: dict = field(default_factory=dict)
    bmd_str: str = "—"
    bmdl_str: str = "—"
    trend_marker: str = ""
    bmd_status: str | None = None
    loel: float | None = None
    noel: float | None = None
    direction: str = ""
    responsive: bool = False
    # Missing-animal footnote data.  Populated when an xlsx study file
    # provides the authoritative animal roster and some animals are absent
    # from the bm2 data (died before terminal sacrifice).
    # Maps dose → number of animals in the xlsx roster that have no data
    # for this endpoint.  Empty dict means no missing animals.
    missing_animals_by_dose: dict[float, int] = field(default_factory=dict)


def _safe_float(val) -> float | None:
    """
    Safely convert a value to float, handling None, string "NaN", and
    Jackson serialization quirks.  Returns None if the value is not a
    valid finite number.
    """
    if val is None:
        return None
    if isinstance(val, str):
        if val.lower() == "nan" or val.lower() == "infinity":
            return None
        try:
            f = float(val)
            return f if f == f else None  # NaN check
        except ValueError:
            return None
    if isinstance(val, (int, float)):
        return val if val == val else None  # NaN check
    return None


def _adaptive_decimals(*values: float | None, min_dp: int = 1, max_dp: int = 4) -> int:
    """
    Choose the number of decimal places for a group of related values
    (typically mean and SE for the same dose group).

    Uses the smallest non-zero absolute value to decide:
      - |v| >= 1    → 1 decimal  (e.g., "234.5 ± 12.3")
      - |v| >= 0.1  → 2 decimals (e.g., "0.34 ± 0.12")
      - |v| >= 0.01 → 3 decimals (e.g., "0.034 ± 0.012")
      - |v| < 0.01  → 4 decimals (e.g., "0.0034 ± 0.0012")

    This keeps large body weights at 1 dp (matching reference format) while
    giving small clinical chemistry values enough precision.
    """
    # Find the smallest non-zero magnitude among all provided values
    smallest = None
    for v in values:
        if v is not None and v != 0:
            mag = abs(v)
            if smallest is None or mag < smallest:
                smallest = mag

    if smallest is None or smallest >= 1.0:
        return min_dp
    elif smallest >= 0.1:
        return max(min_dp, 2)
    elif smallest >= 0.01:
        return max(min_dp, 3)
    else:
        return max_dp


def _classify_bmd_result(
    bmd: float | None,
    bmdl: float | None,
    bmdu: float | None,
    lowest_nonzero_dose: float,
    best_stat_result: dict | None,
) -> str:
    """
    Classify a BMD result into NIEHS reference report bins.

    The NIEHS 5-day study reports use these classifications for apical
    endpoint BMDs (see Appendix D, Table D-1 in the reference report):

      - "failure":  Model did not successfully complete.
      - "NVM":      Nonviable model — completed but failed acceptability.
      - "NR":       Not reportable — BMD below lower limit of extrapolation
                    (<1/3 lowest nonzero dose tested).  BMD is reported as
                    "<1/3 LNZD" and BMDL is not reported.
      - "UREP":     Unreliable estimate of potency — subject matter expert
                    review flag.  Triggered by wide confidence interval
                    (BMDU/BMDL ratio > 40) or other quality concerns.
      - "viable":   Candidate for recommended model without warning.

    Args:
        bmd:   Best BMD value (None if modeling failed).
        bmdl:  Best BMDL value (lower confidence limit).
        bmdu:  Best BMDU value (upper confidence limit).
        lowest_nonzero_dose: Smallest nonzero dose in the experiment.
        best_stat_result: The bestStatResult dict from BMDExpress 3,
                         containing model fit metrics.
    """
    # No BMD at all → failure (model didn't produce a result)
    if bmd is None:
        return "failure"

    # Check model fit quality from bestStatResult
    if best_stat_result:
        success = best_stat_result.get("success")
        # BMDExpress 3 serializes success as string "true" or null
        if success is None or success == "false":
            return "NVM"

        r_squared = _safe_float(best_stat_result.get("rSquared"))
        if r_squared is not None and r_squared < 0.1:
            return "NVM"

    # BMD below lower limit of extrapolation: <1/3 lowest nonzero dose.
    # Reference report: "BMD values derived from viable models that were
    # threefold lower than the lowest nonzero dose tested were reported
    # as <1/3 the lowest nonzero dose tested, and corresponding BMDL
    # values were not reported."
    if lowest_nonzero_dose > 0 and bmd < lowest_nonzero_dose / 3:
        return "NR"

    # Wide confidence interval → unreliable estimate.
    # Reference report: BMDU/BMDL ratio > 40 is flagged.
    if bmdl is not None and bmdu is not None and bmdl > 0:
        ratio = bmdu / bmdl
        if ratio > 40:
            return "UREP"

    # Step function check — if BMD lands below the lowest dose, the
    # model is extrapolating beyond the observed range.
    if best_stat_result:
        step_flag = best_stat_result.get("stepWithBMDLessLowest", False)
        if step_flag:
            return "UREP"

    return "viable"


def _build_bmd_result_lookup(bm2_json: dict) -> dict[tuple[str, str], dict]:
    """
    Build a lookup of BMD/BMDL values and NIEHS classification from
    BMDExpress 3 native bMDResult data.

    BMDExpress 3 stores dose-response modeling results in the bMDResult
    array.  Each entry has probeStatResults with bestBMD/bestBMDL and
    model fit metrics (rSquared, success, confidence bounds).

    This function resolves Jackson @ref pointers, extracts BMD values,
    and classifies each result into NIEHS bins: viable, NVM, NR, UREP,
    or failure.  The lowest nonzero dose (LNZD) from the parent experiment
    is used for the NR classification (BMD < LNZD/3).

    Returns:
        Dict mapping (experiment_name, probe_id) → {
            "bmd": float|None, "bmdl": float|None, "bmdu": float|None,
            "status": str  # "viable", "NVM", "NR", "UREP", "failure"
        }.
    """
    import math

    # Build @ref → object lookups for experiments and probeResponses.
    # Jackson serialization uses integer @ref/@id pairs for object identity.
    ref_to_exp_name: dict[int, str] = {}
    ref_to_probe_id: dict[int, str] = {}
    # Also need @ref → experiment for LNZD calculation
    ref_to_exp: dict[int, dict] = {}
    for exp in bm2_json.get("doseResponseExperiments", []):
        ref = exp.get("@ref")
        if ref is not None:
            ref_to_exp_name[ref] = exp["name"]
            ref_to_exp[ref] = exp
        for pr in exp.get("probeResponses", []):
            pref = pr.get("@ref")
            if pref is not None:
                ref_to_probe_id[pref] = pr.get("probe", {}).get("id", "")

    # Cache LNZD per experiment (lowest nonzero dose tested).
    # Used for NR classification: BMD < LNZD/3 means the model is
    # extrapolating below the observed dose range.
    lnzd_cache: dict[int, float] = {}

    lookup: dict[tuple[str, str], dict] = {}
    for result in bm2_json.get("bMDResult", []):
        exp_ref = result.get("doseResponseExperiment")
        exp_name = ref_to_exp_name.get(exp_ref, "")
        if not exp_name:
            continue

        # Compute LNZD for this experiment (cached per @ref)
        if exp_ref not in lnzd_cache:
            exp = ref_to_exp.get(exp_ref, {})
            doses = [t["dose"] for t in exp.get("treatments", [])]
            nonzero = [d for d in doses if d > 0]
            lnzd_cache[exp_ref] = min(nonzero) if nonzero else 0.0
        lnzd = lnzd_cache[exp_ref]

        for psr in result.get("probeStatResults", []):
            pr_ref = psr.get("probeResponse")
            probe_id = ref_to_probe_id.get(pr_ref, "")
            if not probe_id:
                continue

            bmd = _safe_float(psr.get("bestBMD"))
            bmdl = _safe_float(psr.get("bestBMDL"))
            bmdu = _safe_float(psr.get("bestBMDU"))

            # Get bestStatResult for model fit classification.
            # It can be a dict (inline) or an int (@ref pointer).
            bsr = psr.get("bestStatResult")
            if isinstance(bsr, int):
                bsr = None  # unresolved @ref — treat as no info

            status = _classify_bmd_result(bmd, bmdl, bmdu, lnzd, bsr)

            lookup[(exp_name, probe_id)] = {
                "bmd": bmd,
                "bmdl": bmdl,
                "bmdu": bmdu,
                "status": status,
            }

    return lookup


def build_table_data(
    bm2_json: dict,
    category_lookup: dict,
) -> dict[str, list[TableRow]]:
    """
    Build the table data for all endpoints, organized by sex.

    Three-phase approach:
      Phase 1: Batch all endpoints into a single Java RunPrefilter call
               to get Williams/Dunnett p-values (one JVM launch total).
      Phase 2: For each endpoint, compute Jonckheere in Python, combine
               with Java pairwise results, apply business rules.
      BMD lookup: Uses BMDExpress 3 native bMDResult data (bestBMD/bestBMDL)
               as the primary source, with category_lookup as fallback for
               genomics endpoints.

    Args:
        bm2_json:        Parsed JSON from the .bm2 export.
        category_lookup: Dict from build_category_lookup().

    Returns:
        Dict mapping sex label ("Male" / "Female") → list of TableRow.
        Each list is ordered: body weight endpoints first, then organ
        weight endpoints.
    """
    from apical_stats import (
        _call_java_prefilter, _sig_marker, jonckheere_test,
        ALPHA_STAR, EndpointStats,
    )
    import math

    # Build lookup of BMDExpress 3-native BMD values from bMDResult.
    # This is the primary BMD source for apical endpoints (body weight,
    # organ weight, clinical chemistry, hematology, hormones).  Category
    # analysis (category_lookup) is the fallback for genomics endpoints.
    bmd_result_lookup = _build_bmd_result_lookup(bm2_json)

    # ---- Phase 1: Batch Java prefilter ----
    # Collect all endpoints across all experiments into a single batch,
    # so we only launch one JVM for the entire set (~20 endpoints instead
    # of 20 separate JVM launches = ~15s saved).
    batch_doses: list[float] = []
    batch_probe_ids: list[str] = []
    batch_responses: list[list[float]] = []
    # Track endpoint metadata for reassembly after Java returns
    endpoint_meta: list[dict] = []

    for exp in bm2_json.get("doseResponseExperiments", []):
        exp_name = exp.get("name", "")
        treatments = exp.get("treatments", [])
        exp_doses = [t["dose"] for t in treatments]
        treat_doses = {i: t["dose"] for i, t in enumerate(treatments)}

        # Determine sex from experiment name
        exp_name_lower = exp_name.lower()
        if "female" in exp_name_lower:
            sex = "Female"
        elif "male" in exp_name_lower:
            sex = "Male"
        else:
            sex = "Unknown"

        for pr in exp.get("probeResponses", []):
            probe_name = pr["probe"]["id"]
            responses = pr["responses"]

            # Group responses by dose (raw, unsanitized — keeps None/NaN)
            groups_by_dose_raw: dict[float, list] = {}
            for i, val in enumerate(responses):
                dose = treat_doses[i]
                groups_by_dose_raw.setdefault(dose, []).append(val)

            # Skip endpoints without control
            if 0.0 not in groups_by_dose_raw:
                continue

            # Sanitized groups for Python Jonckheere (needs clean values)
            groups_by_dose = {
                dose: [v for v in vals
                       if v is not None
                       and not (isinstance(v, float) and math.isnan(v))]
                for dose, vals in groups_by_dose_raw.items()
            }

            # For Java batching, use the RAW response vector (preserving
            # the original sample count so all endpoints have the same
            # number of columns).  Replace None with 0.0 for JSON safety —
            # None values are exceedingly rare in apical data and would
            # represent missing observations (dead animals), which don't
            # meaningfully affect Williams/Dunnett on the remaining groups.
            flat_responses = [
                v if v is not None else 0.0
                for v in responses
            ]

            batch_key = f"{exp_name}::{probe_name}"
            batch_probe_ids.append(batch_key)
            batch_responses.append(flat_responses)
            endpoint_meta.append({
                "exp_name": exp_name,
                "probe_name": probe_name,
                "sex": sex,
                "groups_by_dose": groups_by_dose,
                "sorted_doses": sorted(groups_by_dose.keys()),
                "exp_doses": exp_doses,
            })

    # Run Java prefilter — one batch per unique dose design.
    # Different experiments can have different numbers of treatments/replicates
    # (e.g., female_hematology has 37 samples, male_hematology has 30).
    # Group endpoints by their dose vector and batch each group separately.
    java_results: dict[str, dict] = {}
    if batch_probe_ids:
        # Group by dose vector (as tuple for hashability)
        from collections import defaultdict
        dose_groups: dict[tuple, list[int]] = defaultdict(list)
        for idx, meta in enumerate(endpoint_meta):
            dose_key = tuple(meta["exp_doses"])
            dose_groups[dose_key].append(idx)

        for dose_key, indices in dose_groups.items():
            group_doses = list(dose_key)
            group_ids = [batch_probe_ids[i] for i in indices]
            group_responses = [batch_responses[i] for i in indices]

            try:
                # williams_p_cutoff=1.0 → always run Dunnett's for apical
                # endpoints.  Jonckheere (Python) is the trend gatekeeper
                # for apical data, not Williams.  With ~20 endpoints,
                # running Dunnett's on all of them is fast.
                result = _call_java_prefilter(
                    doses=group_doses,
                    probe_ids=group_ids,
                    responses=group_responses,
                    num_permutations=1000,
                    num_threads=4,
                    dunnett_simulations=15000,
                    williams_p_cutoff=1.0,
                )
                for probe_result in result.get("probes", []):
                    java_results[probe_result["probe_id"]] = probe_result
            except (RuntimeError, FileNotFoundError) as e:
                import sys
                print(f"  WARNING: Batch Java prefilter failed for "
                      f"{len(group_ids)} endpoints: {e}", file=sys.stderr)

    # ---- Phase 2: Build EndpointStats from Jonckheere + Java results ----
    tables: dict[str, list[TableRow]] = {}

    for i, meta in enumerate(endpoint_meta):
        exp_name = meta["exp_name"]
        probe_name = meta["probe_name"]
        sex = meta["sex"]
        groups_by_dose = meta["groups_by_dose"]
        sorted_doses = meta["sorted_doses"]

        if sex not in tables:
            tables[sex] = []

        # Descriptive statistics
        control_dose = 0.0
        treatment_doses = [d for d in sorted_doses if d != control_dose]
        n_per_dose = {}
        mean_per_dose = {}
        se_per_dose = {}

        for dose in sorted_doses:
            vals = groups_by_dose[dose]
            n = len(vals)
            if n == 0:
                n_per_dose[dose] = 0
                mean_per_dose[dose] = None
                se_per_dose[dose] = None
                continue
            mean_val = sum(vals) / n
            n_per_dose[dose] = n
            mean_per_dose[dose] = mean_val
            if n > 1:
                variance = sum((v - mean_val) ** 2 for v in vals) / (n - 1)
                se_per_dose[dose] = math.sqrt(variance / n)
            else:
                se_per_dose[dose] = 0.0

        # Jonckheere trend test (stays in Python — NTP convention)
        jonck_p, direction = jonckheere_test(groups_by_dose)
        jonckheere_sig = jonck_p <= ALPHA_STAR
        trend_marker = _sig_marker(jonck_p)

        # Pairwise p-values from Java batch.
        # NaN/null p-values (from zero-variance endpoints where Dunnett's
        # can't compute a test statistic) are treated as non-significant (1.0).
        batch_key = f"{exp_name}::{probe_name}"
        java_probe = java_results.get(batch_key)

        pairwise_p: dict[float, float] = {}
        if java_probe:
            for dose_str, pval in java_probe.get("dunnett_p", {}).items():
                if pval is None or (isinstance(pval, str)):
                    pval = 1.0
                pairwise_p[float(dose_str)] = float(pval)
            pairwise_method = "williams" if jonckheere_sig else "dunnett"
        else:
            pairwise_method = "none"

        pairwise_marker = {d: _sig_marker(p) for d, p in pairwise_p.items()}
        any_pairwise_sig = any(p <= ALPHA_STAR for p in pairwise_p.values())
        report_bmd = jonckheere_sig and any_pairwise_sig

        # Look up BMD/BMDL and NIEHS classification status.
        # BMD modeling and statistical significance are INDEPENDENT concerns:
        #   - BMDExpress 3 runs its own prefilter → produces bMDResult
        #   - Our NTP stats (Jonckheere + Dunnett) → determine LOEL/NOEL
        # The reference report Table 8 shows both side by side.  We always
        # show the BMDExpress 3 classification if a bMDResult exists,
        # regardless of our stat gate.
        #
        # NIEHS classification bins (from reference report Appendix D):
        #   "viable" → report BMD and BMDL values
        #   "NR"     → not reportable (BMD < 1/3 LNZD); show "<LNZD/3"
        #   "NVM"    → nonviable model; show "NVM"
        #   "UREP"   → unreliable estimate of potency; show "UREP"
        #   "failure" → model did not complete; show "NVM"
        #   None     → endpoint not modeled by BMDExpress 3
        bmd_str = "—"
        bmdl_str = "—"
        bmd_status = None

        # Primary source: BMDExpress 3 native bMDResult
        bmd_entry = bmd_result_lookup.get((exp_name, probe_name))
        if bmd_entry:
            status = bmd_entry["status"]
            bmd_status = status

            if status == "viable":
                bmd_val = bmd_entry["bmd"]
                bmdl_val = bmd_entry["bmdl"]
                if bmd_val is not None and bmdl_val is not None:
                    bmd_str = _format_bmd(bmd_val)
                    bmdl_str = _format_bmd(bmdl_val)
                else:
                    bmd_str = "NVM"
                    bmdl_str = "NVM"
                    bmd_status = "NVM"
            elif status == "NR":
                # BMD below extrapolation limit — report as <1/3 LNZD.
                # BMDL is not reported per NIEHS convention.
                lnzd = sorted_doses[1] if len(sorted_doses) > 1 else 0
                bmd_str = f"<{_format_bmd(lnzd / 3)}"
                bmdl_str = "—"
            elif status == "UREP":
                bmd_str = "UREP"
                bmdl_str = "UREP"
            elif status == "NVM":
                bmd_str = "NVM"
                bmdl_str = "NVM"
            else:
                # "failure" or unknown
                bmd_str = "NVM"
                bmdl_str = "NVM"
                bmd_status = "NVM"
        else:
            # Fallback: category analysis (genomics endpoints via GO terms)
            cat = category_lookup.get((exp_name, probe_name))
            if cat:
                try:
                    bmd_val = float(cat["bmd"])
                    bmdl_val = float(cat["bmdl"])
                    bmd_str = _format_bmd(bmd_val)
                    bmdl_str = _format_bmd(bmdl_val)
                    bmd_status = "viable"
                except (ValueError, KeyError, TypeError):
                    pass

        # Build the formatted values for each dose group
        values_by_dose = {}
        n_by_dose = {}

        for dose in sorted_doses:
            n = n_per_dose[dose]
            mean = mean_per_dose[dose]
            se = se_per_dose[dose]
            n_by_dose[dose] = n

            # Format: "mean ± SE" with significance marker
            if dose == control_dose:
                marker = trend_marker
            else:
                marker = pairwise_marker.get(dose, "")

            if se is not None and se > 0:
                # Adaptive decimal places: use more for small values so
                # precision isn't lost (e.g., 0.0034 ± 0.0012 instead of
                # 0.0 ± 0.0).  Large values (≥1) keep 1 decimal place to
                # match the NTP reference format.
                dp = _adaptive_decimals(mean, se)
                values_by_dose[dose] = f"{mean:.{dp}f} ± {se:.{dp}f}{marker}"
            elif mean is not None:
                dp = _adaptive_decimals(mean)
                values_by_dose[dose] = f"{mean:.{dp}f}{marker}"
            else:
                values_by_dose[dose] = "—"

        # Derive LOEL, NOEL, and direction of change for Table 8.
        # LOEL: lowest treatment dose with a significant (p ≤ 0.05) pairwise
        #   difference from control.  Per NIEHS: "the lowest dose demonstrating
        #   a significant (p ≤ 0.05) pairwise difference relative to the
        #   vehicle control group."
        # NOEL: highest dose NOT showing a significant pairwise difference.
        #   Per NIEHS: "the highest dose not showing a significant (p ≤ 0.05)
        #   pairwise difference relative to the vehicle control group."
        # Direction: UP or DOWN based on Jonckheere trend direction.
        loel = None
        noel = None
        for d in treatment_doses:
            if pairwise_p.get(d, 1.0) <= ALPHA_STAR:
                if loel is None or d < loel:
                    loel = d
        if loel is not None:
            # NOEL = highest dose below LOEL that is not significant
            candidates = [d for d in treatment_doses if d < loel]
            noel = max(candidates) if candidates else control_dose
        else:
            # No significant pairwise → NOEL is the highest dose tested
            noel = max(treatment_doses) if treatment_doses else None

        # Direction from Jonckheere: "UP" or "DOWN" string.
        # Only show direction when the Jonckheere trend test is significant,
        # matching the NIEHS reference convention (Table 8 shows direction
        # only for endpoints with a significant monotonic trend).
        dir_str = ""
        if jonckheere_sig and direction in ("UP", "DOWN"):
            dir_str = direction

        row = TableRow(
            label=probe_name,
            values_by_dose=values_by_dose,
            n_by_dose=n_by_dose,
            bmd_str=bmd_str,
            bmdl_str=bmdl_str,
            trend_marker=trend_marker,
            bmd_status=bmd_status,
            loel=loel,
            noel=noel,
            direction=dir_str,
            responsive=report_bmd,
        )
        tables[sex].append(row)

        # Collect dose-response summary stats for optional BMDS modeling.
        # These are the same means/SDs/Ns we computed above, packaged in a
        # format that apical_bmds.run_bmds_for_endpoints() expects.
        # Stored on the TableRow so the orchestrator can extract them without
        # re-parsing the integrated JSON.
        _sorted_doses = sorted(groups_by_dose.keys())
        row._bmds_input = {
            "key": f"{sex}::{probe_name}",
            "doses": _sorted_doses,
            "ns": [n_per_dose.get(d, 0) for d in _sorted_doses],
            "means": [
                mean_per_dose.get(d, 0.0) if mean_per_dose.get(d) is not None else 0.0
                for d in _sorted_doses
            ],
            "stdevs": [
                # Convert SE back to SD: SD = SE * sqrt(n)
                (se_per_dose.get(d, 0.0) or 0.0) * math.sqrt(n_per_dose.get(d, 1))
                for d in _sorted_doses
            ],
        }

    return tables


def annotate_missing_animals(
    domain_tables: dict[str, dict[str, list]],
    xlsx_rosters: dict[str, dict],
    reference_domain: str = "body_weight",
) -> None:
    """
    Annotate TableRows with missing-animal counts by comparing bm2 N counts
    against the xlsx study file roster.

    Dead animals are detected by comparing the bm2 N per dose (survivors
    with measurements) against the xlsx Core Animals roster (full assignment).
    The difference is recorded as missing_animals_by_dose on each TableRow.

    For domains whose xlsx roster has fewer dose groups than the reference
    domain (e.g., clinical chemistry missing 333/1000 mg/kg), the entire
    dose group is flagged as missing.

    Mutates TableRows in place — no return value.

    Args:
        domain_tables:    {domain → {sex → [TableRow, ...]}}, as from _partition_by_domain.
        xlsx_rosters:     From integrated["_meta"]["xlsx_rosters"].
        reference_domain: Domain with the most complete roster (usually body_weight).
    """
    if not xlsx_rosters:
        return

    # Use the reference domain (body weight) to establish the full Core
    # Animals roster — all dose groups and all animals expected to survive
    # to terminal sacrifice.  If no body weight xlsx exists, use the domain
    # with the most dose groups.
    ref_roster = xlsx_rosters.get(reference_domain)
    if not ref_roster:
        # Fallback: domain with most dose groups
        ref_roster = max(
            xlsx_rosters.values(),
            key=lambda r: len(r.get("dose_groups", [])),
            default=None,
        )
    if not ref_roster:
        return

    # Reference: dose (float) → sex → expected Core Animal count.
    # Dose keys may be strings from JSON — normalize to float.
    ref_core_raw = ref_roster.get("core_animals_by_dose_sex", {})
    ref_core = {float(k): v for k, v in ref_core_raw.items()}

    for domain, sex_rows in domain_tables.items():
        # Get this domain's xlsx roster for comparison.
        # Normalize dose keys to float (may be strings from JSON).
        dom_roster = xlsx_rosters.get(domain, {})
        dom_core_raw = dom_roster.get("core_animals_by_dose_sex", {})
        dom_core = {float(k): v for k, v in dom_core_raw.items()}

        for sex, rows in sex_rows.items():
            for row in rows:
                missing = {}
                for dose, actual_n in row.n_by_dose.items():
                    # Expected N: from this domain's xlsx, or from reference
                    # domain if this domain's xlsx doesn't have the dose group
                    dom_sex_aids = dom_core.get(dose, {}).get(sex)
                    ref_sex_aids = ref_core.get(dose, {}).get(sex)

                    if dom_sex_aids is not None:
                        expected_n = len(dom_sex_aids)
                    elif ref_sex_aids is not None:
                        # This domain doesn't have this dose group in its
                        # xlsx — all animals at this dose died before this
                        # endpoint could be measured
                        expected_n = len(ref_sex_aids)
                    else:
                        continue

                    diff = expected_n - actual_n
                    if diff > 0:
                        missing[dose] = diff

                # Also flag dose groups that are in the reference but
                # completely absent from this endpoint's data
                for dose in ref_core:
                    if dose not in row.n_by_dose:
                        sex_aids = ref_core[dose].get(sex)
                        if sex_aids:
                            missing[dose] = len(sex_aids)

                row.missing_animals_by_dose = missing


def backfill_missing_doses(
    domain_tables: dict[str, dict[str, list]],
    xlsx_rosters: dict[str, dict],
    reference_domain: str = "body_weight",
) -> None:
    """
    Ensure every TableRow uses the full study dose list as columns, filling
    in "–" (em-dash) for dose groups where no animals survived to be measured.

    The NIEHS reference report shows ALL dose columns in every domain table,
    even when an entire dose group is absent (all animals died).  Those
    columns display "–" for the value and 0 for N, making it visually clear
    that the study included that dose but no data was collected.

    This function runs AFTER annotate_missing_animals() so that the
    missing_animals_by_dose dict is already populated.  It:
      1. Determines the full dose list from the reference xlsx roster
      2. For each TableRow, adds any missing doses with "–" values and N=0

    Mutates TableRows in place — no return value.

    Args:
        domain_tables:    {domain → {sex → [TableRow, ...]}}, as from _partition_by_domain.
        xlsx_rosters:     From integrated["_meta"]["xlsx_rosters"].
        reference_domain: Domain with the most complete roster (usually body_weight).
    """
    if not xlsx_rosters:
        return

    # Determine the full study dose list from the reference domain's roster.
    # The body_weight xlsx always has every dose group because body weight
    # is measured on all animals regardless of survival.
    ref_roster = xlsx_rosters.get(reference_domain)
    if not ref_roster:
        # Fallback: domain with the most dose groups
        ref_roster = max(
            xlsx_rosters.values(),
            key=lambda r: len(r.get("dose_groups", [])),
            default=None,
        )
    if not ref_roster:
        return

    # Full dose list from the reference roster, sorted ascending.
    # Dose keys may be strings from JSON — normalize to float.
    ref_doses_raw = ref_roster.get("dose_groups", [])
    full_doses = sorted(float(d) for d in ref_doses_raw)
    if not full_doses:
        return

    for domain, sex_rows in domain_tables.items():
        for sex, rows in sex_rows.items():
            for row in rows:
                # Add any doses from the full list that are absent from
                # this row's data.  These are dose groups where no animals
                # survived to have this endpoint measured.
                for dose in full_doses:
                    if dose not in row.values_by_dose:
                        row.values_by_dose[dose] = "–"
                        row.n_by_dose[dose] = 0


def _format_bmd(value: float) -> str:
    """
    Format a BMD or BMDL value for display in the report table.

    Uses fixed-point notation with 3 significant figures, matching the
    precision level in NTP reports.

    Args:
        value: BMD or BMDL value (in dose units, e.g., mg/kg).

    Returns:
        Formatted string, e.g., "8.492" or "0.00679".
    """
    if value == 0:
        return "0"
    # Use 3 significant figures
    return f"{value:.3g}"


# ---------------------------------------------------------------------------
# Narrative generation — auto-generate NTP-style prose from table data
# ---------------------------------------------------------------------------
# These helpers and the main generate_results_narrative() function produce
# the "Results" paragraphs that appear between the section heading and the
# tables in an NTP subchronic study report.  The prototype structure comes
# from the PFHxSAm study report:
#   - Paragraph about body weight findings
#   - Paragraphs about organ weight findings per sex (significant + non-sig)
# Clinical observations (paragraph 1 in real reports) must be user-supplied
# since that data isn't in the .bm2 file.

def _parse_cell_mean(cell_text: str) -> float | None:
    """
    Extract the numeric mean from a formatted table cell value.

    Table cells look like "330.2 ± 5.1**" — this function pulls out
    the leading number (330.2) and ignores the SE and significance markers.

    Args:
        cell_text: Formatted cell string, e.g., "330.2 ± 5.1**".

    Returns:
        The mean as a float, or None if parsing fails.
    """
    if not cell_text or cell_text.strip() == "–":
        return None
    # Match the leading number (possibly negative, with decimal)
    match = re.match(r"(-?\d+\.?\d*)", cell_text.strip())
    if match:
        return float(match.group(1))
    return None


def _lowest_sig_dose(row: TableRow) -> float | None:
    """
    Find the lowest non-control dose group that has pairwise significance.

    Significance is indicated by * or ** appended to the cell value.
    The control dose (0.0) is excluded because its marker represents the
    trend test, not pairwise comparison.

    Args:
        row: A TableRow with values_by_dose containing formatted strings.

    Returns:
        The lowest dose (float) with a pairwise significance marker,
        or None if no treatment dose is significant.
    """
    sorted_doses = sorted(row.values_by_dose.keys())
    for dose in sorted_doses:
        # Skip control — its marker is for the trend test (Jonckheere)
        if dose == 0.0:
            continue
        val = row.values_by_dose.get(dose, "")
        if "*" in val:
            return dose
    return None


def _endpoint_direction(row: TableRow) -> str:
    """
    Determine whether an endpoint increased or decreased relative to control.

    Compares the control group mean to the highest-dose group mean.  This
    is a simplification — real studies might have non-monotonic responses —
    but it matches the NTP report convention of reporting the overall
    direction of the dose-response.

    Args:
        row: A TableRow with values_by_dose.

    Returns:
        "increase" or "decrease".
    """
    sorted_doses = sorted(row.values_by_dose.keys())
    if len(sorted_doses) < 2:
        return "increase"  # fallback

    control_mean = _parse_cell_mean(row.values_by_dose.get(sorted_doses[0], ""))
    highest_mean = _parse_cell_mean(row.values_by_dose.get(sorted_doses[-1], ""))

    if control_mean is None or highest_mean is None:
        return "increase"  # fallback

    return "decrease" if highest_mean < control_mean else "increase"


def _fmt_dose(dose: float) -> str:
    """
    Format a dose value for prose output.

    Drops trailing ".0" for integer-valued doses to match natural writing
    style:  12.0 → "12", 0.15 → "0.15".

    Args:
        dose: Dose value in dose units.

    Returns:
        Clean string representation.
    """
    if dose == int(dose):
        return str(int(dose))
    return str(dose)


def _parse_organ_label(label: str) -> tuple[str, str]:
    """
    Parse an endpoint label into (organ_name, weight_type).

    The .bm2 endpoint labels follow patterns like:
        "Liver Absolute"       → ("Liver", "absolute")
        "Liver Weight Relative" → ("Liver", "relative")
        "R. Kidney Absolute"   → ("R. Kidney", "absolute")
        "Kidney-Left Absolute" → ("Kidney-Left", "absolute")
        "Terminal Body Wt."    → ("", "body_weight")
        "Body Weight"          → ("", "body_weight")
        "Body Weight Gain"     → ("", "body_weight")
        "SD5"                  → ("", "body_weight")   # Study Day 5 terminal wt
        "Alanine aminotransferase" → ("Alanine aminotransferase", "endpoint")

    The organ_name is used to group absolute/relative pairs when
    describing organ weight findings in prose.  The "endpoint" type
    is for clinical pathology or other non-weight endpoints — their
    narrative prose omits the word "weight".

    Args:
        label: The endpoint label string from the TableRow.

    Returns:
        A 2-tuple (organ_name, weight_type) where weight_type is one of:
        "absolute", "relative", "body_weight", or "endpoint".
    """
    lower = label.lower()

    # Body weight endpoints — various naming conventions from BMDExpress
    if "body" in lower and ("wt" in lower or "weight" in lower):
        return ("", "body_weight")

    # "SD5", "SD 5", etc. — Study Day N terminal body weight labels.
    # BMDExpress uses these short codes for the terminal body weight
    # measurement at the end of a 5-day study.
    if re.match(r"^sd\s*\d+$", lower):
        return ("", "body_weight")

    # Organ weight endpoints — last word is "Absolute" or "Relative".
    # Also strip redundant "Weight" from the organ name so we get
    # "Liver" instead of "Liver Weight" (avoids "Liver Weight relative
    # weight" double-weight in prose).
    if lower.endswith("absolute"):
        organ = label[:-len("Absolute")].strip()
        # Strip trailing "Weight" / "Wt" / "Wt." from organ name
        organ = re.sub(r"\s+(?:Weight|Wt\.?)$", "", organ, flags=re.IGNORECASE)
        return (organ, "absolute")
    elif lower.endswith("relative"):
        organ = label[:-len("Relative")].strip()
        organ = re.sub(r"\s+(?:Weight|Wt\.?)$", "", organ, flags=re.IGNORECASE)
        return (organ, "relative")

    # Fallback: clinical pathology or other non-weight endpoints.
    # Marked as "endpoint" so the narrative generator can omit the
    # word "weight" from prose (e.g., "ALT was significantly increased"
    # rather than "ALT absolute weight was significantly increased").
    return (label, "endpoint")


def _oxford_comma(items: list[str], conjunction: str = "or") -> str:
    """
    Join a list of strings with an Oxford comma.

    Follows the NTP report convention:
        []            → ""
        ["a"]         → "a"
        ["a", "b"]    → "a or b"
        ["a", "b", "c"] → "a, b, or c"

    Args:
        items:       List of strings to join.
        conjunction: The conjunction word ("or", "and").

    Returns:
        The joined string with Oxford comma formatting.
    """
    if len(items) == 0:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} {conjunction} {items[1]}"
    return ", ".join(items[:-1]) + f", {conjunction} {items[-1]}"


def generate_results_narrative(
    table_data: dict[str, list[TableRow]],
    compound_name: str,
    dose_unit: str = "mg/kg",
    start_table_num: int = 1,
) -> list[str]:
    """
    Auto-generate NTP-style results narrative paragraphs from table data.

    Produces the prose that appears between the section heading and the
    statistical tables in an NTP report.  Follows the structure of the
    PFHxSAm prototype:

      1. Body weight paragraph (significant or not, per sex)
      2+ Organ weight paragraphs per sex (significant endpoints with
         BMD/BMDL, then non-significant endpoint list)

    Clinical observations (animal condition, mortality) are NOT generated
    here because that data isn't in the .bm2 file — those are user-supplied.

    Args:
        table_data:      Dict mapping sex label ("Male"/"Female") → list of
                         TableRow, as returned by build_table_data().
        compound_name:   Chemical name for use in prose.
        dose_unit:       Dose unit string (e.g., "mg/kg").
        start_table_num: Starting table number for cross-references.

    Returns:
        A list of paragraph strings, ready for insertion into the .docx
        or display in the UI textarea.
    """
    paragraphs = []

    # ---------------------------------------------------------------
    # Body weight paragraph — look for body weight rows across sexes
    # ---------------------------------------------------------------
    # Collect body weight findings for each sex to build a single paragraph
    bw_findings = []  # list of per-sex description strings
    table_num = start_table_num  # track which table number each sex maps to

    # We need the table number for each sex to reference in prose
    sex_table_num = {}
    tnum = start_table_num
    for sex in ["Male", "Female"]:
        if sex in table_data and table_data[sex]:
            sex_table_num[sex] = tnum
            tnum += 1

    for sex in ["Male", "Female"]:
        rows = table_data.get(sex, [])
        bw_rows = [r for r in rows if _parse_organ_label(r.label)[1] == "body_weight"]

        for bw_row in bw_rows:
            tbl_ref = sex_table_num.get(sex, start_table_num)
            if bw_row.bmd_str == "ND":
                # Not significant — report as such
                bw_findings.append(
                    f"{sex.lower()} rats (Table {tbl_ref})"
                )
            else:
                # Significant — report direction, lowest sig dose, trend, BMD
                direction = _endpoint_direction(bw_row)
                low_dose = _lowest_sig_dose(bw_row)
                # Trend direction label for prose
                trend_dir = "positive" if direction == "increase" else "negative"

                parts = []
                parts.append(
                    f"Terminal body weight was significantly "
                    f"{'increased' if direction == 'increase' else 'decreased'} "
                    f"in {sex.lower()} rats"
                )
                if low_dose is not None:
                    parts.append(
                        f" at ≥{_fmt_dose(low_dose)} {dose_unit}"
                    )
                parts.append(
                    f" with a {trend_dir} trend "
                    f"(Table {tbl_ref}). "
                    f"The BMD and BMDL were {bw_row.bmd_str} and "
                    f"{bw_row.bmdl_str} {dose_unit}, respectively."
                )
                bw_findings.append("".join(parts))

    # Build the body weight paragraph
    if bw_findings:
        # Check if ALL findings are non-significant (just table refs)
        all_nd = all("Table" in f and f.startswith(("male", "female"))
                     for f in bw_findings)
        if all_nd:
            # All non-significant — single sentence
            sex_refs = _oxford_comma(bw_findings, conjunction="or")
            paragraphs.append(
                f"No significant changes in terminal body weight for "
                f"{sex_refs} occurred with exposure to {compound_name}."
            )
        else:
            # At least one significant — join the individual findings
            paragraphs.append(" ".join(bw_findings))

    # ---------------------------------------------------------------
    # Organ weight paragraphs — one block per sex
    # ---------------------------------------------------------------
    for sex in ["Male", "Female"]:
        rows = table_data.get(sex, [])
        if not rows:
            continue

        tbl_ref = sex_table_num.get(sex, start_table_num)

        # Separate non-body-weight rows (organ weights + clinical endpoints)
        organ_rows = [r for r in rows if _parse_organ_label(r.label)[1] != "body_weight"]
        if not organ_rows:
            continue

        # Detect whether this section is organ weights or clinical pathology.
        # If ANY row has "absolute" or "relative" type, it's an organ weight
        # section.  Otherwise it's clinical pathology / generic endpoints.
        all_types = {_parse_organ_label(r.label)[1] for r in organ_rows}
        is_organ_section = bool(all_types & {"absolute", "relative"})

        # Split into significant (BMD reported) and non-significant (BMD = "ND")
        sig_rows = [r for r in organ_rows if r.bmd_str != "ND"]
        nonsig_rows = [r for r in organ_rows if r.bmd_str == "ND"]

        if not sig_rows:
            # No significant findings for this sex
            endpoint_noun = "organ weights" if is_organ_section else "endpoints"
            paragraphs.append(
                f"In {sex.lower()} rats at study termination, there were no "
                f"{endpoint_noun} that exhibited significant trend and pairwise "
                f"comparisons (Table {tbl_ref})."
            )
            continue

        # Group significant rows by organ/endpoint name so we can pair
        # absolute/relative for organ weights, or list individual clinical
        # pathology endpoints
        organ_groups: dict[str, list[TableRow]] = {}
        for r in sig_rows:
            organ_name, wtype = _parse_organ_label(r.label)
            organ_groups.setdefault(organ_name, []).append(r)

        # Build per-organ/endpoint finding descriptions
        sex_findings = []
        for organ_name, group_rows in organ_groups.items():
            # Determine direction from the first row (abs and rel should agree)
            direction = _endpoint_direction(group_rows[0])
            dir_word = "increased" if direction == "increase" else "decreased"
            trend_dir = "positive" if direction == "increase" else "negative"

            # Find lowest pairwise-significant dose across the group
            low_doses = [_lowest_sig_dose(r) for r in group_rows]
            low_doses = [d for d in low_doses if d is not None]
            min_low_dose = min(low_doses) if low_doses else None

            # Determine weight types present in this group
            weight_types = []
            for r in group_rows:
                _, wt = _parse_organ_label(r.label)
                weight_types.append(wt)

            # Collect BMD/BMDL values, qualifying with weight type only
            # for organ weight endpoints (not clinical pathology)
            bmd_parts = []
            for r in group_rows:
                _, wt = _parse_organ_label(r.label)
                if wt in ("absolute", "relative"):
                    bmd_parts.append(
                        f"{r.bmd_str} and {r.bmdl_str} {dose_unit} "
                        f"({wt} weight)"
                    )
                else:
                    bmd_parts.append(
                        f"{r.bmd_str} and {r.bmdl_str} {dose_unit}"
                    )

            # Compose the finding sentence — phrasing differs between
            # organ weight endpoints ("liver relative weight was...") and
            # clinical pathology endpoints ("ALT was...").
            has_weight_types = any(
                wt in ("absolute", "relative") for wt in weight_types
            )
            if has_weight_types:
                # Organ weight: "Liver absolute and relative weight was..."
                wt_type_str = _oxford_comma(
                    sorted(set(wt for wt in weight_types
                               if wt in ("absolute", "relative"))),
                    conjunction="and",
                )
                finding = (
                    f"{organ_name} {wt_type_str} weight was significantly "
                    f"{dir_word}"
                )
            else:
                # Clinical pathology / generic: "ALT was significantly..."
                finding = (
                    f"{organ_name} was significantly {dir_word}"
                )

            if min_low_dose is not None:
                finding += f" at ≥{_fmt_dose(min_low_dose)} {dose_unit}"
            finding += f" with a {trend_dir} trend"

            # Add BMD/BMDL
            if len(bmd_parts) == 1:
                finding += (
                    f". The BMD and BMDL were {bmd_parts[0]}, respectively."
                )
            else:
                finding += (
                    ". The BMD and BMDL were "
                    + "; ".join(bmd_parts)
                    + ", respectively."
                )

            sex_findings.append(finding)

        # Opening line for significant findings
        sex_para = (
            f"In {sex.lower()} rats at study termination (Table {tbl_ref}), "
        )
        sex_para += " ".join(sex_findings)

        # Non-significant endpoints listed at the end
        if nonsig_rows:
            nonsig_labels = [r.label.lower() for r in nonsig_rows]
            nonsig_str = _oxford_comma(nonsig_labels, conjunction="or")
            sex_para += (
                f" Significant trend and pairwise comparisons were not "
                f"observed in {nonsig_str}."
            )

        paragraphs.append(sex_para)

    return paragraphs


# ---------------------------------------------------------------------------
# .docx table generation — build the Word document tables
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text: str, bold: bool = False, size: int = TABLE_FONT_SIZE):
    """
    Set the text of a Word table cell with consistent formatting.

    Clears any existing content, sets font to Calibri at the specified
    size, and optionally makes it bold.  Also sets minimal paragraph
    spacing for compact tables.

    Args:
        cell:  The docx table cell to modify.
        text:  The text to display.
        bold:  Whether to make the text bold.
        size:  Font size in points (default TABLE_FONT_SIZE).
    """
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold


def generate_table(
    doc: Document,
    sex: str,
    rows: list[TableRow],
    table_number: int,
    compound_name: str,
    dose_unit: str = "mg/kg",
    table_caption_template: str = (
        "Summary of Body Weights and Organ Weights "
        "of {sex} Rats Administered {compound} for Five Days"
    ),
) -> None:
    """
    Generate one NTP-style summary table in the Word document.

    Creates a table matching the PFHxSAm prototype format:
      - Title row with table number, compound name, and sex
      - Header row with "Endpoint" + dose columns + BMD/BMDL columns
      - "n" row showing sample sizes per dose group
      - One row per endpoint with mean ± SE and significance markers
      - Footnotes explaining significance markers and statistical methods

    Args:
        doc:                    The Document to add the table to.
        sex:                    "Male" or "Female".
        rows:                   List of TableRow objects for this sex.
        table_number:           Table number for the caption (e.g., 3).
        compound_name:          Chemical name for the caption.
        dose_unit:              Dose unit string (default "mg/kg").
        table_caption_template: Format string for the table caption.
                                Must contain {sex} and {compound} placeholders.
                                Default matches the NTP body/organ weight format.
    """
    if not rows:
        return

    # Get the dose groups from the first row (all rows share the same doses)
    sorted_doses = sorted(rows[0].values_by_dose.keys())

    # Column layout: Endpoint | dose_0 | dose_1 | ... | dose_n | BMD1Std | BMDL1Std
    n_cols = 1 + len(sorted_doses) + 2
    # Rows: header + n row + data rows
    n_rows = 1 + 1 + len(rows)

    # -- Table caption --
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(12)
    caption.paragraph_format.space_after = Pt(4)
    # Format the caption from the template, filling in sex and compound name
    caption_text = table_caption_template.format(
        sex=sex, compound=compound_name,
    )
    run = caption.add_run(f"Table {table_number}. {caption_text}")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.bold = True

    # -- Create the table --
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # -- Header row --
    header_row = table.rows[0]
    _set_cell_text(header_row.cells[0], "Endpoint", bold=True, size=HEADER_FONT_SIZE)

    for i, dose in enumerate(sorted_doses):
        # Format dose column header: "0 mg/kg", "0.15 mg/kg", etc.
        if dose == 0:
            label = f"0 {dose_unit}"
        elif dose == int(dose):
            label = f"{int(dose)} {dose_unit}"
        else:
            label = f"{dose} {dose_unit}"
        _set_cell_text(header_row.cells[1 + i], label, bold=True, size=HEADER_FONT_SIZE)

    _set_cell_text(header_row.cells[-2], f"BMD1Std\n({dose_unit})", bold=True, size=HEADER_FONT_SIZE)
    _set_cell_text(header_row.cells[-1], f"BMDL1Std\n({dose_unit})", bold=True, size=HEADER_FONT_SIZE)

    # -- "n" row (sample sizes per dose group) --
    n_row = table.rows[1]
    _set_cell_text(n_row.cells[0], "n", bold=True, size=TABLE_FONT_SIZE)
    # Use the first data row's n values (they may vary slightly across
    # endpoints if some animals died, so we show the max across endpoints)
    for i, dose in enumerate(sorted_doses):
        # Find max n across all endpoints for this dose
        max_n = max(row.n_by_dose.get(dose, 0) for row in rows)
        _set_cell_text(n_row.cells[1 + i], str(max_n), size=TABLE_FONT_SIZE)
    _set_cell_text(n_row.cells[-2], "NA", size=TABLE_FONT_SIZE)
    _set_cell_text(n_row.cells[-1], "NA", size=TABLE_FONT_SIZE)

    # -- Data rows --
    for ri, row in enumerate(rows):
        tbl_row = table.rows[2 + ri]
        _set_cell_text(tbl_row.cells[0], row.label, bold=True, size=TABLE_FONT_SIZE)

        for i, dose in enumerate(sorted_doses):
            val = row.values_by_dose.get(dose, "–")
            _set_cell_text(tbl_row.cells[1 + i], val, size=TABLE_FONT_SIZE)

        _set_cell_text(tbl_row.cells[-2], row.bmd_str, size=TABLE_FONT_SIZE)
        _set_cell_text(tbl_row.cells[-1], row.bmdl_str, size=TABLE_FONT_SIZE)

    # -- Footnotes --
    footnotes = doc.add_paragraph()
    footnotes.paragraph_format.space_before = Pt(2)
    footnotes.paragraph_format.space_after = Pt(2)

    notes = [
        "* Statistically significant at p ≤ 0.05; ** p ≤ 0.01.",
        "Statistical analysis performed by Jonckheere (trend) and "
        "Williams or Dunnett (pairwise) tests.",
        "BMD1Std and BMDL1Std from BMDExpress 3 model averaging (defined "
        "category analysis).  ND = not determined (trend and/or pairwise "
        "significance criteria not met).",
    ]
    for note in notes:
        run = footnotes.add_run(note + "\n")
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.italic = True

    doc.add_paragraph()  # spacing after table


# ---------------------------------------------------------------------------
# Main entry point — generate the full report
# ---------------------------------------------------------------------------

def build_table_data_from_bm2(
    bm2_path: str,
    bmd_stat: str = "mean",
) -> tuple[dict[str, list[TableRow]], dict, dict]:
    """
    Convenience wrapper: export .bm2 → JSON + categories, run stats, return
    (table_data, category_lookup, bm2_json).

    Handles the full export-and-analyze pipeline so callers don't need
    to manage temp files or know about the BMDExpress Java classes.

    Caching: both the full BMDProject JSON and category analysis are cached
    in LMDB (bm2_cache module).  On first call: run Java exports → store
    in LMDB.  On subsequent calls: LMDB lookup + orjson.loads (no JVM).

    The category export now uses ExportCategories.java which reads
    CategoryAnalysisResults directly from the BMDProject via BMDExpress's
    native Java API — replacing the broken DataCombinerService TSV path.

    Args:
        bm2_path: Path to the BMDExpress 3 .bm2 file.
        bmd_stat: Which BMD aggregate statistic to use for the lookup.
                  One of: "mean", "median", "minimum", "weighted_mean",
                  "fifth_pct", "tenth_pct".  Defaults to "mean".

    Returns:
        A 3-tuple:
          - table_data: Dict mapping sex label ("Male"/"Female") → list of
            TableRow, ready for generate_table() or add_apical_tables_to_doc().
          - category_lookup: Dict from build_category_lookup(), mapping
            (experiment_prefix, endpoint_name) → BMD info dict.
          - bm2_json: The full deserialized BMDProject as a dict.
    """
    print(f"Processing: {bm2_path}")

    # Step 1: Load or export the full BMDProject JSON.
    # LMDB B-tree cache stores deserialized dicts keyed by file path.
    bm2_json = bm2_cache.get_json(bm2_path)
    category_lookup = bm2_cache.get_categories(bm2_path)

    if bm2_json is not None and category_lookup is not None:
        print("  Loaded BMDProject + categories from LMDB cache")
    else:
        # Export BMDProject JSON if not cached
        if bm2_json is None:
            print("  Exporting .bm2 → JSON (ExportBm2)...")
            with tempfile.NamedTemporaryFile(suffix=".json", delete=False) as tmp:
                tmp_json = tmp.name
            # Skip the TSV export — we use ExportCategories instead
            export_bm2(bm2_path, tmp_json, "NONE")
            with open(tmp_json, "rb") as f:
                import orjson
                bm2_json = orjson.loads(f.read())
            bm2_cache.put_json(bm2_path, bm2_json)
            os.unlink(tmp_json)

        # Export category analysis if not cached
        if category_lookup is None:
            print("  Exporting categories (ExportCategories)...")
            with tempfile.NamedTemporaryFile(suffix=".json", delete=False) as tmp:
                tmp_cat = tmp.name
            categories_json = export_categories(bm2_path, tmp_cat)
            # Pass experiment names so the lookup resolves BMDExpress
            # pipeline suffixes (e.g., "_williams_0.05_NOMTC_nofoldfilter")
            # back to raw experiment names used by build_table_data().
            exp_names = [
                exp.get("name", "")
                for exp in bm2_json.get("doseResponseExperiments", [])
            ]
            category_lookup = build_category_lookup(
                categories_json, bmd_stat, experiment_names=exp_names,
            )
            bm2_cache.put_categories(bm2_path, category_lookup)
            os.unlink(tmp_cat)

    print(f"  Category analysis: {len(category_lookup)} endpoints with BMD values")

    # Steps 3-5: Build table data (runs stats + cross-references + business rules).
    # Phase 1 batches all endpoints into a single Java RunPrefilter call
    # for Williams/Dunnett (BMDExpress-native stats via sciome-commons-math).
    # Phase 2 runs Jonckheere in Python and combines with Java results.
    print("  Running NTP statistical tests (Jonckheere + Java Williams/Dunnett)...")
    table_data = build_table_data(bm2_json, category_lookup)

    return table_data, category_lookup, bm2_json


def add_apical_tables_to_doc(
    doc: Document,
    table_data: dict[str, list[TableRow]],
    section_title: str,
    compound_name: str,
    dose_unit: str = "mg/kg",
    table_caption_template: str = (
        "Summary of Body Weights and Organ Weights "
        "of {sex} Rats Administered {compound} for Five Days"
    ),
    start_table_num: int = 1,
    narrative_paragraphs: list[str] | None = None,
) -> int:
    """
    Add apical endpoint tables to an existing Document object.

    This is the docx-building half of generate_report(), extracted so
    the web server can add NTP tables to a report document that already
    contains the background section paragraphs.

    Adds a section heading, optional narrative paragraphs (NTP-style prose
    describing body weight and organ weight findings), then one table per
    sex (Male, Female), each with its own caption, header row, n row, data
    rows, and footnotes.

    Args:
        doc:                    The Document to add tables to.
        table_data:             Dict mapping sex label → list of TableRow,
                                as returned by build_table_data_from_bm2().
        section_title:          Heading text for the section (e.g.,
                                "Animal Condition, Body Weights, and Organ Weights").
        compound_name:          Chemical name for table captions.
        dose_unit:              Dose unit for column headers (default "mg/kg").
        table_caption_template: Format string for per-table captions.  Must
                                contain {sex} and {compound} placeholders.
        start_table_num:        Starting table number (default 1).  Use this
                                when adding multiple .bm2 sections so table
                                numbers continue sequentially.
        narrative_paragraphs:   Optional list of prose paragraphs to insert
                                between the section heading and the first table.
                                These are the auto-generated (and possibly
                                user-edited) NTP-style results narrative.

    Returns:
        The next available table number (so multiple .bm2 files can be
        added sequentially with correct numbering).
    """
    # Section heading — parameterized so we can reuse for clinical pathology, etc.
    doc.add_heading(section_title, level=2)

    # Insert narrative paragraphs (if provided) between heading and tables.
    # These describe body weight and organ weight findings in NTP prose style.
    if narrative_paragraphs:
        for para_text in narrative_paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

    # Generate one table per sex, in a consistent order
    table_num = start_table_num
    for sex in ["Male", "Female"]:
        rows = table_data.get(sex, [])
        if rows:
            generate_table(
                doc, sex, rows, table_num,
                compound_name=compound_name,
                dose_unit=dose_unit,
                table_caption_template=table_caption_template,
            )
            table_num += 1

            # Print summary to console for debugging / CLI usage
            for row in rows:
                sig_doses = sum(
                    1 for d, v in row.values_by_dose.items()
                    if "*" in v and d != 0.0
                )
                print(f"    {sex} / {row.label}: "
                      f"BMD={row.bmd_str}, BMDL={row.bmdl_str}, "
                      f"{sig_doses} sig dose groups")

    return table_num


# ---------------------------------------------------------------------------
# NIEHS report table helpers — BMD Summary, Gene Set BMD, Gene BMD
# ---------------------------------------------------------------------------
# These three functions add new table types to a .docx document, matching
# the structure of NIEHS Report 10 (PFHxSAm, NBK589955).  They are
# composable: the server calls them to assemble the full NIEHS-style
# report alongside the existing apical endpoint tables.

def _format_table_cell(cell, text: str, bold: bool = False, size: int = 8) -> None:
    """
    Set the text and formatting of a single table cell.

    This is a small helper to avoid repeating font-size / font-name /
    bold assignment on every cell.  Matches the formatting style used
    by generate_table() for visual consistency across all report tables.

    Args:
        cell:  A python-docx table cell object.
        text:  The string to display in the cell.
        bold:  Whether the text should be bold (for headers / sex rows).
        size:  Font size in points (default 8, matching TABLE_FONT_SIZE).
    """
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(size)
            run.font.name = "Calibri"
            run.bold = bold


def add_bmd_summary_table_to_doc(
    doc: Document,
    endpoints: list[dict],
    table_num: int,
    dose_unit: str = "mg/kg",
) -> int:
    """
    Add an Apical Endpoint BMD Summary table to the document.

    This table aggregates BMD, BMDL, LOEL, and NOEL across all apical
    endpoints (organ weights + clinical pathology), sorted by BMD from
    low to high.  It provides a single-glance overview of which endpoints
    were most sensitive.

    The table is split by sex: a "Male" header row followed by male
    endpoints, then a "Female" header row followed by female endpoints.
    Within each sex block, endpoints are sorted by BMD ascending.

    Matching NIEHS format:
      Caption: "BMD, BMDL, LOEL, and NOEL Summary for Apical Endpoints,
                Sorted by BMD or LOEL from Low to High."
      Columns: Endpoint | BMD1Std (unit) | BMDL1Std (unit) | LOEL (unit) |
               NOEL (unit) | Direction of Change

    Args:
        doc:        The Document to add the table to.
        endpoints:  List of endpoint dicts, each with keys:
                      endpoint (str), bmd (str/float), bmdl (str/float),
                      loel (str/float), noel (str/float), direction (str),
                      sex (str — "Male" or "Female").
        table_num:  The sequential table number for the caption.
        dose_unit:  Unit string for column headers (default "mg/kg").

    Returns:
        The next available table number (table_num + 1).
    """
    # Section heading
    doc.add_heading("Apical Endpoint BMD Summary", level=3)

    # Table caption — styled as italic paragraph before the table
    caption = (
        f"Table {table_num}. BMD, BMDL, LOEL, and NOEL Summary for "
        f"Apical Endpoints, Sorted by BMD or LOEL from Low to High."
    )
    cap_para = doc.add_paragraph()
    run = cap_para.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.italic = True
    cap_para.paragraph_format.space_after = Pt(4)

    # Column headers
    headers = [
        "Endpoint",
        f"BMD₁Std ({dose_unit})",
        f"BMDL₁Std ({dose_unit})",
        f"LOEL ({dose_unit})",
        f"NOEL ({dose_unit})",
        "Direction of Change",
    ]

    # Split endpoints by sex for the grouped display
    male_eps = sorted(
        [e for e in endpoints if e.get("sex", "").lower() == "male"],
        key=lambda e: float(e.get("bmd", 9999)) if str(e.get("bmd", "ND")) != "ND" else 9999,
    )
    female_eps = sorted(
        [e for e in endpoints if e.get("sex", "").lower() == "female"],
        key=lambda e: float(e.get("bmd", 9999)) if str(e.get("bmd", "ND")) != "ND" else 9999,
    )

    # Calculate total rows: header + male sex row + male data + female sex row + female data
    n_rows = 1  # column header row
    if male_eps:
        n_rows += 1 + len(male_eps)    # "Male" row + data rows
    if female_eps:
        n_rows += 1 + len(female_eps)  # "Female" row + data rows

    if n_rows <= 1:
        # No endpoints to show — add a note and return
        p = doc.add_paragraph()
        run = p.add_run("No apical endpoints with BMD values available.")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.italic = True
        return table_num + 1

    table = doc.add_table(rows=n_rows, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        _format_table_cell(table.rows[0].cells[i], header, bold=True, size=HEADER_FONT_SIZE)

    # Fill data rows, grouped by sex
    row_idx = 1

    for sex_label, sex_eps in [("Male", male_eps), ("Female", female_eps)]:
        if not sex_eps:
            continue

        # Sex header row — merged across all columns, bold
        for i in range(len(headers)):
            _format_table_cell(table.rows[row_idx].cells[i], sex_label if i == 0 else "", bold=True)
        row_idx += 1

        # Individual endpoint rows
        for ep in sex_eps:
            bmd_val = ep.get("bmd", "ND")
            bmdl_val = ep.get("bmdl", "ND")
            loel_val = ep.get("loel", "—")
            noel_val = ep.get("noel", "—")
            direction = ep.get("direction", "")

            # Format numeric values using 3 significant figures (consistent
            # with _format_bmd used in domain tables).  Previous .2f format
            # lost precision for small values (e.g., 0.00679 → 0.01).
            def _fmt_val(v):
                if v is None or str(v) in ("ND", "—", ""):
                    return str(v) if v else "—"
                try:
                    return _format_bmd(float(v))
                except (ValueError, TypeError):
                    return str(v)

            row_data = [
                ep.get("endpoint", ""),
                _fmt_val(bmd_val),
                _fmt_val(bmdl_val),
                _fmt_val(loel_val),
                _fmt_val(noel_val),
                direction,
            ]
            for i, val in enumerate(row_data):
                _format_table_cell(table.rows[row_idx].cells[i], val)
            row_idx += 1

    doc.add_paragraph()  # spacing after table
    return table_num + 1


def add_gene_set_bmd_tables_to_doc(
    doc: Document,
    gene_sets_data: list[dict],
    organ: str,
    sex: str,
    table_num: int,
    dose_unit: str = "mg/kg",
    bmd_stat_label: str = "Median",
) -> int:
    """
    Add a Gene Set (GO Biological Process) BMD table to the document.

    Produces one table for a specific organ × sex combination (e.g.,
    "Liver Male").  The table shows the top GO terms ranked by BMD —
    lower BMD means the gene set was perturbed at a lower dose (more potent).

    The bmd_stat_label controls the column header text (e.g., "Median",
    "5th %ile", "10th %ile") and must match whichever statistic was used
    to populate the bmd/bmdl values in gene_sets_data.

    Matching NIEHS format:
      Caption: "Top 10 {Organ} Gene Ontology Biological Process Gene Sets
                Ranked by Potency of Perturbation, Sorted by Benchmark
                Dose {stat_label} ({Sex})"
      Columns: GO Term | GO ID | BMD {stat} (unit) | BMDL {stat} (unit) |
               # Genes | Direction

    Args:
        doc:             The Document to add the table to.
        gene_sets_data:  List of gene set dicts:
                           [{rank, go_id, go_term, bmd, bmdl,
                             n_genes, genes, direction}, ...]
                         Legacy data may use bmd_median/bmdl_median instead.
        organ:           Organ name for the caption (e.g., "Liver").
        sex:             Sex label for the caption (e.g., "Male").
        table_num:       Sequential table number.
        dose_unit:       Unit string for column headers.
        bmd_stat_label:  Human-readable label for the BMD statistic used
                         (e.g., "Median", "5th %ile").

    Returns:
        The next available table number (table_num + 1).
    """
    # Section heading
    organ_title = organ.title()
    sex_title = sex.title()
    doc.add_heading(f"Gene Set Benchmark Dose Analysis — {organ_title} ({sex_title})", level=3)

    # Caption
    caption = (
        f"Table {table_num}. Top 10 {organ_title} Gene Ontology Biological "
        f"Process Gene Sets Ranked by Potency of Perturbation, Sorted by "
        f"Benchmark Dose {bmd_stat_label} ({sex_title})."
    )
    cap_para = doc.add_paragraph()
    run = cap_para.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.italic = True
    cap_para.paragraph_format.space_after = Pt(4)

    if not gene_sets_data:
        p = doc.add_paragraph()
        run = p.add_run(f"No qualifying gene sets for {organ_title} ({sex_title}).")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.italic = True
        return table_num + 1

    headers = [
        "GO Term",
        "GO ID",
        f"BMD {bmd_stat_label} ({dose_unit})",
        f"BMDL {bmd_stat_label} ({dose_unit})",
        "# Genes",
        "Direction",
    ]

    table = doc.add_table(rows=1 + len(gene_sets_data), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        _format_table_cell(table.rows[0].cells[i], header, bold=True, size=HEADER_FONT_SIZE)

    # Data rows — read "bmd"/"bmdl" first, fall back to legacy "bmd_median"/"bmdl_median"
    for row_idx, gs in enumerate(gene_sets_data, 1):
        bmd_val = gs.get("bmd") if gs.get("bmd") is not None else gs.get("bmd_median")
        bmdl_val = gs.get("bmdl") if gs.get("bmdl") is not None else gs.get("bmdl_median")
        row_data = [
            gs.get("go_term", ""),
            gs.get("go_id", ""),
            f"{bmd_val:.3f}" if bmd_val is not None else "—",
            f"{bmdl_val:.3f}" if bmdl_val is not None else "—",
            str(gs.get("n_genes", 0)),
            gs.get("direction", ""),
        ]
        for i, val in enumerate(row_data):
            _format_table_cell(table.rows[row_idx].cells[i], val)

    doc.add_paragraph()  # spacing
    return table_num + 1


def add_gene_bmd_tables_to_doc(
    doc: Document,
    genes_data: list[dict],
    organ: str,
    sex: str,
    table_num: int,
    dose_unit: str = "mg/kg",
) -> int:
    """
    Add a Gene-level BMD table to the document.

    Produces one table for a specific organ × sex combination showing
    the top 10 individual genes ranked by BMD (most sensitive first).

    Matching NIEHS format:
      Caption: "Top 10 {Organ} Genes Ranked by Potency of Perturbation,
                Sorted by Benchmark Dose Median ({Sex})"
      Columns: Gene Symbol | Gene Name | BMD (unit) | BMDL (unit) |
               Fold Change | Direction

    Args:
        doc:          The Document to add the table to.
        genes_data:   List of gene dicts from rank_genes_by_bmd():
                        [{rank, gene_symbol, full_name, bmd, bmdl, bmdu,
                          direction, fold_change}, ...]
        organ:        Organ name for the caption.
        sex:          Sex label for the caption.
        table_num:    Sequential table number.
        dose_unit:    Unit string for column headers.

    Returns:
        The next available table number (table_num + 1).
    """
    organ_title = organ.title()
    sex_title = sex.title()
    doc.add_heading(f"Gene Benchmark Dose Analysis — {organ_title} ({sex_title})", level=3)

    # Caption
    caption = (
        f"Table {table_num}. Top 10 {organ_title} Genes Ranked by Potency "
        f"of Perturbation, Sorted by Benchmark Dose Median ({sex_title})."
    )
    cap_para = doc.add_paragraph()
    run = cap_para.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.italic = True
    cap_para.paragraph_format.space_after = Pt(4)

    if not genes_data:
        p = doc.add_paragraph()
        run = p.add_run(f"No qualifying genes for {organ_title} ({sex_title}).")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.italic = True
        return table_num + 1

    headers = [
        "Gene Symbol",
        "Gene Name",
        f"BMD ({dose_unit})",
        f"BMDL ({dose_unit})",
        "Fold Change",
        "Direction",
    ]

    table = doc.add_table(rows=1 + len(genes_data), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        _format_table_cell(table.rows[0].cells[i], header, bold=True, size=HEADER_FONT_SIZE)

    # Data rows
    for row_idx, gene in enumerate(genes_data, 1):
        fc = gene.get("fold_change")
        row_data = [
            gene.get("gene_symbol", ""),
            gene.get("full_name", ""),
            f"{gene['bmd']:.3f}" if gene.get("bmd") is not None else "—",
            f"{gene['bmdl']:.3f}" if gene.get("bmdl") is not None else "—",
            f"{fc:.2f}" if fc is not None else "—",
            gene.get("direction", ""),
        ]
        for i, val in enumerate(row_data):
            _format_table_cell(table.rows[row_idx].cells[i], val)

    doc.add_paragraph()  # spacing
    return table_num + 1


def generate_report(
    bm2_path: str,
    output_docx: str = "apical_report.docx",
    compound_name: str = "Test Compound",
    dose_unit: str = "mg/kg",
    section_title: str = "Animal Condition, Body Weights, and Organ Weights",
    table_caption_template: str = (
        "Summary of Body Weights and Organ Weights "
        "of {sex} Rats Administered {compound} for Five Days"
    ),
) -> str:
    """
    Generate the complete report section from a .bm2 file.

    Pipeline:
      1. Export .bm2 → JSON (full structure with dose-response data)
      2. Export .bm2 → TSV (category analysis results with BMD values)
      3. Run NTP statistical tests on each endpoint
      4. Cross-reference with category analysis for BMD/BMDL
      5. Apply business rules (both trend + pairwise required)
      6. Generate .docx with tables split by sex

    Internally delegates to build_table_data_from_bm2() for steps 1-5
    and add_apical_tables_to_doc() for step 6.  The CLI entry point
    calls this function, while the web server calls the helpers directly
    to compose tables into a larger report document.

    Args:
        bm2_path:               Path to the BMDExpress 3 .bm2 file.
        output_docx:            Path for the output Word document.
        compound_name:          Chemical name for table captions.
        dose_unit:              Dose unit for column headers and BMD columns.
        section_title:          Heading text for the report section.
                                Default: "Animal Condition, Body Weights, and
                                Organ Weights".
        table_caption_template: Format string for per-table captions.  Must
                                contain {sex} and {compound} placeholders.
                                Default matches the NTP body/organ weight format.

    Returns:
        Path to the generated .docx file.
    """
    # Steps 1-5: export .bm2 and build table data
    table_data, _category_lookup, _bm2_json = build_table_data_from_bm2(bm2_path)

    # Step 6: Generate .docx using the composable helper
    print("  Generating .docx...")
    doc = Document()

    add_apical_tables_to_doc(
        doc,
        table_data,
        section_title=section_title,
        compound_name=compound_name,
        dose_unit=dose_unit,
        table_caption_template=table_caption_template,
    )

    doc.save(output_docx)
    print(f"\nSaved: {output_docx}")
    return output_docx


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description=(
            "Generate NTP-style apical endpoint report from a .bm2 file. "
            "Works for body/organ weights, clinical pathology, or any other "
            "apical endpoint type — use --section-title and --table-caption "
            "to customize headings."
        ),
    )
    parser.add_argument(
        "bm2_path",
        help="Path to the BMDExpress 3 .bm2 file",
    )
    parser.add_argument(
        "--output", "-o",
        default="apical_report.docx",
        help="Output .docx path (default: apical_report.docx)",
    )
    parser.add_argument(
        "--compound", "-c",
        default="Test Compound",
        help="Compound name for table captions",
    )
    parser.add_argument(
        "--dose-unit", "-u",
        default="mg/kg",
        help="Dose unit for column headers (default: mg/kg)",
    )
    parser.add_argument(
        "--section-title",
        default="Animal Condition, Body Weights, and Organ Weights",
        help=(
            "Section heading in the report "
            "(default: 'Animal Condition, Body Weights, and Organ Weights')"
        ),
    )
    parser.add_argument(
        "--table-caption",
        default=(
            "Summary of Body Weights and Organ Weights "
            "of {sex} Rats Administered {compound} for Five Days"
        ),
        help=(
            "Table caption template with {sex} and {compound} placeholders "
            "(default: 'Summary of Body Weights and Organ Weights of {sex} "
            "Rats Administered {compound} for Five Days')"
        ),
    )

    args = parser.parse_args()
    generate_report(
        args.bm2_path,
        output_docx=args.output,
        compound_name=args.compound,
        dose_unit=args.dose_unit,
        section_title=args.section_title,
        table_caption_template=args.table_caption,
    )
