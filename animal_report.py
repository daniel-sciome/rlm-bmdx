"""
animal_report — Per-animal traceability across file pool tiers.

While file_integrator.py fingerprints files and tracks aggregate counts
(how many animals per dose), this module goes deeper: it reads individual
animal IDs from every file in the pool and cross-references them across
tiers (xlsx → txt/csv → bm2) and domains (body_weight, hematology, etc.).

This solves the problem discovered in the PFHxSAm study (C20022-01):
  - "Selection" column in xlsx separates Core Animals (110) from
    Biosampling Animals (12 PK animals, 500-series IDs)
  - Systematic attrition across tiers is structurally explainable
    (dose group exclusion, biosampling exclusion, QC removal)
  - Without this module, the user can't see *which* animals dropped
    or *why* — only that counts differ

The report consists of five subsections:
  A. Study Design Summary — dose × sex × selection breakdown
  B. Animal Roster — every animal with per-domain tier presence flags
  C. Domain Coverage Matrix — counts per domain per tier
  D. Attrition Analysis — who dropped, why, per domain
  E. Consistency Checks — sex/dose mismatches across domains

Usage:
    from animal_report import build_animal_report, add_animal_report_to_doc

    report = build_animal_report(session_dir, fingerprints)
    next_table_num = add_animal_report_to_doc(doc, report, start_table_num=1)
"""

# ---------------------------------------------------------------------------
# Imports
# ---------------------------------------------------------------------------

import json
import logging
import os
import re
from dataclasses import dataclass, field, asdict

# file_integrator provides domain detection and sex detection helpers
# that we reuse here to avoid duplicating pattern matching logic.
from file_integrator import (
    base_domain,
    detect_domain,
    _detect_sex_from_filename,
    _BM2_SEX_PATTERN,
    _GENE_EXPR_ORGAN_PATTERN,
)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

logger = logging.getLogger(__name__)

# Short domain labels for the compact animal roster table columns.
# Order matches the DOCX table column order (left to right).
DOMAIN_COLUMN_ORDER = [
    "body_weight",
    "organ_weights",
    "clin_chem",
    "hematology",
    "hormones",
    "tissue_conc",
    "clinical_obs",
    "gene_expression",
]

# Human-readable abbreviations for domain column headers in the roster.
# Keep to 2-3 characters to fit in compact table cells.
DOMAIN_SHORT_LABELS = {
    "body_weight": "BW",
    "organ_weights": "OW",
    "clin_chem": "CC",
    "hematology": "Hem",
    "hormones": "Horm",
    "tissue_conc": "TC",
    "clinical_obs": "CO",
    "gene_expression": "GE",
}

# Human-readable full names for domain rows in the coverage matrix.
DOMAIN_FULL_LABELS = {
    "body_weight": "Body Weight",
    "organ_weights": "Organ Weights",
    "clin_chem": "Clinical Chemistry",
    "hematology": "Hematology",
    "hormones": "Hormones",
    "tissue_conc": "Tissue Concentration",
    "clinical_obs": "Clinical Observations",
    "gene_expression": "Gene Expression",
}

# Tier key names used in domain_presence dicts and coverage matrices.
# These match the file type groupings: xlsx is tier 1, txt/csv is tier 2,
# bm2 is tier 3 (same precedence as file_integrator).
TIER_XLSX = "xlsx"
TIER_TXT_CSV = "txt_csv"
TIER_BM2 = "bm2"


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class AnimalRecord:
    """
    One animal's complete profile across all domains and tiers.

    The animal_id is the NTP animal identifier (e.g., "101", "503").
    Sex and dose come from the most authoritative source (xlsx > txt/csv > bm2).
    Selection distinguishes "Core Animals" from "Biosampling Animals" —
    only present in xlsx files that have a Selection column.

    domain_presence tracks where this animal appears:
      {"body_weight": {"xlsx": True, "txt_csv": True, "bm2": True},
       "hematology": {"xlsx": True, "txt_csv": False, "bm2": False}}
    """
    animal_id: str
    sex: str | None = None
    dose: float | None = None
    selection: str | None = None
    # domain → {tier_key: bool} — which tiers include this animal
    domain_presence: dict[str, dict[str, bool]] = field(default_factory=dict)


@dataclass
class DomainAttrition:
    """
    Per-domain tracking of animal dropout between tiers.

    For each domain (body_weight, hematology, etc.), tracks the set of
    animal IDs present at each tier and the animals excluded between tiers.
    Exclusion reasons are classified as:
      - dose_exclusion: animal's entire dose group is absent from the target tier
      - biosampling_exclusion: animal's selection is "Biosampling Animals"
      - qc_exclusion: individual animal removed (not explained by dose or selection)
    """
    domain: str
    xlsx_ids: set[str] = field(default_factory=set)
    txt_csv_ids: set[str] = field(default_factory=set)
    bm2_ids: set[str] = field(default_factory=set)
    excluded_xlsx_to_txt: set[str] = field(default_factory=set)
    excluded_txt_to_bm2: set[str] = field(default_factory=set)
    # reason → [animal_ids] — why each animal was excluded
    exclusion_reasons: dict[str, list[str]] = field(default_factory=dict)


@dataclass
class AnimalReport:
    """
    Complete animal report for a session.

    This is the top-level data structure returned by build_animal_report()
    and consumed by add_animal_report_to_doc() for DOCX generation and
    by the frontend for HTML rendering.

    Fields are grouped by report subsection:
      A. animals — the full roster keyed by animal_id
      B. study_* / dose_* — study design summary
      C. domain_coverage — counts per domain per tier
      D. attrition / completeness — per-domain dropout analysis
      E. consistency_issues — cross-domain mismatches
    """
    # A. Roster — every animal in the study
    animals: dict[str, AnimalRecord] = field(default_factory=dict)

    # B. Study Design — overall structure
    study_number: str | None = None
    total_animals: int = 0
    core_count: int = 0
    biosampling_count: int = 0
    dose_groups: list[float] = field(default_factory=list)
    # dose → {"Male": n, "Female": n} — design matrix
    dose_design: dict[float, dict[str, int]] = field(default_factory=dict)

    # C. Coverage — domain → tier → animal count
    domain_coverage: dict[str, dict[str, int]] = field(default_factory=dict)

    # D. Attrition — domain → DomainAttrition
    attrition: dict[str, DomainAttrition] = field(default_factory=dict)
    # domain → fraction of xlsx animals retained in bm2 (0.0–1.0)
    completeness: dict[str, float] = field(default_factory=dict)

    # E. Consistency — list of {type, animal_id, details} dicts
    consistency_issues: list[dict] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Extraction functions — one per file type
# ---------------------------------------------------------------------------
# These read individual animal records from each file type.  They're
# separate from file_integrator's fingerprint functions because they
# extract *per-animal detail* (animal_id → dose, sex, selection) rather
# than just aggregates (n_animals_by_dose).


def _extract_animals_from_xlsx(path: str) -> dict[str, dict]:
    """
    Extract per-animal records from an NTP xlsx file.

    Opens the xlsx with openpyxl and reads the "Data" sheet.  Finds
    column indices by header name (not position) to handle files that
    have extra or missing columns (e.g., some files lack "Selection").

    Returns a dict of {animal_id: {dose, sex, selection}} with deduplication
    across observation rows (each animal appears many times — once per
    endpoint per study day).

    Args:
        path: Absolute path to the xlsx file on disk.

    Returns:
        Dict mapping animal_id strings to metadata dicts.
        Each metadata dict has keys: dose (float|None), sex (str|None),
        selection (str|None).
    """
    import openpyxl

    animals = {}

    try:
        # read_only=False because NTP xlsx files sometimes lack dimension
        # metadata, causing read-only mode to see only 1 row per sheet.
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception as e:
        logger.warning("Could not open xlsx %s: %s", path, e)
        return animals

    try:
        # Find the "Data" sheet — it's always named "Data" in NTP files,
        # but fall back to the last sheet if not found.
        data_sheet = wb["Data"] if "Data" in wb.sheetnames else wb[wb.sheetnames[-1]]

        # --- Find column indices by header name ---
        # NTP xlsx files have these standard columns:
        #   A: NTP Study Number, B: Concentration, C: Animal ID, D: Sex
        #   E: Selection (sometimes missing), F+: endpoints
        # We find them by name to handle column order variations.
        col_indices = {}
        header_map = {
            "concentration": "dose",
            "animal id": "animal_id",
            "sex": "sex",
            "selection": "selection",
            "ntp study number": "study_number",
        }

        headers = None
        for row_idx, row in enumerate(data_sheet.iter_rows(values_only=True)):
            if row_idx == 0:
                # First row contains column headers.
                # Match each header against known patterns.  Only accept
                # the FIRST column that matches each field_name to avoid
                # endpoint columns like "Mean Cell Hgb Concentration"
                # overwriting the actual "Concentration" (dose) column.
                headers = [str(c).strip().lower() if c is not None else "" for c in row]
                assigned_fields = set()  # track which field_names are already assigned
                for col_idx, header_text in enumerate(headers):
                    for pattern, field_name in header_map.items():
                        if field_name in assigned_fields:
                            continue  # already found this field — skip
                        if pattern in header_text:
                            col_indices[field_name] = col_idx
                            assigned_fields.add(field_name)
                            break
                continue

            # Data rows — extract animal record
            aid_col = col_indices.get("animal_id")
            if aid_col is None or aid_col >= len(row) or row[aid_col] is None:
                continue

            animal_id = str(row[aid_col]).strip()
            if not animal_id:
                continue

            # Only record the first occurrence — all rows for the same
            # animal have identical dose/sex/selection.
            if animal_id in animals:
                continue

            record = {"dose": None, "sex": None, "selection": None}

            dose_col = col_indices.get("dose")
            if dose_col is not None and dose_col < len(row) and row[dose_col] is not None:
                try:
                    record["dose"] = float(row[dose_col])
                except (ValueError, TypeError):
                    pass

            sex_col = col_indices.get("sex")
            if sex_col is not None and sex_col < len(row) and row[sex_col] is not None:
                record["sex"] = str(row[sex_col]).strip().title()

            sel_col = col_indices.get("selection")
            if sel_col is not None and sel_col < len(row) and row[sel_col] is not None:
                record["selection"] = str(row[sel_col]).strip()

            animals[animal_id] = record

    except Exception as e:
        logger.debug("Data sheet parse failed for %s: %s", path, e)

    wb.close()
    return animals


def _extract_animals_from_txt_csv(
    path: str,
    filename: str,
    file_type: str,
) -> dict[str, dict]:
    """
    Extract per-animal records from a BMDExpress-importable txt/csv file.

    NTP txt/csv files are transposed tables:
      Row 1: Animal IDs (tab or comma separated), first cell is a label
      Row 2: Dose concentrations per animal (same column order)
      Row 3+: Endpoint measurements

    Sex comes from the filename (e.g., "male_body_weight.txt") via
    file_integrator's _detect_sex_from_filename().

    Args:
        path:      Absolute path to the txt/csv file on disk.
        filename:  Original filename (needed for sex detection).
        file_type: "txt" or "csv" (determines separator character).

    Returns:
        Dict mapping animal_id strings to metadata dicts.
        Each metadata dict has keys: dose (float|None), sex (str|None).
    """
    animals = {}

    # Detect sex from filename — NTP convention: "male_body_weight.txt"
    sexes = _detect_sex_from_filename(filename)
    sex = sexes[0] if sexes else None

    separator = "," if file_type == "csv" else "\t"

    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            lines = f.readlines()
    except Exception as e:
        logger.warning("Could not read %s: %s", filename, e)
        return animals

    # Strip trailing whitespace and skip empty lines
    lines = [ln.rstrip("\n\r") for ln in lines if ln.strip()]
    if len(lines) < 2:
        return animals

    # Row 1: animal IDs (first cell is a label like "Something" or blank)
    row1_cells = lines[0].split(separator)
    animal_ids = [c.strip() for c in row1_cells[1:] if c.strip()]

    # Row 2: dose concentrations per animal (same column order)
    row2_cells = lines[1].split(separator)
    doses_raw = row2_cells[1:]

    for i, aid in enumerate(animal_ids):
        dose = None
        if i < len(doses_raw):
            try:
                dose = float(doses_raw[i].strip())
            except (ValueError, TypeError):
                pass
        animals[aid] = {"dose": dose, "sex": sex}

    return animals


def _extract_animals_from_bm2(
    path: str,
    bm2_json: dict | None = None,
) -> dict[str, dict]:
    """
    Extract per-animal records from a BMDExpress 3 .bm2 JSON.

    Iterates ALL experiments' treatments[] — not just the first experiment.
    Each experiment contributes animal_ids with dose from treatment.dose
    and sex inferred from the experiment name (e.g., "BodyWeightMale").

    If bm2_json is not provided, tries to load from the LMDB cache.
    Returns an empty dict if the JSON is unavailable.

    Args:
        path:     Absolute path to the .bm2 file on disk.
        bm2_json: Pre-loaded BMDProject dict (optional — avoids cache lookup).

    Returns:
        Dict mapping animal_id strings to metadata dicts.
        Each metadata dict has keys: dose (float|None), sex (str|None).
    """
    animals = {}

    # Try loading from LMDB cache if not provided
    if bm2_json is None:
        try:
            import bm2_cache
            bm2_json = bm2_cache.get_json(path)
        except Exception:
            pass

    if bm2_json is None:
        logger.debug("No cached JSON for %s — cannot extract animals", path)
        return animals

    experiments = bm2_json.get("doseResponseExperiments", [])

    for exp in experiments:
        # Infer sex from experiment name (e.g., "BodyWeightMale" → "Male")
        exp_name = exp.get("name", "")
        sex_matches = _BM2_SEX_PATTERN.findall(exp_name)
        sex = sex_matches[0].title() if sex_matches else None

        for treatment in exp.get("treatments", []):
            animal_name = treatment.get("name", "")
            if not animal_name:
                continue

            dose = None
            raw_dose = treatment.get("dose")
            if raw_dose is not None:
                try:
                    dose = float(raw_dose)
                except (ValueError, TypeError):
                    pass

            # Don't overwrite if already seen — earlier experiments
            # may have more authoritative metadata.
            if animal_name not in animals:
                animals[animal_name] = {"dose": dose, "sex": sex}

    return animals


# ---------------------------------------------------------------------------
# Study number extraction helper
# ---------------------------------------------------------------------------

# Pattern matches NTP study numbers like "C20022-01"
_STUDY_NUMBER_RE = re.compile(r"^[A-Z]\d+-\d+$")


def _extract_study_number_from_xlsx(path: str) -> str | None:
    """
    Extract the NTP study number from an xlsx file's Key sheet or Data sheet.

    Checks column D of the "Key to Column Labels" sheet first (where NTP
    places study-level metadata), then falls back to the first data row's
    column A (NTP Study Number column).

    Args:
        path: Absolute path to the xlsx file.

    Returns:
        Study number string (e.g., "C20022-01") or None.
    """
    import openpyxl

    try:
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception:
        return None

    study_number = None

    # Check "Key to Column Labels" sheet column D
    try:
        key_sheet = wb[wb.sheetnames[0]]
        for row in key_sheet.iter_rows(values_only=True):
            if len(row) >= 4 and row[3] is not None:
                val = str(row[3]).strip()
                if _STUDY_NUMBER_RE.match(val):
                    study_number = val
                    break
    except Exception:
        pass

    # Fallback: check Data sheet column A
    if study_number is None:
        try:
            data_sheet = wb["Data"] if "Data" in wb.sheetnames else wb[wb.sheetnames[-1]]
            for row_idx, row in enumerate(data_sheet.iter_rows(values_only=True)):
                if row_idx == 0:
                    continue  # skip header
                if row[0] is not None:
                    val = str(row[0]).strip()
                    if _STUDY_NUMBER_RE.match(val):
                        study_number = val
                        break
        except Exception:
            pass

    wb.close()
    return study_number


# ---------------------------------------------------------------------------
# Main orchestrator
# ---------------------------------------------------------------------------


def build_animal_report(
    session_dir: str,
    fingerprints: dict,
) -> AnimalReport:
    """
    Build a complete animal report by reading all files in the session.

    This is the main entry point.  For each fingerprinted file, it calls
    the appropriate _extract_animals_from_* function to get per-animal
    detail, then builds a unified roster, computes study design summary,
    domain coverage, attrition analysis, and consistency checks.

    The fingerprints dict maps file_id → FileFingerprint (from
    file_integrator.fingerprint_file).  Only the domain and file_type
    fields are needed to route extraction; the actual animal data is
    read fresh from the files on disk.

    Args:
        session_dir:   Path to the session directory (e.g., sessions/DTXSID50469320).
        fingerprints:  Dict of file_id → FileFingerprint from _pool_fingerprints.

    Returns:
        AnimalReport with all five subsections populated.
    """
    report = AnimalReport()
    files_dir = os.path.join(session_dir, "files")

    # --- Phase 1: Extract per-animal data from every file ---
    # We organize extractions by (domain, tier) so we can build the
    # domain_presence map and attrition analysis.

    # domain → tier → {animal_id: {dose, sex, selection?}}
    domain_tier_animals: dict[str, dict[str, dict[str, dict]]] = {}

    study_number = None

    for file_id, fp in fingerprints.items():
        # fp can be a FileFingerprint dataclass or a dict (from JSON restore).
        # Normalize to dict access for compatibility.
        if hasattr(fp, "filename"):
            filename = fp.filename
            file_type = fp.file_type
            domain = fp.domain
        else:
            filename = fp.get("filename", "")
            file_type = fp.get("file_type", "")
            domain = fp.get("domain")

        if not domain:
            continue  # skip files with no detected domain

        # Normalize to base domain (strip _tox_study / _inferred suffix)
        # so both tox_study and inferred files contribute to the same
        # conceptual domain column in the animal roster table.
        domain = base_domain(domain)

        path = os.path.join(files_dir, filename)
        if not os.path.exists(path):
            continue

        # Determine tier from file_type
        if file_type == "xlsx":
            tier = TIER_XLSX
        elif file_type in ("txt", "csv"):
            tier = TIER_TXT_CSV
        elif file_type == "bm2":
            tier = TIER_BM2
        else:
            continue

        # Extract animals from this file
        if file_type == "xlsx":
            file_animals = _extract_animals_from_xlsx(path)
            # Also extract study number from xlsx (most authoritative source)
            if study_number is None:
                study_number = _extract_study_number_from_xlsx(path)
        elif file_type in ("txt", "csv"):
            file_animals = _extract_animals_from_txt_csv(path, filename, file_type)
        elif file_type == "bm2":
            # Try to get bm2_json from LMDB cache
            file_animals = _extract_animals_from_bm2(path)
        else:
            continue

        if not file_animals:
            continue

        # Store in domain_tier_animals structure
        if domain not in domain_tier_animals:
            domain_tier_animals[domain] = {}

        if tier not in domain_tier_animals[domain]:
            domain_tier_animals[domain][tier] = {}

        # Merge animals from this file into the domain×tier bucket.
        # Multiple files can contribute to the same domain×tier
        # (e.g., male_body_weight.txt + female_body_weight.txt).
        domain_tier_animals[domain][tier].update(file_animals)

    report.study_number = study_number

    # --- Phase 2: Build unified animal roster ---
    # Merge all extracted animals into a single roster.  xlsx is
    # authoritative for dose/sex/selection; txt/csv and bm2 fill gaps.

    all_animals: dict[str, AnimalRecord] = {}

    for domain, tiers in domain_tier_animals.items():
        for tier, animals_dict in tiers.items():
            for aid, meta in animals_dict.items():
                if aid not in all_animals:
                    all_animals[aid] = AnimalRecord(animal_id=aid)

                record = all_animals[aid]

                # Update sex/dose/selection from most authoritative source.
                # xlsx > txt_csv > bm2 (check if current value is None
                # before overwriting with a less authoritative source).
                if tier == TIER_XLSX:
                    # xlsx is always authoritative — overwrite
                    if meta.get("sex"):
                        record.sex = meta["sex"]
                    if meta.get("dose") is not None:
                        record.dose = meta["dose"]
                    if meta.get("selection"):
                        record.selection = meta["selection"]
                elif tier == TIER_TXT_CSV:
                    # txt/csv fills in if xlsx didn't provide
                    if record.sex is None and meta.get("sex"):
                        record.sex = meta["sex"]
                    if record.dose is None and meta.get("dose") is not None:
                        record.dose = meta["dose"]
                elif tier == TIER_BM2:
                    # bm2 is least authoritative
                    if record.sex is None and meta.get("sex"):
                        record.sex = meta["sex"]
                    if record.dose is None and meta.get("dose") is not None:
                        record.dose = meta["dose"]

                # Track domain presence for this animal
                if domain not in record.domain_presence:
                    record.domain_presence[domain] = {
                        TIER_XLSX: False,
                        TIER_TXT_CSV: False,
                        TIER_BM2: False,
                    }
                record.domain_presence[domain][tier] = True

    report.animals = all_animals
    report.total_animals = len(all_animals)

    # --- Phase 3: Study design summary ---
    # Count core vs biosampling animals and build dose × sex design matrix.

    core_ids = set()
    biosamp_ids = set()
    dose_design: dict[float, dict[str, int]] = {}

    for aid, record in all_animals.items():
        sel = record.selection
        if sel and "biosamp" in sel.lower():
            biosamp_ids.add(aid)
        else:
            core_ids.add(aid)

        if record.dose is not None and record.sex:
            dose = record.dose
            if dose not in dose_design:
                dose_design[dose] = {}
            sex = record.sex
            dose_design[dose][sex] = dose_design[dose].get(sex, 0) + 1

    report.core_count = len(core_ids)
    report.biosampling_count = len(biosamp_ids)
    report.dose_groups = sorted(dose_design.keys())
    report.dose_design = {dose: dose_design[dose] for dose in sorted(dose_design.keys())}

    # --- Phase 4: Domain coverage matrix ---
    # For each domain, count how many animals appear at each tier.

    for domain in sorted(domain_tier_animals.keys()):
        tiers = domain_tier_animals[domain]
        report.domain_coverage[domain] = {
            TIER_XLSX: len(tiers.get(TIER_XLSX, {})),
            TIER_TXT_CSV: len(tiers.get(TIER_TXT_CSV, {})),
            TIER_BM2: len(tiers.get(TIER_BM2, {})),
        }

    # --- Phase 5: Attrition analysis ---
    # For each domain, compute who dropped between tiers and classify why.

    # Build lookup tables for exclusion reason classification:
    # - animal_id → dose (for dose exclusion detection)
    # - animal_id → selection (for biosampling exclusion detection)
    animal_doses = {aid: rec.dose for aid, rec in all_animals.items()}
    animal_selections = {aid: rec.selection for aid, rec in all_animals.items()}

    for domain, tiers in domain_tier_animals.items():
        xlsx_ids = set(tiers.get(TIER_XLSX, {}).keys())
        txt_csv_ids = set(tiers.get(TIER_TXT_CSV, {}).keys())
        bm2_ids = set(tiers.get(TIER_BM2, {}).keys())

        attrition = DomainAttrition(
            domain=domain,
            xlsx_ids=xlsx_ids,
            txt_csv_ids=txt_csv_ids,
            bm2_ids=bm2_ids,
        )

        # Compute excluded sets between tiers
        attrition.excluded_xlsx_to_txt = xlsx_ids - txt_csv_ids if xlsx_ids and txt_csv_ids else set()
        attrition.excluded_txt_to_bm2 = txt_csv_ids - bm2_ids if txt_csv_ids and bm2_ids else set()

        # Classify exclusion reasons for xlsx→txt drops
        _classify_exclusions(
            attrition,
            attrition.excluded_xlsx_to_txt,
            xlsx_ids,
            txt_csv_ids,
            animal_doses,
            animal_selections,
            "xlsx_to_txt",
        )

        # Classify exclusion reasons for txt→bm2 drops
        _classify_exclusions(
            attrition,
            attrition.excluded_txt_to_bm2,
            txt_csv_ids,
            bm2_ids,
            animal_doses,
            animal_selections,
            "txt_to_bm2",
        )

        # Compute completeness — ratio of bm2 animals to xlsx animals.
        # If a tier is empty, use the next available tier as denominator.
        denominator = len(xlsx_ids) or len(txt_csv_ids) or 1
        numerator = len(bm2_ids) or len(txt_csv_ids) or len(xlsx_ids)
        report.completeness[domain] = round(numerator / denominator, 3) if denominator else 1.0

        report.attrition[domain] = attrition

    # --- Phase 6: Consistency checks ---
    # Look for animals that have different sex or dose across domains.

    report.consistency_issues = _check_consistency(domain_tier_animals)

    return report


def _classify_exclusions(
    attrition: DomainAttrition,
    excluded_ids: set[str],
    source_ids: set[str],
    target_ids: set[str],
    animal_doses: dict[str, float | None],
    animal_selections: dict[str, str | None],
    transition_label: str,
) -> None:
    """
    Classify why each excluded animal was dropped between tiers.

    Three categories:
      1. dose_exclusion — the animal's entire dose group is absent from
         the target tier (e.g., 333/1000 mg/kg groups excluded from txt)
      2. biosampling_exclusion — the animal's selection is "Biosampling Animals"
         and the target tier doesn't include biosampling animals
      3. qc_exclusion — individual animal removed for QC reasons (everything
         else — the animal's dose group is present, it's a core animal,
         but it was still excluded individually)

    Results are stored in attrition.exclusion_reasons as a dict mapping
    "{transition_label}_{reason}" → list of animal_id strings.

    Args:
        attrition:         DomainAttrition object to update.
        excluded_ids:      Set of animal IDs that were excluded.
        source_ids:        Animal IDs in the source tier (for dose group computation).
        target_ids:        Animal IDs in the target tier (for dose group comparison).
        animal_doses:      Global animal_id → dose lookup.
        animal_selections: Global animal_id → selection lookup.
        transition_label:  "xlsx_to_txt" or "txt_to_bm2" (used in reason keys).
    """
    if not excluded_ids:
        return

    # Find which dose groups are present in the target tier
    target_doses = {
        animal_doses.get(aid)
        for aid in target_ids
        if animal_doses.get(aid) is not None
    }

    dose_excluded = []
    biosamp_excluded = []
    qc_excluded = []

    for aid in sorted(excluded_ids):
        dose = animal_doses.get(aid)
        selection = animal_selections.get(aid) or ""

        if dose is not None and dose not in target_doses:
            # Entire dose group is missing from target tier
            dose_excluded.append(aid)
        elif "biosamp" in selection.lower():
            # Animal is a biosampling animal not included in downstream analysis
            biosamp_excluded.append(aid)
        else:
            # Individual removal — likely QC (blood sample quality, outlier, etc.)
            qc_excluded.append(aid)

    if dose_excluded:
        attrition.exclusion_reasons[f"{transition_label}_dose_exclusion"] = dose_excluded
    if biosamp_excluded:
        attrition.exclusion_reasons[f"{transition_label}_biosampling_exclusion"] = biosamp_excluded
    if qc_excluded:
        attrition.exclusion_reasons[f"{transition_label}_qc_exclusion"] = qc_excluded


def _check_consistency(
    domain_tier_animals: dict[str, dict[str, dict[str, dict]]],
) -> list[dict]:
    """
    Check for sex/dose mismatches for the same animal across domains.

    An animal should have the same sex and dose regardless of which domain
    or tier it appears in.  If domain A says animal 101 is Male at dose 4
    but domain B says it's Female at dose 12, that's a consistency issue.

    Args:
        domain_tier_animals: The domain → tier → {animal_id: metadata} structure.

    Returns:
        List of issue dicts with keys: type, animal_id, details.
    """
    issues = []

    # Build per-animal: {animal_id: [(domain, tier, sex, dose), ...]}
    animal_observations: dict[str, list[tuple[str, str, str | None, float | None]]] = {}

    for domain, tiers in domain_tier_animals.items():
        for tier, animals_dict in tiers.items():
            for aid, meta in animals_dict.items():
                if aid not in animal_observations:
                    animal_observations[aid] = []
                animal_observations[aid].append((
                    domain, tier, meta.get("sex"), meta.get("dose"),
                ))

    for aid, observations in sorted(animal_observations.items()):
        # Check sex consistency
        sexes = {obs[2] for obs in observations if obs[2] is not None}
        if len(sexes) > 1:
            issues.append({
                "type": "sex_mismatch",
                "animal_id": aid,
                "details": {
                    "sexes_found": sorted(sexes),
                    "sources": [
                        {"domain": obs[0], "tier": obs[1], "sex": obs[2]}
                        for obs in observations if obs[2] is not None
                    ],
                },
            })

        # Check dose consistency
        doses = {obs[3] for obs in observations if obs[3] is not None}
        if len(doses) > 1:
            issues.append({
                "type": "dose_mismatch",
                "animal_id": aid,
                "details": {
                    "doses_found": sorted(doses),
                    "sources": [
                        {"domain": obs[0], "tier": obs[1], "dose": obs[3]}
                        for obs in observations if obs[3] is not None
                    ],
                },
            })

    return issues


# ---------------------------------------------------------------------------
# Serialization helpers
# ---------------------------------------------------------------------------


def report_to_dict(report: AnimalReport) -> dict:
    """
    Convert an AnimalReport to a JSON-serializable dict.

    Handles the set→list conversion needed for DomainAttrition fields
    (JSON doesn't support sets).  The animals dict uses AnimalRecord
    dataclasses which are converted via dataclasses.asdict().

    Args:
        report: The AnimalReport to serialize.

    Returns:
        JSON-safe dict suitable for json.dumps().
    """
    animals_dict = {}
    for aid, rec in report.animals.items():
        animals_dict[aid] = asdict(rec)

    attrition_dict = {}
    for domain, att in report.attrition.items():
        attrition_dict[domain] = {
            "domain": att.domain,
            "xlsx_ids": sorted(att.xlsx_ids),
            "txt_csv_ids": sorted(att.txt_csv_ids),
            "bm2_ids": sorted(att.bm2_ids),
            "excluded_xlsx_to_txt": sorted(att.excluded_xlsx_to_txt),
            "excluded_txt_to_bm2": sorted(att.excluded_txt_to_bm2),
            "exclusion_reasons": att.exclusion_reasons,
        }

    # Convert dose_design keys from float to string for JSON
    dose_design_str = {}
    for dose, sexes in report.dose_design.items():
        dose_design_str[str(dose)] = sexes

    return {
        "animals": animals_dict,
        "study_number": report.study_number,
        "total_animals": report.total_animals,
        "core_count": report.core_count,
        "biosampling_count": report.biosampling_count,
        "dose_groups": report.dose_groups,
        "dose_design": dose_design_str,
        "domain_coverage": report.domain_coverage,
        "attrition": attrition_dict,
        "completeness": report.completeness,
        "consistency_issues": report.consistency_issues,
    }


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------


def add_animal_report_to_doc(
    doc,
    report: AnimalReport,
    start_table_num: int = 1,
) -> int:
    """
    Add the animal report section to a python-docx Document.

    Generates five subsections as tables and narrative text:
      Table A — Study Design Summary (dose × sex count table)
      Table B — Animal Roster (compact per-animal × per-domain matrix)
      Table C — Domain Coverage Matrix (counts per tier)
      Table D — Attrition Analysis (narrative with exclusion grouping)
      Table E — Consistency Check Results (bullet list or "none")

    Args:
        doc:             python-docx Document object.
        report:          AnimalReport with all subsections populated.
        start_table_num: Sequential table number to start from (continues
                         from previous report sections for correct numbering).

    Returns:
        Next table number (for subsequent sections to continue from).
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table_num = start_table_num

    # --- Section heading ---
    doc.add_heading("Animal Report", level=2)

    # ===== Table A — Study Design Summary =====
    # Narrative paragraph describing the study enrollment
    study_label = report.study_number or "this study"
    sel_detail = ""
    if report.biosampling_count > 0:
        sel_detail = (
            f" ({report.core_count} core, "
            f"{report.biosampling_count} biosampling)"
        )

    summary_text = (
        f"Study {study_label} enrolled {report.total_animals} animals"
        f"{sel_detail} across {len(report.dose_groups)} dose groups."
    )
    p = doc.add_paragraph()
    run = p.add_run(summary_text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)

    # Dose × Sex count table
    if report.dose_design:
        # Collect all sexes that appear in any dose group
        all_sexes = sorted({
            sex for sexes in report.dose_design.values()
            for sex in sexes.keys()
        })

        headers = ["Dose (mg/kg)"] + all_sexes + ["Total"]
        n_rows = len(report.dose_groups)
        table = doc.add_table(rows=1 + n_rows, cols=len(headers))
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Caption
        caption = doc.add_paragraph()
        run = caption.add_run(f"Table {table_num}. Study design — animals per dose group and sex.")
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.italic = True
        caption.paragraph_format.space_after = Pt(4)
        # Move caption before the table (insert it before the table element)
        table._element.addprevious(caption._element)

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for r in para.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.name = "Calibri"

        # Data rows
        for row_idx, dose in enumerate(report.dose_groups):
            row = table.rows[1 + row_idx]
            # Dose value
            row.cells[0].text = str(int(dose)) if dose == int(dose) else str(dose)
            total = 0
            for sex_idx, sex in enumerate(all_sexes):
                count = report.dose_design.get(dose, {}).get(sex, 0)
                row.cells[1 + sex_idx].text = str(count)
                total += count
            row.cells[len(headers) - 1].text = str(total)

            # Format all cells in this row
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(9)
                        r.font.name = "Calibri"

        table_num += 1

    # ===== Table B — Animal Roster =====
    # Compact table with one row per animal and domain columns showing
    # tier presence: "XTB" = present in xlsx, txt, bm2; "X--" = xlsx only.

    if report.animals:
        # Determine which domains actually have data (skip empty columns)
        active_domains = [
            d for d in DOMAIN_COLUMN_ORDER
            if d in report.domain_coverage
        ]

        headers = ["Animal ID", "Sex", "Dose", "Selection"] + [
            DOMAIN_SHORT_LABELS.get(d, d[:3]) for d in active_domains
        ]

        # Sort animals numerically (IDs like "101", "503" should sort by number)
        sorted_aids = sorted(
            report.animals.keys(),
            key=lambda x: (int(x) if x.isdigit() else float('inf'), x),
        )

        n_rows = len(sorted_aids)
        table = doc.add_table(rows=1 + n_rows, cols=len(headers))
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Caption
        caption = doc.add_paragraph()
        run = caption.add_run(
            f"Table {table_num}. Animal roster — tier presence per domain. "
            f"X=xlsx, T=txt/csv, B=bm2."
        )
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.italic = True
        caption.paragraph_format.space_after = Pt(4)
        table._element.addprevious(caption._element)

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for r in para.runs:
                    r.bold = True
                    r.font.size = Pt(7)
                    r.font.name = "Calibri"

        # Data rows
        for row_idx, aid in enumerate(sorted_aids):
            rec = report.animals[aid]
            row = table.rows[1 + row_idx]

            row.cells[0].text = aid
            row.cells[1].text = rec.sex or ""
            row.cells[2].text = (
                str(int(rec.dose)) if rec.dose is not None and rec.dose == int(rec.dose)
                else str(rec.dose) if rec.dose is not None
                else ""
            )
            row.cells[3].text = rec.selection or ""

            # Domain columns — show tier presence as "XTB" pattern
            for col_idx, domain in enumerate(active_domains):
                presence = rec.domain_presence.get(domain, {})
                x = "X" if presence.get(TIER_XLSX) else "-"
                t = "T" if presence.get(TIER_TXT_CSV) else "-"
                b = "B" if presence.get(TIER_BM2) else "-"
                row.cells[4 + col_idx].text = f"{x}{t}{b}"

            # Format — compact 7pt font
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(7)
                        r.font.name = "Calibri"

        table_num += 1

    # ===== Table C — Domain Coverage Matrix =====
    if report.domain_coverage:
        headers_c = ["Domain", "xlsx", "txt/csv", "bm2", "Drop (X→T)", "Drop (T→B)"]
        n_domains = len(report.domain_coverage)
        table = doc.add_table(rows=1 + n_domains, cols=len(headers_c))
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Caption
        caption = doc.add_paragraph()
        run = caption.add_run(
            f"Table {table_num}. Domain coverage matrix — animal counts per tier."
        )
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.italic = True
        caption.paragraph_format.space_after = Pt(4)
        table._element.addprevious(caption._element)

        # Header
        for i, h in enumerate(headers_c):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for r in para.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.name = "Calibri"

        # Rows — one per domain
        for row_idx, domain in enumerate(sorted(report.domain_coverage.keys())):
            cov = report.domain_coverage[domain]
            att = report.attrition.get(domain)
            row = table.rows[1 + row_idx]

            row.cells[0].text = DOMAIN_FULL_LABELS.get(domain, domain)
            row.cells[1].text = str(cov.get(TIER_XLSX, 0))
            row.cells[2].text = str(cov.get(TIER_TXT_CSV, 0))
            row.cells[3].text = str(cov.get(TIER_BM2, 0))
            row.cells[4].text = str(len(att.excluded_xlsx_to_txt)) if att else "—"
            row.cells[5].text = str(len(att.excluded_txt_to_bm2)) if att else "—"

            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(9)
                        r.font.name = "Calibri"

        table_num += 1

    # ===== Table D — Attrition Analysis =====
    if report.attrition:
        doc.add_heading("Attrition Analysis", level=3)

        for domain in sorted(report.attrition.keys()):
            att = report.attrition[domain]
            domain_label = DOMAIN_FULL_LABELS.get(domain, domain)
            completeness_pct = report.completeness.get(domain, 1.0) * 100

            # Domain header with completeness
            p = doc.add_paragraph()
            run = p.add_run(f"{domain_label}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run = p.add_run(f" — completeness: {completeness_pct:.0f}%")
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(2)

            # List exclusion reasons
            has_exclusions = False
            for reason_key, animal_ids in sorted(att.exclusion_reasons.items()):
                if not animal_ids:
                    continue
                has_exclusions = True

                # Parse reason label from key like "xlsx_to_txt_dose_exclusion"
                reason_label = reason_key.replace("xlsx_to_txt_", "").replace("txt_to_bm2_", "")
                transition = "xlsx→txt" if "xlsx_to_txt" in reason_key else "txt→bm2"
                reason_display = reason_label.replace("_", " ").title()

                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(
                    f"{transition} {reason_display}: "
                    f"{len(animal_ids)} animals "
                    f"({', '.join(animal_ids[:10])}"
                )
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                if len(animal_ids) > 10:
                    run = p.add_run(f", ... +{len(animal_ids) - 10} more")
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                run = p.add_run(")")
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                p.paragraph_format.space_after = Pt(2)

            if not has_exclusions:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run("No exclusions detected.")
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                p.paragraph_format.space_after = Pt(2)

    # ===== Table E — Consistency Check Results =====
    doc.add_heading("Consistency Checks", level=3)

    if report.consistency_issues:
        for issue in report.consistency_issues:
            issue_type = issue.get("type", "unknown")
            aid = issue.get("animal_id", "?")
            details = issue.get("details", {})

            if issue_type == "sex_mismatch":
                text = (
                    f"Animal {aid}: sex mismatch — "
                    f"found {', '.join(details.get('sexes_found', []))} "
                    f"across domains"
                )
            elif issue_type == "dose_mismatch":
                text = (
                    f"Animal {aid}: dose mismatch — "
                    f"found {', '.join(str(d) for d in details.get('doses_found', []))} "
                    f"across domains"
                )
            else:
                text = f"Animal {aid}: {issue_type}"

            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(2)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No inconsistencies detected.")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(6)

    return table_num
