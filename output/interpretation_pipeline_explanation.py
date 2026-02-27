"""
Generates a Word document (.docx) from the interpretation pipeline explanation
that was given in a previous Claude Code session.

This script uses python-docx to create a formatted document with headings,
paragraphs, tables, and styled text that mirrors the original markdown response.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def build_document():
    """
    Builds and returns a python-docx Document containing the full
    interpretation pipeline explanation.
    """
    doc = Document()

    # -- Document styles --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # -- Title --
    title = doc.add_heading('How the Interpretation Document Task Was Achieved', level=1)

    doc.add_paragraph(
        'The interpretation document task is a multi-stage pipeline that transforms '
        'BMDExpress dose-response gene expression data (a CSV of gene symbols + benchmark '
        'dose values) into a comprehensive, literature-grounded toxicogenomics interpretation '
        'report. Here\u2019s how it works:'
    )

    # -------------------------------------------------------------------
    # Pipeline Overview
    # -------------------------------------------------------------------
    doc.add_heading('Pipeline Overview', level=2)
    doc.add_paragraph(
        'The main engine is interpret.py (1,797 lines), which orchestrates 6 steps:'
    )

    # -- Step 1 --
    doc.add_heading('1. Load & Parse Input CSV', level=3)
    doc.add_paragraph(
        'Reads a BMD CSV file (gene symbol, BMD value, direction up/down), '
        'normalizes gene symbols, and extracts the gene list.'
    )

    # -- Step 2 --
    doc.add_heading('2. Structured Enrichment Analysis', level=3)
    doc.add_paragraph(
        'Performs statistical enrichment using Fisher\u2019s exact test with '
        'Benjamini-Hochberg FDR correction:'
    )
    bullets_2 = [
        ('Pathway enrichment', 'KEGG/Reactome pathways over-represented in the gene set'),
        ('GO term enrichment', 'Biological processes enriched'),
        ('BMD ordering', 'Median/min BMD per pathway, showing dose-response progression'),
        ('Organ signature', 'Which organs are enriched among responsive genes'),
        ('Literature context', 'Genes, papers, claims, and evidence levels from the knowledge base'),
    ]
    for bold_part, rest in bullets_2:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(f' \u2014 {rest}')

    # -- Step 3 --
    doc.add_heading('3. Format Structured Context', level=3)
    doc.add_paragraph(
        'Assembles all enrichment results into a text block (top 20 pathways, top 20 GO terms, '
        'organ scores, per-gene literature claims, top papers) that fits within LLM context windows.'
    )

    # -- Step 4 --
    doc.add_heading('4. Multi-Model Narrative Generation', level=3)
    doc.add_paragraph(
        'Sends the structured context + a detailed interpretation prompt to '
        'multiple LLMs in parallel:'
    )
    models = [
        ('Local Ollama', 'qwen2.5:14b on two GPUs (RTX 3080 Ti + RX 6900 XT via SSH)'),
        ('Anthropic API', 'claude-sonnet-4-6, claude-opus-4-6, claude-haiku-4-5'),
    ]
    for bold_part, rest in models:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(f': {rest}')

    doc.add_paragraph(
        'Each model runs 3 times \u2192 15 independent narratives total.'
    )

    doc.add_paragraph('Each narrative covers 6 sections:')
    sections = [
        'Biological Response Narrative (phase-by-phase dose progression)',
        'Organ-Level Prediction',
        'Mechanism of Action (molecular initiating event + adverse outcome pathway)',
        'Protective vs. Adverse Responses',
        'Literature Support (citations from the KB)',
        'Confidence Assessment',
    ]
    for i, s in enumerate(sections, 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')

    # -- Step 5 --
    doc.add_heading('5. Concordance Analysis', level=3)
    doc.add_paragraph(
        'A designated concordance model (claude-sonnet-4-6) receives all 15 narratives and identifies:'
    )
    concordance_items = [
        'Where models agree (high-confidence findings)',
        'Where models diverge (conflicting interpretations)',
        'Model-specific unique insights',
        'Overall concordance ratings per section',
    ]
    for item in concordance_items:
        doc.add_paragraph(item, style='List Bullet')

    # -- Step 6 --
    doc.add_heading('6. Export', level=3)
    doc.add_paragraph(
        'Outputs to the output/ directory as both Markdown and Word documents (.docx):'
    )
    doc.add_paragraph(
        'Per-model narratives, combined narratives, and concordance report.',
        style='List Bullet'
    )

    # -------------------------------------------------------------------
    # Knowledge Base
    # -------------------------------------------------------------------
    doc.add_heading('Knowledge Base (bmdx.duckdb)', level=2)
    doc.add_paragraph(
        'The pipeline queries a DuckDB database with 9 normalized tables built from:'
    )
    kb_items = [
        ('2,315 papers', 'from citation graph crawls across 7 sources (citegraph.py)'),
        ('2,840 GO terms', 'from BMDExpress annotations'),
        ('1,377 genes', '(161 consensus, 149 moderate evidence)'),
        ('4,431 claims', 'extracted from abstracts by LLMs (extract.py)'),
        ('2,860 pathway mappings', '(KEGG + Reactome, via pathway_enrich.py)'),
        ('2,224 citation edges', 'for graph traversal'),
    ]
    for bold_part, rest in kb_items:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(f' {rest}')

    # -------------------------------------------------------------------
    # Key Design Choices
    # -------------------------------------------------------------------
    doc.add_heading('Key Design Choices', level=2)
    choices = [
        ('Structured-then-narrative',
         'Mathematical enrichment first provides grounded statistics; LLM synthesis '
         'then adds domain expertise and contextual interpretation.'),
        ('Multi-model consensus',
         '5 models \u00d7 3 runs identifies robust findings (supported by 4+ models) '
         'vs. model-specific hallucinations.'),
        ('Recursive decomposition',
         'Rather than stuffing all 2,315 papers into context, the system programmatically '
         'queries the KB to retrieve only relevant evidence per gene/pathway \u2014 staying '
         'within context limits.'),
        ('Dose-ordered response',
         'Pathways ranked by median BMD reveal a progression from adaptive (low dose) '
         'to adverse (high dose) responses.'),
        ('Literature traceability',
         'Every interpretation is tied back to specific papers, genes, and claims in the '
         'knowledge base.'),
    ]
    for bold_part, rest in choices:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(f': {rest}')

    # -------------------------------------------------------------------
    # Supporting Scripts table
    # -------------------------------------------------------------------
    doc.add_heading('Supporting Scripts', level=2)

    table_data = [
        ('Script', 'Role'),
        ('interpret.py', 'Main interpretation engine'),
        ('build_db.py', 'Builds bmdx.duckdb from crawl data'),
        ('citegraph.py', 'Citation graph crawler'),
        ('extract.py', 'LLM extraction from abstracts (genes, organs, claims)'),
        ('go_gene_map.py', 'GO-gene cross-reference mapping'),
        ('pathway_enrich.py', 'KEGG + Reactome enrichment queries'),
        ('fulltext.py', 'Full-text retrieval from PMC, arXiv, S2, Unpaywall'),
    ]

    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (col1, col2) in enumerate(table_data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            # Bold header row
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    return doc


if __name__ == '__main__':
    doc = build_document()
    out_path = os.path.join(os.path.dirname(__file__), 'interpretation_pipeline_explanation.docx')
    doc.save(out_path)
    print(f'Saved to {out_path}')
