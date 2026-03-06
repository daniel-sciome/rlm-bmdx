"""
FastAPI server for 5dToxReport.

Provides a web interface and REST API for generating structured toxicology
reports for "5 Day Genomic Dose Response in Sprague-Dawley Rats" studies.
The report follows the NIEHS Report 10 structure: Background, Materials and
Methods, Results (apical endpoints, BMD summary, gene set BMD analysis, gene
BMD analysis), and Summary.

Session persistence: approved sections are saved as JSON files under
sessions/{dtxsid}/, keyed by the chemical's DTXSID.  When a DTXSID is
resolved, any previously-saved session is auto-restored in the UI.

Version history: every Approve archives the previous version into
sessions/{dtxsid}/history/{section_key}/ with a timestamped filename.
Users can browse past versions and non-destructively restore any of them
(restoring creates a new version with the old content).

Endpoints:
  GET  /                                            — Serve the web/index.html UI
  POST /api/resolve                                 — Resolve a chemical identifier to all IDs
  POST /api/generate                                — Gather data + generate 6-paragraph background
  POST /api/upload-bm2                              — Upload .bm2 files for apical endpoint analysis
  POST /api/upload-csv                              — Upload gene-level BMD CSV for transcriptomic analysis
  POST /api/upload-zip                              — Upload a .zip archive; extract and register .bm2, .csv, .txt, .xlsx files
  POST /api/process-bm2                             — Process a .bm2 file and return table data as JSON
  POST /api/process-genomics                        — Rank GO gene sets and genes by BMD from CSV data
  GET  /api/session/{dtxsid}/bmd-summary            — Auto-derive BMD/BMDL/LOEL/NOEL summary from approved .bm2 sections
  POST /api/generate-methods                        — LLM-generate Materials and Methods from study params
  POST /api/generate-summary                        — LLM-generate Summary synthesizing all approved sections
  POST /api/export-docx                             — Export full report in NIEHS section order
  POST /api/export-pdf                              — Export full report as tagged PDF/UA-1
  GET  /api/session/{dtxsid}                        — Load a previously-saved session (all section types)
  POST /api/session/approve                         — Approve (save) a report section
  POST /api/session/unapprove                       — Unapprove (delete) a report section
  GET  /api/session/{dtxsid}/history/{section_key}  — Version history for a section
  POST /api/session/{dtxsid}/restore                — Restore a past version as current
  GET  /api/style-profile                           — Retrieve the global style profile
  DELETE /api/style-profile/{idx}                   — Delete a style rule by index

The server uses Server-Sent Events (SSE) during /api/generate to stream
progress updates to the browser in real time (e.g., "Querying ATSDR...",
"Generating with Claude...").

Usage:
    python background_server.py                   # start on port 8000
    python background_server.py --port 8080       # custom port
    python background_server.py --host 0.0.0.0    # listen on all interfaces
"""

import asyncio
import json
import logging
import os
import re
import shutil
import sys
import tempfile
import traceback
import uuid
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import orjson
from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import (
    FileResponse, HTMLResponse, JSONResponse, Response, StreamingResponse,
)
from starlette.staticfiles import StaticFiles

import bm2_cache
from chem_resolver import ChemicalIdentity, resolve_chemical
from data_gatherer import BackgroundData, gather_all
from background_writer import generate_background
from file_integrator import (
    FileFingerprint,
    ValidationReport,
    fingerprint_file,
    validate_pool,
    lightweight_validate,
)
from pool_integrator import integrate_pool
from animal_report import (
    build_animal_report,
    report_to_dict,
    add_animal_report_to_doc,
)

# AnthropicEndpoint wraps the Claude API — used here for style rule extraction
# (haiku model for speed/cost, ~$0.001 per extraction call).
# rank_go_sets_by_bmd / rank_genes_by_bmd produce the data for the Gene Set
# and Gene BMD Analysis report sections (NIEHS format).
# ToxKBQuerier provides read-only access to the bmdx.duckdb knowledge base.
# load_dose_response parses and validates gene-level BMD CSV files.
from interpret import (
    AnthropicEndpoint,
    ToxKBQuerier,
    fetch_gene_descriptions,
    fetch_go_descriptions,
    load_dose_response,
    rank_go_sets_by_bmd,
    rank_genes_by_bmd,
)

import pandas as pd

logger = logging.getLogger(__name__)

# Import .docx building utilities from the existing codebase
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from build_docx import add_heading, add_para, fmt

# Import the composable helpers from apical_report for .bm2 processing
# and the new NIEHS table builders for BMD summary + genomics sections
from apical_report import (
    build_table_data_from_bm2,
    add_apical_tables_to_doc,
    generate_results_narrative,
    add_bmd_summary_table_to_doc,
    add_gene_set_bmd_tables_to_doc,
    add_gene_bmd_tables_to_doc,
)


# ---------------------------------------------------------------------------
# FastAPI app
# ---------------------------------------------------------------------------

app = FastAPI(
    title="5dToxReport",
    description=(
        "Generate toxicology reports for 5-day genomic dose-response studies. "
        "Includes background sections and NTP-style apical endpoint tables."
    ),
    version="2.0.0",
)

# ---------------------------------------------------------------------------
# Server-side session storage for uploaded .bm2 files
# ---------------------------------------------------------------------------
# Maps bm2_id (UUID string) → dict with filename, temp_path, and table_data
# (table_data is populated after /api/process-bm2 is called).
# This is an in-memory store; files live in a temp directory per upload.
_bm2_uploads: dict[str, dict] = {}

# Maps csv_id (UUID string) → dict with filename, temp_path, and parsed DataFrame.
# Used for gene-level BMD CSV uploads (transcriptomic data for the Gene Set
# and Gene BMD Analysis report sections).
_csv_uploads: dict[str, dict] = {}

# Maps file_id (UUID string) → dict with filename, temp_path, type.
# Used for raw dose-response experimental data (.csv, .txt, .xlsx) extracted
# from zip archives.  These are BMDExpress-importable input data, not
# gene-level BMD results.  The client references these by ID in the file pool.
_data_uploads: dict[str, dict] = {}

# Maps dtxsid → { file_id → FileFingerprint } for cross-validation.
# Populated when files are fingerprinted (on upload or validation request).
# Persisted to disk as validation_report.json per session directory.
_pool_fingerprints: dict[str, dict[str, FileFingerprint]] = {}

# Maps dtxsid → merged BMDProject dict from pool integration.
# Populated by /api/pool/integrate/{dtxsid} and persisted to
# sessions/{dtxsid}/integrated.json for cross-session restore.
_integrated_pool: dict[str, dict] = {}


# ---------------------------------------------------------------------------
# Session persistence helpers
# ---------------------------------------------------------------------------
# Approved sections are persisted to disk as JSON files under sessions/{dtxsid}/.
# This allows the user to close the browser, restart the server, and pick up
# exactly where they left off — the UI auto-restores on DTXSID resolution.

SESSIONS_DIR = Path(__file__).parent / "sessions"

# Global style profile — stores learned writing style rules extracted from
# user edits.  Lives at sessions/_style_profile.json (not per-chemical,
# because writing style transcends individual chemicals).
STYLE_PROFILE_PATH = SESSIONS_DIR / "_style_profile.json"

# Maximum number of style rules to retain — when full, the lowest-confidence
# rule is evicted to make room for new ones.
MAX_STYLE_RULES = 30

# Valid categories for style rules — used for validation when merging
STYLE_CATEGORIES = {
    "terminology", "grammar", "phrasing", "structure", "formatting", "tone",
}


def _load_style_profile() -> dict:
    """
    Load the global style profile from disk, returning an empty structure if
    the file doesn't exist yet.

    The profile has three top-level keys:
      - version (int): schema version for future migrations
      - updated_at (str): ISO 8601 timestamp of last modification
      - rules (list[dict]): ordered list of learned style rules, each with:
          rule (str), category (str), confidence (int), first_seen, last_seen
    """
    if STYLE_PROFILE_PATH.exists():
        try:
            return json.loads(STYLE_PROFILE_PATH.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            # Corrupted file — start fresh rather than crash
            pass
    return {"version": 1, "updated_at": _now_iso(), "rules": []}


def _save_style_profile(profile: dict) -> None:
    """
    Write the style profile to disk.  Creates the sessions/ directory if
    needed (same dir used for per-chemical session storage).
    """
    SESSIONS_DIR.mkdir(parents=True, exist_ok=True)
    profile["updated_at"] = _now_iso()
    STYLE_PROFILE_PATH.write_text(
        json.dumps(profile, indent=2, default=str), encoding="utf-8",
    )


def _merge_rules(profile: dict, new_rules: list[dict]) -> dict:
    """
    Merge newly-extracted style rules into the existing profile.

    For each new rule:
      - If it matches an existing rule (exact string match on the 'rule' key),
        increment the existing rule's confidence and update last_seen.
      - If it's genuinely new, add it with confidence=1.
      - If adding pushes the total past MAX_STYLE_RULES, evict the rule with
        the lowest confidence (oldest last_seen breaks ties).

    The LLM extraction prompt is given the existing rules to avoid semantic
    duplicates, so exact-match dedup here is a safety net — most real dedup
    happens at extraction time.

    Returns the modified profile (also mutates in place for convenience).
    """
    existing_rules = profile.get("rules", [])
    now = _now_iso()

    for nr in new_rules:
        rule_text = nr.get("rule", "").strip()
        category = nr.get("category", "phrasing").strip().lower()
        if not rule_text:
            continue
        # Default to "phrasing" if the LLM returns an unknown category
        if category not in STYLE_CATEGORIES:
            category = "phrasing"

        # Check for exact duplicate by rule text
        matched = False
        for er in existing_rules:
            if er["rule"].strip().lower() == rule_text.lower():
                # Reinforce: bump confidence and update timestamp
                er["confidence"] = er.get("confidence", 1) + 1
                er["last_seen"] = now
                matched = True
                break

        if not matched:
            existing_rules.append({
                "rule": rule_text,
                "category": category,
                "confidence": 1,
                "first_seen": now,
                "last_seen": now,
            })

    # Evict lowest-confidence rules if over the cap.
    # Sort by confidence ascending, then by last_seen ascending (oldest first)
    # so we drop the least-reinforced, least-recently-seen rules.
    if len(existing_rules) > MAX_STYLE_RULES:
        existing_rules.sort(
            key=lambda r: (r.get("confidence", 1), r.get("last_seen", "")),
        )
        existing_rules = existing_rules[-MAX_STYLE_RULES:]

    profile["rules"] = existing_rules
    return profile


def _extract_and_merge_style_rules(original: str, edited: str) -> int:
    """
    Extract writing style rules by comparing original LLM text to user edits,
    then merge them into the global style profile.

    This runs in a background thread (called via run_in_executor) so it
    doesn't block the approve response.  Uses Claude Haiku for speed and
    cost (~$0.001 per call).

    Args:
        original: The original text generated by the LLM
        edited: The user's edited version of the same text

    Returns:
        Number of new rules extracted and merged (0 if none found or on error)
    """
    try:
        profile = _load_style_profile()
        existing_rule_strings = [r["rule"] for r in profile.get("rules", [])]

        # Build the extraction prompt — tells the LLM to compare the two
        # versions and find deliberate style preferences.  Existing rules are
        # included so the LLM can avoid re-extracting duplicates.
        existing_rules_json = json.dumps(existing_rule_strings, indent=2)
        prompt = f"""Compare the ORIGINAL and EDITED versions of this scientific/toxicology text.
The editor made deliberate style changes. Extract specific, reusable writing
style rules that the editor is consistently applying.

Focus on these categories:
- terminology: preferred word choices and technical terms
- grammar: comma usage, voice (active/passive), tense preferences
- phrasing: preferred sentence constructions, transition patterns
- structure: paragraph organization, how information is ordered
- formatting: citation style, abbreviation conventions
- tone: formality level, hedging language, precision

EXISTING RULES (already learned — do NOT re-extract these):
{existing_rules_json}

ORIGINAL TEXT:
{original}

EDITED TEXT:
{edited}

Return ONLY a JSON array of new rules not already covered above:
[{{"rule": "description of the style preference", "category": "terminology|grammar|phrasing|structure|formatting|tone"}}]

If no new rules are evident, return an empty array: []"""

        # Use Haiku for speed and cost — style rule extraction doesn't need
        # the full reasoning power of Sonnet/Opus.
        # _llm_generate_json handles endpoint creation, fence stripping,
        # and JSON parsing in one call.
        new_rules = _llm_generate_json(
            "style-rule-extractor",
            prompt,
            system=(
                "You are a writing style analyst. Compare original and edited text "
                "to identify deliberate style preferences. Output ONLY valid JSON."
            ),
            model="claude-haiku-4-5-20251001",
            max_tokens=2048,
        )

        if not isinstance(new_rules, list) or len(new_rules) == 0:
            logger.info("Style extraction found no new rules")
            return 0

        # Merge extracted rules into the profile and save
        count_before = len(profile.get("rules", []))
        _merge_rules(profile, new_rules)
        _save_style_profile(profile)
        count_after = len(profile.get("rules", []))

        new_count = count_after - count_before
        logger.info(
            "Style learning: extracted %d rules, %d new (total: %d)",
            len(new_rules), max(new_count, 0), count_after,
        )
        return len(new_rules)

    except Exception:
        # Style learning is non-critical — log the error but don't crash
        # the approve flow
        logger.error("Style rule extraction failed:\n%s", traceback.format_exc())
        return 0


def _session_dir(dtxsid: str) -> Path:
    """
    Return the session directory for a given DTXSID, creating it if needed.

    Each chemical gets its own directory under sessions/ (e.g.,
    sessions/DTXSID6020430/).  The directory is created on first approve.
    """
    d = SESSIONS_DIR / dtxsid
    d.mkdir(parents=True, exist_ok=True)
    return d


def _bm2_slug(filename: str) -> str:
    """
    Slugify a .bm2 filename for use as a JSON key / filename stem.

    Strips the compound prefix (everything before the first hyphen), lowercases,
    replaces spaces/non-alphanum with hyphens, and strips the .bm2 extension.

    Example:
        'P3MP-Organ and Body Weights.bm2' → 'organ-and-body-weights'
        'P3MP-Clinical Pathology.bm2'     → 'clinical-pathology'
    """
    # Remove .bm2 extension
    stem = filename.rsplit(".bm2", 1)[0]
    # Drop the compound prefix before the first hyphen (e.g. "P3MP-")
    if "-" in stem:
        stem = stem.split("-", 1)[1]
    # Lowercase, replace non-alphanumeric runs with single hyphens, strip edges
    slug = re.sub(r"[^a-z0-9]+", "-", stem.lower()).strip("-")
    return slug


def _now_iso() -> str:
    """Return the current UTC time as an ISO 8601 string."""
    return datetime.now(timezone.utc).isoformat()


def _save_section(dtxsid: str, section_key: str, data: dict) -> None:
    """
    Write data as JSON to sessions/{dtxsid}/{section_key}.json, archiving the
    previous version first so we keep full history.

    Version history layout:
        sessions/{dtxsid}/history/{section_key}/{safe_timestamp}.json
    The current file ({section_key}.json) is always the latest version.
    Before overwriting, the existing current is copied into history/ using
    its approved_at timestamp as the filename.

    Also stamps the data with an incrementing "version" number (v1, v2, ...)
    and updates meta.json's updated_at timestamp.
    """
    d = _session_dir(dtxsid)
    current_path = d / f"{section_key}.json"
    history_dir = d / "history" / section_key

    # --- Archive the current version before overwriting (if it exists) ---
    # This preserves every previously-approved version as a timestamped file
    # in the history/ subdirectory.  First-ever approve has no file to archive.
    if current_path.exists():
        existing = json.loads(current_path.read_text(encoding="utf-8"))
        # Use the existing file's approved_at as the archive filename
        # so timestamps reflect when that version was actually approved
        ts = existing.get("approved_at", _now_iso())
        # Replace colons with hyphens so the filename is filesystem-safe
        # (ISO 8601 timestamps contain colons, e.g. "2026-03-02T19:23:59+00:00")
        safe_ts = ts.replace(":", "-")
        history_dir.mkdir(parents=True, exist_ok=True)
        (history_dir / f"{safe_ts}.json").write_text(
            json.dumps(existing, indent=2, default=str), encoding="utf-8",
        )

    # --- Compute the version number for this new save ---
    # Count existing history files (previous versions) and add 1 for the
    # new version.  First approve = 0 history files → version 1.
    version_count = len(list(history_dir.glob("*.json"))) if history_dir.exists() else 0
    data["version"] = version_count + 1

    # --- Write the new version as the canonical current file ---
    current_path.write_text(
        json.dumps(data, indent=2, default=str), encoding="utf-8",
    )

    # Touch meta.json's updated_at
    meta_path = d / "meta.json"
    if meta_path.exists():
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
    else:
        meta = {"dtxsid": dtxsid, "created_at": _now_iso()}
    meta["updated_at"] = _now_iso()
    meta_path.write_text(json.dumps(meta, indent=2), encoding="utf-8")


def _delete_section(dtxsid: str, section_key: str) -> None:
    """
    Remove sessions/{dtxsid}/{section_key}.json if it exists.

    Called when the user clicks "Try Again" to unapprove a section.
    The .bm2 file in files/ is kept — it's still useful for reprocessing.
    """
    d = SESSIONS_DIR / dtxsid
    section_path = d / f"{section_key}.json"
    if section_path.exists():
        section_path.unlink()


# ---------------------------------------------------------------------------
# GET / — serve the web UI
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse)
async def serve_ui():
    """
    Serve the main web UI (web/index.html).

    The HTML file contains the full single-page application with chemical
    ID form, .bm2 upload area, output panels, and copy/export buttons.
    """
    html_path = Path(__file__).parent / "web" / "index.html"
    if not html_path.exists():
        return HTMLResponse(
            "<h1>Error</h1><p>web/index.html not found</p>",
            status_code=404,
        )
    return HTMLResponse(html_path.read_text(encoding="utf-8"))


# ---------------------------------------------------------------------------
# POST /api/resolve — chemical identity resolution
# ---------------------------------------------------------------------------

@app.post("/api/resolve")
async def api_resolve(request: Request):
    """
    Resolve a chemical identifier to all known identifiers.

    Input JSON:
      {"identifier": "95-50-1", "id_type": "auto"}

    Returns the full ChemicalIdentity as JSON, with all resolved fields
    (name, CASRN, DTXSID, CID, EC number, formula, etc.).
    """
    body = await request.json()
    identifier = body.get("identifier", "").strip()
    id_type = body.get("id_type", "auto")

    if not identifier:
        return JSONResponse(
            {"error": "No identifier provided"},
            status_code=400,
        )

    # Run resolution in a thread pool to avoid blocking the event loop
    # (it makes multiple HTTP requests to PubChem and CTX)
    loop = asyncio.get_running_loop()
    identity = await loop.run_in_executor(
        None, resolve_chemical, identifier, id_type, "",
    )

    return JSONResponse(identity.to_dict())


# ---------------------------------------------------------------------------
# POST /api/generate — full pipeline: gather data + generate background
# ---------------------------------------------------------------------------

@app.post("/api/generate")
async def api_generate(request: Request):
    """
    Run the full background generation pipeline.

    Input JSON:
      {"identity": {...ChemicalIdentity fields...}, "use_ollama": false, "model": ""}

    Returns a streaming SSE response with progress updates, followed by
    the final generated background as JSON.

    SSE events:
      event: progress   data: {"message": "Querying ATSDR ToxProfiles..."}
      event: complete   data: {"paragraphs": [...], "references": [...], ...}
      event: error      data: {"error": "..."}
    """
    body = await request.json()
    identity_dict = body.get("identity", {})
    use_ollama = body.get("use_ollama", False)
    model = body.get("model", "")

    if not identity_dict.get("name") and not identity_dict.get("casrn"):
        return JSONResponse(
            {"error": "Identity must include at least a name or CASRN"},
            status_code=400,
        )

    # Reconstruct ChemicalIdentity from the JSON dict
    identity = ChemicalIdentity.from_dict(identity_dict)

    # Use SSE to stream progress updates
    async def event_stream():
        progress_messages = []

        def progress_callback(msg: str):
            """Collect progress messages from the data gathering step."""
            progress_messages.append(msg)

        try:
            # Step 1: Gather data (with progress callback)
            loop = asyncio.get_running_loop()

            # Yield initial progress
            yield _sse_event("progress", {"message": "Starting data gathering..."})

            # Run gather_all in a thread pool (it makes blocking HTTP calls)
            bg_data = await loop.run_in_executor(
                None, gather_all, identity, progress_callback,
            )

            # Yield all progress messages collected during gathering
            for msg in progress_messages:
                yield _sse_event("progress", {"message": msg})

            # Step 2: Load learned style preferences (if any) and generate
            # background with LLM.  Style rules are injected into the prompt
            # so the LLM writes in the user's preferred style from the start.
            style_rules = []
            profile = _load_style_profile()
            if profile.get("rules"):
                style_rules = [r["rule"] for r in profile["rules"]]

            if style_rules:
                yield _sse_event("progress", {
                    "message": f"Applying {len(style_rules)} learned style preference{'s' if len(style_rules) != 1 else ''}..."
                })

            yield _sse_event("progress", {
                "message": f"Generating background with {'Ollama' if use_ollama else 'Claude'}..."
            })

            result = await loop.run_in_executor(
                None, generate_background, bg_data, use_ollama, model,
                style_rules or None,
            )

            # Step 3: Return the complete result
            # Include the raw data for the export endpoint
            result["raw_data"] = bg_data.to_dict()
            result["notes"] = bg_data.notes

            yield _sse_event("complete", result)

        except Exception as e:
            yield _sse_event("error", {"error": str(e)})

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # disable nginx buffering
        },
    )


def _sse_event(event_type: str, data: dict) -> str:
    """Format a Server-Sent Event message."""
    json_str = json.dumps(data, default=str)
    return f"event: {event_type}\ndata: {json_str}\n\n"


def _js_dose_key(dose: float) -> str:
    """
    Convert a dose float to the same string JavaScript's String() produces.

    Python's str(1.0) gives "1.0" but JavaScript's String(1) gives "1".
    This mismatch causes dict key lookups to fail in the browser when the
    JS code does row.values[String(dose)].  We fix it by dropping the
    trailing ".0" for integer-valued floats, matching JS behavior.

    Examples:
        _js_dose_key(0.0)   → "0"
        _js_dose_key(0.15)  → "0.15"
        _js_dose_key(1.0)   → "1"
        _js_dose_key(300.0) → "300"
    """
    if dose == int(dose):
        return str(int(dose))
    return str(dose)


def _serialize_table_rows(table_data: dict) -> dict:
    """
    Convert a {sex: [TableRow, ...]} dict to JSON-friendly nested dicts.

    Each TableRow has values_by_dose, n_by_dose, bmd_str, bmdl_str, and
    trend_marker attributes.  Dose float keys are converted via _js_dose_key()
    to match JavaScript's String(number) behavior.

    Used by /api/process-bm2 and /api/process-integrated to serialize
    the NTP stats pipeline output for the browser.

    Args:
        table_data: Dict mapping sex label ("Male", "Female") to lists of
                    TableRow objects from apical_report.build_table_data().

    Returns:
        Dict mapping sex label to lists of JSON-serializable row dicts.
    """
    tables_json = {}
    for sex, rows in table_data.items():
        tables_json[sex] = []
        for row in rows:
            sorted_doses = sorted(row.values_by_dose.keys())
            tables_json[sex].append({
                "label": row.label,
                "doses": sorted_doses,
                "values": {_js_dose_key(d): v for d, v in row.values_by_dose.items()},
                "n": {_js_dose_key(d): n for d, n in row.n_by_dose.items()},
                "bmd": row.bmd_str,
                "bmdl": row.bmdl_str,
                "trend_marker": row.trend_marker,
            })
    return tables_json


def _strip_markdown_fences(text: str) -> str:
    """
    Remove markdown code fences (```json ... ```) from LLM responses.

    Claude sometimes wraps JSON output in markdown code blocks even when
    instructed to return raw JSON.  This strips those fences so the
    response can be parsed as JSON.

    Args:
        text: Raw LLM response text, possibly wrapped in code fences.

    Returns:
        The text with leading ```json and trailing ``` removed, if present.
    """
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned


def _llm_generate_json(
    name: str,
    prompt: str,
    system: str,
    *,
    model: str = "claude-sonnet-4-6",
    max_tokens: int = 8192,
    temperature: float = 0.2,
) -> Any:
    """
    Synchronous helper: call Claude → strip markdown fences → parse JSON.

    Centralizes the repeated pattern of creating an AnthropicEndpoint,
    generating text, stripping markdown code fences, and JSON-parsing the
    result.  Can be called directly in sync code (e.g. inside a thread-pool
    worker) or wrapped in run_in_executor for async endpoints.

    Args:
        name:        Logical name for the endpoint (appears in logs / billing).
        prompt:      The user-turn prompt to send.
        system:      The system prompt.
        model:       Claude model ID (default: claude-sonnet-4-6).
        max_tokens:  Max tokens for the response.
        temperature: Sampling temperature.

    Returns:
        Parsed JSON (dict or list) from the LLM response.

    Raises:
        ValueError:  If the LLM returns an empty response.
        json.JSONDecodeError: If the response isn't valid JSON after fence stripping.
    """
    endpoint = AnthropicEndpoint(
        name=name,
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
    )
    response = endpoint.generate(prompt, system=system)
    if not response:
        raise ValueError(f"LLM '{name}' returned empty response")
    return json.loads(_strip_markdown_fences(response))


async def _llm_generate_json_async(
    name: str,
    prompt: str,
    system: str,
    **kwargs,
) -> Any:
    """
    Async wrapper around _llm_generate_json — runs the blocking LLM call
    in a thread-pool executor so it doesn't block the event loop.

    Accepts the same keyword arguments as _llm_generate_json (model,
    max_tokens, temperature).

    Returns:
        Parsed JSON (dict or list) from the LLM response.
    """
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(
        None,
        lambda: _llm_generate_json(name, prompt, system, **kwargs),
    )


def _safe_filename(name: str) -> str:
    """
    Sanitize a chemical name for use as a filename.

    Replaces non-alphanumeric characters (except spaces, hyphens, and
    underscores) with underscores.  Used when building download filenames
    for exported .docx and .pdf reports.

    Args:
        name: The chemical name (e.g., "1,2-Dichlorobenzene").

    Returns:
        A filesystem-safe string (e.g., "1_2-Dichlorobenzene").
    """
    return "".join(c if c.isalnum() or c in " -_" else "_" for c in name)


def _ensure_fingerprints(dtxsid: str, force: bool = False) -> dict:
    """
    Ensure fingerprints are populated for a session's file pool.

    Checks the in-memory _pool_fingerprints cache first.  If empty (e.g.,
    after a server restart), re-fingerprints all files from the session's
    files/ directory by scanning _bm2_uploads, _data_uploads, and the
    filesystem.

    This logic was previously duplicated in api_pool_validate() and
    api_generate_animal_report().

    Args:
        dtxsid: The DTXSID identifying the session.
        force:  If True, clear existing fingerprints and re-scan from disk.
                Used by the validation endpoint which always wants a fresh scan.

    Returns:
        The fingerprint dict {file_id: FileFingerprint} for this session.
    """
    fps = _pool_fingerprints.get(dtxsid, {})
    if fps and not force:
        return fps

    # Re-fingerprint all files from disk
    files_dir = _session_dir(dtxsid) / "files"
    if not files_dir.exists():
        return {}

    _pool_fingerprints[dtxsid] = {}
    fingerprinted: set[str] = set()

    # 1. Fingerprint files registered in _bm2_uploads
    for fid, entry in _bm2_uploads.items():
        path = entry.get("temp_path", "")
        if path and os.path.exists(path) and str(files_dir) in path:
            bm2_json = entry.get("bm2_json")
            _fingerprint_and_store(fid, entry["filename"], path, "bm2", dtxsid, bm2_json)
            fingerprinted.add(entry["filename"])

    # 2. Fingerprint files registered in _data_uploads
    for fid, entry in _data_uploads.items():
        path = entry.get("temp_path", "")
        if path and os.path.exists(path) and str(files_dir) in path:
            _fingerprint_and_store(fid, entry["filename"], path, entry["type"], dtxsid)
            fingerprinted.add(entry["filename"])

    # 3. Scan files/ directory for anything not yet registered
    for data_file in sorted(files_dir.iterdir()):
        if not data_file.is_file() or data_file.name in fingerprinted:
            continue
        ext = data_file.suffix.lower().lstrip(".")
        if ext not in ("xlsx", "txt", "csv", "bm2"):
            continue
        fid = f"scan-{data_file.name}"
        _fingerprint_and_store(fid, data_file.name, str(data_file), ext, dtxsid)

    return _pool_fingerprints.get(dtxsid, {})


# ---------------------------------------------------------------------------
# POST /api/upload-bm2 — upload .bm2 files for apical endpoint analysis
# ---------------------------------------------------------------------------

@app.post("/api/upload-bm2")
async def api_upload_bm2(request: Request, files: list[UploadFile] = File(...)):
    """
    Accept one or more .bm2 file uploads.

    Saves each file to the session's files/ directory (if a dtxsid query
    parameter is provided) so it persists across page reloads, or to a
    temp directory as a fallback.  Returns a JSON list of metadata objects:
    {id, filename}.  The id is a UUID that other endpoints use to
    reference the uploaded file.
    """
    # If the client provides a DTXSID, save directly to the session's
    # files/ directory so the file survives page reloads and server
    # restarts — no need to wait for approve to persist.
    dtxsid = request.query_params.get("dtxsid", "")
    persist_dir = None
    if dtxsid:
        persist_dir = _session_dir(dtxsid) / "files"
        persist_dir.mkdir(exist_ok=True)

    results = []

    for upload in files:
        bm2_id = str(uuid.uuid4())
        safe_filename = os.path.basename(upload.filename or "upload.bm2")
        content = await upload.read()

        if persist_dir:
            # Save to the session's files/ directory for persistence
            file_path = str(persist_dir / safe_filename)
            with open(file_path, "wb") as f:
                f.write(content)
        else:
            # Fallback: save to a temp directory (lost on restart)
            tmp_dir = tempfile.mkdtemp(prefix="bm2_")
            file_path = os.path.join(tmp_dir, safe_filename)
            with open(file_path, "wb") as f:
                f.write(content)

        _bm2_uploads[bm2_id] = {
            "filename": safe_filename,
            "temp_path": file_path,
            "table_data": None,   # populated by /api/process-bm2 or preview
            "bm2_json": None,     # populated on first preview/process
            "narrative": None,    # populated by /api/process-bm2
        }

        # Fingerprint the uploaded file for cross-validation.
        # bm2 fingerprints are partial at this stage (no JSON yet) because
        # we haven't run the expensive Java export — domain is detected
        # from the filename only.  Full metadata fills in on first preview.
        validation_issues = []
        if dtxsid:
            fp = _fingerprint_and_store(bm2_id, safe_filename, file_path, "bm2", dtxsid)
            validation_issues = _run_lightweight_validation(fp, dtxsid)

        results.append({
            "id": bm2_id,
            "filename": safe_filename,
            "validation_issues": validation_issues,
        })

    return JSONResponse(results)


# ---------------------------------------------------------------------------
# POST /api/process-bm2 — analyze a .bm2 file and return table data as JSON
# ---------------------------------------------------------------------------

@app.post("/api/process-bm2")
async def api_process_bm2(request: Request):
    """
    Process an uploaded .bm2 file through the NTP statistical pipeline.

    Input JSON:
      {
        "bm2_id": "<uuid from upload>",
        "section_title": "Animal Condition, Body Weights, and Organ Weights",
        "table_caption_template": "Summary of ... {sex} ... {compound} ...",
        "compound_name": "PFHxSAm",
        "dose_unit": "mg/kg"
      }

    Calls build_table_data_from_bm2() from apical_report.py, which exports
    the .bm2 via the BMDExpress Java CLI and runs Jonckheere → Williams/Dunnett
    tests.

    Returns JSON with the table data for HTML preview:
      {
        "bm2_id": "...",
        "tables": {
          "Male": [
            {
              "label": "Terminal Body Wt.",
              "doses": [0.0, 0.15, 0.5, 1.5, 5.0],
              "values": {"0.0": "330.2 ± 5.1", ...},
              "n": {"0.0": 5, ...},
              "bmd": "8.49", "bmdl": "4.23", "trend_marker": "**"
            }, ...
          ],
          "Female": [...]
        }
      }
    """
    body = await request.json()
    bm2_id = body.get("bm2_id", "")
    compound_name = body.get("compound_name", "Test Compound")
    dose_unit = body.get("dose_unit", "mg/kg")

    upload = _bm2_uploads.get(bm2_id)
    if not upload:
        return JSONResponse(
            {"error": f"Unknown bm2_id: {bm2_id}"},
            status_code=404,
        )

    bm2_path = upload["temp_path"]
    if not os.path.exists(bm2_path):
        return JSONResponse(
            {"error": f"Uploaded file no longer exists: {upload['filename']}"},
            status_code=410,
        )

    try:
        # Run the full export-and-analyze pipeline in a thread pool
        # because it spawns Java subprocesses (blocking I/O)
        loop = asyncio.get_running_loop()
        table_data, _cat_lookup, bm2_json = await loop.run_in_executor(
            None, build_table_data_from_bm2, bm2_path,
        )

        # Cache the table_data and full JSON so /api/export-docx and
        # /api/preview can reuse them without re-running the Java export
        upload["table_data"] = table_data
        upload["bm2_json"] = bm2_json

        # Generate the NTP-style results narrative from the table data.
        # This produces paragraphs describing body weight and organ weight
        # findings that the user can edit in the UI before export.
        narrative = generate_results_narrative(
            table_data, compound_name, dose_unit,
        )
        upload["narrative"] = narrative

        # Serialize TableRow objects to JSON-friendly dicts
        tables_json = _serialize_table_rows(table_data)

        return JSONResponse({
            "bm2_id": bm2_id,
            "tables": tables_json,
            "narrative": narrative,  # list of auto-generated paragraph strings
        })

    except Exception as e:
        return JSONResponse(
            {"error": f"Processing failed: {e}"},
            status_code=500,
        )


# ---------------------------------------------------------------------------
# POST /api/upload-csv — upload gene-level BMD CSV files for genomics analysis
# ---------------------------------------------------------------------------

@app.post("/api/upload-csv")
async def api_upload_csv(files: list[UploadFile] = File(...)):
    """
    Accept one or more gene-level BMD CSV file uploads.

    Each CSV represents one organ × sex combination (e.g., "Liver Male").
    Required columns: gene_symbol, bmd.
    Optional columns: bmdl, bmdu, direction, fold_change, best_model,
                      full_name, gof_p.

    The uploaded CSV is saved to a temp directory and parsed via
    interpret.py's load_dose_response() to validate structure and
    normalize gene symbols.

    Returns a JSON list of metadata objects:
      [{id, filename, row_count, columns_found}, ...]
    """
    results = []

    for upload in files:
        csv_id = str(uuid.uuid4())

        # Save the uploaded file to a temp directory so we can parse it
        tmp_dir = tempfile.mkdtemp(prefix="csv_")
        safe_filename = os.path.basename(upload.filename or "upload.csv")
        tmp_path = os.path.join(tmp_dir, safe_filename)

        with open(tmp_path, "wb") as f:
            content = await upload.read()
            f.write(content)

        try:
            # Validate and parse the CSV using the existing dose-response loader.
            # This normalizes gene symbols and ensures bmd is numeric.
            df = load_dose_response(tmp_path)

            _csv_uploads[csv_id] = {
                "filename": safe_filename,
                "temp_path": tmp_path,
                "temp_dir": tmp_dir,
                "df": df,  # parsed DataFrame for fast reuse
            }

            results.append({
                "id": csv_id,
                "filename": safe_filename,
                "row_count": len(df),
                "columns_found": list(df.columns),
            })
        except Exception as e:
            # Clean up temp dir on parse failure
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return JSONResponse(
                {"error": f"CSV validation failed for {safe_filename}: {e}"},
                status_code=400,
            )

    return JSONResponse(results)


# ---------------------------------------------------------------------------
# POST /api/upload-zip — extract a .zip archive and register individual files
# ---------------------------------------------------------------------------
# Accepts a single .zip file, extracts its contents to temp directories,
# and registers each file with the appropriate in-memory store (_bm2_uploads
# or _csv_uploads) just as if they had been uploaded individually via
# /api/upload-bm2 or /api/upload-csv.
#
# Supported file types inside the zip:
#   .bm2  — registered as bm2 uploads (same as /api/upload-bm2)
#   .csv  — parsed via load_dose_response() and registered as csv uploads
#   .txt  — registered as generic text files (stored in uploadedFiles on client)
#   .xlsx — registered as generic spreadsheet files (stored in uploadedFiles)
#
# Files in subdirectories within the zip are extracted with their basename
# only (flattened).  __MACOSX/ entries and hidden files (starting with .)
# are silently skipped.

# Valid extensions that we extract from zip archives.  Anything else is
# silently skipped to avoid polluting the file pool with irrelevant files.
_ZIP_VALID_EXTENSIONS = {".bm2", ".csv", ".txt", ".xlsx"}


@app.post("/api/upload-zip")
async def api_upload_zip(request: Request, file: UploadFile = File(...)):
    """
    Accept a single .zip archive upload, extract its contents, and register
    each recognized file type as if it had been uploaded individually.

    .bm2 files are BMDExpress output — they go into _bm2_uploads for
    downstream processing via /api/process-bm2.

    .csv, .txt, .xlsx are raw dose-response experimental data (animal IDs
    across columns, dose concentrations in row 2, measured endpoints in
    subsequent rows — the format BMDExpress imports as doseResponseExperiment).
    These are stored as generic file pool entries.

    Returns a JSON object with:
      bm2_files:   [{id, filename}, ...]            — same shape as /api/upload-bm2
      other_files: [{id, filename, type}, ...]       — csv, txt, and xlsx raw data
      skipped:     [filename, ...]                   — files with unrecognized extensions
    """
    # Save the uploaded zip to a temp file so zipfile can read it
    zip_tmp_dir = tempfile.mkdtemp(prefix="zip_")
    zip_path = os.path.join(zip_tmp_dir, "upload.zip")

    with open(zip_path, "wb") as f:
        content = await file.read()
        f.write(content)

    # Validate that it's actually a zip file
    if not zipfile.is_zipfile(zip_path):
        shutil.rmtree(zip_tmp_dir, ignore_errors=True)
        return JSONResponse(
            {"error": "The uploaded file is not a valid .zip archive"},
            status_code=400,
        )

    # If the client provides a DTXSID, persist .bm2 files to the session's
    # files/ directory immediately so they survive page reloads.
    dtxsid = request.query_params.get("dtxsid", "")
    persist_dir = None
    if dtxsid:
        persist_dir = _session_dir(dtxsid) / "files"
        persist_dir.mkdir(exist_ok=True)

    bm2_results = []
    other_results = []  # csv, txt, xlsx — raw data files
    skipped = []

    with zipfile.ZipFile(zip_path, "r") as zf:
        for member in zf.namelist():
            # Skip directories, __MACOSX metadata, and hidden files
            if member.endswith("/"):
                continue
            basename = os.path.basename(member)
            if not basename or basename.startswith(".") or "__MACOSX" in member:
                continue

            # Check file extension
            _, ext = os.path.splitext(basename)
            ext = ext.lower()
            if ext not in _ZIP_VALID_EXTENSIONS:
                skipped.append(basename)
                continue

            safe_filename = os.path.basename(basename)

            if persist_dir:
                # Persist to session files/ directory so the file survives
                # page reloads and server restarts — applies to ALL file
                # types (.bm2, .csv, .txt, .xlsx), not just .bm2.
                file_path = str(persist_dir / safe_filename)
                with zf.open(member) as src, open(file_path, "wb") as dst:
                    shutil.copyfileobj(src, dst)
                tmp_dir = None  # no temp dir needed — file is in session dir
            else:
                # No DTXSID provided — extract to a temp directory (lost on restart)
                tmp_dir = tempfile.mkdtemp(prefix=f"{ext.lstrip('.')}_")
                file_path = os.path.join(tmp_dir, safe_filename)
                with zf.open(member) as src, open(file_path, "wb") as dst:
                    shutil.copyfileobj(src, dst)

            file_id = str(uuid.uuid4())

            if ext == ".bm2":
                # Register in _bm2_uploads — same as /api/upload-bm2.
                # .bm2 files are BMDExpress output (benchmark dose results),
                # which need the special _bm2_uploads store because
                # downstream endpoints (/api/process-bm2) reference them.
                _bm2_uploads[file_id] = {
                    "filename": safe_filename,
                    "temp_path": file_path,
                    "table_data": None,
                    "bm2_json": None,     # populated on first preview/process
                    "narrative": None,
                }
                # Fingerprint for cross-validation
                v_issues = []
                if dtxsid:
                    fp = _fingerprint_and_store(file_id, safe_filename, file_path, "bm2", dtxsid)
                    v_issues = _run_lightweight_validation(fp, dtxsid)
                bm2_results.append({
                    "id": file_id,
                    "filename": safe_filename,
                    "validation_issues": v_issues,
                })

            else:
                # .csv, .txt, and .xlsx — raw dose-response experimental
                # data (animal IDs × endpoints, tab- or comma-separated)
                # or spreadsheet data.  These are *input* data suitable for
                # importing into BMDExpress, NOT gene-level BMD output.
                # We store them as generic file pool entries; the client
                # tracks them in uploadedFiles and the user chooses which
                # file to use for each report section.
                file_type = ext.lstrip(".")  # "csv", "txt", or "xlsx"

                # Store internally so downstream endpoints can access
                # the extracted file by ID if needed (e.g., /api/preview).
                _data_uploads[file_id] = {
                    "filename": safe_filename,
                    "temp_path": file_path,
                    "type": file_type,
                }

                # Fingerprint for cross-validation
                v_issues = []
                if dtxsid:
                    fp = _fingerprint_and_store(file_id, safe_filename, file_path, file_type, dtxsid)
                    v_issues = _run_lightweight_validation(fp, dtxsid)

                # Only send safe metadata to the client (no server paths)
                other_results.append({
                    "id": file_id,
                    "filename": safe_filename,
                    "type": file_type,
                    "validation_issues": v_issues,
                })

    # Clean up the zip temp file (individual extracted files are in their
    # own temp dirs and persist until the server shuts down)
    shutil.rmtree(zip_tmp_dir, ignore_errors=True)

    # Collect all validation issues from individual file fingerprints
    # into a flat list for the response — the client shows these as
    # immediate feedback after zip extraction.
    all_validation_issues = []
    for r in bm2_results + other_results:
        all_validation_issues.extend(r.get("validation_issues", []))

    return JSONResponse({
        "bm2_files": bm2_results,
        "other_files": other_results,
        "skipped": skipped,
        "validation_issues": all_validation_issues,
    })


# ---------------------------------------------------------------------------
# POST /api/process-genomics — run gene set and gene BMD ranking
# ---------------------------------------------------------------------------

@app.post("/api/process-genomics")
async def api_process_genomics(request: Request):
    """
    Process an uploaded gene-level BMD CSV to produce the data for the
    Gene Set BMD Analysis and Gene BMD Analysis report sections.

    Takes a previously-uploaded CSV (via csv_id from /api/upload-csv)
    and runs:
      1. NIEHS quality filtering (fold change, goodness-of-fit, BMDU/BMDL)
      2. GO Biological Process gene set ranking by median BMD
      3. Individual gene ranking by BMD

    Input JSON:
      {
        "csv_id": "uuid-from-upload",
        "organ": "liver",
        "sex": "male",
        "compound_name": "PFHxSAm",
        "dose_unit": "mg/kg"
      }

    Returns JSON:
      {
        "gene_sets": [{rank, go_id, go_term, bmd_median, bmdl_median,
                       n_genes, genes, direction}, ...],
        "top_genes": [{rank, gene_symbol, full_name, bmd, bmdl, bmdu,
                       direction, fold_change}, ...],
        "total_responsive_genes": N,
        "organ": "liver",
        "sex": "male"
      }
    """
    body = await request.json()
    csv_id = body.get("csv_id", "")
    organ = body.get("organ", "")
    sex = body.get("sex", "")

    if not csv_id:
        return JSONResponse(
            {"error": "csv_id is required"},
            status_code=400,
        )

    csv_upload = _csv_uploads.get(csv_id)
    if not csv_upload:
        return JSONResponse(
            {"error": f"Unknown csv_id: {csv_id}"},
            status_code=404,
        )

    df = csv_upload.get("df")
    if df is None or df.empty:
        return JSONResponse(
            {"error": "No valid data in the uploaded CSV"},
            status_code=400,
        )

    try:
        # Run gene set and gene ranking in a thread pool because
        # ToxKBQuerier makes DB queries that could be slow for large datasets
        loop = asyncio.get_running_loop()

        def _run_genomics():
            # Open a read-only connection to the knowledge base for GO term
            # lookups.  The KB is the bmdx.duckdb file in the project root.
            db_path = str(Path(__file__).parent / "bmdx.duckdb")
            kb = ToxKBQuerier(db_path)
            try:
                gene_sets = rank_go_sets_by_bmd(df, kb, top_n=10)
                top_genes = rank_genes_by_bmd(df, top_n=10)

                # Fetch descriptions for the NIEHS dense 9pt blocks.
                # GO definitions from EBI QuickGO, gene summaries from
                # MyGene.info (human orthologs for full descriptions).
                go_ids = [gs["go_id"] for gs in gene_sets]
                go_descs = fetch_go_descriptions(go_ids, kb=kb)

                gene_syms = [g["gene_symbol"] for g in top_genes]
                gene_descs = fetch_gene_descriptions(gene_syms)

                return gene_sets, top_genes, go_descs, gene_descs
            finally:
                kb.close()

        gene_sets, top_genes, go_descs, gene_descs = await loop.run_in_executor(
            None, _run_genomics,
        )

        return JSONResponse({
            "gene_sets": gene_sets,
            "top_genes": top_genes,
            "go_descriptions": go_descs,
            "gene_descriptions": gene_descs,
            "total_responsive_genes": len(df),
            "organ": organ,
            "sex": sex,
        })

    except Exception as e:
        return JSONResponse(
            {"error": f"Genomics processing failed: {e}"},
            status_code=500,
        )


# ---------------------------------------------------------------------------
# GET /api/session/{dtxsid}/bmd-summary — auto-derive apical endpoint BMD summary
# ---------------------------------------------------------------------------

@app.get("/api/session/{dtxsid}/bmd-summary")
async def api_bmd_summary(dtxsid: str):
    """
    Auto-derive an Apical Endpoint BMD Summary from all approved .bm2 sections.

    Iterates all bm2_*.json files in the session directory, extracts every
    endpoint where BMD ≠ "ND", and computes LOEL (lowest dose with a
    significance marker) and NOEL (highest dose below LOEL without a marker).

    This data is already present in the approved table JSON — we just need
    to aggregate it across sections into one sorted summary.

    Returns JSON:
      {
        "endpoints": [
          {
            "endpoint": "Liver Relative",
            "bmd": "4.23",
            "bmdl": "2.10",
            "loel": "5.0",
            "noel": "1.5",
            "direction": "Increased",
            "sex": "Male"
          }, ...
        ],
        "sorted_by": "bmd_asc"
      }
    """
    d = SESSIONS_DIR / dtxsid
    if not d.exists():
        return JSONResponse(
            {"error": f"No session found for {dtxsid}"},
            status_code=404,
        )

    # Gather all approved bm2_*.json files
    endpoints = []
    for f in sorted(d.glob("bm2_*.json")):
        try:
            section = json.loads(f.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue

        tables_json = section.get("tables_json", {})

        # Each table is keyed by sex ("Male", "Female")
        for sex, rows in tables_json.items():
            for row in rows:
                bmd_str = row.get("bmd", "ND")
                bmdl_str = row.get("bmdl", "ND")

                # Skip endpoints without a valid BMD (not dose-responsive
                # by the NTP dual-significance criterion)
                if bmd_str == "ND":
                    continue

                label = row.get("label", "")
                # Skip the "n" row (sample sizes, not an endpoint)
                if label.lower() == "n":
                    continue

                # Derive LOEL and NOEL from significance markers in the
                # values_by_dose dict.  Each value string may end with
                # "*" or "**" if the dose group showed a significant
                # pairwise difference vs. control.
                values = row.get("values", {})
                doses = row.get("doses", [])

                # Sort doses numerically (they may be strings)
                try:
                    sorted_doses = sorted(
                        [float(d) for d in doses if float(d) > 0],
                    )
                except (ValueError, TypeError):
                    sorted_doses = []

                # Find LOEL: lowest treatment dose with a significance marker
                loel = None
                for dose in sorted_doses:
                    dose_key = _js_dose_key(dose)
                    val = values.get(dose_key, "")
                    if "*" in str(val):
                        loel = dose
                        break

                # Find NOEL: highest dose below LOEL without a marker.
                # If LOEL is the lowest dose, NOEL is None (below tested range).
                noel = None
                if loel is not None:
                    for dose in sorted_doses:
                        if dose >= loel:
                            break
                        dose_key = _js_dose_key(dose)
                        val = values.get(dose_key, "")
                        if "*" not in str(val):
                            noel = dose

                # Determine direction from the trend marker.
                # The trend_marker field has "*" or "**" for significant
                # Jonckheere trend.  To get direction, we check if the
                # highest-dose mean is greater or less than control mean.
                trend_marker = row.get("trend_marker", "")
                direction = ""
                if trend_marker:
                    # Try to parse control and highest-dose values to
                    # determine direction of change
                    try:
                        control_key = _js_dose_key(0.0)
                        control_val = values.get(control_key, "")
                        # Extract numeric part (before ±)
                        control_num = float(control_val.split("±")[0].replace("*", "").strip())
                        if sorted_doses:
                            high_key = _js_dose_key(sorted_doses[-1])
                            high_val = values.get(high_key, "")
                            high_num = float(high_val.split("±")[0].replace("*", "").strip())
                            direction = "Increased" if high_num > control_num else "Decreased"
                    except (ValueError, IndexError, AttributeError):
                        direction = ""

                endpoints.append({
                    "endpoint": label,
                    "bmd": bmd_str,
                    "bmdl": bmdl_str,
                    "loel": loel,
                    "noel": noel,
                    "direction": direction,
                    "sex": sex,
                })

    # Sort by BMD ascending (numeric sort; "ND" already filtered out)
    endpoints.sort(
        key=lambda e: float(e["bmd"]) if e["bmd"] != "ND" else 9999,
    )

    return JSONResponse({
        "endpoints": endpoints,
        "sorted_by": "bmd_asc",
    })


# ---------------------------------------------------------------------------
# GET /api/session/{dtxsid} — load a previously saved session
# ---------------------------------------------------------------------------

@app.get("/api/session/{dtxsid}")
async def api_session_load(dtxsid: str):
    """
    Load a previously-saved session for a given DTXSID.

    Looks up sessions/{dtxsid}/ on disk and returns all saved section data
    including the original sections (meta, identity, background, bm2) and
    the new NIEHS sections (methods, bmd_summary, genomics, summary).

    Returns JSON:
      {
        "exists": true/false,
        "meta": {...} or null,
        "identity": {...} or null,
        "background": {...} or null,
        "methods": {...} or null,
        "bm2_sections": { "organ-and-body-weights": {...}, ... },
        "bmd_summary": {...} or null,
        "genomics_sections": { "liver_male": {...}, ... },
        "summary": {...} or null
      }
    """
    d = SESSIONS_DIR / dtxsid
    if not d.exists():
        return JSONResponse({"exists": False})

    # Helper: read a JSON file if it exists, else return None
    def _read_json(name: str):
        p = d / name
        if p.exists():
            return json.loads(p.read_text(encoding="utf-8"))
        return None

    # Gather all bm2_*.json files into a dict keyed by slug.
    # Also re-register any persisted .bm2 files into _bm2_uploads so they
    # are available for preview, reprocessing, and export without re-upload.
    # The returned section data includes a server_file_id that the client
    # can use as a real upload ID (not a synthetic "file-restored-*" key).
    bm2_sections = {}
    files_dir = d / "files"
    for f in sorted(d.glob("bm2_*.json")):
        # Filename is e.g. "bm2_organ-and-body-weights.json"
        slug = f.stem.removeprefix("bm2_")
        section = json.loads(f.read_text(encoding="utf-8"))

        # Re-register the .bm2 file in _bm2_uploads if it exists on disk.
        # This makes the file fully functional (preview, process, export)
        # without needing the user to re-upload it.
        original_filename = section.get("filename", "")
        bm2_on_disk = files_dir / original_filename if original_filename else None
        if bm2_on_disk and bm2_on_disk.exists():
            file_id = str(uuid.uuid4())
            _bm2_uploads[file_id] = {
                "filename": original_filename,
                "temp_path": str(bm2_on_disk),
                "table_data": None,   # will be populated on first preview/process
                "bm2_json": None,     # will be populated on first preview/process
            }
            # Tell the client what server-side file_id to use
            section["server_file_id"] = file_id

        bm2_sections[slug] = section

    # Discover files in files/ that don't have a corresponding approved
    # section yet (uploaded but not yet processed/approved).  Register them
    # in the appropriate server-side store (_bm2_uploads or _data_uploads)
    # and return them as "pending_files" so the client can add them to the
    # file pool on restore.
    pending_files = []
    # Build a set of filenames that already have approved sections
    approved_filenames = {
        sec.get("filename", "") for sec in bm2_sections.values()
    }
    if files_dir.exists():
        for data_file in sorted(files_dir.iterdir()):
            if not data_file.is_file():
                continue
            ext = data_file.suffix.lower()

            if ext == ".bm2":
                if data_file.name in approved_filenames:
                    continue  # already handled above as an approved section
                # Unapproved .bm2 file — register in _bm2_uploads
                file_id = str(uuid.uuid4())
                _bm2_uploads[file_id] = {
                    "filename": data_file.name,
                    "temp_path": str(data_file),
                    "table_data": None,
                    "bm2_json": None,
                }
                pending_files.append({
                    "id": file_id,
                    "filename": data_file.name,
                    "type": "bm2",
                })
                # Fingerprint the .bm2 file so _pool_fingerprints is
                # populated for methods context extraction and validation.
                # This is cheap (~10ms from LMDB cache, no Java export).
                _fingerprint_and_store(
                    file_id, data_file.name, str(data_file), "bm2", dtxsid,
                )

            elif ext in (".csv", ".txt", ".xlsx"):
                # Raw data file (dose-response experimental data or
                # spreadsheet).  Register in _data_uploads so the
                # preview endpoint can serve it.
                file_type = ext.lstrip(".")
                file_id = str(uuid.uuid4())
                _data_uploads[file_id] = {
                    "filename": data_file.name,
                    "temp_path": str(data_file),
                    "type": file_type,
                }
                pending_files.append({
                    "id": file_id,
                    "filename": data_file.name,
                    "type": file_type,
                })
                # Fingerprint the data file so endpoints, sexes, and
                # dose groups are available for methods context extraction.
                _fingerprint_and_store(
                    file_id, data_file.name, str(data_file), file_type, dtxsid,
                )
            # Skip other file types (e.g., leftover pickle sidecars)

    # Gather all genomics_*.json files into a dict keyed by organ_sex
    # (e.g., "liver_male", "kidney_female")
    genomics_sections = {}
    for f in sorted(d.glob("genomics_*.json")):
        slug = f.stem.removeprefix("genomics_")
        genomics_sections[slug] = json.loads(f.read_text(encoding="utf-8"))

    # Load cached validation report and precedence decisions if they exist.
    # These are generated by POST /api/pool/validate/{dtxsid} and
    # POST /api/pool/resolve respectively, and persisted to disk so
    # they survive page reloads without re-running validation.
    validation_report = _read_json("validation_report.json")
    precedence = _read_json("precedence.json")

    return JSONResponse({
        "exists": True,
        "meta": _read_json("meta.json"),
        "identity": _read_json("identity.json"),
        "background": _read_json("background.json"),
        "methods": _read_json("methods.json"),
        "bm2_sections": bm2_sections,
        "pending_files": pending_files,
        "animal_report": _read_json("animal_report.json"),
        "bmd_summary": _read_json("bmd_summary.json"),
        "genomics_sections": genomics_sections,
        "summary": _read_json("summary.json"),
        "validation_report": validation_report,
        "precedence": precedence,
    })


# ---------------------------------------------------------------------------
# POST /api/session/approve — approve (save) a report section
# ---------------------------------------------------------------------------

@app.post("/api/session/approve")
async def api_session_approve(request: Request):
    """
    Approve a report section and persist it to disk.

    Input JSON:
      {
        "dtxsid": "DTXSID6020430",
        "identity": {ChemicalIdentity dict},   // saved on first approve
        "section_type": "background" | "bm2" | "methods" | "bmd_summary"
                        | "genomics" | "summary",
        "data": {
          // For background: paragraphs, references, model_used, notes,
          //                 original_paragraphs, original_references
          // For bm2: filename, section_title, table_caption, compound_name,
          //          dose_unit, narrative, tables_json, original_narrative
          // For methods/summary: paragraphs, original_paragraphs
          // For bmd_summary: endpoints
          // For genomics: organ, sex, gene_sets, top_genes
        }
      }

    Saves the section data with an approved_at timestamp.  Also saves/updates
    identity.json and meta.json on every approve call.

    For bm2 sections, the .bm2 file is copied from /tmp to
    sessions/{dtxsid}/files/ so it survives server restarts.

    Style learning: if the user edited the generated text before approving,
    the original and edited versions are compared to extract writing style
    rules.  Extraction runs asynchronously in a background thread so the
    approve response is immediate.
    """
    # Valid section types — the original two plus the new NIEHS sections
    VALID_SECTION_TYPES = {
        "background", "bm2", "methods", "bmd_summary", "genomics", "summary",
    }

    body = await request.json()
    dtxsid = body.get("dtxsid", "")
    identity = body.get("identity")
    section_type = body.get("section_type", "")
    data = body.get("data", {})

    if not dtxsid:
        return JSONResponse({"error": "dtxsid is required"}, status_code=400)
    if section_type not in VALID_SECTION_TYPES:
        return JSONResponse(
            {"error": f"section_type must be one of: {', '.join(sorted(VALID_SECTION_TYPES))}"},
            status_code=400,
        )

    # Ensure the session directory exists
    d = _session_dir(dtxsid)

    # Save / update identity.json on every approve so it stays current
    if identity:
        (d / "identity.json").write_text(
            json.dumps(identity, indent=2, default=str), encoding="utf-8",
        )

    # Save / update meta.json with chemical name and CASRN
    meta_path = d / "meta.json"
    if meta_path.exists():
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
    else:
        meta = {"dtxsid": dtxsid, "created_at": _now_iso()}
    meta["updated_at"] = _now_iso()
    if identity:
        meta["name"] = identity.get("name", "")
        meta["casrn"] = identity.get("casrn", "")
    meta_path.write_text(json.dumps(meta, indent=2), encoding="utf-8")

    # Stamp the data with the approval time
    data["approved_at"] = _now_iso()

    if section_type == "background":
        _save_section(dtxsid, "background", data)
    elif section_type == "bm2":
        # Build the section key from the original .bm2 filename
        filename = data.get("filename", "")
        slug = _bm2_slug(filename) if filename else ""
        if not slug:
            return JSONResponse(
                {"error": "bm2 data must include a filename"},
                status_code=400,
            )
        section_key = f"bm2_{slug}"
        _save_section(dtxsid, section_key, data)

        # Copy the .bm2 file from /tmp into the session's files/ directory
        # so it persists across server restarts.  (Processed data is cached
        # in LMDB via bm2_cache, so no sidecar files need copying.)
        bm2_id = data.get("bm2_id", "")
        upload = _bm2_uploads.get(bm2_id)
        if upload and os.path.exists(upload["temp_path"]):
            files_dir = d / "files"
            files_dir.mkdir(exist_ok=True)
            src = upload["temp_path"]
            dest = files_dir / filename
            if not dest.exists():
                shutil.copy2(src, dest)

    elif section_type == "methods":
        # Materials and Methods — single section, saved as methods.json
        _save_section(dtxsid, "methods", data)

    elif section_type == "bmd_summary":
        # Apical Endpoint BMD Summary — auto-derived, saved as bmd_summary.json
        _save_section(dtxsid, "bmd_summary", data)

    elif section_type == "genomics":
        # Genomics section — one file per organ × sex combination.
        # The section_key is built from organ and sex (e.g., "genomics_liver_male").
        organ = data.get("organ", "").lower().replace(" ", "_")
        sex = data.get("sex", "").lower().replace(" ", "_")
        if not organ or not sex:
            return JSONResponse(
                {"error": "genomics data must include organ and sex"},
                status_code=400,
            )
        section_key = f"genomics_{organ}_{sex}"
        _save_section(dtxsid, section_key, data)

        # Copy uploaded CSV into files/ for future reprocessing
        csv_id = data.get("csv_id", "")
        csv_upload = _csv_uploads.get(csv_id)
        if csv_upload and os.path.exists(csv_upload["temp_path"]):
            files_dir = d / "files"
            files_dir.mkdir(exist_ok=True)
            dest = files_dir / csv_upload["filename"]
            if not dest.exists():
                shutil.copy2(csv_upload["temp_path"], dest)

    elif section_type == "summary":
        # Summary section — saved as summary.json
        _save_section(dtxsid, "summary", data)

    # --- Style learning: detect edits and extract rules ---
    # The client sends the original LLM-generated text alongside the
    # (possibly edited) approved text.  If they differ, the user made
    # deliberate style changes — we extract rules from those changes
    # in a background thread to avoid blocking the approve response.
    user_edited = False

    if section_type == "background":
        original_paragraphs = data.get("original_paragraphs")
        edited_paragraphs = data.get("paragraphs")
        if original_paragraphs and edited_paragraphs:
            user_edited = original_paragraphs != edited_paragraphs
            if user_edited:
                original_text = "\n\n".join(original_paragraphs)
                edited_text = "\n\n".join(edited_paragraphs)

    elif section_type == "bm2":
        original_narrative = data.get("original_narrative", "")
        edited_narrative = data.get("narrative", "")
        if original_narrative and edited_narrative:
            user_edited = original_narrative != edited_narrative
            if user_edited:
                original_text = original_narrative
                edited_text = edited_narrative

    elif section_type in ("methods", "summary"):
        # Methods and summary sections use the same paragraph-based
        # style learning as background — compare original vs. edited paragraphs
        original_paragraphs = data.get("original_paragraphs")
        edited_paragraphs = data.get("paragraphs")
        if original_paragraphs and edited_paragraphs:
            user_edited = original_paragraphs != edited_paragraphs
            if user_edited:
                original_text = "\n\n".join(original_paragraphs)
                edited_text = "\n\n".join(edited_paragraphs)

    if user_edited:
        # Fire-and-forget: extract style rules in a background thread.
        # We don't await the result — the user gets their approve response
        # immediately, and the rules are saved asynchronously.
        loop = asyncio.get_running_loop()
        loop.run_in_executor(
            None, _extract_and_merge_style_rules, original_text, edited_text,
        )
        logger.info("Style learning triggered for %s/%s", dtxsid, section_type)

    # Read back the saved version number so the UI can display "v1", "v2", etc.
    # The version was stamped onto `data` by _save_section() before writing.
    version = data.get("version", 1)

    return JSONResponse({
        "ok": True,
        "section_type": section_type,
        "user_edited": user_edited,
        "version": version,
    })


# ---------------------------------------------------------------------------
# POST /api/session/unapprove — unapprove (delete) a report section
# ---------------------------------------------------------------------------

@app.post("/api/session/unapprove")
async def api_session_unapprove(request: Request):
    """
    Unapprove a report section by deleting its JSON file from disk.

    Input JSON:
      {
        "dtxsid": "DTXSID6020430",
        "section_type": "background" | "bm2" | "methods" | "bmd_summary"
                        | "genomics" | "summary",
        "bm2_slug": "organ-and-body-weights",  // required for bm2 only
        "organ": "liver",                       // required for genomics only
        "sex": "male"                           // required for genomics only
      }

    The .bm2 file and CSV in files/ are NOT deleted — they're still useful
    for reprocessing without re-uploading.
    """
    body = await request.json()
    dtxsid = body.get("dtxsid", "")
    section_type = body.get("section_type", "")
    bm2_slug = body.get("bm2_slug", "")

    if not dtxsid:
        return JSONResponse({"error": "dtxsid is required"}, status_code=400)

    if section_type == "background":
        _delete_section(dtxsid, "background")
    elif section_type == "bm2":
        if not bm2_slug:
            return JSONResponse(
                {"error": "bm2_slug is required for bm2 sections"},
                status_code=400,
            )
        _delete_section(dtxsid, f"bm2_{bm2_slug}")
    elif section_type == "methods":
        _delete_section(dtxsid, "methods")
    elif section_type == "bmd_summary":
        _delete_section(dtxsid, "bmd_summary")
    elif section_type == "genomics":
        organ = body.get("organ", "").lower().replace(" ", "_")
        sex = body.get("sex", "").lower().replace(" ", "_")
        if not organ or not sex:
            return JSONResponse(
                {"error": "organ and sex are required for genomics sections"},
                status_code=400,
            )
        _delete_section(dtxsid, f"genomics_{organ}_{sex}")
    elif section_type == "summary":
        _delete_section(dtxsid, "summary")
    else:
        return JSONResponse(
            {"error": f"Unknown section_type: {section_type}"},
            status_code=400,
        )

    return JSONResponse({"ok": True})


# ---------------------------------------------------------------------------
# GET /api/session/{dtxsid}/history/{section_key} — version history for a section
# ---------------------------------------------------------------------------

@app.get("/api/session/{dtxsid}/history/{section_key}")
async def api_session_history(dtxsid: str, section_key: str, version: int = 0):
    """
    Return version history for an approved section.

    Lists all past versions (from the history/ subdirectory) plus the current
    version.  Each entry includes the version number, approved_at timestamp,
    and whether it's the current version.

    If the `version` query parameter is provided (e.g. ?version=1), returns
    the full JSON content of that specific version instead of the list.

    URL pattern:
        GET /api/session/{dtxsid}/history/{section_key}         → version list
        GET /api/session/{dtxsid}/history/{section_key}?version=1 → full content

    Response (list mode):
    {
      "section_key": "background",
      "current_version": 3,
      "versions": [
        {"version": 1, "approved_at": "2026-03-02T19:23:59...", "is_current": false},
        {"version": 2, "approved_at": "2026-03-02T19:45:12...", "is_current": false},
        {"version": 3, "approved_at": "2026-03-02T20:01:33...", "is_current": true}
      ]
    }
    """
    d = SESSIONS_DIR / dtxsid
    current_path = d / f"{section_key}.json"
    history_dir = d / "history" / section_key

    # The current file must exist — otherwise there's no history to show
    if not current_path.exists():
        return JSONResponse(
            {"error": f"No approved '{section_key}' section found"},
            status_code=404,
        )

    current_data = json.loads(current_path.read_text(encoding="utf-8"))
    current_version = current_data.get("version", 1)

    # --- Collect archived versions from history/ ---
    # Each file is named {safe_timestamp}.json where colons in the ISO
    # timestamp have been replaced with hyphens.
    archived_versions = []
    if history_dir.exists():
        for f in sorted(history_dir.glob("*.json")):
            try:
                v = json.loads(f.read_text(encoding="utf-8"))
                archived_versions.append(v)
            except (json.JSONDecodeError, OSError):
                # Skip corrupt history files
                continue

    # --- If a specific version was requested, return its full content ---
    if version > 0:
        # Check archived versions first
        for v in archived_versions:
            if v.get("version") == version:
                return JSONResponse(v)
        # Check if it's the current version
        if current_version == version:
            return JSONResponse(current_data)
        return JSONResponse(
            {"error": f"Version {version} not found"}, status_code=404,
        )

    # --- List mode: return metadata for all versions ---
    versions = []
    for v in archived_versions:
        versions.append({
            "version": v.get("version", 0),
            "approved_at": v.get("approved_at", ""),
            "is_current": False,
        })
    # Add the current version as the last entry
    versions.append({
        "version": current_version,
        "approved_at": current_data.get("approved_at", ""),
        "is_current": True,
    })

    return JSONResponse({
        "section_key": section_key,
        "current_version": current_version,
        "versions": versions,
    })


# ---------------------------------------------------------------------------
# POST /api/session/{dtxsid}/restore — restore a past version as current
# ---------------------------------------------------------------------------

@app.post("/api/session/{dtxsid}/restore")
async def api_session_restore(dtxsid: str, request: Request):
    """
    Restore a past version of a section by re-saving it as the new current.

    This is non-destructive: restoring v1 does NOT delete v2 or v3.  Instead
    it creates a new version (v4) whose content matches v1.  The old current
    is archived to history/ first, as usual.

    Input JSON:
      {"section_key": "background", "version": 1}

    The restored data gets a fresh approved_at timestamp and an incremented
    version number.
    """
    body = await request.json()
    section_key = body.get("section_key", "")
    target_version = body.get("version", 0)

    if not section_key or not target_version:
        return JSONResponse(
            {"error": "section_key and version are required"}, status_code=400,
        )

    d = SESSIONS_DIR / dtxsid
    current_path = d / f"{section_key}.json"
    history_dir = d / "history" / section_key

    if not current_path.exists():
        return JSONResponse(
            {"error": f"No approved '{section_key}' section found"},
            status_code=404,
        )

    # Find the requested version in history files
    target_data = None
    if history_dir.exists():
        for f in sorted(history_dir.glob("*.json")):
            try:
                v = json.loads(f.read_text(encoding="utf-8"))
                if v.get("version") == target_version:
                    target_data = v
                    break
            except (json.JSONDecodeError, OSError):
                continue

    # Also check if they're restoring the current version (no-op but valid)
    if target_data is None:
        current_data = json.loads(current_path.read_text(encoding="utf-8"))
        if current_data.get("version") == target_version:
            return JSONResponse({
                "ok": True,
                "message": "Already the current version",
                "version": target_version,
            })
        return JSONResponse(
            {"error": f"Version {target_version} not found"}, status_code=404,
        )

    # Stamp the restored data with a fresh approval timestamp.
    # _save_section() will archive the old current and assign a new version number.
    target_data["approved_at"] = _now_iso()
    # Remove the old version number — _save_section() will compute the new one
    target_data.pop("version", None)

    _save_section(dtxsid, section_key, target_data)

    return JSONResponse({
        "ok": True,
        "version": target_data.get("version", 1),
        "section_key": section_key,
    })


# ---------------------------------------------------------------------------
# POST /api/generate-methods — generate Materials and Methods section
# ---------------------------------------------------------------------------

@app.post("/api/generate-methods")
async def api_generate_methods(request: Request):
    """
    Generate a structured Materials and Methods section for the 5dToxReport.

    Replicates the NIEHS Report 10 M&M structure with 6 major sections,
    10+ conditional subsections, and Table 1 (genomics sample counts).

    The approach is hybrid data + LLM:
      1. Extract study metadata from fingerprints, animal report, and .bm2
         analysisInfo.notes (doses, sample sizes, BMDExpress params, etc.)
      2. Build a structured LLM prompt that asks for prose per subsection key
      3. Parse the LLM's JSON response into a MethodsReport with heading hierarchy
      4. Subsections are CONDITIONAL — only included when the relevant data domain
         exists in the file pool (e.g., Transcriptomics only if gene_expression)

    Input JSON:
      {
        "identity": {ChemicalIdentity dict},
        "study_params": {"vehicle": "...", "route": "...", "duration_days": 5, "species": "..."},
        "animal_report": {optional animal_report.json dict}
      }
      All other study metadata (dose groups, sample sizes, endpoints, BMD params)
      is extracted automatically from the file pool fingerprints and .bm2 caches.

    Returns JSON:
      {
        "sections": [{heading, level, key, paragraphs, table}, ...],
        "context": {MethodsContext fields},
        "section_key": "methods",
        "model_used": "claude-sonnet-4-6"
      }
    """
    from methods_report import (
        MethodsReport,
        MethodsSection,
        build_methods_prompt,
        build_subsection_skeleton,
        build_table1_data,
        extract_methods_context,
    )

    body = await request.json()
    identity = body.get("identity", {})
    study_params = body.get("study_params", {})
    animal_report_data = body.get("animal_report")
    dtxsid = identity.get("dtxsid", "")

    # --- Backwards compatibility: accept old flat fields too ---
    # The old frontend sent vehicle, route, etc. as top-level fields.
    # Merge them into study_params if present.
    for key in ("vehicle", "route", "duration_days", "species"):
        if key in body and key not in study_params:
            study_params[key] = body[key]

    # --- Collect fingerprints from the server's pool cache ---
    fingerprints = {}
    if dtxsid and dtxsid in _pool_fingerprints:
        for fid, fp in _pool_fingerprints[dtxsid].items():
            # Convert FileFingerprint to dict for extract_methods_context
            if hasattr(fp, "__dataclass_fields__"):
                fingerprints[fid] = {
                    k: getattr(fp, k) for k in fp.__dataclass_fields__
                }
            else:
                fingerprints[fid] = fp

    # --- Collect .bm2 JSON caches for BMDExpress metadata extraction ---
    # Each .bm2 file's analysisInfo.notes contains the BMDExpress version,
    # BMDS version, models fit, BMR type, etc.
    bm2_jsons = {}
    if dtxsid:
        session_files_dir = SESSIONS_DIR / dtxsid / "files"
        if session_files_dir.exists():
            for bm2_path in session_files_dir.glob("*.bm2"):
                try:
                    cached = bm2_cache.get_json(str(bm2_path))
                    if cached:
                        bm2_jsons[bm2_path.stem] = cached
                except Exception:
                    pass

    # --- Load animal report from session if not provided in request ---
    if not animal_report_data and dtxsid:
        ar_path = SESSIONS_DIR / dtxsid / "animal_report.json"
        if ar_path.exists():
            try:
                animal_report_data = json.loads(ar_path.read_text())
            except Exception:
                pass

    # --- Extract structured context from all data sources ---
    ctx = extract_methods_context(
        identity=identity,
        fingerprints=fingerprints,
        animal_report=animal_report_data,
        study_params=study_params,
        bm2_jsons=bm2_jsons,
    )

    # --- Build the structured LLM prompt ---
    system, prompt = build_methods_prompt(ctx)

    try:
        # Call Claude and parse the JSON response — keyed by subsection key,
        # e.g. {"study_design": "paragraph text", "dose_selection": "..."}
        subsection_texts = await _llm_generate_json_async(
            "methods-generator", prompt, system,
        )

        # --- Assemble into MethodsReport ---
        skeleton = build_subsection_skeleton(ctx)
        sections = []
        for key, heading, level in skeleton:
            text = subsection_texts.get(key, "")
            if not text:
                continue
            # Split multi-paragraph strings on double newlines
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            sections.append(MethodsSection(
                heading=heading,
                level=level,
                key=key,
                paragraphs=paragraphs,
            ))

        # --- Add Table 1 data to the report context ---
        table1 = build_table1_data(ctx)

        report = MethodsReport(sections=sections, context=ctx)
        report_dict = report.to_dict()

        # Include Table 1 data separately for the frontend to render
        if table1:
            report_dict["table1"] = table1

        report_dict["section_key"] = "methods"
        report_dict["model_used"] = "claude-sonnet-4-6"

        return JSONResponse(report_dict)

    except json.JSONDecodeError as e:
        # If the LLM didn't return valid JSON, try to salvage as flat paragraphs
        logger.warning("Methods LLM response was not valid JSON: %s", e)
        # Fall back: treat the entire response as a single paragraph per line
        paragraphs = [p.strip() for p in response.split("\n\n") if p.strip()]
        sections = [MethodsSection(
            heading="Materials and Methods",
            level=3,
            key="fallback",
            paragraphs=paragraphs,
        )]
        report = MethodsReport(sections=sections, context=ctx)
        report_dict = report.to_dict()
        report_dict["section_key"] = "methods"
        report_dict["model_used"] = "claude-sonnet-4-6"
        report_dict["warning"] = "LLM response was not structured JSON; content placed in single section"
        return JSONResponse(report_dict)

    except Exception as e:
        logger.exception("Methods generation failed")
        return JSONResponse(
            {"error": f"Methods generation failed: {e}"},
            status_code=500,
        )


# ---------------------------------------------------------------------------
# GET /api/methods-context/{dtxsid} — preview extracted M&M context data
# ---------------------------------------------------------------------------

@app.get("/api/methods-context/{dtxsid}")
async def api_methods_context(dtxsid: str):
    """
    Return the extracted MethodsContext for a DTXSID without running the LLM.

    This is a read-only inspection endpoint that lets the user verify what
    data the system has extracted from the file pool (fingerprints, animal
    report, .bm2 analysisInfo) before generating M&M prose.  The response
    is a JSON object with all MethodsContext fields plus the subsection
    skeleton (which headings will be included) and Table 1 data.

    Used by the "Preview Data" button in the Methods section UI.
    """
    from methods_report import (
        extract_methods_context,
        build_subsection_skeleton,
        build_table1_data,
    )

    # --- Collect fingerprints ---
    fingerprints = {}
    if dtxsid in _pool_fingerprints:
        for fid, fp in _pool_fingerprints[dtxsid].items():
            if hasattr(fp, "__dataclass_fields__"):
                fingerprints[fid] = {
                    k: getattr(fp, k) for k in fp.__dataclass_fields__
                }
            else:
                fingerprints[fid] = fp

    # --- Collect .bm2 JSON caches ---
    bm2_jsons = {}
    session_files_dir = SESSIONS_DIR / dtxsid / "files"
    if session_files_dir.exists():
        for bm2_path in session_files_dir.glob("*.bm2"):
            try:
                cached = bm2_cache.get_json(str(bm2_path))
                if cached:
                    bm2_jsons[bm2_path.stem] = cached
            except Exception:
                pass

    # --- Load animal report ---
    animal_report_data = None
    ar_path = SESSIONS_DIR / dtxsid / "animal_report.json"
    if ar_path.exists():
        try:
            animal_report_data = json.loads(ar_path.read_text())
        except Exception:
            pass

    # --- Load identity ---
    identity = {}
    id_path = SESSIONS_DIR / dtxsid / "identity.json"
    if id_path.exists():
        try:
            identity = json.loads(id_path.read_text())
        except Exception:
            pass

    # --- Extract context ---
    ctx = extract_methods_context(
        identity=identity,
        fingerprints=fingerprints,
        animal_report=animal_report_data,
        bm2_jsons=bm2_jsons,
    )

    # --- Build response ---
    result = ctx.to_dict()

    # Add the subsection skeleton so the user can see which headings
    # will be generated (and which conditional ones are active/skipped)
    skeleton = build_subsection_skeleton(ctx)
    result["_subsection_skeleton"] = [
        {"key": k, "heading": h, "level": lvl}
        for k, h, lvl in skeleton
    ]

    # Add Table 1 data
    table1 = build_table1_data(ctx)
    if table1:
        result["_table1"] = table1

    # Add fingerprint summary so user can see what data sources fed the context
    fp_summary = {}
    for fid, fp in fingerprints.items():
        _get = fp.get if isinstance(fp, dict) else lambda k, d=None: getattr(fp, k, d)
        domain = _get("domain")
        if domain:
            if domain not in fp_summary:
                fp_summary[domain] = []
            fp_summary[domain].append({
                "file_id": fid,
                "filename": _get("filename", fid),
                "tier": _get("tier"),
                "sexes": _get("sexes", []),
                "endpoint_count": len(_get("endpoint_names", [])),
                "dose_groups": _get("dose_groups", []),
            })
    result["_fingerprint_summary"] = fp_summary

    return JSONResponse(result)


# ---------------------------------------------------------------------------
# POST /api/generate-summary — generate report Summary section
# ---------------------------------------------------------------------------

@app.post("/api/generate-summary")
async def api_generate_summary(request: Request):
    """
    Generate a Summary section that synthesizes all approved report sections.

    Reads all approved sections from the session, builds a context block
    describing the key findings, and sends it to Claude to produce a
    NIEHS-style summary.

    The summary references:
      - Most sensitive apical endpoint (from BMD summary)
      - Most sensitive gene set (from genomics)
      - Most sensitive gene (from genomics)
      - Comparison across hierarchy levels (gene < gene set < apical)

    Input JSON:
      {
        "dtxsid": "DTXSID...",
        "identity": {ChemicalIdentity dict}
      }

    Returns JSON:
      {
        "paragraphs": ["Summary paragraph 1...", "Summary paragraph 2...", ...],
        "section_key": "summary",
        "model_used": "claude-..."
      }
    """
    body = await request.json()
    dtxsid = body.get("dtxsid", "")
    identity = body.get("identity", {})

    if not dtxsid:
        return JSONResponse(
            {"error": "dtxsid is required"},
            status_code=400,
        )

    compound_name = identity.get("name", "the test chemical")

    # Gather context from all approved sections in the session
    d = SESSIONS_DIR / dtxsid
    context_parts = []

    # Background — extract a brief summary
    bg_path = d / "background.json"
    if bg_path.exists():
        try:
            bg = json.loads(bg_path.read_text(encoding="utf-8"))
            bg_paras = bg.get("paragraphs", [])
            if bg_paras:
                context_parts.append(
                    "=== BACKGROUND (first paragraph) ===\n"
                    + bg_paras[0][:500]
                )
        except (json.JSONDecodeError, OSError):
            pass

    # Apical endpoint findings — summarize from bm2 sections
    for f in sorted(d.glob("bm2_*.json")):
        try:
            section = json.loads(f.read_text(encoding="utf-8"))
            narrative = section.get("narrative", "")
            if isinstance(narrative, list):
                narrative = " ".join(narrative)
            if narrative:
                section_name = f.stem.removeprefix("bm2_").replace("-", " ").title()
                context_parts.append(
                    f"=== APICAL RESULTS: {section_name} ===\n"
                    + narrative[:800]
                )
        except (json.JSONDecodeError, OSError):
            continue

    # BMD summary — if available
    bmd_path = d / "bmd_summary.json"
    if bmd_path.exists():
        try:
            bmd_data = json.loads(bmd_path.read_text(encoding="utf-8"))
            eps = bmd_data.get("endpoints", [])
            if eps:
                lines = ["=== BMD SUMMARY (sorted by BMD) ==="]
                for ep in eps[:10]:
                    lines.append(
                        f"  {ep.get('endpoint', '')}: BMD={ep.get('bmd', 'ND')}, "
                        f"BMDL={ep.get('bmdl', 'ND')}, {ep.get('sex', '')}, "
                        f"{ep.get('direction', '')}"
                    )
                context_parts.append("\n".join(lines))
        except (json.JSONDecodeError, OSError):
            pass

    # Genomics findings — summarize from genomics_*.json files
    for f in sorted(d.glob("genomics_*.json")):
        try:
            genomics = json.loads(f.read_text(encoding="utf-8"))
            organ = genomics.get("organ", "")
            sex = genomics.get("sex", "")
            gene_sets = genomics.get("gene_sets", [])
            top_genes = genomics.get("top_genes", [])

            lines = [f"=== GENOMICS: {organ.title()} {sex.title()} ==="]
            if gene_sets:
                lines.append("Top gene sets by BMD:")
                for gs in gene_sets[:5]:
                    lines.append(
                        f"  {gs.get('go_term', '')}: median BMD={gs.get('bmd_median', '')}, "
                        f"{gs.get('n_genes', 0)} genes, {gs.get('direction', '')}"
                    )
            if top_genes:
                lines.append("Top genes by BMD:")
                for g in top_genes[:5]:
                    lines.append(
                        f"  {g.get('gene_symbol', '')}: BMD={g.get('bmd', '')}, "
                        f"{g.get('direction', '')}"
                    )
            context_parts.append("\n".join(lines))
        except (json.JSONDecodeError, OSError):
            continue

    if not context_parts:
        return JSONResponse(
            {"error": "No approved sections found to summarize"},
            status_code=400,
        )

    context_block = "\n\n".join(context_parts)

    prompt = f"""Based on the following approved report sections for {compound_name},
generate a Summary section in the style of an NIEHS/NTP 5-day study technical report.

{context_block}

---

Generate 3-4 summary paragraphs covering:
1. Overview — briefly restate the study design and the chemical tested
2. Key Apical Findings — which endpoints were most sensitive (lowest BMD), in which sex, and in what direction
3. Key Genomic Findings (if available) — which gene sets and genes were most sensitive, what biological processes they represent
4. Concordance — compare sensitivity across biological levels (gene < gene set < apical endpoint). Note whether transcriptomic changes occurred at lower doses than apical effects (as expected).

Return ONLY a JSON array of paragraph strings: ["paragraph1", "paragraph2", ...]"""

    system = (
        "You are a toxicology report writer specializing in NTP/NIEHS-style "
        "technical reports. Synthesize findings across biological levels "
        "(molecular, pathway, organism) into a coherent summary. Return ONLY "
        "valid JSON with no markdown formatting."
    )

    try:
        paragraphs = await _llm_generate_json_async(
            "summary-generator", prompt, system,
            max_tokens=4096,
        )
        if not isinstance(paragraphs, list):
            paragraphs = [str(paragraphs)]

        return JSONResponse({
            "paragraphs": paragraphs,
            "section_key": "summary",
            "model_used": "claude-sonnet-4-6",
        })

    except Exception as e:
        return JSONResponse(
            {"error": f"Summary generation failed: {e}"},
            status_code=500,
        )


# ---------------------------------------------------------------------------
# POST /api/generate-genomics-narrative — LLM-generate narrative for genomics
# ---------------------------------------------------------------------------
#
# Generates 1–2 paragraphs each for:
#   - Gene Set BMD Analysis (which biological processes were most sensitive)
#   - Gene BMD Analysis (which individual genes were most sensitive)
#
# These narratives appear above the data tables in the NIEHS report.
# The endpoint is called separately from /api/process-genomics because
# the LLM call is slower than the table computation, and the user may
# want to review the tables before generating narrative.
# ---------------------------------------------------------------------------

@app.post("/api/generate-genomics-narrative")
async def api_generate_genomics_narrative(request: Request):
    """
    Generate narrative paragraphs for the genomics Results section.

    Takes gene set and gene ranking data (from /api/process-genomics)
    and produces LLM-generated narrative for each subsection.

    Input JSON:
      {
        "identity": {ChemicalIdentity fields},
        "organ": "liver",
        "sex": "male",
        "gene_sets": [{go_id, go_term, bmd_median, n_genes, direction}, ...],
        "top_genes": [{gene_symbol, bmd, bmdl, fold_change, direction}, ...],
        "total_responsive_genes": 150,
        "dose_unit": "mg/kg"
      }

    Returns JSON:
      {
        "gene_set_narrative": ["paragraph1", "paragraph2"],
        "gene_narrative": ["paragraph1", "paragraph2"],
        "model_used": "claude-sonnet-4-6"
      }
    """
    body = await request.json()
    identity = body.get("identity", {})
    organ = body.get("organ", "")
    sex = body.get("sex", "")
    gene_sets = body.get("gene_sets", [])
    top_genes = body.get("top_genes", [])
    total_responsive = body.get("total_responsive_genes", 0)
    dose_unit = body.get("dose_unit", "mg/kg")

    compound = identity.get("name", "the test article")

    # --- Build gene set table as text for the prompt ---
    gs_lines = []
    for gs in gene_sets[:10]:
        gs_lines.append(
            f"  {gs.get('go_term', '')} (GO:{gs.get('go_id', '')}): "
            f"median BMD = {gs.get('bmd_median', 'N/A')} {dose_unit}, "
            f"{gs.get('n_genes', 0)} genes, direction = {gs.get('direction', 'N/A')}"
        )
    gs_table = "\n".join(gs_lines) if gs_lines else "(no gene sets)"

    # --- Build gene table as text for the prompt ---
    gene_lines = []
    for g in top_genes[:10]:
        gene_lines.append(
            f"  {g.get('gene_symbol', '')}: "
            f"BMD = {g.get('bmd', 'N/A')} {dose_unit}, "
            f"BMDL = {g.get('bmdl', 'N/A')} {dose_unit}, "
            f"fold change = {g.get('fold_change', 'N/A')}, "
            f"direction = {g.get('direction', 'N/A')}"
        )
    gene_table = "\n".join(gene_lines) if gene_lines else "(no genes)"

    # --- Load style rules for consistent voice ---
    style_rules = ""
    try:
        profile = _load_style_profile()
        rules = profile.get("rules", [])
        if rules:
            style_rules = "\n\nApply these writing style preferences:\n" + "\n".join(
                f"- {r['rule']}" for r in rules[:10]
            )
    except Exception:
        pass  # Style learning is optional

    prompt = f"""Generate narrative paragraphs for the genomics Results section of an
NIEHS/NTP 5-day study technical report on {compound}.

The study examined gene expression in the {organ} of {sex} Sprague Dawley rats.
A total of {total_responsive} genes had significant dose-responsive changes.

=== GENE SET BENCHMARK DOSE ANALYSIS ===
Top gene sets ranked by median BMD (most sensitive first):
{gs_table}

=== GENE BENCHMARK DOSE ANALYSIS ===
Top individual genes ranked by BMD (most sensitive first):
{gene_table}

Return a JSON object with two keys:
1. "gene_set_narrative": An array of 1–2 paragraphs summarizing the gene set BMD analysis.
   Note which biological processes were perturbed at the lowest doses, the predominant
   direction of perturbation, and the number of responsive gene sets.

2. "gene_narrative": An array of 1–2 paragraphs summarizing the individual gene BMD analysis.
   Note which genes were most sensitive, the direction and magnitude of their response,
   and any notable patterns in the top genes.

Use the passive voice and formal scientific register matching NIEHS report style.
Do NOT include table data in the narrative — the tables are presented separately.
{style_rules}

Return ONLY valid JSON, no markdown formatting."""

    system = (
        "You are a toxicology report writer specializing in NTP/NIEHS-style "
        "technical reports. Write concise, data-driven narrative for the genomics "
        "Results section. Return ONLY valid JSON with no markdown formatting."
    )

    try:
        result = await _llm_generate_json_async(
            "genomics-narrative-generator", prompt, system,
            max_tokens=4096,
        )

        # Normalize: ensure both keys are arrays of strings
        gs_narr = result.get("gene_set_narrative", [])
        gene_narr = result.get("gene_narrative", [])
        if isinstance(gs_narr, str):
            gs_narr = [gs_narr]
        if isinstance(gene_narr, str):
            gene_narr = [gene_narr]

        return JSONResponse({
            "gene_set_narrative": gs_narr,
            "gene_narrative": gene_narr,
            "model_used": "claude-sonnet-4-6",
        })

    except Exception as e:
        return JSONResponse(
            {"error": f"Genomics narrative generation failed: {e}"},
            status_code=500,
        )


# ---------------------------------------------------------------------------
# GET /api/style-profile — retrieve the global style profile
# ---------------------------------------------------------------------------

@app.get("/api/style-profile")
async def api_style_profile():
    """
    Return the global style profile (learned writing preferences).

    Returns the full profile JSON including version, updated_at, and rules
    array.  If no profile exists yet, returns an empty structure with zero
    rules so the client can always expect the same shape.
    """
    profile = _load_style_profile()
    return JSONResponse(profile)


# ---------------------------------------------------------------------------
# DELETE /api/style-profile/{idx} — delete a specific style rule by index
# ---------------------------------------------------------------------------

@app.delete("/api/style-profile/{idx}")
async def api_delete_style_rule(idx: int):
    """
    Delete a style rule at the given index from the global profile.

    The index is 0-based and corresponds to the rule's position in the
    rules array.  After deletion, the profile is rewritten to disk.

    Returns the updated profile so the client can re-render immediately.
    """
    profile = _load_style_profile()
    rules = profile.get("rules", [])

    if idx < 0 or idx >= len(rules):
        return JSONResponse(
            {"error": f"Rule index {idx} out of range (0..{len(rules) - 1})"},
            status_code=404,
        )

    removed = rules.pop(idx)
    _save_style_profile(profile)
    logger.info("Deleted style rule #%d: %s", idx, removed.get("rule", ""))

    return JSONResponse(profile)


# ---------------------------------------------------------------------------
# GET /api/preview/{file_id} — preview uploaded file contents
# ---------------------------------------------------------------------------
# Returns a JSON payload shaped for the browser's file-preview modal.
# The shape varies by file type so the front end can switch on `type`:
#   "bm2_json"  — processed .bm2: includes tables_json + narrative
#   "bm2_raw"   — unprocessed .bm2: just a "not yet processed" message
#   "table"     — .csv or .txt: first 50 rows parsed into headers + rows
#   "info"      — .xlsx or unknown: filename and size only

@app.get("/api/preview/{file_id}")
async def api_preview_file(file_id: str):
    """
    Return preview-ready content for an uploaded file.

    Lookup order: check _bm2_uploads first (apical endpoint .bm2 files),
    then _data_uploads (raw dose-response .csv/.txt/.xlsx from zip archives).
    Returns 404 if the file_id is not found in either store.

    Response shape depends on file type — see module-level comment above.
    """
    # --- Check .bm2 uploads first ---
    bm2_entry = _bm2_uploads.get(file_id)
    if bm2_entry:
        filename = bm2_entry["filename"]
        bm2_json = bm2_entry.get("bm2_json")

        if bm2_json is None:
            # In-memory cache is empty.  Try the LMDB B-tree cache first —
            # the bm2_cache module stores deserialized BMDProject dicts via
            # LMDB + orjson.  After first access, reads are memory-mapped
            # (OS page cache) + orjson.loads (~10ms for 5 MB).  This is the
            # fast path for previewing a file after page reload.
            bm2_path = bm2_entry["temp_path"]
            if not os.path.exists(bm2_path):
                return JSONResponse(
                    {"error": f"Uploaded file no longer exists: {filename}"},
                    status_code=410,
                )

            # Fast path — LMDB B-tree lookup.  Only loads the BMDProject
            # JSON (skips category analysis, stats, and narrative).
            loop = asyncio.get_running_loop()
            bm2_json = await loop.run_in_executor(
                None, bm2_cache.get_json, bm2_path,
            )
            if bm2_json is not None:
                bm2_entry["bm2_json"] = bm2_json

            if bm2_json is None:
                # Not in LMDB — run the full export-and-analyze pipeline.
                # This launches Java (slow) but also populates table_data
                # and narrative, so later process-bm2 calls are instant.
                # The pipeline stores the result in LMDB automatically.
                try:
                    table_data, _cat_lookup, bm2_json = await loop.run_in_executor(
                        None, build_table_data_from_bm2, bm2_path,
                    )
                    bm2_entry["table_data"] = table_data
                    bm2_entry["bm2_json"] = bm2_json

                    # Generate a default narrative (will be overwritten with
                    # real compound name / dose unit when the user processes
                    # via the section builder)
                    narrative = generate_results_narrative(
                        table_data, "Test Compound", "mg/kg",
                    )
                    bm2_entry["narrative"] = narrative
                except Exception as e:
                    return JSONResponse(
                        {"error": f"Processing failed: {e}"},
                        status_code=500,
                    )

        # Return the full BMDProject JSON — contains all 7 top-level lists:
        #   doseResponseExperiments, bMDResult, categoryAnalysisResults,
        #   oneWayANOVAResults, williamsTrendResults, curveFitPrefilterResults,
        #   oriogenResults.
        # The JSON tree viewer in the modal lets users expand/collapse any
        # section, so they can inspect BMD model fits, category analysis
        # results, etc. — not just the dose-response tables.
        # Use orjson directly instead of JSONResponse (stdlib json.dumps).
        # For a 5-10 MB BMDProject dict, orjson.dumps is ~10x faster
        # (~15ms vs ~200ms), making the second preview feel instant.
        payload = {"type": "bm2_json", "filename": filename, "data": bm2_json or {}}
        return Response(
            content=orjson.dumps(payload),
            media_type="application/json",
        )

    # --- Check data uploads (.csv, .txt, .xlsx from zip extraction) ---
    data_entry = _data_uploads.get(file_id)
    if data_entry:
        filename = data_entry["filename"]
        temp_path = data_entry["temp_path"]
        file_type = data_entry.get("type", "")

        # Tabular files: parse first 50 rows into headers + rows arrays
        if file_type in ("csv", "txt"):
            try:
                separator = "," if file_type == "csv" else "\t"
                with open(temp_path, "r", encoding="utf-8", errors="replace") as f:
                    all_lines = f.readlines()

                # Strip trailing newlines and skip completely empty lines
                all_lines = [ln.rstrip("\n\r") for ln in all_lines if ln.strip()]
                total_rows = max(len(all_lines) - 1, 0)  # minus header row

                headers = all_lines[0].split(separator) if all_lines else []
                # Return up to 50 data rows (excluding the header)
                rows = [
                    line.split(separator)
                    for line in all_lines[1:51]
                ]

                return JSONResponse({
                    "type": "table",
                    "filename": filename,
                    "headers": headers,
                    "rows": rows,
                    "total_rows": total_rows,
                })
            except Exception as e:
                logger.warning("Preview parse failed for %s: %s", filename, e)
                return JSONResponse({
                    "type": "info",
                    "filename": filename,
                    "error": f"Could not parse file: {e}",
                })

        # Excel files: parse each worksheet with openpyxl and return
        # headers + first 50 rows per sheet for tabular preview.
        if file_type == "xlsx":
            try:
                import openpyxl
                # Note: read_only=True is avoided here because some xlsx files
                # (e.g., NTP/NIEHS exports) lack dimension metadata, causing
                # openpyxl's read-only mode to see only 1 row per sheet.
                wb = openpyxl.load_workbook(temp_path, data_only=True)
                sheets = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    rows_iter = ws.iter_rows(values_only=True)
                    # First row is treated as column headers
                    first_row = next(rows_iter, None)
                    headers = [str(c) if c is not None else "" for c in first_row] if first_row else []
                    # Collect up to 50 data rows for the preview
                    data_rows = []
                    for i, row in enumerate(rows_iter):
                        if i >= 50:
                            break
                        data_rows.append([str(c) if c is not None else "" for c in row])
                    # Total row count (subtract 1 for the header row)
                    total = max(ws.max_row - 1, 0) if ws.max_row else 0
                    sheets.append({
                        "name": sheet_name,
                        "headers": headers,
                        "rows": data_rows,
                        "total_rows": total,
                    })
                wb.close()
                return JSONResponse({
                    "type": "xlsx_table",
                    "filename": filename,
                    "sheets": sheets,
                })
            except Exception as e:
                logger.warning("XLSX preview failed for %s: %s", filename, e)
                return JSONResponse({
                    "type": "info",
                    "filename": filename,
                    "error": f"Could not parse xlsx: {e}",
                })

        # Unknown type fallback
        return JSONResponse({
            "type": "info",
            "filename": filename,
            "message": f"Preview not available for .{file_type} files.",
        })

    # --- Also check _csv_uploads (gene-level BMD CSVs) ---
    csv_entry = _csv_uploads.get(file_id)
    if csv_entry:
        filename = csv_entry["filename"]
        temp_path = csv_entry["temp_path"]
        try:
            with open(temp_path, "r", encoding="utf-8", errors="replace") as f:
                all_lines = f.readlines()
            all_lines = [ln.rstrip("\n\r") for ln in all_lines if ln.strip()]
            total_rows = max(len(all_lines) - 1, 0)
            headers = all_lines[0].split(",") if all_lines else []
            rows = [line.split(",") for line in all_lines[1:51]]
            return JSONResponse({
                "type": "table",
                "filename": filename,
                "headers": headers,
                "rows": rows,
                "total_rows": total_rows,
            })
        except Exception as e:
            return JSONResponse({
                "type": "info",
                "filename": filename,
                "error": f"Could not parse file: {e}",
            })

    # Not found in any store
    return JSONResponse(
        {"error": f"File not found: {file_id}"},
        status_code=404,
    )


# ---------------------------------------------------------------------------
# POST /api/export-docx — export full report to .docx
# ---------------------------------------------------------------------------

@app.post("/api/export-docx")
async def api_export_docx(request: Request):
    """
    Export the full 5dToxReport to a .docx file.

    Builds a document in NIEHS report order:
      1. Title + subtitle
      2. Background section (paragraphs + references)
      3. Materials and Methods (if approved)
      4. Results heading
         4a. Apical endpoint tables (from .bm2 files)
         4b. Apical Endpoint BMD Summary table (if approved)
         4c. Gene Set BMD Analysis tables (if genomics approved)
         4d. Gene BMD Analysis tables (if genomics approved)
      5. Summary (if approved)
      6. References

    Input JSON:
      {
        "paragraphs": ["paragraph 1...", ...],
        "references": ["[1] Ref text...", ...],
        "chemical_name": "1,2-Dichlorobenzene",
        "apical_sections": [...],
        "methods_paragraphs": ["Methods paragraph 1...", ...],
        "bmd_summary_endpoints": [{endpoint, bmd, ...}, ...],
        "genomics_sections": [
          {"organ": "liver", "sex": "male",
           "gene_sets": [...], "top_genes": [...], "dose_unit": "mg/kg"},
          ...
        ],
        "summary_paragraphs": ["Summary paragraph 1...", ...]
      }

    All new fields are optional — if omitted, those sections are skipped
    (backwards-compatible with the original two-section report).

    Returns the .docx file as a downloadable attachment.
    """
    body = await request.json()
    paragraphs = body.get("paragraphs", [])
    references = body.get("references", [])
    chemical_name = body.get("chemical_name", "Chemical")
    apical_sections = body.get("apical_sections", [])

    # New NIEHS sections (all optional)
    # methods_data is the new structured format (dict with sections + context).
    # methods_paragraphs is the legacy flat format (list of strings).
    # We accept either for backwards compatibility.
    methods_data = body.get("methods_data")
    methods_paragraphs = body.get("methods_paragraphs", [])
    animal_report_data = body.get("animal_report")
    bmd_summary_endpoints = body.get("bmd_summary_endpoints", [])
    genomics_sections = body.get("genomics_sections", [])
    summary_paragraphs = body.get("summary_paragraphs", [])

    # At least the background or one results section must be present
    if not paragraphs and not apical_sections:
        return JSONResponse(
            {"error": "No paragraphs or apical sections provided"},
            status_code=400,
        )

    # Build the .docx document
    doc = Document()

    # ===== 1. Title + subtitle =====
    doc.add_heading(
        f"5 Day Genomic Dose Response: {chemical_name}",
        level=1,
    )
    add_para(doc,
        "5 Day Genomic Dose Response in Sprague-Dawley Rats",
        italic=True, size=12,
    )

    # ===== 2. Background section =====
    if paragraphs:
        doc.add_heading("Background", level=2)
        for para_text in paragraphs:
            p = doc.add_paragraph()
            _add_text_with_superscript_refs(p, para_text)
            p.paragraph_format.space_after = Pt(6)

    # Sequential table numbering across all report sections
    next_table_num = 1

    # ===== 3. Materials and Methods (NEW — if approved) =====
    # Supports two formats:
    #   - Structured (methods_data): dict with "sections" and "context" from
    #     the new NIEHS-style generation.  Uses add_methods_to_doc() for full
    #     heading hierarchy, conditional subsections, and Table 1.
    #   - Legacy (methods_paragraphs): flat list of paragraph strings from the
    #     old 5-paragraph generation.  Rendered as simple body text.
    if methods_data and methods_data.get("sections"):
        from methods_report import MethodsReport, add_methods_to_doc
        report = MethodsReport.from_dict(methods_data)
        next_table_num = add_methods_to_doc(doc, report, start_table_num=next_table_num)
    elif methods_paragraphs:
        doc.add_heading("Materials and Methods", level=2)
        for para_text in methods_paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

    # ===== 3b. Animal Report (NEW — if approved) =====
    # Inserted between Materials & Methods and Results because it documents
    # the study's animal roster and data coverage — context needed before
    # reading the dose-response results.
    #
    # animal_report_data is a dict matching report_to_dict() output from
    # animal_report.py.  We reconstruct an AnimalReport object from it
    # to reuse the add_animal_report_to_doc() DOCX builder.
    if animal_report_data:
        from animal_report import AnimalReport, AnimalRecord
        # Reconstruct AnimalReport from serialized dict
        ar = AnimalReport()
        ar.study_number = animal_report_data.get("study_number")
        ar.total_animals = animal_report_data.get("total_animals", 0)
        ar.core_count = animal_report_data.get("core_count", 0)
        ar.biosampling_count = animal_report_data.get("biosampling_count", 0)
        ar.dose_groups = animal_report_data.get("dose_groups", [])
        # Reconstruct dose_design — keys were stringified for JSON
        raw_dd = animal_report_data.get("dose_design", {})
        ar.dose_design = {float(k): v for k, v in raw_dd.items()}
        ar.domain_coverage = animal_report_data.get("domain_coverage", {})
        ar.completeness = animal_report_data.get("completeness", {})
        ar.consistency_issues = animal_report_data.get("consistency_issues", [])
        # Reconstruct animal records
        for aid, rec_dict in animal_report_data.get("animals", {}).items():
            ar.animals[aid] = AnimalRecord(
                animal_id=aid,
                sex=rec_dict.get("sex"),
                dose=rec_dict.get("dose"),
                selection=rec_dict.get("selection"),
                domain_presence=rec_dict.get("domain_presence", {}),
            )
        # Reconstruct attrition objects
        from animal_report import DomainAttrition
        for domain, att_dict in animal_report_data.get("attrition", {}).items():
            ar.attrition[domain] = DomainAttrition(
                domain=domain,
                xlsx_ids=set(att_dict.get("xlsx_ids", [])),
                txt_csv_ids=set(att_dict.get("txt_csv_ids", [])),
                bm2_ids=set(att_dict.get("bm2_ids", [])),
                excluded_xlsx_to_txt=set(att_dict.get("excluded_xlsx_to_txt", [])),
                excluded_txt_to_bm2=set(att_dict.get("excluded_txt_to_bm2", [])),
                exclusion_reasons=att_dict.get("exclusion_reasons", {}),
            )
        next_table_num = add_animal_report_to_doc(doc, ar, start_table_num=next_table_num)

    # ===== 4. Results =====
    # Add a "Results" heading if we have any results sections
    has_results = bool(apical_sections or bmd_summary_endpoints or genomics_sections)
    if has_results:
        doc.add_heading("Results", level=2)

    # --- 4a. Apical endpoint tables (existing .bm2 tables) ---
    for section in apical_sections:
        bm2_id = section.get("bm2_id", "")
        section_title = section.get(
            "section_title",
            "Animal Condition, Body Weights, and Organ Weights",
        )
        table_caption = section.get(
            "table_caption_template",
            "Summary of Body Weights and Organ Weights "
            "of {sex} Rats Administered {compound} for Five Days",
        )
        compound = section.get("compound_name", chemical_name)
        dose_unit = section.get("dose_unit", "mg/kg")

        upload = _bm2_uploads.get(bm2_id)
        if not upload:
            continue

        table_data = upload.get("table_data")
        if table_data is None:
            try:
                loop = asyncio.get_running_loop()
                table_data, _, bm2_json = await loop.run_in_executor(
                    None, build_table_data_from_bm2, upload["temp_path"],
                )
                upload["table_data"] = table_data
                upload["bm2_json"] = bm2_json
            except Exception:
                continue

        narrative_paras = section.get("narrative_paragraphs")
        if narrative_paras is None and upload:
            narrative_paras = upload.get("narrative")

        next_table_num = add_apical_tables_to_doc(
            doc,
            table_data,
            section_title=section_title,
            compound_name=compound,
            dose_unit=dose_unit,
            table_caption_template=table_caption,
            start_table_num=next_table_num,
            narrative_paragraphs=narrative_paras,
        )

    # --- 4b. Apical Endpoint BMD Summary table (NEW) ---
    if bmd_summary_endpoints:
        dose_unit = body.get("dose_unit", "mg/kg")
        next_table_num = add_bmd_summary_table_to_doc(
            doc,
            bmd_summary_endpoints,
            table_num=next_table_num,
            dose_unit=dose_unit,
        )

    # --- 4c & 4d. Gene Set BMD and Gene BMD tables (NEW) ---
    for gs_section in genomics_sections:
        organ = gs_section.get("organ", "")
        sex = gs_section.get("sex", "")
        gene_sets = gs_section.get("gene_sets", [])
        top_genes = gs_section.get("top_genes", [])
        dose_unit = gs_section.get("dose_unit", "mg/kg")

        # Gene Set BMD Analysis table (4c)
        if gene_sets:
            next_table_num = add_gene_set_bmd_tables_to_doc(
                doc, gene_sets, organ, sex,
                table_num=next_table_num,
                dose_unit=dose_unit,
            )

        # Gene BMD Analysis table (4d)
        if top_genes:
            next_table_num = add_gene_bmd_tables_to_doc(
                doc, top_genes, organ, sex,
                table_num=next_table_num,
                dose_unit=dose_unit,
            )

    # ===== 5. Summary (NEW — if approved) =====
    if summary_paragraphs:
        doc.add_heading("Summary", level=2)
        for para_text in summary_paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

    # ===== 6. References (moved to end for NIEHS ordering) =====
    if references:
        doc.add_heading("References", level=2)
        for ref_line in references:
            add_para(doc, ref_line, size=10)

    # Save to a temporary file and return as download
    safe_name = _safe_filename(chemical_name)
    filename = f"5dToxReport_{safe_name}.docx"

    tmp = tempfile.NamedTemporaryFile(
        delete=False, suffix=".docx", prefix="5dtox_",
    )
    doc.save(tmp.name)
    tmp.close()

    return FileResponse(
        tmp.name,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# POST /api/export-pdf — export full report to tagged PDF/UA-1
# ---------------------------------------------------------------------------

@app.post("/api/export-pdf")
async def api_export_pdf(request: Request):
    """
    Export the full 5dToxReport to a PDF/UA-1 compliant PDF file.

    Accepts the same JSON payload as /api/export-docx for consistency.
    Marshals the data into the Typst template schema and compiles it
    using the Typst compiler (embedded Rust binary via typst-py).

    The output PDF has:
      - Full StructTreeRoot with H1-H3, P, Table/TH/TD/TR tags
      - PDF/UA-1 identifier in XMP metadata
      - /Lang set to "en" in the document catalog
      - /MarkInfo << /Marked true >>
      - Proper heading hierarchy matching NIEHS Report 10
      - Running header and page numbers marked as Artifacts

    Returns the PDF file as a downloadable attachment.
    """
    from report_pdf import build_report_pdf, marshal_export_data

    body = await request.json()

    # Resolve table_data for apical sections that reference uploaded .bm2 files.
    # The web UI may send bm2_id references instead of inline table_data,
    # so we need to look up the cached data from the upload store.
    for sec in body.get("apical_sections", []):
        if "table_data" not in sec or not sec["table_data"]:
            bm2_id = sec.get("bm2_id", "")
            upload = _bm2_uploads.get(bm2_id)
            if upload and upload.get("table_data"):
                sec["table_data"] = upload["table_data"]

    # Also inject narrative from server cache if not provided inline
    for sec in body.get("apical_sections", []):
        if not sec.get("narrative_paragraphs"):
            bm2_id = sec.get("bm2_id", "")
            upload = _bm2_uploads.get(bm2_id)
            if upload and upload.get("narrative"):
                sec["narrative_paragraphs"] = upload["narrative"]

    try:
        # Marshal the DOCX-format payload into the Typst template schema
        report_data = marshal_export_data(body)

        # Compile to PDF/UA-1 — sub-second for typical report sizes
        pdf_bytes = build_report_pdf(report_data)
    except Exception as e:
        logging.exception("PDF export failed")
        return JSONResponse(
            {"error": f"PDF generation failed: {e}"},
            status_code=500,
        )

    # Write to a temp file for FileResponse
    chemical_name = body.get("chemical_name", "Chemical")
    safe_name = _safe_filename(chemical_name)
    filename = f"5dToxReport_{safe_name}.pdf"

    tmp = tempfile.NamedTemporaryFile(
        delete=False, suffix=".pdf", prefix="5dtox_",
    )
    tmp.write(pdf_bytes)
    tmp.close()

    return FileResponse(
        tmp.name,
        filename=filename,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# GET /api/export-pdf-scaffold — generate a scaffold PDF showing all sections
# ---------------------------------------------------------------------------

@app.get("/api/export-pdf-scaffold")
async def api_export_pdf_scaffold(
    chemical_name: str = "Test Article",
    casrn: str = "000-00-0",
    dtxsid: str = "DTXSID0000000",
):
    """
    Generate a complete scaffold PDF with placeholder content in every section.

    This endpoint produces a full NIEHS Report 10-structured PDF with all
    sections populated by clearly-marked placeholder text (wrapped in angle
    quotes).  The purpose is to show the exact page flow, typography, table
    layout, landscape pages, and pagination that the final report will have.

    Every template code path is exercised: title page, roman-numeral front
    matter (foreword, TOC, tables list, about, peer review, pub details,
    acknowledgments, abstract), arabic body pages (background, M&M with
    Table 1, results with landscape dose-response tables, internal dose
    portrait table, BMD summary sex-grouped table, genomics gene set
    and gene tables with GO/gene descriptions), summary, and references.

    Query parameters allow customizing the chemical identity on the
    title page and throughout the document:
      ?chemical_name=Perfluorohexanesulfonamide&casrn=41997-13-1&dtxsid=DTXSID50469320

    Returns the PDF as a downloadable attachment.
    """
    from report_pdf import build_report_pdf, scaffold_report_data

    try:
        data = scaffold_report_data(
            chemical_name=chemical_name,
            casrn=casrn,
            dtxsid=dtxsid,
        )
        pdf_bytes = build_report_pdf(data)
    except Exception as e:
        logging.exception("Scaffold PDF generation failed")
        return JSONResponse(
            {"error": f"Scaffold PDF generation failed: {e}"},
            status_code=500,
        )

    # Clean filename from chemical name
    safe_name = _safe_filename(chemical_name)
    filename = f"5dToxReport_Scaffold_{safe_name}.pdf"

    tmp = tempfile.NamedTemporaryFile(
        delete=False, suffix=".pdf", prefix="5dtox_scaffold_",
    )
    tmp.write(pdf_bytes)
    tmp.close()

    return FileResponse(
        tmp.name,
        filename=filename,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _add_text_with_superscript_refs(paragraph, text: str) -> None:
    """
    Add text to a paragraph, converting [N] reference markers to superscript.

    Splits the text on [N] patterns, adds normal text as regular runs and
    reference numbers as superscript runs. Uses Calibri 11pt as the base font.
    """
    import re
    from docx.oxml.ns import qn

    # Split on reference markers like [1], [2,3], [1-3], etc.
    parts = re.split(r'(\[\d+(?:[,\-–]\d+)*\])', text)

    for part in parts:
        if re.match(r'\[\d+(?:[,\-–]\d+)*\]', part):
            # This is a reference marker — make it superscript
            run = paragraph.add_run(part)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.superscript = True
        else:
            # Normal text
            if part:
                run = paragraph.add_run(part)
                run.font.size = Pt(11)
                run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# File pool fingerprinting and validation
# ---------------------------------------------------------------------------
# These endpoints integrate with file_integrator.py to provide structural
# cross-validation of uploaded files.  Every file in the pool gets a
# "fingerprint" (extracted metadata: doses, animals, endpoints, domain)
# which enables detection of conflicts, coverage gaps, and redundancy.


def _fingerprint_and_store(
    file_id: str,
    filename: str,
    path: str,
    file_type: str,
    dtxsid: str,
    bm2_json: dict | None = None,
) -> FileFingerprint:
    """
    Fingerprint a single file and store the result in _pool_fingerprints.

    Called on upload (both direct and zip extraction) and on session load.
    The fingerprint is stored in the dtxsid-keyed pool so it's available
    for lightweight_validate() on subsequent uploads and for full
    validate_pool() when the user clicks "Validate & Integrate".

    Args:
        file_id:   UUID from upload.
        filename:  Original filename.
        path:      Absolute path to the file on disk.
        file_type: "xlsx", "txt", "csv", or "bm2".
        dtxsid:    The DTXSID session this file belongs to.
        bm2_json:  Pre-loaded BMDProject dict (optional, for bm2 files).

    Returns:
        The created FileFingerprint.
    """
    ts_added = datetime.now(tz=timezone.utc).isoformat()
    fp = fingerprint_file(file_id, filename, path, file_type, ts_added, bm2_json)

    if dtxsid:
        if dtxsid not in _pool_fingerprints:
            _pool_fingerprints[dtxsid] = {}
        _pool_fingerprints[dtxsid][file_id] = fp

    return fp


def _run_lightweight_validation(
    fp: FileFingerprint,
    dtxsid: str,
) -> list[dict]:
    """
    Run lightweight validation on a new file against the existing pool.

    Returns a list of issue dicts (may be empty).  Called after fingerprinting
    a newly uploaded file to give immediate feedback.

    Args:
        fp:      Fingerprint of the newly added file.
        dtxsid:  The DTXSID session this file belongs to.

    Returns:
        List of validation issue dicts for JSON serialization.
    """
    if not dtxsid or dtxsid not in _pool_fingerprints:
        return []
    existing = {
        fid: efp for fid, efp in _pool_fingerprints[dtxsid].items()
        if fid != fp.file_id
    }
    from dataclasses import asdict
    issues = lightweight_validate(fp, existing)
    return [asdict(issue) for issue in issues]


@app.post("/api/pool/validate/{dtxsid}")
async def api_pool_validate(dtxsid: str):
    """
    Run full cross-validation on a session's file pool.

    Fingerprints (or re-fingerprints) every file in the session's files/
    directory, then runs all validation checks: coverage, dose consistency,
    animal counts, sex coverage, and redundancy detection.

    Returns a ValidationReport as JSON with:
      - coverage_matrix: domain → tier → file_id(s)
      - issues: list of { severity, domain, issue_type, message, ... }
      - is_complete: whether all domains have full tier coverage

    Saves the report to sessions/{dtxsid}/validation_report.json for
    persistence across page reloads.
    """
    # Re-fingerprint all files from disk to catch any out-of-band changes
    # (e.g., files added manually or by other processes).
    files_dir = _session_dir(dtxsid) / "files"
    if not files_dir.exists():
        return JSONResponse({
            "error": "No files directory found for this session",
        }, status_code=404)

    # Force a full re-scan of all files in the session
    fps = _ensure_fingerprints(dtxsid, force=True)
    report = validate_pool(dtxsid, fps)

    # Persist the report to disk
    report_dict = {
        "dtxsid": report.dtxsid,
        "run_at": report.run_at,
        "file_count": report.file_count,
        "fingerprints": report.fingerprints,
        "issues": report.issues,
        "coverage_matrix": report.coverage_matrix,
        "is_complete": report.is_complete,
    }
    report_path = _session_dir(dtxsid) / "validation_report.json"
    report_path.write_text(
        json.dumps(report_dict, indent=2, default=str),
        encoding="utf-8",
    )

    return Response(
        content=orjson.dumps(report_dict),
        media_type="application/json",
    )


@app.post("/api/pool/resolve")
async def api_pool_resolve(request: Request):
    """
    Record a user's precedence decision for a specific validation conflict.

    When the validation report shows an error (e.g., dose group mismatch),
    the user picks which file is authoritative.  This endpoint persists
    that decision to sessions/{dtxsid}/precedence.json so it survives
    page reloads.

    Input JSON:
      {
        "dtxsid": "DTXSID50469320",
        "issue_index": 0,
        "chosen_file_id": "abc123-..."
      }

    Returns { "ok": true } on success.
    """
    body = await request.json()
    dtxsid = body.get("dtxsid", "")
    issue_index = body.get("issue_index")
    chosen_file_id = body.get("chosen_file_id", "")

    if not dtxsid or issue_index is None or not chosen_file_id:
        return JSONResponse(
            {"error": "dtxsid, issue_index, and chosen_file_id are required"},
            status_code=400,
        )

    # Load existing precedence decisions
    precedence_path = _session_dir(dtxsid) / "precedence.json"
    precedence: list[dict] = []
    if precedence_path.exists():
        try:
            precedence = json.loads(precedence_path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, Exception):
            precedence = []

    # Record the new decision
    precedence.append({
        "issue_index": issue_index,
        "chosen_file_id": chosen_file_id,
        "resolved_at": datetime.now(tz=timezone.utc).isoformat(),
    })

    # Persist to disk
    precedence_path.write_text(
        json.dumps(precedence, indent=2),
        encoding="utf-8",
    )

    return JSONResponse({"ok": True})


# ---------------------------------------------------------------------------
# POST /api/pool/integrate/{dtxsid} — merge files into unified BMDProject
# ---------------------------------------------------------------------------
# After validation + conflict resolution, this endpoint calls integrate_pool()
# to select the best file per domain and merge all dose-response data into a
# single BMDProject JSON.  The integrated structure is the unified source of
# truth that drives all downstream section generation (tables + narratives).


@app.post("/api/pool/integrate/{dtxsid}")
async def api_pool_integrate(dtxsid: str):
    """
    Merge all pool files into a unified BMDProject JSON.

    Reads fingerprints from _pool_fingerprints, coverage_matrix from the
    persisted validation_report.json, and precedence decisions from
    precedence.json.  Calls integrate_pool() to select the best file per
    domain and produce the merged structure.

    The result is stored both in-memory (_integrated_pool) and on disk
    (sessions/{dtxsid}/integrated.json) for session restore.

    Returns the full integrated BMDProject JSON, including a _meta block
    with provenance: which file was chosen for each domain and why.
    """
    session_dir = _session_dir(dtxsid)
    files_dir = session_dir / "files"
    if not files_dir.exists():
        return JSONResponse(
            {"error": "No files directory found for this session"},
            status_code=404,
        )

    # Load fingerprints — prefer in-memory, fall back to validation_report.json
    fps = _pool_fingerprints.get(dtxsid, {})
    if not fps:
        report_path = session_dir / "validation_report.json"
        if report_path.exists():
            try:
                report = json.loads(report_path.read_text(encoding="utf-8"))
                fps = report.get("fingerprints", {})
            except (json.JSONDecodeError, Exception):
                pass

    if not fps:
        return JSONResponse(
            {"error": "No fingerprints found — run validation first"},
            status_code=400,
        )

    # Load the coverage matrix from the validation report
    report_path = session_dir / "validation_report.json"
    coverage_matrix: dict = {}
    if report_path.exists():
        try:
            report = json.loads(report_path.read_text(encoding="utf-8"))
            coverage_matrix = report.get("coverage_matrix", {})
        except (json.JSONDecodeError, Exception):
            pass

    if not coverage_matrix:
        return JSONResponse(
            {"error": "No coverage matrix found — run validation first"},
            status_code=400,
        )

    # Load user precedence decisions (may be empty if no conflicts resolved)
    precedence_path = session_dir / "precedence.json"
    precedence: list[dict] = []
    if precedence_path.exists():
        try:
            precedence = json.loads(precedence_path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, Exception):
            pass

    # Run integration in a thread pool — xlsx parsing uses openpyxl (blocking I/O)
    loop = asyncio.get_running_loop()
    try:
        integrated = await loop.run_in_executor(
            None,
            integrate_pool,
            dtxsid,
            str(session_dir),
            fps,
            coverage_matrix,
            precedence,
        )
    except Exception as e:
        logger.exception("Pool integration failed for %s", dtxsid)
        return JSONResponse(
            {"error": f"Integration failed: {e}"},
            status_code=500,
        )

    # Cache in memory for the process-integrated endpoint
    _integrated_pool[dtxsid] = integrated

    return Response(
        content=orjson.dumps(integrated),
        media_type="application/json",
    )


# ---------------------------------------------------------------------------
# GET /api/integrated/{dtxsid} — stream the full cached integrated.json
# ---------------------------------------------------------------------------
# Serves the integrated.json directly from disk via FileResponse (streaming).
# Does NOT re-run integration — just reads the cached file.  Used by the
# "Preview Data" modal which streams it via Oboe.js for progressive rendering.


@app.get("/api/integrated/{dtxsid}")
async def api_integrated_full(dtxsid: str):
    """
    Stream the full integrated BMDProject JSON from disk.

    Returns the cached integrated.json via FileResponse (chunked streaming)
    so the browser can parse it progressively with Oboe.js.  If no cached
    file exists, returns 404 — the caller should trigger integration first.
    """
    integrated_path = _session_dir(dtxsid) / "integrated.json"
    if not integrated_path.exists():
        return JSONResponse(
            {"error": "No integrated data found — run integration first"},
            status_code=404,
        )
    return FileResponse(
        path=str(integrated_path),
        media_type="application/json",
        filename="integrated.json",
    )


# ---------------------------------------------------------------------------
# GET /api/integrated-summary/{dtxsid} — lightweight summary of integrated data
# ---------------------------------------------------------------------------
# Returns metadata, experiment list, and probe counts without the full 60MB+
# experiment response arrays.  Used by the Data tab previewer on session restore.


@app.get("/api/integrated-summary/{dtxsid}")
async def api_integrated_summary(dtxsid: str):
    """
    Return a lightweight summary of the integrated BMDProject.

    Loads from in-memory cache or disk, then extracts only the metadata
    and per-experiment names/probe counts — NOT the full response arrays.
    """
    integrated = _integrated_pool.get(dtxsid)
    if integrated is None:
        integrated_path = _session_dir(dtxsid) / "integrated.json"
        if integrated_path.exists():
            try:
                integrated = json.loads(integrated_path.read_text(encoding="utf-8"))
                _integrated_pool[dtxsid] = integrated
            except (json.JSONDecodeError, Exception):
                pass

    if not integrated:
        return JSONResponse(
            {"error": "No integrated data found"},
            status_code=404,
        )

    meta = integrated.get("_meta", {})
    experiments = integrated.get("doseResponseExperiments", [])
    bmd_results = integrated.get("bMDResult", [])
    cat_results = integrated.get("categoryAnalysisResults", [])

    # Build experiment summaries (name + probe count only — no response data)
    exp_summaries = []
    for exp in experiments:
        exp_summaries.append({
            "name": exp.get("name", ""),
            "probe_count": len(exp.get("probeResponses", [])),
        })

    return JSONResponse({
        "_meta": meta,
        "experiment_count": len(experiments),
        "experiments": exp_summaries,
        "bmd_result_count": len(bmd_results),
        "category_analysis_count": len(cat_results),
    })


# ---------------------------------------------------------------------------
# POST /api/process-integrated/{dtxsid} — generate sections from integrated data
# ---------------------------------------------------------------------------
# Replaces per-file process-bm2 calls for the auto-processing flow.
# Takes the unified BMDProject JSON and produces section cards (tables +
# narratives) for every domain, ready for the UI to display.


@app.post("/api/process-integrated/{dtxsid}")
async def api_process_integrated(dtxsid: str, request: Request):
    """
    Process the integrated BMDProject JSON into section cards with tables
    and narratives for each apical endpoint domain.

    Input JSON:
      {
        "compound_name": "PFHxSAm",
        "dose_unit": "mg/kg"
      }

    Loads the integrated JSON from _integrated_pool (in-memory) or from
    sessions/{dtxsid}/integrated.json (disk fallback).  Calls
    build_table_data() to run NTP stats on all experiments, then partitions
    the results by domain for the UI's section cards.

    Returns:
      {
        "sections": [
          {
            "domain": "body_weight",
            "title": "Body Weight",
            "tables_json": {"Male": [...], "Female": [...]},
            "narrative": ["paragraph1", "paragraph2", ...]
          },
          ...
        ]
      }
    """
    body = await request.json()
    compound_name = body.get("compound_name", "Test Compound")
    dose_unit = body.get("dose_unit", "mg/kg")

    # Load integrated data — prefer in-memory, fall back to disk
    integrated = _integrated_pool.get(dtxsid)
    if integrated is None:
        integrated_path = _session_dir(dtxsid) / "integrated.json"
        if integrated_path.exists():
            try:
                integrated = json.loads(integrated_path.read_text(encoding="utf-8"))
                _integrated_pool[dtxsid] = integrated
            except (json.JSONDecodeError, Exception):
                pass

    if not integrated:
        return JSONResponse(
            {"error": "No integrated data found — run integration first"},
            status_code=400,
        )

    try:
        # Restore category lookup from the serialized pipe-separated keys.
        # integrate_pool() stored this as _category_lookup with "prefix|endpoint"
        # string keys; we restore them to (prefix, endpoint) tuple keys that
        # build_table_data() expects.
        from apical_report import build_table_data

        flat_cat = integrated.get("_category_lookup", {})
        cat_lookup: dict[tuple[str, str], dict] = {
            tuple(k.split("|", 1)): v
            for k, v in flat_cat.items()
        }

        loop = asyncio.get_running_loop()

        # Filter out gene expression experiments before running NTP stats.
        # Gene expression .bm2 data has thousands of probes — running Dunnett's
        # test on each would be extremely slow and isn't meaningful for clinical
        # endpoints.  Genomics is handled separately by export_genomics().
        from file_integrator import _BM2_DOMAIN_MAP, detect_domain

        meta = integrated.get("_meta", {})
        source_files = meta.get("source_files", {})
        ge_source = source_files.get("gene_expression")
        ge_exp_names = set()
        if ge_source:
            # Gene expression experiments have names starting with the organ
            # (e.g., "Liver_PFHxSAm_Male_No0") — identify them by checking
            # which experiments DON'T match any clinical domain prefix.
            for exp in integrated.get("doseResponseExperiments", []):
                exp_name = exp.get("name", "")
                exp_lower = exp_name.lower().replace("_", "")
                matched = False
                for prefix in _BM2_DOMAIN_MAP:
                    clean = exp_lower.replace("female", "").replace("male", "").strip()
                    if clean.startswith(prefix) or prefix.startswith(clean):
                        matched = True
                        break
                if not matched:
                    ge_exp_names.add(exp_name)

        # Build a filtered copy without gene expression experiments for NTP stats
        if ge_exp_names:
            apical_integrated = {
                **integrated,
                "doseResponseExperiments": [
                    exp for exp in integrated.get("doseResponseExperiments", [])
                    if exp.get("name", "") not in ge_exp_names
                ],
            }
            logger.info(
                "Filtered %d gene expression experiments from NTP stats pipeline",
                len(ge_exp_names),
            )
        else:
            apical_integrated = integrated

        # Run the NTP stats pipeline on clinical endpoint experiments only.
        # This is pure Python (no JVM) and typically takes <1s.
        table_data = await loop.run_in_executor(
            None, build_table_data, apical_integrated, cat_lookup,
        )

        # --- Partition table rows by domain ---
        # build_table_data() returns {"Male": [TableRow, ...], "Female": [...]}.
        # We need to split these into per-domain sections so the UI can create
        # separate section cards for body weight, organ weights, etc.
        #
        # Strategy: look at the experiment names in the integrated data to build
        # a mapping of endpoint_name → domain, then partition the table rows.

        # Build experiment_name → domain mapping.
        # Strategy: use _meta.source_files to know which experiment names
        # belong to which domain.  Each source file contributed experiments
        # whose names we can map back.  Also use detect_domain() on the
        # experiment name itself as fallback.
        exp_name_to_domain: dict[str, str] = {}

        # Use the filtered (apical-only) experiments for domain mapping —
        # gene expression experiments were already excluded above.
        for exp in apical_integrated.get("doseResponseExperiments", []):
            exp_name = exp.get("name", "")
            exp_lower = exp_name.lower()

            # Strip sex suffix/prefix for matching.
            # IMPORTANT: strip "female" BEFORE "male" — "female" contains
            # "male" as a substring, so stripping "male" first leaves "fe".
            stripped = exp_lower.replace("female", "").replace("male", "").replace("_", "").strip()

            domain_for_exp = None
            for prefix, dom in _BM2_DOMAIN_MAP.items():
                if stripped.startswith(prefix) or prefix.startswith(stripped):
                    domain_for_exp = dom
                    break

            # Fallback: try detect_domain() which uses regex patterns.
            # This handles abbreviated names like "clin_chem" that don't
            # match the full BM2 prefix "clinicalchemistry".
            if not domain_for_exp:
                domain_for_exp = detect_domain(exp_name, "bm2", 0)

            # Last resort: check if experiment name overlaps with source domain keys
            if not domain_for_exp:
                for dom in source_files:
                    dom_key = dom.replace("_", "")
                    if dom_key in exp_lower.replace("_", ""):
                        domain_for_exp = dom
                        break

            if domain_for_exp:
                exp_name_to_domain[exp_name] = domain_for_exp

        # Build endpoint → domain map using the experiment mapping.
        # Each probe/endpoint in an experiment inherits that experiment's domain.
        endpoint_domain_map: dict[str, str] = {}
        for exp in apical_integrated.get("doseResponseExperiments", []):
            exp_name = exp.get("name", "")
            dom = exp_name_to_domain.get(exp_name)
            if dom:
                for pr in exp.get("probeResponses", []):
                    probe_id = pr.get("probe", {}).get("id", "")
                    if probe_id:
                        # Key by (sex, probe_id) to avoid collisions when
                        # the same endpoint name appears in different domains
                        # (unlikely but possible for generic names like "Day")
                        endpoint_domain_map[(exp_name, probe_id)] = dom

        # Partition TableRows by domain, preserving sex grouping.
        # build_table_data() groups by sex and uses probe_name as the label.
        # We need to match back to the (exp_name, probe_id) key.
        #
        # Since build_table_data doesn't preserve the experiment name on
        # TableRow, we build a secondary map: (sex, probe_name) → domain.
        sex_probe_domain: dict[tuple[str, str], str] = {}
        for (exp_name, probe_id), dom in endpoint_domain_map.items():
            sex = "Female" if "female" in exp_name.lower() else \
                  "Male" if "male" in exp_name.lower() else "Unknown"
            sex_probe_domain[(sex, probe_id)] = dom

        # Structure: {domain: {sex: [TableRow, ...]}}
        domain_tables: dict[str, dict[str, list]] = {}
        for sex, rows in table_data.items():
            for row in rows:
                dom = sex_probe_domain.get((sex, row.label), "unknown")
                domain_tables.setdefault(dom, {}).setdefault(sex, []).append(row)

        # Human-readable domain titles for section headers
        _DOMAIN_TITLES = {
            "body_weight":    "Body Weight",
            "organ_weights":  "Organ Weights",
            "clin_chem":      "Clinical Chemistry",
            "hematology":     "Hematology",
            "hormones":       "Hormones",
            "tissue_conc":    "Tissue Concentration",
            "clinical_obs":   "Clinical Observations",
        }

        # Build sections array: one per domain that has data
        sections = []
        for dom, sex_rows in sorted(domain_tables.items()):
            # Serialize TableRow objects to JSON-friendly dicts
            tables_json = _serialize_table_rows(sex_rows)

            # Generate narrative for this domain's data only
            narrative = generate_results_narrative(
                sex_rows, compound_name, dose_unit,
            )

            sections.append({
                "domain": dom,
                "title": _DOMAIN_TITLES.get(dom, dom.replace("_", " ").title()),
                "tables_json": tables_json,
                "narrative": narrative,
            })

        # --- Gene expression genomics (from integrated .bm2) ---
        # If the integration included gene_expression, extract per-gene BMD
        # and GO BP category results directly from the .bm2 using the
        # BMDExpress 3 Java API.  This replaces the old CSV-based workflow.
        genomics_sections = {}
        meta = integrated.get("_meta", {})
        ge_source = meta.get("source_files", {}).get("gene_expression")
        if ge_source and ge_source.get("tier") == "bm2":
            ge_filename = ge_source.get("filename", "")
            ge_path = str(_session_dir(dtxsid) / "files" / ge_filename)

            if os.path.exists(ge_path):
                from apical_report import export_genomics
                import tempfile

                # Run the Java export in a thread pool (JVM startup ~0.5s)
                tmp_json = tempfile.NamedTemporaryFile(
                    delete=False, suffix=".json", prefix="genomics_",
                )
                tmp_json.close()

                try:
                    ge_result = await loop.run_in_executor(
                        None, export_genomics, ge_path, tmp_json.name,
                    )

                    # Reshape into the format the UI expects: keyed by organ_sex
                    for exp in ge_result.get("experiments", []):
                        organ = exp.get("organ", "unknown").lower()
                        sex = exp.get("sex", "unknown").lower()
                        key = f"{organ}_{sex}"

                        # Sort genes by BMD ascending (lowest = most sensitive).
                        # Java serializes NaN/Infinity as strings — coerce to float.
                        def _safe_float(val, default=float("inf")):
                            if val is None:
                                return default
                            try:
                                v = float(val)
                                # NaN sorts inconsistently — treat as infinity
                                return default if v != v else v
                            except (TypeError, ValueError):
                                return default

                        genes = sorted(
                            exp.get("genes", []),
                            key=lambda g: _safe_float(g.get("bmd")),
                        )

                        # Sort GO terms by bmd_median ascending
                        go_bp = sorted(
                            exp.get("go_bp", []),
                            key=lambda g: _safe_float(g.get("bmd_median")),
                        )

                        genomics_sections[key] = {
                            "organ": organ,
                            "sex": sex,
                            "total_probes": exp.get("total_probes", 0),
                            "total_responsive_genes": len(genes),
                            "gene_sets": [
                                {
                                    "rank": i + 1,
                                    "go_id": g["go_id"],
                                    "go_term": g["go_term"],
                                    "bmd_median": g.get("bmd_median"),
                                    "bmdl_median": g.get("bmdl_median"),
                                    "n_genes": g.get("n_passed", 0),
                                    "direction": g.get("direction", ""),
                                    "fishers_p": g.get("fishers_two_tail"),
                                    "genes": g.get("gene_symbols", ""),
                                }
                                for i, g in enumerate(go_bp[:20])
                            ],
                            "top_genes": [
                                {
                                    "rank": i + 1,
                                    "gene_symbol": g["gene_symbol"],
                                    "bmd": g.get("bmd"),
                                    "bmdl": g.get("bmdl"),
                                    "bmdu": g.get("bmdu"),
                                    "direction": g.get("direction", ""),
                                    "fold_change": g.get("fold_change"),
                                    "r_squared": g.get("r_squared"),
                                }
                                for i, g in enumerate(genes[:20])
                            ],
                        }
                finally:
                    os.unlink(tmp_json.name)

        return JSONResponse({
            "sections": sections,
            "genomics_sections": genomics_sections,
        })

    except Exception as e:
        logger.exception("Processing integrated data failed for %s", dtxsid)
        return JSONResponse(
            {"error": f"Processing failed: {e}"},
            status_code=500,
        )


# ---------------------------------------------------------------------------
# POST /api/generate-animal-report/{dtxsid} — per-animal traceability report
# ---------------------------------------------------------------------------
# This endpoint reads every file in the session's pool and traces individual
# animals across tiers (xlsx → txt/csv → bm2) and domains (body_weight,
# hematology, etc.).  The result documents study design, animal roster,
# domain coverage, attrition patterns, and cross-domain consistency.
#
# Unlike /api/pool/validate (which checks aggregate structural consistency),
# this goes to per-animal granularity: which specific animals dropped between
# tiers, and why (dose group exclusion, biosampling exclusion, QC removal).


@app.post("/api/generate-animal-report/{dtxsid}")
async def api_generate_animal_report(dtxsid: str):
    """
    Generate a per-animal traceability report for a session's file pool.

    Reads all fingerprinted files from disk, extracts per-animal data
    (animal_id → dose, sex, selection), and cross-references across
    tiers and domains.  Persists the result to
    sessions/{dtxsid}/animal_report.json.

    Requires fingerprints to exist (from prior /api/pool/validate call).
    If no fingerprints are cached, re-fingerprints all files first.

    Returns the full AnimalReport as JSON.
    """
    session_path = _session_dir(dtxsid)
    files_dir = session_path / "files"

    if not files_dir.exists():
        return JSONResponse(
            {"error": "No files directory found for this session"},
            status_code=404,
        )

    # Ensure we have fingerprints — re-fingerprint if the pool is empty.
    # This can happen if the server restarted since the last validation.
    fps = _ensure_fingerprints(dtxsid)

    if not fps:
        return JSONResponse(
            {"error": "No fingerprinted files found — upload files first"},
            status_code=400,
        )

    # Build the animal report in a thread executor to avoid blocking
    # the event loop (xlsx/bm2 parsing can take a few seconds).
    loop = asyncio.get_running_loop()
    try:
        report = await loop.run_in_executor(
            None,
            build_animal_report,
            str(session_path),
            fps,
        )
    except Exception as e:
        logger.exception("Failed to build animal report for %s", dtxsid)
        return JSONResponse(
            {"error": f"Animal report generation failed: {e}"},
            status_code=500,
        )

    # Serialize and persist to disk
    report_dict = report_to_dict(report)
    report_path = session_path / "animal_report.json"
    report_path.write_text(
        json.dumps(report_dict, indent=2, default=str),
        encoding="utf-8",
    )

    return Response(
        content=orjson.dumps(report_dict),
        media_type="application/json",
    )


# ---------------------------------------------------------------------------
# Static file serving — CSS, JS, and other assets under web/
# ---------------------------------------------------------------------------
# Mounted AFTER all explicit @app routes so that named endpoints (like GET /)
# take priority.  The StaticFiles handler is a catch-all that serves anything
# inside the web/ directory (style.css, js/state.js, js/utils.js, js/main.js,
# images, etc.) with correct MIME types and caching headers.
app.mount("/", StaticFiles(directory=Path(__file__).parent / "web"), name="static")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import uvicorn

    host = "127.0.0.1"
    port = 9000

    # Parse CLI args
    args = sys.argv[1:]
    i = 0
    while i < len(args):
        if args[i] == "--host" and i + 1 < len(args):
            host = args[i + 1]
            i += 2
        elif args[i] == "--port" and i + 1 < len(args):
            port = int(args[i + 1])
            i += 2
        else:
            i += 1

    import webbrowser

    print(f"Starting 5dToxReport on http://{host}:{port}")

    # Open the browser after a short delay so the server is ready
    import threading
    threading.Timer(1.0, webbrowser.open, args=[f"http://{host}:{port}"]).start()

    uvicorn.run(app, host=host, port=port)
