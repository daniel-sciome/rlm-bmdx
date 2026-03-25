"""
Dose-response interpretation engine for bmdx.

Takes gene-level BMD results (from BMDExpress) and produces a grounded,
literature-backed narrative by querying the bmdx knowledge base.

Usage:
    python interpret.py <dose_response.csv> [--db bmdx.duckdb] [--output output/interpretation.md] [--fdr 0.05]
"""

import argparse
import logging
import os
import sys
import time
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from pathlib import Path

import duckdb
import pandas as pd
import requests
from scipy.stats import fisher_exact

from extract import OllamaEndpoint, LOCAL_OLLAMA, REMOTE_OLLAMA, normalize_gene, parse_json_response


# ---------------------------------------------------------------------------
# AnthropicEndpoint — Claude API endpoint
# ---------------------------------------------------------------------------

@dataclass
class AnthropicEndpoint:
    name: str
    model: str          # e.g. "claude-sonnet-4-6"
    max_tokens: int = 4096
    temperature: float = 0.3

    def generate(self, prompt: str, system: str = "",
                 temperature: float | None = None) -> str:
        import anthropic
        temp = temperature if temperature is not None else self.temperature
        client = anthropic.Anthropic()  # reads ANTHROPIC_API_KEY from env
        messages = [{"role": "user", "content": prompt}]
        kwargs = {
            "model": self.model,
            "max_tokens": self.max_tokens,
            "messages": messages,
            "temperature": temp,
        }
        if system:
            kwargs["system"] = system
        response = client.messages.create(**kwargs)
        return response.content[0].text if response.content else ""

    def is_available(self) -> bool:
        return bool(os.environ.get("ANTHROPIC_API_KEY"))


# ---------------------------------------------------------------------------
# NarrativeRun — result of a single LLM narrative generation
# ---------------------------------------------------------------------------

@dataclass
class NarrativeRun:
    model: str
    endpoint_name: str
    run_index: int
    narrative: str
    elapsed_seconds: float


# ---------------------------------------------------------------------------
# ToxKBQuerier — read-only DB query helpers
# ---------------------------------------------------------------------------

class ToxKBQuerier:
    """Wraps bmdx.duckdb with typed query methods."""

    def __init__(self, db_path: str = "bmdx.duckdb"):
        self.con = duckdb.connect(db_path, read_only=True)

    def close(self):
        self.con.close()

    def gene_pathways(self, gene: str) -> list[dict]:
        rows = self.con.execute(
            "SELECT pathway_db, pathway_id, pathway_name, species "
            "FROM pathways WHERE gene_symbol = ?",
            [gene],
        ).fetchall()
        return [
            {"pathway_db": r[0], "pathway_id": r[1],
             "pathway_name": r[2], "species": r[3]}
            for r in rows
        ]

    def gene_go_terms(self, gene: str) -> list[dict]:
        rows = self.con.execute(
            "SELECT g.go_id, t.go_term, t.cluster_id "
            "FROM gene_go_terms g JOIN go_terms t ON g.go_id = t.go_id "
            "WHERE g.gene_symbol = ?",
            [gene],
        ).fetchall()
        return [
            {"go_id": r[0], "go_term": r[1], "cluster_id": r[2]}
            for r in rows
        ]

    def gene_organs(self, gene: str) -> list[str]:
        row = self.con.execute(
            "SELECT organs FROM genes WHERE gene_symbol = ?",
            [gene],
        ).fetchone()
        if row and row[0]:
            return list(row[0])
        return []

    def gene_papers(self, gene: str) -> list[dict]:
        rows = self.con.execute(
            "SELECT p.paper_id, p.title, p.year, p.citation_count "
            "FROM papers p JOIN paper_genes pg ON p.paper_id = pg.paper_id "
            "WHERE pg.gene_symbol = ?",
            [gene],
        ).fetchall()
        return [
            {"paper_id": r[0], "title": r[1], "year": r[2],
             "citation_count": r[3]}
            for r in rows
        ]

    def gene_claims(self, gene: str) -> list[dict]:
        rows = self.con.execute(
            "SELECT pc.claim, p.title, p.year "
            "FROM paper_claims pc "
            "JOIN paper_genes pg ON pc.paper_id = pg.paper_id "
            "JOIN papers p ON pc.paper_id = p.paper_id "
            "WHERE pg.gene_symbol = ?",
            [gene],
        ).fetchall()
        return [{"claim": r[0], "paper_title": r[1], "year": r[2]} for r in rows]

    def pathway_genes(self, pathway_name: str) -> list[str]:
        rows = self.con.execute(
            "SELECT DISTINCT gene_symbol FROM pathways WHERE pathway_name = ?",
            [pathway_name],
        ).fetchall()
        return [r[0] for r in rows]

    def all_pathway_gene_counts(self) -> dict[str, int]:
        rows = self.con.execute(
            "SELECT pathway_name, COUNT(DISTINCT gene_symbol) "
            "FROM pathways GROUP BY pathway_name",
        ).fetchall()
        return {r[0]: r[1] for r in rows}

    def all_go_term_gene_counts(self) -> dict[str, int]:
        rows = self.con.execute(
            "SELECT go_id, COUNT(DISTINCT gene_symbol) "
            "FROM gene_go_terms GROUP BY go_id",
        ).fetchall()
        return {r[0]: r[1] for r in rows}

    def total_gene_count(self) -> int:
        row = self.con.execute(
            "SELECT COUNT(DISTINCT gene_symbol) FROM genes",
        ).fetchone()
        return row[0] if row else 0

    def total_pathway_gene_count(self) -> int:
        row = self.con.execute(
            "SELECT COUNT(DISTINCT gene_symbol) FROM pathways",
        ).fetchone()
        return row[0] if row else 0

    def total_go_gene_count(self) -> int:
        row = self.con.execute(
            "SELECT COUNT(DISTINCT gene_symbol) FROM gene_go_terms",
        ).fetchone()
        return row[0] if row else 0

    def gene_evidence(self, gene: str) -> dict:
        row = self.con.execute(
            "SELECT evidence, mention_count FROM genes WHERE gene_symbol = ?",
            [gene],
        ).fetchone()
        if row:
            return {"evidence": row[0], "mention_count": row[1]}
        return {"evidence": None, "mention_count": 0}

    def go_term_name(self, go_id: str) -> str:
        row = self.con.execute(
            "SELECT go_term FROM go_terms WHERE go_id = ?",
            [go_id],
        ).fetchone()
        return row[0] if row else go_id

    def all_organ_counts(self) -> dict[str, int]:
        """Count how many genes are annotated to each organ in the KB."""
        rows = self.con.execute(
            "SELECT organ, COUNT(*) FROM ("
            "  SELECT UNNEST(organs) AS organ FROM genes "
            "  WHERE organs IS NOT NULL"
            ") GROUP BY organ",
        ).fetchall()
        return {r[0]: r[1] for r in rows}


# ---------------------------------------------------------------------------
# BH FDR correction
# ---------------------------------------------------------------------------

def benjamini_hochberg(pvalues: list[float]) -> list[float]:
    """Benjamini-Hochberg FDR correction. Returns adjusted p-values."""
    n = len(pvalues)
    if n == 0:
        return []
    indexed = sorted(enumerate(pvalues), key=lambda x: x[1])
    adjusted = [0.0] * n
    prev = 1.0
    for rank_minus_1 in range(n - 1, -1, -1):
        orig_idx, pval = indexed[rank_minus_1]
        rank = rank_minus_1 + 1
        adj = min(prev, pval * n / rank)
        adjusted[orig_idx] = adj
        prev = adj
    return adjusted


# ---------------------------------------------------------------------------
# Pathway enrichment (Fisher's exact test)
# ---------------------------------------------------------------------------

def enrich_pathways(
    responsive_genes: list[str],
    kb: ToxKBQuerier,
    fdr_cutoff: float = 0.05,
) -> list[dict]:
    """
    Over-representation analysis for pathways using Fisher's exact test.

    For each pathway with at least one responsive gene, build a 2x2 table:
                      In pathway    Not in pathway
    Responsive           a              b
    Not responsive       c              d

    Background = all genes in the KB (genes table).
    """
    responsive_set = set(responsive_genes)
    n_responsive = len(responsive_set)

    pathway_counts = kb.all_pathway_gene_counts()
    bg_total = kb.total_gene_count()

    # Map each responsive gene to its pathways
    gene_to_pathways: dict[str, set[str]] = {}
    for gene in responsive_set:
        pathways = kb.gene_pathways(gene)
        if pathways:
            gene_to_pathways[gene] = {p["pathway_name"] for p in pathways}

    # Count responsive genes per pathway
    pathway_responsive: dict[str, list[str]] = {}
    for gene, pnames in gene_to_pathways.items():
        for pname in pnames:
            pathway_responsive.setdefault(pname, []).append(gene)

    results = []
    pvals = []

    for pathway_name, genes_in_pathway in pathway_responsive.items():
        a = len(genes_in_pathway)
        if a < 2:
            continue  # skip singletons
        b = n_responsive - a
        pathway_size = pathway_counts.get(pathway_name, 0)
        c = pathway_size - a
        d = bg_total - pathway_size - b
        if c < 0:
            c = 0
        if d < 0:
            d = 0

        _, pval = fisher_exact([[a, b], [c, d]], alternative="greater")

        results.append({
            "pathway_name": pathway_name,
            "overlap_genes": sorted(genes_in_pathway),
            "overlap_count": a,
            "pathway_size": pathway_size,
            "pvalue": pval,
        })
        pvals.append(pval)

    # BH correction
    if pvals:
        fdrs = benjamini_hochberg(pvals)
        for i, r in enumerate(results):
            r["fdr"] = fdrs[i]
    else:
        for r in results:
            r["fdr"] = 1.0

    # Filter and sort
    results = [r for r in results if r["fdr"] < fdr_cutoff]
    results.sort(key=lambda x: x["pvalue"])
    return results


# ---------------------------------------------------------------------------
# GO term enrichment (Fisher's exact test)
# ---------------------------------------------------------------------------

def enrich_go_terms(
    responsive_genes: list[str],
    kb: ToxKBQuerier,
    fdr_cutoff: float = 0.05,
) -> list[dict]:
    """
    Over-representation analysis for GO terms using Fisher's exact test.
    Same logic as enrich_pathways but against gene_go_terms table.
    """
    responsive_set = set(responsive_genes)
    n_responsive = len(responsive_set)

    go_counts = kb.all_go_term_gene_counts()
    bg_total = kb.total_gene_count()

    # Map each responsive gene to its GO terms
    gene_to_go: dict[str, set[str]] = {}
    for gene in responsive_set:
        terms = kb.gene_go_terms(gene)
        if terms:
            gene_to_go[gene] = {t["go_id"] for t in terms}

    # Count responsive genes per GO term
    go_responsive: dict[str, list[str]] = {}
    for gene, go_ids in gene_to_go.items():
        for go_id in go_ids:
            go_responsive.setdefault(go_id, []).append(gene)

    results = []
    pvals = []

    for go_id, genes_in_term in go_responsive.items():
        a = len(genes_in_term)
        if a < 2:
            continue
        b = n_responsive - a
        term_size = go_counts.get(go_id, 0)
        c = term_size - a
        d = bg_total - term_size - b
        if c < 0:
            c = 0
        if d < 0:
            d = 0

        _, pval = fisher_exact([[a, b], [c, d]], alternative="greater")

        go_name = kb.go_term_name(go_id)
        results.append({
            "go_id": go_id,
            "go_term": go_name,
            "overlap_genes": sorted(genes_in_term),
            "overlap_count": a,
            "term_size": term_size,
            "pvalue": pval,
        })
        pvals.append(pval)

    if pvals:
        fdrs = benjamini_hochberg(pvals)
        for i, r in enumerate(results):
            r["fdr"] = fdrs[i]
    else:
        for r in results:
            r["fdr"] = 1.0

    results = [r for r in results if r["fdr"] < fdr_cutoff]
    results.sort(key=lambda x: x["pvalue"])
    return results


# ---------------------------------------------------------------------------
# BMD-ranked GO gene sets and individual genes (NIEHS report tables)
# ---------------------------------------------------------------------------
# These two functions produce the data for the "Gene Set Benchmark Dose
# Analysis" and "Gene Benchmark Dose Analysis" report sections, matching
# the methodology described in NIEHS Report 10 (PFHxSAm, NBK589955).
#
# The NIEHS filtering criteria are:
#   - Fold change magnitude > 2  (|log2FC| not used — raw fold change)
#   - Goodness-of-fit p-value > 0.1  (model fits the data adequately)
#   - BMDU/BMDL ratio ≤ 40  (BMD estimate is precise enough to be useful)
#
# For gene sets: a GO term is "active" if it has ≥ 3 genes passing all
# filters.  The set's BMD is the median of its member genes' BMDs.
# For individual genes with multiple probes: the median BMD is used.

def _apply_niehs_gene_filters(df: pd.DataFrame) -> pd.DataFrame:
    """
    Apply the standard NIEHS quality filters to a gene-level BMD dataframe.

    These filters remove low-quality or imprecise BMD estimates, keeping
    only genes whose dose-response modeling meets minimum quality standards.
    The same filter set is used for both gene-level and gene-set-level
    ranking, ensuring consistency between the two report tables.

    Filters applied (in order):
      1. Fold change magnitude > 2  — genes must show a biologically
         meaningful change, not just a statistically detectable one.
      2. Goodness-of-fit p > 0.1   — the fitted dose-response model must
         adequately describe the data (p ≤ 0.1 suggests poor fit).
      3. BMDU/BMDL ratio ≤ 40      — the confidence interval around the
         BMD must be reasonably narrow; ratios above 40 indicate the BMD
         is too uncertain to rank on.

    Args:
        df: DataFrame with columns gene_symbol, bmd, and optionally
            fold_change, gof_p (goodness-of-fit p-value), bmdu, bmdl.

    Returns:
        Filtered copy of df with only rows passing all applicable filters.
        Filters are skipped if the corresponding column is absent (e.g.,
        if no fold_change column exists, that filter is not applied).
    """
    filtered = df.copy()

    # Filter 1: fold change magnitude > 2
    # The NIEHS report uses absolute fold change (not log2), so a value
    # of 2 means the gene is at least 2x up or 0.5x down vs. control.
    if "fold_change" in filtered.columns:
        fc = pd.to_numeric(filtered["fold_change"], errors="coerce")
        filtered = filtered[fc.abs() > 2].copy()

    # Filter 2: goodness-of-fit p-value > 0.1
    # A low p-value means the model is a poor fit to the observed data.
    # We keep only genes where the model fits well (p > 0.1).
    if "gof_p" in filtered.columns:
        gof = pd.to_numeric(filtered["gof_p"], errors="coerce")
        filtered = filtered[gof > 0.1].copy()

    # Filter 3: BMDU/BMDL ratio ≤ 40
    # This measures the width of the BMD confidence interval.  Ratios
    # above 40 mean the upper and lower bounds are so far apart that
    # the point estimate is essentially meaningless.
    if "bmdu" in filtered.columns and "bmdl" in filtered.columns:
        bmdu = pd.to_numeric(filtered["bmdu"], errors="coerce")
        bmdl = pd.to_numeric(filtered["bmdl"], errors="coerce")
        # Avoid division by zero — if BMDL is 0, the ratio is infinite
        ratio = bmdu / bmdl.replace(0, float("nan"))
        filtered = filtered[ratio <= 40].copy()

    return filtered


def rank_go_sets_by_bmd(
    df: pd.DataFrame,
    kb: ToxKBQuerier,
    top_n: int = 10,
) -> list[dict]:
    """
    Rank GO Biological Process gene sets by the potency of perturbation.

    This produces the data for the "Gene Set Benchmark Dose Analysis" table
    in the NIEHS report.  For each GO term that has enough passing genes,
    we compute the median BMD across its member genes — lower median BMD
    means the gene set was perturbed at a lower dose, i.e., more potent.

    The algorithm:
      1. Apply NIEHS quality filters to the gene-level data.
      2. For each GO term in the knowledge base, find which of the
         filtered genes are annotated to it.
      3. Keep only GO terms with ≥ 3 passing genes (the "active" threshold).
      4. Compute median BMD and median BMDL for each active GO term.
      5. Determine the predominant direction (up/down) of the member genes.
      6. Sort by median BMD ascending and return the top N.

    Args:
        df:    Gene-level BMD DataFrame.  Required columns: gene_symbol, bmd.
               Optional: bmdl, fold_change, direction, gof_p, bmdu.
        kb:    ToxKBQuerier connected to bmdx.duckdb — used to look up
               which genes belong to which GO terms.
        top_n: How many top gene sets to return (default 10, matching NIEHS).

    Returns:
        List of dicts, each representing one GO gene set, sorted by
        median BMD ascending:
          [{rank, go_id, go_term, bmd_median, bmdl_median, n_genes,
            genes, direction}, ...]
    """
    # Step 1: filter genes to only those passing NIEHS quality criteria
    filtered = _apply_niehs_gene_filters(df)

    if filtered.empty:
        return []

    # Build lookup dicts from the filtered data for fast access
    gene_bmd = dict(zip(filtered["gene_symbol"], filtered["bmd"]))
    gene_bmdl = {}
    if "bmdl" in filtered.columns:
        gene_bmdl = dict(zip(
            filtered["gene_symbol"],
            pd.to_numeric(filtered["bmdl"], errors="coerce"),
        ))
    gene_dir = {}
    if "direction" in filtered.columns:
        gene_dir = dict(zip(filtered["gene_symbol"], filtered["direction"]))

    # Step 2: for each filtered gene, find its GO term annotations
    # Build a reverse map: go_id → [list of passing genes]
    go_to_genes: dict[str, list[str]] = {}
    for gene in filtered["gene_symbol"].unique():
        terms = kb.gene_go_terms(gene)
        for t in terms:
            go_to_genes.setdefault(t["go_id"], []).append(gene)

    # Step 3 & 4: compute stats for each active GO term (≥ 3 genes)
    results = []
    for go_id, genes in go_to_genes.items():
        if len(genes) < 3:
            continue  # skip terms with too few passing genes

        # Median BMD across member genes
        bmds = [gene_bmd[g] for g in genes if g in gene_bmd]
        if not bmds:
            continue
        bmd_median = float(pd.Series(bmds).median())

        # Median BMDL (if available)
        bmdls = [gene_bmdl[g] for g in genes if g in gene_bmdl and pd.notna(gene_bmdl[g])]
        bmdl_median = float(pd.Series(bmdls).median()) if bmdls else None

        # Step 5: predominant direction of member genes
        dirs = [str(gene_dir.get(g, "")).lower() for g in genes]
        up_count = sum(1 for d in dirs if d == "up")
        down_count = sum(1 for d in dirs if d == "down")
        if up_count > down_count:
            direction = "Up"
        elif down_count > up_count:
            direction = "Down"
        elif up_count == 0 and down_count == 0:
            direction = "N/A"
        else:
            direction = "Mixed"

        go_term_name = kb.go_term_name(go_id)

        results.append({
            "go_id": go_id,
            "go_term": go_term_name,
            "bmd_median": round(bmd_median, 3),
            "bmdl_median": round(bmdl_median, 3) if bmdl_median is not None else None,
            "n_genes": len(genes),
            "genes": sorted(genes),
            "direction": direction,
        })

    # Step 6: sort by median BMD ascending (most potent first) and rank
    results.sort(key=lambda x: x["bmd_median"])
    for i, r in enumerate(results[:top_n], 1):
        r["rank"] = i

    return results[:top_n]


def rank_genes_by_bmd(
    df: pd.DataFrame,
    top_n: int = 10,
) -> list[dict]:
    """
    Rank individual genes by the potency of perturbation (BMD).

    This produces the data for the "Gene Benchmark Dose Analysis" table
    in the NIEHS report.  Genes with the lowest BMD were perturbed at
    the lowest dose, indicating they are the most sensitive responders.

    For genes with multiple probes (multiple rows in the CSV), the median
    BMD across probes is used — this matches the NIEHS methodology where
    probe-level results are collapsed to gene-level by taking the median.

    The algorithm:
      1. Apply NIEHS quality filters.
      2. Group by gene_symbol and take the median BMD (and median BMDL,
         fold_change) across probes for the same gene.
      3. Sort by median BMD ascending.
      4. Return the top N with full metadata.

    Args:
        df:    Gene-level BMD DataFrame.  Required columns: gene_symbol, bmd.
               Optional: bmdl, bmdu, fold_change, direction, best_model,
               full_name, gof_p.
        top_n: How many top genes to return (default 10, matching NIEHS).

    Returns:
        List of dicts, each representing one gene, sorted by BMD ascending:
          [{rank, gene_symbol, full_name, bmd, bmdl, bmdu, direction,
            fold_change}, ...]
    """
    # Step 1: apply NIEHS quality filters
    filtered = _apply_niehs_gene_filters(df)

    if filtered.empty:
        return []

    # Step 2: collapse multiple probes per gene to median values.
    # We group by gene_symbol and aggregate numeric columns with median,
    # and non-numeric columns (direction, full_name) by taking the first.
    agg_spec = {"bmd": "median"}

    # Add optional numeric columns to the aggregation
    for col in ("bmdl", "bmdu", "fold_change"):
        if col in filtered.columns:
            filtered[col] = pd.to_numeric(filtered[col], errors="coerce")
            agg_spec[col] = "median"

    # Non-numeric columns: keep the first value per gene
    for col in ("direction", "full_name", "best_model"):
        if col in filtered.columns:
            agg_spec[col] = "first"

    grouped = filtered.groupby("gene_symbol", as_index=False).agg(agg_spec)

    # Step 3: sort by BMD ascending (most sensitive gene first)
    grouped = grouped.sort_values("bmd").reset_index(drop=True)

    # Step 4: build result dicts with full metadata
    results = []
    for i, row in grouped.head(top_n).iterrows():
        entry = {
            "rank": len(results) + 1,
            "gene_symbol": row["gene_symbol"],
            "full_name": row.get("full_name", "") or "",
            "bmd": round(float(row["bmd"]), 3),
            "bmdl": round(float(row["bmdl"]), 3) if "bmdl" in row and pd.notna(row.get("bmdl")) else None,
            "bmdu": round(float(row["bmdu"]), 3) if "bmdu" in row and pd.notna(row.get("bmdu")) else None,
            "direction": row.get("direction", "N/A") or "N/A",
            "fold_change": round(float(row["fold_change"]), 2) if "fold_change" in row and pd.notna(row.get("fold_change")) else None,
        }
        results.append(entry)

    return results


# ---------------------------------------------------------------------------
# GO term definition fetcher — EBI QuickGO API
# ---------------------------------------------------------------------------
#
# The NIEHS Report 10 includes dense 9pt blocks of GO term descriptions
# after each gene set table.  Each entry shows the bold GO ID + term name,
# followed by the Gene Ontology biological process definition.
#
# We fetch definitions from the EBI QuickGO REST API (no API key needed,
# supports batch queries up to 200 IDs).  A module-level cache avoids
# re-fetching the same terms when processing multiple organ×sex CSVs.
# ---------------------------------------------------------------------------

# Module-level caches — persist for the lifetime of the server process.
# Keys are GO IDs (str) and gene symbols (str), respectively.
_go_description_cache: dict[str, dict] = {}
_gene_description_cache: dict[str, dict] = {}

logger = logging.getLogger(__name__)


def fetch_go_descriptions(
    go_ids: list[str],
    kb: ToxKBQuerier | None = None,
) -> list[dict]:
    """
    Fetch GO term definitions from the EBI QuickGO REST API.

    For each GO ID, returns a dict with go_id, name, and definition.
    Results are cached in a module-level dict so repeated calls for the
    same term (e.g., across liver_male and liver_female) are free.

    The QuickGO endpoint accepts up to 200 comma-separated IDs in a single
    GET request.  We typically only need 10 (top gene sets), so one call
    suffices.

    Args:
        go_ids: List of GO IDs to fetch (e.g., ["GO:0006355", "GO:0007165"]).
        kb:     Optional ToxKBQuerier — used as fallback if the API fails.
                Provides the GO term name from the local DuckDB, even though
                it doesn't have the full definition text.

    Returns:
        List of dicts in the same order as go_ids:
          [{"go_id": "GO:0006355",
            "name": "regulation of DNA-templated transcription",
            "definition": "Any process that modulates the frequency..."}, ...]
        If the API fails for a term, definition will be empty string.
    """
    if not go_ids:
        return []

    # Separate cached vs. uncached IDs
    uncached = [gid for gid in go_ids if gid not in _go_description_cache]

    if uncached:
        try:
            # QuickGO REST API — batch query for term definitions
            # Docs: https://www.ebi.ac.uk/QuickGO/api/index.html
            ids_param = ",".join(uncached)
            url = f"https://www.ebi.ac.uk/QuickGO/services/ontology/go/terms/{ids_param}"
            resp = requests.get(
                url,
                headers={"Accept": "application/json"},
                timeout=15,
            )
            resp.raise_for_status()
            data = resp.json()

            # Parse the response — each result has id, name, definition.text
            for result in data.get("results", []):
                gid = result.get("id", "")
                name = result.get("name", "")
                defn = ""
                if result.get("definition"):
                    defn = result["definition"].get("text", "")
                _go_description_cache[gid] = {
                    "go_id": gid,
                    "name": name,
                    "definition": defn,
                }

        except Exception as e:
            # API failure is non-fatal — we'll fall back to local names
            logger.warning("QuickGO API failed: %s", e)

    # Build the result list, using cache hits and local fallback for misses
    results = []
    for gid in go_ids:
        if gid in _go_description_cache:
            results.append(_go_description_cache[gid])
        else:
            # Fallback: use the GO term name from local DuckDB if available
            name = ""
            if kb:
                name = kb.go_term_name(gid)
            entry = {"go_id": gid, "name": name, "definition": ""}
            _go_description_cache[gid] = entry
            results.append(entry)

    return results


# ---------------------------------------------------------------------------
# Gene description fetcher — MyGene.info API
# ---------------------------------------------------------------------------
#
# The NIEHS Report 10 includes dense 9pt blocks of gene descriptions
# after each gene table.  Each entry shows the bold italicized gene symbol,
# followed by a functional description (typically from NCBI Gene / UniProt).
#
# We use the MyGene.info API (https://mygene.info/) which:
#   - Requires no API key
#   - Supports batch POST queries (up to 1000 genes at once)
#   - Returns name (short) and summary (full functional description)
#   - Has generous rate limits (1000 req/s)
# ---------------------------------------------------------------------------

def fetch_gene_descriptions(
    gene_symbols: list[str],
    species: str = "rat",
) -> list[dict]:
    """
    Fetch gene functional descriptions from the MyGene.info API.

    For each gene symbol, returns a dict with gene_symbol, name (short),
    and description (full functional summary).  Results are cached in a
    module-level dict.

    Strategy: gene summaries (the full paragraph-length descriptions used
    in NIEHS Report 10's dense 9pt blocks) are only available for human
    genes in NCBI Gene.  Since rat gene symbols map to human orthologs
    with the same symbol (case-insensitive), we:
      1. Query rat genes → get the short `name` field
      2. Query human orthologs (uppercased symbol) → get the `summary` field
      3. Merge: rat name + human summary per gene

    Uses the MyGene.info batch POST /query endpoint (up to 1000 genes/request).

    Args:
        gene_symbols: List of gene symbols (e.g., ["Nfe2l2", "Cyp1a1"]).
        species:      Species name for the rat query (default "rat").

    Returns:
        List of dicts in the same order as gene_symbols:
          [{"gene_symbol": "Nfe2l2",
            "name": "NFE2 like bZIP transcription factor 2",
            "description": "This gene encodes a transcription factor..."}, ...]
        If the API fails for a gene, name and description will be empty.
    """
    if not gene_symbols:
        return []

    # Separate cached vs. uncached symbols
    uncached = [g for g in gene_symbols if g not in _gene_description_cache]

    if uncached:
        # Step 1: fetch rat gene names (short description)
        rat_names = _mygene_batch_query(uncached, species, "symbol,name")

        # Step 2: fetch human orthologs for full summaries.
        # Rat symbols like "Nfe2l2" map to human "NFE2L2" (uppercased).
        human_symbols = [g.upper() for g in uncached]
        human_data = _mygene_batch_query(human_symbols, "human", "symbol,name,summary")

        # Step 3: merge rat name + human summary per gene
        for i, symbol in enumerate(uncached):
            rat_info = rat_names.get(symbol, {})
            human_info = human_data.get(symbol.upper(), {})

            _gene_description_cache[symbol] = {
                "gene_symbol": symbol,
                # Prefer rat name (species-specific); fall back to human
                "name": rat_info.get("name", "") or human_info.get("name", ""),
                # Summary only available from human orthologs
                "description": human_info.get("summary", ""),
            }

    # Build the result list, using cache hits and empty fallback for misses
    results = []
    for symbol in gene_symbols:
        if symbol in _gene_description_cache:
            results.append(_gene_description_cache[symbol])
        else:
            entry = {"gene_symbol": symbol, "name": "", "description": ""}
            _gene_description_cache[symbol] = entry
            results.append(entry)

    return results


def _mygene_batch_query(
    symbols: list[str],
    species: str,
    fields: str,
) -> dict[str, dict]:
    """
    Batch query MyGene.info for gene metadata.

    Uses the POST /v3/query endpoint which accepts multiple gene symbols
    and returns results in a single request.

    Args:
        symbols: List of gene symbols to query.
        species: Species name ("rat", "human", etc.).
        fields:  Comma-separated field list (e.g., "symbol,name,summary").

    Returns:
        Dict mapping query symbol → {name, summary, symbol, ...}.
        Symbols that weren't found are omitted from the result.
    """
    if not symbols:
        return {}

    try:
        resp = requests.post(
            "https://mygene.info/v3/query",
            json={
                "q": symbols,
                "scopes": "symbol",
                "species": species,
                "fields": fields,
            },
            timeout=15,
        )
        resp.raise_for_status()
        results_raw = resp.json()

        # Map query string → result dict (skip "notfound" entries)
        out = {}
        for item in results_raw:
            if isinstance(item, dict) and not item.get("notfound"):
                query = item.get("query", "")
                out[query] = item
        return out

    except Exception as e:
        logger.warning("MyGene.info API failed for %s: %s", species, e)
        return {}


# ---------------------------------------------------------------------------
# BMD-ordered pathway narrative
# ---------------------------------------------------------------------------

def compute_bmd_ordering(
    df: pd.DataFrame,
    enriched_pathways: list[dict],
) -> list[dict]:
    """
    For each enriched pathway, compute median/min BMD across its genes
    and summarize direction. Returns list sorted by median BMD ascending.
    """
    gene_bmd = dict(zip(df["gene_symbol"], df["bmd"]))
    gene_dir = {}
    if "direction" in df.columns:
        gene_dir = dict(zip(df["gene_symbol"], df["direction"]))

    ordered = []
    for pw in enriched_pathways:
        genes = pw["overlap_genes"]
        bmds = [gene_bmd[g] for g in genes if g in gene_bmd]
        if not bmds:
            continue

        dirs = [gene_dir.get(g, "unknown") for g in genes if g in gene_bmd]
        up_count = sum(1 for d in dirs if str(d).lower() == "up")
        down_count = sum(1 for d in dirs if str(d).lower() == "down")
        if up_count > down_count:
            direction = "mostly UP"
        elif down_count > up_count:
            direction = "mostly DOWN"
        else:
            direction = "mixed"

        ordered.append({
            "pathway": pw["pathway_name"],
            "median_bmd": round(float(pd.Series(bmds).median()), 3),
            "mean_bmd": round(float(pd.Series(bmds).mean()), 3),
            "min_bmd": round(float(min(bmds)), 3),
            "gene_count": len(genes),
            "direction": direction,
            "genes": genes,
            "pvalue": pw["pvalue"],
            "fdr": pw["fdr"],
        })

    ordered.sort(key=lambda x: x["median_bmd"])
    return ordered


# ---------------------------------------------------------------------------
# Organ signature
# ---------------------------------------------------------------------------

def compute_organ_signature(
    responsive_genes: list[str],
    kb: ToxKBQuerier,
) -> dict[str, dict]:
    """
    Score organs by enrichment of responsive genes relative to KB background.
    Returns {organ: {score, count, genes}} sorted by score descending.
    """
    # Count organ occurrences across responsive genes
    organ_genes: dict[str, list[str]] = {}
    for gene in responsive_genes:
        organs = kb.gene_organs(gene)
        for organ in organs:
            organ_genes.setdefault(organ, []).append(gene)

    # Background frequencies
    bg_counts = kb.all_organ_counts()
    total_genes_in_kb = kb.total_gene_count()
    n_responsive = len(responsive_genes)

    results = {}
    for organ, genes in organ_genes.items():
        count = len(genes)
        bg_freq = bg_counts.get(organ, 1) / max(total_genes_in_kb, 1)
        observed_freq = count / max(n_responsive, 1)
        enrichment = observed_freq / bg_freq if bg_freq > 0 else 0.0
        results[organ] = {
            "score": round(enrichment, 2),
            "count": count,
            "genes": sorted(genes),
        }

    # Sort by score descending
    return dict(sorted(results.items(), key=lambda x: x[1]["score"], reverse=True))


# ---------------------------------------------------------------------------
# Structured analysis result
# ---------------------------------------------------------------------------

@dataclass
class AnalysisResult:
    """Structured results from dose-response analysis."""
    df: pd.DataFrame
    responsive_genes: list
    bmd_min: float
    bmd_max: float
    n_up: int
    n_down: int
    pw_enriched: list
    go_enriched: list
    bmd_ordered: list
    organ_sig: dict
    gene_literature: list   # [{gene, evidence, mention_count, n_papers, organs, claims}]
    top_papers: list        # [{paper_id, title, year, citation_count, genes}]
    fdr_cutoff: float = 0.05
    context_text: str = ""
    llm_narrative: str = ""
    narrative_runs: list = field(default_factory=list)
    concordance_text: str = ""
    concordance_model: str = ""


# ---------------------------------------------------------------------------
# Data gathering (separated from formatting)
# ---------------------------------------------------------------------------

def analyze(
    df: pd.DataFrame,
    kb: ToxKBQuerier,
    fdr_cutoff: float = 0.05,
) -> AnalysisResult:
    """
    Run all analyses on dose-response data and return structured results.
    """
    responsive_genes = list(df["gene_symbol"])
    bmd_min = float(df["bmd"].min())
    bmd_max = float(df["bmd"].max())

    n_up = n_down = 0
    if "direction" in df.columns:
        n_up = int((df["direction"].str.lower() == "up").sum())
        n_down = int((df["direction"].str.lower() == "down").sum())

    print("  Running pathway enrichment...")
    pw_enriched = enrich_pathways(responsive_genes, kb, fdr_cutoff)
    print(f"  {len(pw_enriched)} significant pathways (FDR < {fdr_cutoff})")

    print("  Running GO term enrichment...")
    go_enriched = enrich_go_terms(responsive_genes, kb, fdr_cutoff)
    print(f"  {len(go_enriched)} significant GO terms (FDR < {fdr_cutoff})")

    print("  Computing BMD ordering...")
    bmd_ordered = compute_bmd_ordering(df, pw_enriched)

    print("  Computing organ signature...")
    organ_sig = compute_organ_signature(responsive_genes, kb)

    # Per-gene literature context
    print("  Gathering per-gene literature...")
    gene_literature = []
    gene_priority = []
    for gene in responsive_genes:
        ev = kb.gene_evidence(gene)
        n_papers = len(kb.gene_papers(gene))
        gene_priority.append((gene, ev.get("evidence"), ev.get("mention_count") or 0, n_papers))

    evidence_order = {"consensus": 0, "moderate": 1, "single": 2, None: 3}
    gene_priority.sort(key=lambda x: (evidence_order.get(x[1], 3), -x[2]))

    for gene, evidence, mention_count, n_papers in gene_priority[:15]:
        organs = kb.gene_organs(gene)
        claims = kb.gene_claims(gene)
        seen_claims = set()
        unique_claims = []
        for c in claims:
            claim_text = c["claim"][:200]
            if claim_text not in seen_claims:
                seen_claims.add(claim_text)
                unique_claims.append(c)
        gene_literature.append({
            "gene": gene,
            "evidence": evidence,
            "mention_count": mention_count,
            "n_papers": n_papers,
            "organs": organs,
            "claims": unique_claims[:3],
        })

    # Multi-gene papers
    print("  Finding multi-gene papers...")
    paper_gene_overlap: dict[str, dict] = {}
    for gene in responsive_genes:
        for paper in kb.gene_papers(gene):
            pid = paper["paper_id"]
            if pid not in paper_gene_overlap:
                paper_gene_overlap[pid] = {
                    "title": paper["title"],
                    "year": paper["year"],
                    "citation_count": paper["citation_count"] or 0,
                    "genes": [],
                }
            paper_gene_overlap[pid]["genes"].append(gene)

    top_papers = [
        {"paper_id": pid, **info, "genes": sorted(set(info["genes"]))}
        for pid, info in paper_gene_overlap.items()
        if len(info["genes"]) >= 2
    ]
    top_papers.sort(key=lambda x: x["citation_count"], reverse=True)
    top_papers = top_papers[:10]

    result = AnalysisResult(
        df=df,
        responsive_genes=responsive_genes,
        bmd_min=bmd_min,
        bmd_max=bmd_max,
        n_up=n_up,
        n_down=n_down,
        pw_enriched=pw_enriched,
        go_enriched=go_enriched,
        bmd_ordered=bmd_ordered,
        organ_sig=organ_sig,
        gene_literature=gene_literature,
        top_papers=top_papers,
        fdr_cutoff=fdr_cutoff,
    )
    result.context_text = format_context_text(result)
    return result


def format_context_text(result: AnalysisResult) -> str:
    """Format structured analysis results into a text block for the LLM."""
    lines = []
    fdr_cutoff = result.fdr_cutoff

    # --- Summary ---
    lines.append("=== DOSE-RESPONSE SUMMARY ===")
    lines.append(f"{len(result.responsive_genes)} responsive genes, BMD range {result.bmd_min:.3g}-{result.bmd_max:.3g}")
    if result.n_up or result.n_down:
        lines.append(f"Direction: {result.n_up} up, {result.n_down} down")
    lines.append("")

    # --- Pathway Enrichment ---
    lines.append(f"=== PATHWAY ENRICHMENT (FDR < {fdr_cutoff}) ===")
    if result.pw_enriched:
        for i, pw in enumerate(result.pw_enriched[:20], 1):
            lines.append(
                f"{i}. {pw['pathway_name']} "
                f"(p={pw['pvalue']:.2e}, FDR={pw['fdr']:.2e}, "
                f"{pw['overlap_count']}/{pw['pathway_size']} genes)"
            )
            lines.append(f"   Genes: {', '.join(pw['overlap_genes'])}")
    else:
        lines.append("No significantly enriched pathways at this FDR threshold.")
    lines.append("")

    # --- GO Term Enrichment ---
    lines.append(f"=== GO TERM ENRICHMENT (FDR < {fdr_cutoff}, top 20) ===")
    if result.go_enriched:
        for i, gt in enumerate(result.go_enriched[:20], 1):
            lines.append(
                f"{i}. {gt['go_term']} [{gt['go_id']}] "
                f"(p={gt['pvalue']:.2e}, FDR={gt['fdr']:.2e}, "
                f"{gt['overlap_count']}/{gt['term_size']} genes)"
            )
            lines.append(f"   Genes: {', '.join(gt['overlap_genes'])}")
    else:
        lines.append("No significantly enriched GO terms at this FDR threshold.")
    lines.append("")

    # --- BMD Ordering ---
    lines.append("=== DOSE-ORDERED RESPONSE ===")
    if result.bmd_ordered:
        for entry in result.bmd_ordered:
            lines.append(
                f"BMD {entry['min_bmd']:.3g}-{entry['mean_bmd']:.3g}: "
                f"{entry['pathway']} "
                f"(median BMD {entry['median_bmd']}, "
                f"{entry['gene_count']} genes, {entry['direction']})"
            )
    else:
        lines.append("No pathway-level BMD ordering available.")
    lines.append("")

    # --- Organ Signature ---
    lines.append("=== ORGAN SIGNATURE ===")
    if result.organ_sig:
        for organ, info in result.organ_sig.items():
            genes_str = ", ".join(info["genes"][:8])
            if len(info["genes"]) > 8:
                genes_str += f", ... (+{len(info['genes']) - 8} more)"
            lines.append(
                f"{organ.title()}: {info['score']}x enriched "
                f"({genes_str})"
            )
    else:
        lines.append("No organ annotations found for responsive genes.")
    lines.append("")

    # --- Literature Context ---
    lines.append("=== LITERATURE CONTEXT ===")
    for gl in result.gene_literature:
        ev_label = f"{gl['evidence']} gene" if gl['evidence'] else "not in KB"
        organs_str = ", ".join(gl['organs']) if gl['organs'] else "not annotated"
        lines.append(f"## {gl['gene']} ({gl['n_papers']} papers, {ev_label})")
        lines.append(f"  Organs: {organs_str}")
        if gl['claims']:
            lines.append("  Key claims:")
            for c in gl['claims']:
                claim_text = c["claim"][:200]
                year = c.get("year", "")
                title_short = (c.get("paper_title") or "")[:60]
                lines.append(f'  - "{claim_text}" ({title_short}, {year})')
        lines.append("")

    # --- Top Relevant Papers ---
    lines.append("=== TOP RELEVANT PAPERS ===")
    for paper in result.top_papers:
        genes_str = ", ".join(paper["genes"])
        lines.append(
            f"- {paper['title']} ({paper['year']}, "
            f"cited {paper['citation_count']}x) "
            f"[genes: {genes_str}]"
        )

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Build interpretation context (convenience wrapper)
# ---------------------------------------------------------------------------

def build_interpretation_context(
    df: pd.DataFrame,
    kb: ToxKBQuerier,
    fdr_cutoff: float = 0.05,
) -> str:
    """
    Assemble all KB evidence into a structured text block for the LLM.
    """
    return analyze(df, kb, fdr_cutoff).context_text


# ---------------------------------------------------------------------------
# build_genomics_interpretation — bridge from web app genomics_sections format
# ---------------------------------------------------------------------------

def build_genomics_interpretation(
    genomics_section: dict,
    db_path: str = "bmdx.duckdb",
    fdr_cutoff: float = 0.05,
) -> dict:
    """
    Run the full interpretation pipeline on one organ×sex genomics section.

    Converts a genomics_section dict (as produced by _extract_genomics() in
    pool_orchestrator.py) into a DataFrame and runs the analyze() pipeline
    against bmdx.duckdb to produce pathway enrichment, GO enrichment,
    BMD-ordered pathways, organ signatures, and per-gene literature evidence.

    Uses 'all_genes' when available (full responsive gene list, better for
    Fisher's exact enrichment tests), falls back to 'top_genes' for backward
    compatibility with cached sessions that predate the all_genes addition.

    Args:
        genomics_section: One entry from genomics_sections dict, containing
            at minimum 'top_genes' (list of gene dicts with gene_symbol, bmd,
            bmdl, direction, fold_change) and optionally 'all_genes'.
        db_path: Path to the bmdx.duckdb knowledge base file.
        fdr_cutoff: FDR threshold for pathway/GO enrichment significance.

    Returns:
        dict with keys:
            context_text: Formatted multi-section text block for the LLM prompt
                (~200 lines covering pathway enrichment, GO enrichment,
                BMD ordering, organ signature, and literature context).
            analysis_result: Serializable subset of the AnalysisResult —
                pw_enriched, go_enriched, bmd_ordered, organ_sig,
                gene_literature lists/dicts for caching and inspection.
    """
    # Prefer all_genes (full list) over top_genes (truncated to 20) because
    # Fisher's exact test needs the complete responsive gene set to compute
    # meaningful enrichment p-values. top_genes alone would under-power the
    # test and miss real pathway signals.
    gene_list = genomics_section.get("all_genes") or genomics_section.get("top_genes", [])

    if not gene_list:
        return {
            "context_text": "(No responsive genes available for enrichment analysis.)",
            "analysis_result": {},
        }

    # Build the DataFrame that analyze() expects: gene_symbol, bmd, bmdl,
    # direction, fold_change columns.  Missing values are fine — analyze()
    # handles NaN BMDs gracefully for genes that passed prefilter but had
    # model failures.
    df = pd.DataFrame(gene_list)

    # analyze() requires at minimum a 'gene_symbol' column.  Rename if the
    # key happens to differ (defensive — current format already uses gene_symbol).
    if "gene_symbol" not in df.columns and "symbol" in df.columns:
        df = df.rename(columns={"symbol": "gene_symbol"})

    kb = ToxKBQuerier(db_path)
    try:
        result = analyze(df, kb, fdr_cutoff)
    finally:
        # ToxKBQuerier opens a duckdb connection — ensure it's closed even
        # if analyze() raises, to avoid leaking file handles.
        kb.close()

    # Return the formatted context text plus a serializable snapshot of the
    # raw enrichment results.  The snapshot enables caching without re-running
    # the full pipeline and lets the frontend inspect enrichment details.
    return {
        "context_text": result.context_text,
        "analysis_result": {
            "n_responsive": len(result.responsive_genes),
            "bmd_range": [result.bmd_min, result.bmd_max],
            "n_up": result.n_up,
            "n_down": result.n_down,
            "pw_enriched": result.pw_enriched,
            "go_enriched": result.go_enriched,
            "bmd_ordered": result.bmd_ordered,
            "organ_sig": result.organ_sig,
            "gene_literature": result.gene_literature,
        },
    }


# ---------------------------------------------------------------------------
# LLM Synthesis
# ---------------------------------------------------------------------------

INTERPRETATION_SYSTEM = """You are a toxicogenomics expert analyst. You interpret dose-response
gene expression data (BMD analysis results) to predict organism-level and organ-level effects.

Your interpretations are grounded in the scientific literature and biological pathway knowledge
provided to you. You cite specific genes, pathways, and papers to support your conclusions.

You write in a clear, professional scientific style suitable for a regulatory toxicology report."""

INTERPRETATION_PROMPT = """Based on the following dose-response analysis results and literature context,
generate a comprehensive interpretation report.

{context}

---

Please generate a detailed interpretation covering these sections:

## 1. Biological Response Narrative
What biological processes are activated or suppressed at each dose level?
Order your narrative from lowest to highest dose, describing the progression of effects.

## 2. Organ-Level Prediction
Which organs are most likely affected by this exposure? Explain why based on the
gene expression patterns, organ-specific gene annotations, and literature evidence.

## 3. Mechanism of Action
Based on the enriched pathways and GO terms, what is the likely mechanism of toxicity?
Describe the molecular initiating event and key events in the adverse outcome pathway.

## 4. Protective vs. Adverse Responses
Distinguish between adaptive/protective responses (e.g., NRF2 antioxidant defense,
DNA repair, UPR) and responses indicating damage (e.g., apoptosis, inflammation,
fibrosis). At what dose does the transition from adaptive to adverse likely occur?

## 5. Literature Support
For each major conclusion, cite the specific papers and claims from the knowledge base
that support it. Note where your conclusions agree with or extend the literature.

## 6. Confidence Assessment
Rate your confidence in each major conclusion. Note which genes have strong literature
support (consensus genes) vs. those with limited evidence. Identify any novel findings
not well-represented in the existing literature."""


CONCORDANCE_SYSTEM = """You are a meta-analyst comparing multiple independent AI-generated \
toxicogenomics interpretations. Your task is to identify where different models agree, \
disagree, or produce unique insights. You are rigorous, balanced, and cite specific \
model outputs to support your conclusions."""

CONCORDANCE_PROMPT = """Below are {n_narratives} independent toxicogenomics interpretation \
narratives generated by {n_models} different AI models ({n_runs} runs each). Each narrative \
covers the same dose-response dataset but was generated independently.

{narratives_block}

---

Please perform a concordance analysis comparing these narratives across all 6 sections. \
For each section, identify what the models agree on, where they diverge, and any unique insights.

## 1. Biological Response Narrative — Concordance
- **Agreement**: What biological processes/dose-response patterns do most models identify?
- **Divergence**: Where do models disagree on processes, ordering, or significance?
- **Concordance rating**: (Strong / Moderate / Weak agreement)

## 2. Organ-Level Prediction — Concordance
- **Agreement**: Which organ predictions are consistent across models?
- **Divergence**: Where do organ predictions differ?
- **Concordance rating**: (Strong / Moderate / Weak agreement)

## 3. Mechanism of Action — Concordance
- **Agreement**: What molecular mechanisms are consistently identified?
- **Divergence**: Where do mechanistic explanations differ?
- **Concordance rating**: (Strong / Moderate / Weak agreement)

## 4. Protective vs. Adverse Responses — Concordance
- **Agreement**: Do models agree on the adaptive-to-adverse transition dose?
- **Divergence**: Where do models differ on protective vs. adverse classification?
- **Concordance rating**: (Strong / Moderate / Weak agreement)

## 5. Literature Support — Concordance
- **Agreement**: Which papers/claims are cited consistently across models?
- **Divergence**: Are there contradictory literature interpretations?
- **Concordance rating**: (Strong / Moderate / Weak agreement)

## 6. Confidence Assessment — Concordance
- **Agreement**: Do models agree on which findings are high vs. low confidence?
- **Divergence**: Where do confidence assessments differ?
- **Concordance rating**: (Strong / Moderate / Weak agreement)

## High-Confidence Findings
Claims or conclusions supported by 4 or more of the {n_models} models. These represent \
the most robust findings from this analysis.

## Divergent Findings
Claims where models meaningfully disagree. Explain the nature of each disagreement.

## Model-Specific Observations
Unique insights contributed by individual models that others did not identify. \
Note which model produced each observation.

## Overall Concordance Summary
Provide a brief overall assessment of how well the models agree, which sections show \
the strongest consensus, and which would benefit from further investigation."""


def synthesize_interpretation(
    context: str,
    endpoint: OllamaEndpoint | None = None,
) -> str:
    """Send the assembled context to the LLM for narrative synthesis."""
    if endpoint is None:
        endpoint = LOCAL_OLLAMA

    if not endpoint.is_available():
        print(f"  [WARNING] Ollama endpoint {endpoint.name} not available.")
        print("  Skipping LLM synthesis — returning structured context only.")
        return (
            "# Dose-Response Interpretation\n\n"
            "*LLM synthesis unavailable — structured analysis below.*\n\n"
            "```\n" + context + "\n```"
        )

    print(f"  Generating interpretation via {endpoint.name} ({endpoint.model})...")
    prompt = INTERPRETATION_PROMPT.format(context=context)

    # Use higher token limit for interpretation
    payload = {
        "model": endpoint.model,
        "prompt": prompt,
        "system": INTERPRETATION_SYSTEM,
        "stream": False,
        "options": {
            "temperature": 0.3,
            "num_predict": 4096,
        },
    }

    try:
        r = requests.post(
            f"{endpoint.url}/api/generate",
            json=payload,
            timeout=300,
        )
        r.raise_for_status()
        response = r.json().get("response", "")
        if response:
            return f"# Dose-Response Interpretation\n\n{response}"
        else:
            return (
                "# Dose-Response Interpretation\n\n"
                "*LLM returned empty response — structured analysis below.*\n\n"
                "```\n" + context + "\n```"
            )
    except Exception as e:
        print(f"  [ERROR] LLM synthesis failed: {e}")
        return (
            "# Dose-Response Interpretation\n\n"
            f"*LLM synthesis failed ({e}) — structured analysis below.*\n\n"
            "```\n" + context + "\n```"
        )


# ---------------------------------------------------------------------------
# Multi-model narrative generation
# ---------------------------------------------------------------------------

# Model-to-endpoint mapping
MODEL_ENDPOINTS = {
    "qwen2.5:14b": LOCAL_OLLAMA,
    "gemma2:9b": OllamaEndpoint(
        name="remote-6900xt",
        url="http://localhost:11435",
        model="gemma2:9b",
        weight=1,
    ),
    # Claude models resolved dynamically via AnthropicEndpoint
}


def _resolve_endpoint(model: str) -> OllamaEndpoint | AnthropicEndpoint:
    """Return the appropriate endpoint for a model name."""
    if model in MODEL_ENDPOINTS:
        return MODEL_ENDPOINTS[model]
    if model.startswith("claude-"):
        return AnthropicEndpoint(name=f"anthropic-{model}", model=model)
    # Fall back: try local Ollama with the model name overridden
    ep = OllamaEndpoint(
        name=LOCAL_OLLAMA.name,
        url=LOCAL_OLLAMA.url,
        model=model,
        weight=LOCAL_OLLAMA.weight,
    )
    return ep


def _ensure_model_pulled(endpoint: OllamaEndpoint) -> None:
    """Check if model is available on an Ollama endpoint; pull if missing."""
    try:
        r = requests.get(f"{endpoint.url}/api/tags", timeout=10)
        r.raise_for_status()
        models = [m["name"] for m in r.json().get("models", [])]
        # Ollama tags may include `:latest` suffix
        target = endpoint.model
        if any(target == m or target == m.split(":")[0] or m.startswith(target) for m in models):
            print(f"  [{endpoint.name}] Model {target} already present.")
            return
    except requests.RequestException as e:
        print(f"  [WARNING] Cannot check models on {endpoint.name}: {e}")
        return

    print(f"  [{endpoint.name}] Pulling {target}...")
    try:
        r = requests.post(
            f"{endpoint.url}/api/pull",
            json={"name": target, "stream": False},
            timeout=600,
        )
        r.raise_for_status()
        print(f"  [{endpoint.name}] Pull complete for {target}.")
    except requests.RequestException as e:
        print(f"  [WARNING] Pull failed for {target} on {endpoint.name}: {e}")


def generate_narrative(
    context: str,
    endpoint: OllamaEndpoint | AnthropicEndpoint,
    model_override: str | None = None,
    temperature: float = 0.3,
) -> tuple[str, float]:
    """
    Generate a single narrative from the given endpoint.
    Returns (narrative_text, elapsed_seconds).
    """
    prompt = INTERPRETATION_PROMPT.format(context=context)
    start = time.time()

    if isinstance(endpoint, AnthropicEndpoint):
        text = endpoint.generate(prompt, system=INTERPRETATION_SYSTEM,
                                 temperature=temperature)
    elif isinstance(endpoint, OllamaEndpoint):
        # Temporarily override model if needed
        original_model = endpoint.model
        if model_override:
            endpoint.model = model_override
        payload = {
            "model": endpoint.model,
            "prompt": prompt,
            "system": INTERPRETATION_SYSTEM,
            "stream": False,
            "options": {
                "temperature": temperature,
                "num_predict": 4096,
            },
        }
        try:
            r = requests.post(
                f"{endpoint.url}/api/generate",
                json=payload,
                timeout=300,
            )
            r.raise_for_status()
            text = r.json().get("response", "")
        except requests.RequestException as e:
            text = f"[ERROR] Generation failed: {e}"
        finally:
            endpoint.model = original_model
    else:
        text = "[ERROR] Unknown endpoint type"

    elapsed = time.time() - start
    return text, elapsed


def _sanitize_model_name(model: str) -> str:
    """Convert model name to a safe filename component (e.g. qwen2.5:14b -> qwen2.5-14b)."""
    return model.replace(":", "-").replace("/", "-")


def _write_per_model_output(
    result: AnalysisResult,
    runs: list[NarrativeRun],
    model: str,
    base_docx: str | None,
    base_md: str,
) -> None:
    """Write per-model .docx and .md files as each model's runs complete."""
    safe = _sanitize_model_name(model)
    md_path = Path(base_md)
    model_md = md_path.parent / f"{md_path.stem}_{safe}{md_path.suffix}"

    # Build markdown content
    lines = [f"# Narratives: {model}\n"]
    for run in runs:
        lines.append(f"## Run {run.run_index + 1} ({run.elapsed_seconds:.1f}s)\n")
        lines.append(run.narrative)
        lines.append("\n---\n")
    md_text = "\n".join(lines)
    model_md.parent.mkdir(parents=True, exist_ok=True)
    model_md.write_text(md_text)
    print(f"  Saved {model_md}")

    if base_docx:
        docx_path = Path(base_docx)
        model_docx = docx_path.parent / f"{docx_path.stem}_{safe}{docx_path.suffix}"
        _build_per_model_docx(result, runs, model, str(model_docx))


def _build_per_model_docx(
    result: AnalysisResult,
    runs: list[NarrativeRun],
    model: str,
    output_path: str,
) -> None:
    """Build a per-model .docx with the model's narrative runs."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from build_docx import add_heading, add_para

    doc = Document()
    title = doc.add_heading(f"Narratives: {model}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in runs:
        add_heading(doc, f"Run {run.run_index + 1} ({run.elapsed_seconds:.1f}s)", level=1)
        _render_narrative(doc, run.narrative)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  Saved {output_path}")


def generate_all_narratives(
    context: str,
    models: list[str],
    n_runs: int = 3,
    base_docx: str | None = None,
    base_md: str = "output/interpretation.md",
    result: AnalysisResult | None = None,
) -> list[NarrativeRun]:
    """
    Generate narratives from multiple models in parallel.
    Each model runs sequentially for n_runs, but models run concurrently.
    Per-model outputs are written as each model completes.
    """
    all_runs: list[NarrativeRun] = []
    lock = __import__("threading").Lock()

    # Pre-flight: ensure Ollama models are pulled
    for model in models:
        ep = _resolve_endpoint(model)
        if isinstance(ep, OllamaEndpoint):
            _ensure_model_pulled(ep)

    def _run_model(model: str) -> list[NarrativeRun]:
        ep = _resolve_endpoint(model)

        # Check availability
        if not ep.is_available():
            print(f"  [WARNING] {model} endpoint not available, skipping.")
            return []

        runs = []
        for i in range(n_runs):
            print(f"  [{model}] Starting run {i + 1}/{n_runs}...")
            text, elapsed = generate_narrative(context, ep, temperature=0.3)
            run = NarrativeRun(
                model=model,
                endpoint_name=ep.name,
                run_index=i,
                narrative=text,
                elapsed_seconds=elapsed,
            )
            runs.append(run)
            print(f"  [{model}] Run {i + 1} done in {elapsed:.1f}s")

        # Write per-model output as soon as this model finishes
        if runs and result is not None:
            _write_per_model_output(result, runs, model, base_docx, base_md)

        with lock:
            all_runs.extend(runs)
        return runs

    # One thread per model — all run concurrently
    with ThreadPoolExecutor(max_workers=len(models)) as executor:
        futures = {executor.submit(_run_model, m): m for m in models}
        for future in as_completed(futures):
            model = futures[future]
            try:
                future.result()
            except Exception as e:
                print(f"  [ERROR] {model} failed: {e}")

    return all_runs


# ---------------------------------------------------------------------------
# Concordance analysis
# ---------------------------------------------------------------------------

def run_concordance_analysis(
    narrative_runs: list[NarrativeRun],
    concordance_model: str = "claude-sonnet-4-6",
) -> tuple[str, str]:
    """
    Analyze concordance across narrative runs.
    Returns (concordance_text, model_used).
    """
    # Group runs by model
    from collections import OrderedDict
    model_runs: dict[str, list[NarrativeRun]] = OrderedDict()
    for run in narrative_runs:
        model_runs.setdefault(run.model, []).append(run)

    models = list(model_runs.keys())
    n_models = len(models)
    n_runs = max((len(runs) for runs in model_runs.values()), default=0)

    # Build narratives block with clear delimiters
    blocks = []
    for run in narrative_runs:
        blocks.append(
            f"### Model: {run.model} — Run {run.run_index + 1}\n\n"
            f"{run.narrative}"
        )
    narratives_block = "\n\n---\n\n".join(blocks)

    prompt = CONCORDANCE_PROMPT.format(
        n_narratives=len(narrative_runs),
        n_models=n_models,
        n_runs=n_runs,
        narratives_block=narratives_block,
    )

    endpoint = _resolve_endpoint(concordance_model)
    if not endpoint.is_available():
        return (
            f"[Concordance analysis unavailable — {concordance_model} endpoint not reachable]",
            concordance_model,
        )

    print(f"  Concordance model: {concordance_model} via {endpoint.name}")
    start = time.time()

    if isinstance(endpoint, AnthropicEndpoint):
        # Concordance needs more tokens than a single narrative (15 narratives → long output)
        old_max = endpoint.max_tokens
        endpoint.max_tokens = 16384
        text = endpoint.generate(prompt, system=CONCORDANCE_SYSTEM, temperature=0.3)
        endpoint.max_tokens = old_max
    elif isinstance(endpoint, OllamaEndpoint):
        payload = {
            "model": endpoint.model,
            "prompt": prompt,
            "system": CONCORDANCE_SYSTEM,
            "stream": False,
            "options": {
                "temperature": 0.3,
                "num_predict": 4096,
            },
        }
        try:
            r = requests.post(
                f"{endpoint.url}/api/generate",
                json=payload,
                timeout=600,
            )
            r.raise_for_status()
            text = r.json().get("response", "")
        except requests.RequestException as e:
            text = f"[ERROR] Concordance generation failed: {e}"
    else:
        text = "[ERROR] Unknown endpoint type for concordance"

    elapsed = time.time() - start
    print(f"  Concordance analysis complete in {elapsed:.1f}s")
    return text, concordance_model


# ---------------------------------------------------------------------------
# Word document export
# ---------------------------------------------------------------------------

def _render_narrative(doc, text: str) -> None:
    """Parse LLM narrative (markdown-ish) into Word document elements."""
    from build_docx import (
        add_heading as _h, add_para as _p, add_bullet as _b,
        add_bold_lead as _bl,
    )

    if not text:
        _p(doc, "No LLM narrative available.", italic=True)
        return

    in_code_block = False
    for line in text.split("\n"):
        stripped = line.strip()

        # Toggle code blocks — skip their contents
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue

        if not stripped:
            continue

        # Headings (check longer prefixes first)
        if stripped.startswith("### "):
            _h(doc, stripped[4:], level=3)
        elif stripped.startswith("## "):
            _h(doc, stripped[3:], level=2)
        elif stripped.startswith("# "):
            # Skip the top-level title we already rendered
            if "Dose-Response Interpretation" in stripped:
                continue
            _h(doc, stripped[2:], level=2)
        # Bullets
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _b(doc, stripped[2:])
        # Bold-lead patterns: **Key:** rest
        elif stripped.startswith("**") and "**" in stripped[2:]:
            end = stripped.index("**", 2)
            bold_part = stripped[2:end]
            rest = stripped[end + 2:].lstrip(": ")
            if rest:
                _bl(doc, bold_part + ": ", rest)
            else:
                _p(doc, bold_part, bold=True)
        else:
            _p(doc, stripped)


def build_interpretation_docx(result: AnalysisResult, output_path: str) -> None:
    """Render analysis results into a formatted Word document."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from build_docx import (
        add_heading, add_para, add_bold_lead, add_bullet, add_table,
    )

    doc = Document()

    # --- Title Page ---
    title = doc.add_heading("Dose-Response Interpretation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"{len(result.responsive_genes)} genes | "
        f"BMD range {result.bmd_min:.3g}\u2013{result.bmd_max:.3g} | "
        f"FDR < {result.fdr_cutoff}"
    )
    run.font.size = Pt(12)
    run.italic = True
    run.font.name = "Calibri"

    # --- Summary ---
    add_heading(doc, "Summary", level=1)
    add_bold_lead(doc, "Gene count: ", str(len(result.responsive_genes)))
    add_bold_lead(doc, "BMD range: ", f"{result.bmd_min:.3g} \u2013 {result.bmd_max:.3g}")
    if result.n_up or result.n_down:
        add_bold_lead(doc, "Direction: ",
                      f"{result.n_up} up-regulated, {result.n_down} down-regulated")
    add_bold_lead(doc, "Enriched pathways: ", str(len(result.pw_enriched)))
    add_bold_lead(doc, "Enriched GO terms: ", str(len(result.go_enriched)))

    # --- LLM Narrative ---
    if not result.narrative_runs:
        add_heading(doc, "Interpretation", level=1)
        _render_narrative(doc, result.llm_narrative)

    # --- Individual Narratives (multi-model mode) ---
    if result.narrative_runs:
        add_heading(doc, "Individual Narratives", level=1)
        # Group runs by model
        from collections import OrderedDict
        model_runs: dict[str, list[NarrativeRun]] = OrderedDict()
        for run in result.narrative_runs:
            model_runs.setdefault(run.model, []).append(run)

        for model, runs in model_runs.items():
            add_heading(doc, model, level=2)
            for run in sorted(runs, key=lambda r: r.run_index):
                add_heading(doc, f"Run {run.run_index + 1} ({run.elapsed_seconds:.1f}s)", level=3)
                _render_narrative(doc, run.narrative)

        # --- Generation Summary Table ---
        add_heading(doc, "Generation Summary", level=1)
        summary_rows = []
        for run in result.narrative_runs:
            summary_rows.append([
                run.model,
                str(run.run_index + 1),
                run.endpoint_name,
                f"{run.elapsed_seconds:.1f}s",
            ])
        add_table(doc,
                  ["Model", "Run", "Endpoint", "Time"],
                  summary_rows,
                  col_widths=[5, 1.5, 4, 2.5])

    # --- Concordance Analysis ---
    if result.concordance_text:
        add_heading(doc, "Concordance Analysis", level=1)
        add_para(doc, f"Analyzed by {result.concordance_model}", italic=True)
        _render_narrative(doc, result.concordance_text)

    # --- Pathway Enrichment Table ---
    if result.pw_enriched:
        add_heading(doc, "Pathway Enrichment", level=1)
        add_para(doc,
                 f"Significantly enriched pathways (FDR < {result.fdr_cutoff}), "
                 "ranked by p-value.",
                 italic=True)
        pw_rows = []
        for i, pw in enumerate(result.pw_enriched[:20], 1):
            pw_rows.append([
                str(i),
                pw["pathway_name"],
                f"{pw['pvalue']:.2e}",
                f"{pw['fdr']:.2e}",
                f"{pw['overlap_count']}/{pw['pathway_size']}",
                ", ".join(pw["overlap_genes"]),
            ])
        add_table(doc,
                  ["Rank", "Pathway", "p-value", "FDR", "Overlap", "Genes"],
                  pw_rows,
                  col_widths=[1.2, 5, 2, 2, 1.8, 5])

    # --- GO Term Enrichment Table ---
    if result.go_enriched:
        add_heading(doc, "GO Term Enrichment", level=1)
        add_para(doc,
                 f"Significantly enriched GO terms (FDR < {result.fdr_cutoff}), "
                 "ranked by p-value.",
                 italic=True)
        go_rows = []
        for i, gt in enumerate(result.go_enriched[:20], 1):
            go_rows.append([
                str(i),
                f"{gt['go_term']} [{gt['go_id']}]",
                f"{gt['pvalue']:.2e}",
                f"{gt['fdr']:.2e}",
                f"{gt['overlap_count']}/{gt['term_size']}",
                ", ".join(gt["overlap_genes"]),
            ])
        add_table(doc,
                  ["Rank", "GO Term", "p-value", "FDR", "Overlap", "Genes"],
                  go_rows,
                  col_widths=[1.2, 5, 2, 2, 1.8, 5])

    # --- BMD-Ordered Response Table ---
    if result.bmd_ordered:
        add_heading(doc, "BMD-Ordered Response", level=1)
        add_para(doc,
                 "Pathways ordered by median benchmark dose "
                 "(lowest = most sensitive).",
                 italic=True)
        bmd_rows = []
        for entry in result.bmd_ordered:
            bmd_rows.append([
                f"{entry['min_bmd']:.3g}\u2013{entry['mean_bmd']:.3g}",
                entry["pathway"],
                str(entry["median_bmd"]),
                ", ".join(entry["genes"]),
                entry["direction"],
            ])
        add_table(doc,
                  ["BMD Range", "Pathway", "Median BMD", "Genes", "Direction"],
                  bmd_rows,
                  col_widths=[2.5, 5, 2.5, 5, 2])

    # --- Organ Signature Table ---
    if result.organ_sig:
        add_heading(doc, "Organ Signature", level=1)
        add_para(doc,
                 "Organs ranked by enrichment score relative to "
                 "knowledge base background.",
                 italic=True)
        organ_rows = []
        for organ, info in result.organ_sig.items():
            genes_str = ", ".join(info["genes"][:10])
            if len(info["genes"]) > 10:
                genes_str += f" (+{len(info['genes']) - 10} more)"
            organ_rows.append([
                organ.title(),
                f"{info['score']}x",
                str(info["count"]),
                genes_str,
            ])
        add_table(doc,
                  ["Organ", "Enrichment", "Gene Count", "Genes"],
                  organ_rows,
                  col_widths=[3, 2.5, 2.5, 9])

    # --- Literature Context ---
    if result.gene_literature:
        add_heading(doc, "Literature Context", level=1)
        for gl in result.gene_literature:
            ev_label = (f"{gl['evidence']} gene"
                        if gl['evidence'] else "not in KB")
            add_heading(doc,
                        f"{gl['gene']} ({gl['n_papers']} papers, {ev_label})",
                        level=2)
            if gl['organs']:
                add_bold_lead(doc, "Organs: ", ", ".join(gl['organs']))
            if gl['claims']:
                for c in gl['claims']:
                    claim_text = c["claim"][:200]
                    year = c.get("year", "")
                    title_short = (c.get("paper_title") or "")[:60]
                    add_bullet(doc,
                               f" ({title_short}, {year})",
                               bold_prefix=f"\"{claim_text}\"")

    # --- Top Papers Table ---
    if result.top_papers:
        add_heading(doc, "Top Multi-Gene Papers", level=1)
        add_para(doc,
                 "Papers mentioning 2+ responsive genes, "
                 "ranked by citation count.",
                 italic=True)
        paper_rows = []
        for paper in result.top_papers:
            paper_rows.append([
                paper["title"] or "",
                str(paper["year"] or ""),
                str(paper["citation_count"]),
                ", ".join(paper["genes"]),
            ])
        add_table(doc,
                  ["Title", "Year", "Citations", "Genes"],
                  paper_rows,
                  col_widths=[7, 1.5, 2, 6.5])

    # Save
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Word document saved to {output_path}")


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def load_dose_response(csv_path: str) -> pd.DataFrame:
    """Load and validate dose-response CSV."""
    df = pd.read_csv(csv_path)

    # Normalize column names
    df.columns = [c.strip().lower() for c in df.columns]

    # Validate required columns
    if "gene_symbol" not in df.columns:
        raise ValueError("CSV must have a 'gene_symbol' column")
    if "bmd" not in df.columns:
        raise ValueError("CSV must have a 'bmd' column")

    # Normalize gene symbols
    df["gene_symbol"] = df["gene_symbol"].apply(normalize_gene)
    df = df[df["gene_symbol"] != ""].reset_index(drop=True)

    # Ensure BMD is numeric
    df["bmd"] = pd.to_numeric(df["bmd"], errors="coerce")
    df = df.dropna(subset=["bmd"]).reset_index(drop=True)

    return df


def interpret(
    csv_path: str,
    db_path: str = "bmdx.duckdb",
    output_path: str = "output/interpretation.md",
    docx_path: str | None = None,
    fdr_cutoff: float = 0.05,
    models: list[str] | None = None,
    n_runs: int = 3,
    concordance_model: str = "claude-sonnet-4-6",
) -> str:
    """Full interpretation pipeline."""
    print(f"Loading dose-response data from {csv_path}...")
    df = load_dose_response(csv_path)
    print(f"  {len(df)} genes loaded (BMD range: {df['bmd'].min():.3g} - {df['bmd'].max():.3g})")

    print(f"Connecting to knowledge base: {db_path}")
    kb = ToxKBQuerier(db_path)

    # Check how many genes are in the KB
    in_kb = sum(1 for g in df["gene_symbol"] if kb.gene_evidence(g)["evidence"] is not None)
    print(f"  {in_kb}/{len(df)} genes found in KB")

    print("Building interpretation context...")
    result = analyze(df, kb, fdr_cutoff)

    # --- Multi-model mode ---
    if models:
        print(f"Multi-model mode: {len(models)} models x {n_runs} runs = {len(models) * n_runs} narratives")
        result.narrative_runs = generate_all_narratives(
            context=result.context_text,
            models=models,
            n_runs=n_runs,
            base_docx=docx_path,
            base_md=output_path,
            result=result,
        )
        print(f"  {len(result.narrative_runs)} narratives generated.")

        # Run concordance analysis
        if result.narrative_runs:
            print("Running concordance analysis...")
            result.concordance_text, result.concordance_model = run_concordance_analysis(
                result.narrative_runs,
                concordance_model=concordance_model,
            )

        # Write combined "all" outputs
        md_path = Path(output_path)
        all_md = md_path.parent / f"{md_path.stem}_all{md_path.suffix}"
        lines = ["# All Narratives\n"]
        for run in result.narrative_runs:
            lines.append(f"## {run.model} — Run {run.run_index + 1} ({run.elapsed_seconds:.1f}s)\n")
            lines.append(run.narrative)
            lines.append("\n---\n")
        # Insert concordance before appendix
        if result.concordance_text:
            lines.append(f"\n# Concordance Analysis\n\n*Analyzed by {result.concordance_model}*\n")
            lines.append(result.concordance_text)
            lines.append("\n---\n")
        lines.append("\n# Appendix: Structured Analysis\n\n```\n" + result.context_text + "\n```\n")
        all_md_text = "\n".join(lines)
        all_md.parent.mkdir(parents=True, exist_ok=True)
        all_md.write_text(all_md_text)
        print(f"Combined markdown saved to {all_md}")

        # Write separate concordance output files
        if result.concordance_text:
            conc_md = md_path.parent / f"{md_path.stem}_concordance{md_path.suffix}"
            conc_md_text = (
                f"# Concordance Analysis\n\n"
                f"*Analyzed by {result.concordance_model}*\n\n"
                f"{result.concordance_text}\n"
            )
            conc_md.write_text(conc_md_text)
            print(f"Concordance markdown saved to {conc_md}")

            if docx_path:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from build_docx import add_heading, add_para

                conc_doc = Document()
                title = conc_doc.add_heading("Concordance Analysis", level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_para(conc_doc, f"Analyzed by {result.concordance_model}", italic=True)
                _render_narrative(conc_doc, result.concordance_text)
                docx_p = Path(docx_path)
                conc_docx = docx_p.parent / f"{docx_p.stem}_concordance{docx_p.suffix}"
                conc_docx.parent.mkdir(parents=True, exist_ok=True)
                conc_doc.save(str(conc_docx))
                print(f"Concordance docx saved to {conc_docx}")

        if docx_path:
            docx_p = Path(docx_path)
            all_docx = docx_p.parent / f"{docx_p.stem}_all{docx_p.suffix}"
            build_interpretation_docx(result, str(all_docx))

        kb.close()
        return all_md_text

    # --- Single-model mode (backward compatible) ---
    print("Running LLM synthesis...")
    result.llm_narrative = synthesize_interpretation(result.context_text)

    # Build Word document if requested
    if docx_path:
        build_interpretation_docx(result, docx_path)

    # Write markdown output
    report = result.llm_narrative
    report += "\n\n---\n\n# Appendix: Structured Analysis\n\n```\n" + result.context_text + "\n```\n"

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(report)
    print(f"Markdown saved to {output_path}")

    kb.close()
    return report


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def concordance_only(
    all_md_path: str = "output/interpretation_all.md",
    concordance_model: str = "claude-sonnet-4-6",
    output_dir: str | None = None,
) -> str:
    """
    Run concordance analysis on existing narratives without re-generating them.

    Parses the combined all-narratives markdown to reconstruct NarrativeRun
    objects, then feeds them to run_concordance_analysis().
    """
    import re as _re

    all_md = Path(all_md_path)
    if not all_md.exists():
        print(f"Error: {all_md_path} not found")
        print("Run the full pipeline first to generate narratives.")
        return ""

    text = all_md.read_text()

    # Parse "## model_name — Run N (Xs)" headers and their narrative bodies
    # Pattern: ## <model> — Run <N> (<elapsed>s)
    header_pattern = _re.compile(
        r'^## (.+?) — Run (\d+) \((\d+(?:\.\d+)?)s\)\s*$',
        _re.MULTILINE,
    )
    matches = list(header_pattern.finditer(text))
    if not matches:
        print(f"Error: no narrative headers found in {all_md_path}")
        print("Expected format: '## model_name — Run N (Xs)'")
        return ""

    runs: list[NarrativeRun] = []
    for i, m in enumerate(matches):
        model = m.group(1)
        run_index = int(m.group(2)) - 1  # 0-indexed
        elapsed = float(m.group(3))

        # Narrative body is everything between this header and the next (or end)
        start = m.end()
        if i + 1 < len(matches):
            end = matches[i + 1].start()
        else:
            # Stop at concordance section or appendix if present
            end_marker = _re.search(r'^# (?:Concordance|Appendix)', text[start:], _re.MULTILINE)
            end = start + end_marker.start() if end_marker else len(text)

        narrative = text[start:end].strip().rstrip("-").strip()
        runs.append(NarrativeRun(
            model=model,
            endpoint_name="(loaded from file)",
            run_index=run_index,
            narrative=narrative,
            elapsed_seconds=elapsed,
        ))

    print(f"Loaded {len(runs)} narratives from {all_md_path}")
    models_found = sorted(set(r.model for r in runs))
    for model in models_found:
        n = sum(1 for r in runs if r.model == model)
        print(f"  {model}: {n} runs")

    # Run concordance
    print(f"\nRunning concordance analysis with {concordance_model}...")
    conc_text, conc_model = run_concordance_analysis(runs, concordance_model)

    if not conc_text or conc_text.startswith("["):
        print(f"Concordance failed: {conc_text}")
        return conc_text

    # Write output
    out_dir = Path(output_dir) if output_dir else all_md.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    conc_md_path = out_dir / "interpretation_concordance.md"
    conc_md_text = (
        f"# Concordance Analysis\n\n"
        f"*Analyzed by {conc_model}*\n\n"
        f"{conc_text}\n"
    )
    conc_md_path.write_text(conc_md_text)
    print(f"\nConcordance markdown saved to {conc_md_path}")

    # Also write docx if python-docx is available
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from build_docx import add_heading, add_para

        conc_doc = Document()
        title = conc_doc.add_heading("Concordance Analysis", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_para(conc_doc, f"Analyzed by {conc_model}", italic=True)
        _render_narrative(conc_doc, conc_text)
        conc_docx_path = out_dir / "interpretation_concordance.docx"
        conc_doc.save(str(conc_docx_path))
        print(f"Concordance docx saved to {conc_docx_path}")
    except ImportError:
        pass

    return conc_text


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Interpret dose-response BMD results against the bmdx knowledge base",
    )
    sub = parser.add_subparsers(dest="command")

    # Default: full pipeline (no subcommand, uses --csv or positional)
    parser.add_argument("csv", nargs="?", default=None, help="Path to gene-level BMD results CSV")
    parser.add_argument("--csv", dest="csv_flag", default=None, help="Path to gene-level BMD results CSV (alternative to positional)")
    parser.add_argument("--db", default="bmdx.duckdb", help="Path to bmdx.duckdb")
    parser.add_argument("--output", default="output/interpretation.md", help="Output markdown path")
    parser.add_argument("--docx", default=None, help="Output Word document path (.docx)")
    parser.add_argument("--fdr", type=float, default=0.05, help="FDR cutoff for enrichment")
    parser.add_argument(
        "--models", nargs="+", default=None,
        help="Models for multi-narrative generation (e.g. qwen2.5:14b gemma2:9b claude-sonnet-4-6)",
    )
    parser.add_argument(
        "--runs", type=int, default=3,
        help="Number of narrative runs per model (default: 3)",
    )
    parser.add_argument(
        "--concordance-model", default="claude-sonnet-4-6",
        help="Model for concordance analysis (default: claude-sonnet-4-6)",
    )

    # Subcommand: concordance-only
    conc_parser = sub.add_parser(
        "concordance",
        help="Run concordance analysis on existing narratives",
    )
    conc_parser.add_argument(
        "--input", default="output/interpretation_all.md",
        help="Path to combined all-narratives markdown (default: output/interpretation_all.md)",
    )
    conc_parser.add_argument(
        "--concordance-model", default="claude-sonnet-4-6",
        help="Model for concordance analysis (default: claude-sonnet-4-6)",
    )
    conc_parser.add_argument(
        "--output-dir", default=None,
        help="Output directory for concordance files (default: same as input)",
    )

    args = parser.parse_args()

    if args.command == "concordance":
        concordance_only(
            all_md_path=args.input,
            concordance_model=args.concordance_model,
            output_dir=args.output_dir,
        )
    else:
        csv_path = args.csv_flag or args.csv
        if not csv_path:
            parser.error("csv argument is required for the full pipeline")
        interpret(
            csv_path,
            db_path=args.db,
            output_path=args.output,
            docx_path=args.docx,
            fdr_cutoff=args.fdr,
            models=args.models,
            n_runs=args.runs,
            concordance_model=args.concordance_model,
        )
